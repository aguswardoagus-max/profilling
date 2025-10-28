# Mapping Profiling Backend
# Advanced Network Mapping for Intelligence Analysis
# Supports: Terrorism, Social-Cultural, Political, Defense, Economic, Ideology Networks

import os
import json
import logging
from datetime import datetime
from typing import Dict, List, Any, Optional
from pathlib import Path
import google.generativeai as genai
from flask import Blueprint, request, jsonify, render_template
import networkx as nx

# Configure matplotlib to use non-interactive backend to avoid tkinter issues
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend
# Disable interactive mode and warnings
import warnings
warnings.filterwarnings('ignore', category=UserWarning, module='matplotlib')
import matplotlib.pyplot as plt
plt.ioff()  # Turn off interactive mode
import matplotlib.patches as mpatches
from io import BytesIO
import base64
import tempfile
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Create blueprint
mapping_bp = Blueprint('mapping', __name__)

# Gemini AI Configuration
GEMINI_API_KEY = 'AIzaSyDxxD5ZYEsW1Zeo4RiPcM_zEf2bvG8WF1A'
genai.configure(api_key=GEMINI_API_KEY)

try:
    mapping_model = genai.GenerativeModel('gemini-1.5-pro')
except:
    try:
        mapping_model = genai.GenerativeModel('gemini-1.5-flash')
    except:
        mapping_model = None
        logger.warning("Tidak bisa menginisialisasi model Gemini AI untuk mapping")

# Network Categories
NETWORK_CATEGORIES = {
    'terrorism': {
        'name': 'Jaringan Teroris',
        'color': '#FF4444',
        'description': 'Pemetaan jaringan teroris dan aktivitas ekstremis',
        'keywords': ['teroris', 'ekstremis', 'radikal', 'bom', 'serangan', 'jihad', 'ISIS', 'Al-Qaeda']
    },
    'social_cultural': {
        'name': 'Sosial Budaya',
        'color': '#44AA44',
        'description': 'Pemetaan jaringan sosial dan budaya',
        'keywords': ['sosial', 'budaya', 'masyarakat', 'komunitas', 'organisasi', 'kelompok']
    },
    'political': {
        'name': 'Politik',
        'color': '#4444FF',
        'description': 'Pemetaan jaringan politik dan kekuasaan',
        'keywords': ['politik', 'partai', 'pemerintah', 'kekuasaan', 'kepemimpinan', 'demokrasi']
    },
    'defense': {
        'name': 'Hankam',
        'color': '#FF8800',
        'description': 'Pemetaan jaringan pertahanan dan keamanan',
        'keywords': ['hankam', 'militer', 'polisi', 'keamanan', 'pertahanan', 'intelijen']
    },
    'economic': {
        'name': 'Ekonomi',
        'color': '#8844FF',
        'description': 'Pemetaan jaringan ekonomi dan bisnis',
        'keywords': ['ekonomi', 'bisnis', 'perdagangan', 'keuangan', 'investasi', 'perusahaan']
    },
    'ideology': {
        'name': 'Ideologi',
        'color': '#FF44AA',
        'description': 'Pemetaan jaringan ideologi dan keyakinan',
        'keywords': ['ideologi', 'keyakinan', 'agama', 'filosofi', 'pandangan', 'prinsip']
    }
}

class NetworkMapper:
    """Advanced Network Mapping Class with AI Integration"""
    
    def __init__(self):
        self.graph = nx.Graph()
        self.categories = NETWORK_CATEGORIES
        self.ai_model = mapping_model
        
    def analyze_network_with_ai(self, data: Dict[str, Any], category: str) -> Dict[str, Any]:
        """Analyze network data using AI Gemini"""
        try:
            if not self.ai_model:
                return self._fallback_analysis(data, category)
            
            category_info = self.categories.get(category, {})
            category_name = category_info.get('name', category)
            
            # Prepare data for AI analysis
            persons = data.get('persons', [])
            relationships = data.get('relationships', [])
            
            # Create comprehensive prompt for AI analysis
            prompt = f"""
Sebagai analis intelijen senior dengan pengalaman 20 tahun dalam pemetaan jaringan {category_name}, analisis data berikut untuk mengidentifikasi pola, koneksi, dan ancaman potensial.

KATEGORI ANALISIS: {category_name}
DESKRIPSI: {category_info.get('description', '')}

DATA PERSONS ({len(persons)} orang):
{self._format_persons_data(persons)}

DATA RELATIONSHIPS ({len(relationships)} hubungan):
{self._format_relationships_data(relationships)}

INSTRUKSI ANALISIS:
1. Identifikasi pola koneksi yang mencurigakan atau penting
2. Tentukan tingkat risiko dan prioritas setiap node
3. Analisis struktur jaringan dan hierarki
4. Identifikasi key players dan influencers
5. Deteksi potential threats dan vulnerabilities
6. Berikan rekomendasi strategis

FORMAT OUTPUT (JSON):
{{
    "network_analysis": {{
        "total_nodes": number,
        "total_connections": number,
        "network_density": number,
        "key_players": [
            {{"name": "string", "importance_score": number, "role": "string"}}
        ],
        "threat_level": "LOW/MEDIUM/HIGH/CRITICAL",
        "vulnerabilities": ["string"],
        "recommendations": ["string"]
    }},
    "categorized_nodes": {{
        "high_priority": ["string"],
        "medium_priority": ["string"],
        "low_priority": ["string"]
    }},
    "connection_patterns": [
        {{"pattern_type": "string", "description": "string", "risk_level": "string"}}
    ],
    "strategic_insights": [
        {{"insight": "string", "confidence": number, "action_required": boolean}}
    ]
}}
"""
            
            # Get AI response
            response = self.ai_model.generate_content(prompt)
            ai_analysis = json.loads(response.text)
            
            return ai_analysis
            
        except Exception as e:
            logger.error(f"Error in AI network analysis: {e}")
            return self._fallback_analysis(data, category)
    
    def _fallback_analysis(self, data: Dict[str, Any], category: str) -> Dict[str, Any]:
        """Fallback analysis when AI is not available"""
        persons = data.get('persons', [])
        relationships = data.get('relationships', [])
        
        return {
            "network_analysis": {
                "total_nodes": len(persons),
                "total_connections": len(relationships),
                "network_density": len(relationships) / max(len(persons) * (len(persons) - 1) / 2, 1),
                "key_players": [{"name": p.get('name', 'Unknown'), "importance_score": 0.5, "role": "Unknown"} for p in persons[:5]],
                "threat_level": "MEDIUM",
                "vulnerabilities": ["Data incomplete"],
                "recommendations": ["Gather more data for better analysis"]
            },
            "categorized_nodes": {
                "high_priority": [p.get('name', 'Unknown') for p in persons[:3]],
                "medium_priority": [p.get('name', 'Unknown') for p in persons[3:6]],
                "low_priority": [p.get('name', 'Unknown') for p in persons[6:]]
            },
            "connection_patterns": [
                {"pattern_type": "Basic", "description": "Standard connections", "risk_level": "LOW"}
            ],
            "strategic_insights": [
                {"insight": "Requires further investigation", "confidence": 0.3, "action_required": True}
            ]
        }
    
    def _format_persons_data(self, persons: List[Dict]) -> str:
        """Format persons data for AI analysis"""
        formatted = []
        for i, person in enumerate(persons, 1):
            formatted.append(f"{i}. {person.get('name', 'Unknown')} - {person.get('role', 'Unknown')} - {person.get('location', 'Unknown')}")
        return "\n".join(formatted)
    
    def _format_relationships_data(self, relationships: List[Dict]) -> str:
        """Format relationships data for AI analysis"""
        formatted = []
        for i, rel in enumerate(relationships, 1):
            formatted.append(f"{i}. {rel.get('from', 'Unknown')} -> {rel.get('to', 'Unknown')} ({rel.get('type', 'Unknown')})")
        return "\n".join(formatted)
    
    def generate_network_visualization(self, data: Dict[str, Any], category: str, ai_analysis: Dict[str, Any]) -> str:
        """Generate network visualization as base64 image"""
        try:
            # Create new graph
            G = nx.Graph()
            
            # Add nodes
            persons = data.get('persons', [])
            for person in persons:
                G.add_node(person.get('name', 'Unknown'), 
                          role=person.get('role', 'Unknown'),
                          location=person.get('location', 'Unknown'))
            
            # Add edges
            relationships = data.get('relationships', [])
            for rel in relationships:
                G.add_edge(rel.get('from', 'Unknown'), 
                          rel.get('to', 'Unknown'),
                          relationship_type=rel.get('type', 'Unknown'))
            
            # Create visualization
            plt.figure(figsize=(12, 8))
            pos = nx.spring_layout(G, k=3, iterations=50)
            
            # Get category color
            category_info = self.categories.get(category, {})
            base_color = category_info.get('color', '#444444')
            
            # Color nodes based on priority
            high_priority = ai_analysis.get('categorized_nodes', {}).get('high_priority', [])
            medium_priority = ai_analysis.get('categorized_nodes', {}).get('medium_priority', [])
            
            node_colors = []
            for node in G.nodes():
                if node in high_priority:
                    node_colors.append('#FF4444')  # Red for high priority
                elif node in medium_priority:
                    node_colors.append('#FFAA44')  # Orange for medium priority
                else:
                    node_colors.append(base_color)  # Base color for low priority
            
            # Draw network
            nx.draw(G, pos, 
                   node_color=node_colors,
                   node_size=1000,
                   font_size=8,
                   font_weight='bold',
                   with_labels=True,
                   edge_color='#888888',
                   width=2)
            
            # Add title
            plt.title(f"Network Mapping - {category_info.get('name', category)}", 
                     fontsize=16, fontweight='bold', pad=20)
            
            # Add legend
            legend_elements = [
                mpatches.Patch(color='#FF4444', label='High Priority'),
                mpatches.Patch(color='#FFAA44', label='Medium Priority'),
                mpatches.Patch(color=base_color, label='Low Priority')
            ]
            plt.legend(handles=legend_elements, loc='upper right')
            
            # Save to base64
            buffer = BytesIO()
            plt.savefig(buffer, format='png', dpi=300, bbox_inches='tight')
            buffer.seek(0)
            image_base64 = base64.b64encode(buffer.getvalue()).decode()
            plt.close()
            
            return f"data:image/png;base64,{image_base64}"
            
        except Exception as e:
            logger.error(f"Error generating network visualization: {e}")
            # Ensure plot is closed even on error
            try:
                plt.close()
            except:
                pass
            return None
    
    def generate_ai_report(self, data: Dict[str, Any], category: str, ai_analysis: Dict[str, Any]) -> str:
        """Generate AI-powered comprehensive report"""
        try:
            if not self.ai_model:
                return self._generate_fallback_report(data, category, ai_analysis)
            
            category_info = self.categories.get(category, {})
            category_name = category_info.get('name', category)
            
            # Prepare comprehensive data for AI report generation
            persons = data.get('persons', [])
            relationships = data.get('relationships', [])
            
            prompt = f"""
Sebagai analis intelijen senior, buatkan laporan komprehensif tentang pemetaan jaringan {category_name}.

KATEGORI: {category_name}
TANGGAL ANALISIS: {datetime.now().strftime('%d %B %Y')}

DATA YANG DIANALISIS:
- Jumlah Person: {len(persons)}
- Jumlah Hubungan: {len(relationships)}

HASIL ANALISIS AI:
{json.dumps(ai_analysis, indent=2)}

BUATKAN LAPORAN LENGKAP DENGAN STRUKTUR:

1. EXECUTIVE SUMMARY
   - Ringkasan eksekutif tentang temuan utama
   - Tingkat ancaman dan prioritas
   - Rekomendasi strategis utama

2. ANALISIS JARINGAN
   - Struktur jaringan dan hierarki
   - Key players dan peran mereka
   - Pola koneksi yang mencurigakan

3. ASSESSMENT RISIKO
   - Tingkat ancaman keseluruhan
   - Vulnerabilities yang teridentifikasi
   - Potensi dampak operasional

4. REKOMENDASI STRATEGIS
   - Tindakan prioritas
   - Monitoring dan surveillance
   - Koordinasi dengan instansi terkait

5. KESIMPULAN
   - Kesimpulan utama
   - Next steps
   - Timeline implementasi

Gunakan bahasa profesional dan teknis yang sesuai untuk laporan intelijen.
"""
            
            response = self.ai_model.generate_content(prompt)
            return response.text
            
        except Exception as e:
            logger.error(f"Error generating AI report: {e}")
            return self._generate_fallback_report(data, category, ai_analysis)
    
    def _generate_fallback_report(self, data: Dict[str, Any], category: str, ai_analysis: Dict[str, Any]) -> str:
        """Generate fallback report when AI is not available"""
        category_info = self.categories.get(category, {})
        category_name = category_info.get('name', category)
        
        return f"""
LAPORAN PEMETAAN JARINGAN {category_name.upper()}
Tanggal: {datetime.now().strftime('%d %B %Y')}

1. EXECUTIVE SUMMARY
   - Analisis jaringan {category_name} telah dilakukan
   - Total {len(data.get('persons', []))} person dan {len(data.get('relationships', []))} hubungan dianalisis
   - Tingkat ancaman: {ai_analysis.get('network_analysis', {}).get('threat_level', 'UNKNOWN')}

2. ANALISIS JARINGAN
   - Struktur jaringan telah dipetakan
   - Key players teridentifikasi: {len(ai_analysis.get('categorized_nodes', {}).get('high_priority', []))} orang
   - Density jaringan: {ai_analysis.get('network_analysis', {}).get('network_density', 0):.2f}

3. ASSESSMENT RISIKO
   - Vulnerabilities: {', '.join(ai_analysis.get('network_analysis', {}).get('vulnerabilities', ['Tidak teridentifikasi']))}
   - Rekomendasi: {', '.join(ai_analysis.get('network_analysis', {}).get('recommendations', ['Perlu investigasi lebih lanjut']))}

4. REKOMENDASI STRATEGIS
   - Monitoring berkelanjutan diperlukan
   - Koordinasi dengan instansi terkait
   - Update data secara berkala

5. KESIMPULAN
   - Analisis awal telah selesai
   - Perlu investigasi lebih mendalam
   - Monitoring berkelanjutan disarankan
"""

# Global mapper instance
network_mapper = NetworkMapper()

# API Endpoints
@mapping_bp.route('/mapping')
def mapping_dashboard():
    """Mapping dashboard page"""
    return render_template('mapping.html', categories=NETWORK_CATEGORIES)

@mapping_bp.route('/api/mapping/categories', methods=['GET'])
def get_categories():
    """Get available network categories"""
    return jsonify({
        'success': True,
        'categories': NETWORK_CATEGORIES
    })

@mapping_bp.route('/api/mapping/analyze', methods=['POST'])
def analyze_network():
    """Analyze network with AI"""
    try:
        data = request.get_json()
        
        # Validate required fields
        if not data.get('persons') or not data.get('category'):
            return jsonify({'error': 'Persons data and category are required'}), 400
        
        # Analyze network with AI
        ai_analysis = network_mapper.analyze_network_with_ai(data, data['category'])
        
        # Generate network visualization
        visualization = network_mapper.generate_network_visualization(data, data['category'], ai_analysis)
        
        # Generate AI report
        ai_report = network_mapper.generate_ai_report(data, data['category'], ai_analysis)
        
        return jsonify({
            'success': True,
            'analysis': ai_analysis,
            'visualization': visualization,
            'report': ai_report,
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        logger.error(f"Error in network analysis: {e}")
        return jsonify({'error': f'Analysis failed: {str(e)}'}), 500

@mapping_bp.route('/api/mapping/export/pdf', methods=['POST'])
def export_pdf():
    """Export mapping analysis to PDF"""
    try:
        data = request.get_json()
        
        # Create PDF
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        
        # Get styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1
        )
        
        # Build PDF content
        story = []
        
        # Title
        story.append(Paragraph("LAPORAN PEMETAAN JARINGAN", title_style))
        story.append(Spacer(1, 12))
        
        # Add AI report content
        report_content = data.get('report', '')
        if report_content:
            # Split report into paragraphs
            paragraphs = report_content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para.strip(), styles['Normal']))
                    story.append(Spacer(1, 12))
        
        # Build PDF
        doc.build(story)
        
        # Get PDF content
        buffer.seek(0)
        pdf_content = buffer.getvalue()
        buffer.close()
        
        # Convert to base64
        pdf_base64 = base64.b64encode(pdf_content).decode()
        
        return jsonify({
            'success': True,
            'pdf_content': f"data:application/pdf;base64,{pdf_base64}",
            'filename': f"mapping_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        })
        
    except Exception as e:
        logger.error(f"Error exporting PDF: {e}")
        return jsonify({'error': f'PDF export failed: {str(e)}'}), 500

@mapping_bp.route('/api/mapping/export/word', methods=['POST'])
def export_word():
    """Export mapping analysis to Word document"""
    try:
        data = request.get_json()
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('LAPORAN PEMETAAN JARINGAN', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        doc.add_paragraph(f"Tanggal: {datetime.now().strftime('%d %B %Y')}")
        doc.add_paragraph("")
        
        # Add AI report content
        report_content = data.get('report', '')
        if report_content:
            # Split report into paragraphs
            paragraphs = report_content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Convert to base64
        word_base64 = base64.b64encode(buffer.getvalue()).decode()
        
        return jsonify({
            'success': True,
            'word_content': f"data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{word_base64}",
            'filename': f"mapping_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        })
        
    except Exception as e:
        logger.error(f"Error exporting Word: {e}")
        return jsonify({'error': f'Word export failed: {str(e)}'}), 500

@mapping_bp.route('/api/mapping/preview', methods=['POST'])
def preview_data():
    """Preview mapping data"""
    try:
        data = request.get_json()
        
        # Basic validation
        persons = data.get('persons', [])
        relationships = data.get('relationships', [])
        category = data.get('category', '')
        
        # Create preview data
        preview = {
            'category': category,
            'category_info': NETWORK_CATEGORIES.get(category, {}),
            'total_persons': len(persons),
            'total_relationships': len(relationships),
            'persons_preview': persons[:10],  # Show first 10 persons
            'relationships_preview': relationships[:10],  # Show first 10 relationships
            'has_more_persons': len(persons) > 10,
            'has_more_relationships': len(relationships) > 10
        }
        
        return jsonify({
            'success': True,
            'preview': preview
        })
        
    except Exception as e:
        logger.error(f"Error previewing data: {e}")
        return jsonify({'error': f'Preview failed: {str(e)}'}), 500

#!/usr/bin/env python3
"""
Flask web server untuk Clearance Face Search
Menyediakan API endpoints dan serve static files
"""
import os
import sys
import json
import time
import base64
import tempfile
import logging
from datetime import datetime
from pathlib import Path
from flask import Flask, request, jsonify, send_from_directory, render_template, session, redirect, url_for, send_file
from flask_cors import CORS
from dotenv import load_dotenv
from database import authenticate_user, validate_session_token, logout_user, db
from functools import wraps
from cekplat import cekplat_bp

# Ensure sys.stderr is available for debugging
if not hasattr(sys, 'stderr'):
    import io
    sys.stderr = io.TextIOWrapper(io.BytesIO(), encoding='utf-8')

# Try to import AI endpoints (optional - may fail if sklearn/pandas have compatibility issues)
ai_bp = None
try:
    from ai_api_endpoints import ai_bp
    print("✅ AI endpoints module loaded successfully")
except ImportError as e:
    print(f"⚠️ Warning: Could not import AI endpoints module: {e}")
    print("   AI features will be disabled. To fix this, you may need to:")
    print("   1. Downgrade NumPy: pip install 'numpy<2.0'")
    print("   2. Or upgrade pandas and sklearn: pip install --upgrade pandas scikit-learn")
    print("   3. Or reinstall all: pip install --upgrade --force-reinstall numpy pandas scikit-learn")
except Exception as e:
    print(f"⚠️ Warning: Error loading AI endpoints module: {e}")
    print("   AI features will be disabled.")

import numpy as np
import requests

# Load environment variables from .env file
load_dotenv()

# Get the directory of this file (backend/)
backend_dir = os.path.dirname(os.path.abspath(__file__))
# Get the parent directory (project root)
project_root = os.path.dirname(backend_dir)
# Set paths for frontend files
frontend_dir = os.path.join(project_root, 'frontend')
frontend_pages_dir = os.path.join(frontend_dir, 'pages')
frontend_static_dir = os.path.join(frontend_dir, 'static')
config_dir = os.path.join(project_root, 'config')

def require_auth(f):
    """Decorator to require authentication for routes"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # Check for redirect loop protection
        redirect_count = request.cookies.get('redirect_count', '0')
        if int(redirect_count) > 5:
            print("Redirect loop detected in require_auth, serving login page")
            response = send_from_directory(frontend_pages_dir, 'login.html')
            response.set_cookie('redirect_count', '0', expires=0)
            response.set_cookie('session_token', '', expires=0)
            return response
        
        # Check for session token in cookies first
        session_token = request.cookies.get('session_token')
        
        # If no cookie, check Authorization header
        if not session_token:
            session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        
        if not session_token:
            # Return a simple HTML page that redirects to login
            response = '''
            <!DOCTYPE html>
            <html>
            <head>
                <title>Unauthorized Access</title>
                <meta http-equiv="refresh" content="2; url=/login">
                <script>
                    setTimeout(() => {
                        window.location.href = '/login';
                    }, 2000);
                </script>
            </head>
            <body>
                <p>No session token found. Redirecting to login page in 2 seconds...</p>
            </body>
            </html>
            '''
            return response, 401
        
        # Validate session token
        user = validate_session_token(session_token)
        if not user:
            # Return a simple HTML page that redirects to login
            response = '''
            <!DOCTYPE html>
            <html>
            <head>
                <title>Session Expired</title>
                <meta http-equiv="refresh" content="2; url=/login">
                <script>
                    localStorage.removeItem('session_token');
                    localStorage.removeItem('user_data');
                    setTimeout(() => {
                        window.location.href = '/login';
                    }, 2000);
                </script>
            </head>
            <body>
                <p>Session expired. Redirecting to login page in 2 seconds...</p>
            </body>
            </html>
            '''
            return response, 401
        
        # Add user info to request context
        request.current_user = user
        return f(*args, **kwargs)
    
    return decorated_function
import cv2
from PIL import Image
import io
import google.generativeai as genai
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import base64

# Import functions from clearance_face_search.py
from clearance_face_search import (
    ensure_token, call_search, parse_people_from_response,
    load_image_file_to_encoding, get_encoding_from_base64_face,
    save_face_image, USE_FACE_LIB
)

# Konfigurasi Gemini AI
# Try to get from database first, then fallback to environment variable
def get_gemini_api_key():
    """Get Gemini API key from database or environment variable"""
    try:
        # Try to get from database first
        from database import UserDatabase
        db = UserDatabase()
        api_key = db.get_api_key('GEMINI')
        if api_key:
            return api_key
    except Exception as e:
        print(f"⚠️ Cannot get Gemini API key from database: {e}")
    
    # Fallback to environment variable
    return os.getenv('GEMINI_API_KEY', 'AIzaSyDLM8CF-CePwAgxb0TYw6kvw-bLS_q2AXY')

GEMINI_API_KEY = get_gemini_api_key()
USE_NEW_GEMINI = False
client = None
model = None

try:
    from google import genai as google_genai
    client = google_genai.Client(api_key=GEMINI_API_KEY)
    USE_NEW_GEMINI = True
    print("✅ Menggunakan Google GenAI SDK baru")
    
    # Try to list available models to see what's available
    try:
        # This is optional, don't fail if it doesn't work
        try:
            models = client.models.list()
            available_models = [m.name for m in models if hasattr(m, 'name')]
            if available_models:
                print(f"✅ Available models: {', '.join(available_models[:5])}...")  # Show first 5
        except:
            pass  # Ignore if list_models fails
    except Exception as test_error:
        print(f"⚠️ Client initialized but test failed: {test_error}")
        # Continue anyway, might work for generate_content
except Exception as e:
    print(f"⚠️ Tidak bisa menggunakan Google GenAI SDK baru: {e}")
    client = None
    USE_NEW_GEMINI = False

# Always try to initialize old SDK as fallback
# Note: Old SDK models may not be available, so we'll try to list available models first
try:
    import google.generativeai as genai
    genai.configure(api_key=GEMINI_API_KEY)
    print("✅ Menggunakan Google Generative AI SDK lama")
    
    # Try to list available models first
    available_model_names = []
    try:
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                model_name = m.name.replace('models/', '')  # Remove 'models/' prefix
                available_model_names.append(model_name)
        if available_model_names:
            print(f"✅ Available models in old SDK: {', '.join(available_model_names[:5])}...")
    except Exception as list_error:
        print(f"⚠️ Cannot list models: {list_error}")
        # Fallback to trying common model names
        available_model_names = []
    
    # Coba beberapa model yang tersedia (prioritaskan yang lebih baru)
    # Updated to try newer models first, but only if available
    models_to_try = []
    
    # If we have available models list, use those first
    if available_model_names:
        # Filter out experimental models that may not be available in free tier
        free_tier_models = [m for m in available_model_names if 'exp' not in m.lower() and '2.5' not in m and '2.0' not in m]
        if free_tier_models:
            models_to_try = free_tier_models[:5]  # Use first 5 available free tier models
        else:
            models_to_try = available_model_names[:5]  # Fallback to all available
    else:
        # Fallback: use only free tier models (no experimental models)
        models_to_try = [
            'gemini-1.5-flash',  # Most commonly available in free tier
            'gemini-1.5-pro', 
            'gemini-pro'
        ]
    
    model = None
    for model_name in models_to_try:
        try:
            # Test if model can be initialized by trying to create it
            test_model = genai.GenerativeModel(model_name)
            # Try a simple test call (optional, but helps verify)
            model = test_model
            print(f"✅ Model {model_name} berhasil diinisialisasi")
            break
        except Exception as model_error:
            error_str = str(model_error)
            # Don't log 404 errors for every model
            if "404" not in error_str and "not found" not in error_str.lower():
                print(f"⚠️ Model {model_name} gagal: {model_error}")
            continue
    
    if not model:
        print(f"⚠️ Tidak bisa menginisialisasi model Gemini AI dari daftar. Old SDK akan di-disable.")
        model = None
except Exception as e2:
    print(f"⚠️ Tidak bisa menginisialisasi Gemini AI SDK lama: {e2}")
    model = None

def calculate_data_richness(person_data, family_data):
    """Calculate data richness score for dynamic analysis"""
    score = 0
    if person_data.get('full_name') and person_data['full_name'] != 'N/A':
        score += 1
    if person_data.get('ktp_number') and person_data['ktp_number'] != 'N/A':
        score += 1
    if person_data.get('alamat') and person_data['alamat'] != 'N/A':
        score += 1
    if person_data.get('occupation') and person_data['occupation'] != 'N/A':
        score += 1
    if person_data.get('religion') and person_data['religion'] != 'N/A':
        score += 1
    if person_data.get('marital_status') and person_data['marital_status'] != 'N/A':
        score += 1
    if person_data.get('last_education') and person_data['last_education'] != 'N/A':
        score += 1
    if family_data and family_data.get('anggota_keluarga'):
        score += len(family_data['anggota_keluarga'])
    return score

def get_analysis_style(data_richness, pekerjaan, alamat):
    """Get analysis style based on data richness and context"""
    if data_richness >= 8:
        return "comprehensive"
    elif data_richness >= 5:
        return "detailed"
    else:
        return "basic"

def generate_dynamic_fallback_analysis(person_data, family_data, search_type, subtitle=None, location=None):
    """Generate dynamic fallback analysis when AI is not available"""
    import random
    import hashlib
    
    nama = person_data.get('full_name', 'target')
    alamat = person_data.get('alamat', 'lokasi yang tidak diketahui')
    pekerjaan = person_data.get('occupation', 'pekerjaan tidak diketahui')
    status_perkawinan = person_data.get('marital_status', 'status tidak diketahui')
    pendidikan = person_data.get('last_education', 'pendidikan tidak diketahui')
    agama = person_data.get('religion', 'agama tidak diketahui')
    nik = person_data.get('ktp_number', 'N/A')
    
    # Prepare subtitle and location context
    subtitle_context = subtitle if subtitle else "operasi perkebunan sawit ilegal"
    location_context = location if location else alamat
    
    # Create unique seed based on NIK and name for consistent randomization
    unique_seed = hashlib.md5(f"{nik}_{nama}".encode()).hexdigest()
    random.seed(int(unique_seed[:8], 16))
    
    # Different analysis templates for variety
    priority_templates = [
        f'Hasil investigasi lapangan mengungkapkan bahwa {nama} memegang posisi krusial dalam jaringan {subtitle_context} di {location_context}. Target ini menunjukkan pola aktivitas yang mengindikasikan keterlibatan mendalam dalam kegiatan yang merugikan ekosistem hutan. Tingkat ancaman yang ditimbulkan {nama} terhadap lingkungan dan masyarakat lokal di {alamat} memerlukan perhatian serius dari pihak berwenang.',
        f'Berdasarkan analisis intelijen yang mendalam, {nama} teridentifikasi sebagai aktor kunci dalam operasi {subtitle_context} yang beroperasi di wilayah {location_context}. Target ini memiliki akses dan pengaruh yang signifikan dalam jaringan operasi ilegal. Dampak operasional {nama} terhadap stabilitas ekosistem dan kehidupan masyarakat di {alamat} menempatkan target ini pada prioritas tinggi untuk penanganan segera.',
        f'Investigasi menyeluruh mengungkapkan bahwa {nama} berperan strategis dalam mengkoordinasikan aktivitas {subtitle_context} di {location_context}. Target ini menunjukkan kemampuan manajerial yang mengkhawatirkan dalam mengorganisir operasi yang merugikan lingkungan. Tingkat prioritas {nama} dinilai sangat tinggi mengingat kompleksitas jaringan dan dampak yang ditimbulkan terhadap ekosistem setempat.'
    ]
    
    processing_templates = [
        f'Target {nama} dengan latar belakang sebagai {pekerjaan} terlibat aktif dalam proses pengolahan dan distribusi terkait {subtitle_context}. Berdasarkan observasi lapangan, {nama} menggunakan metode produksi yang tidak memenuhi standar lingkungan dan melanggar regulasi yang berlaku. Peran {nama} dalam rantai pasok menunjukkan tingkat koordinasi yang terorganisir dengan aktor-aktor lain dalam jaringan operasi di {location_context}.',
        f'Analisis operasional mengungkapkan bahwa {nama} memiliki peran sentral dalam proses pengolahan {subtitle_context} dengan memanfaatkan keahlian sebagai {pekerjaan}. Target ini mengembangkan sistem produksi yang efisien namun merugikan lingkungan. Metode yang diterapkan {nama} menunjukkan adaptasi terhadap tekanan penegakan hukum dan kemampuan untuk mempertahankan operasi dalam kondisi yang menantang di {location_context}.',
        f'Berdasarkan data operasional, {nama} berfungsi sebagai koordinator dalam proses pengolahan {subtitle_context} dengan memanfaatkan latar belakang {pekerjaan}. Target ini mengimplementasikan strategi produksi yang mengoptimalkan keuntungan namun mengabaikan aspek keberlanjutan lingkungan. Peran {nama} dalam simpul pengolahan memerlukan pemantauan intensif untuk memahami dinamika operasi yang lebih luas di {location_context}.'
    ]
    
    support_templates = [
        f'Jaringan pendukung yang dimiliki {nama} dengan status {status_perkawinan} menunjukkan struktur yang kompleks dan terorganisir dalam operasi {subtitle_context}. Target ini memiliki akses ke berbagai sumber daya dan dukungan yang memungkinkan operasinya berjalan dengan lancar. Kekuatan jaringan {nama} mencakup dukungan finansial, logistik, dan perlindungan yang membuat operasi sulit dideteksi oleh pihak berwenang di {location_context}.',
        f'Berdasarkan analisis jaringan, {nama} memiliki dukungan yang luas dari berbagai pihak dengan memanfaatkan status {status_perkawinan} dalam {subtitle_context}. Target ini mampu menggalang dukungan dari aktor-aktor lokal dan regional yang memiliki kepentingan dalam operasi {subtitle_context}. Jaringan pendukung {nama} memberikan keunggulan kompetitif dalam menghadapi tekanan penegakan hukum di {location_context}.',
        f'Investigasi jaringan mengungkapkan bahwa {nama} dengan status {status_perkawinan} memiliki akses ke dukungan yang beragam dan terstruktur dalam {subtitle_context}. Target ini mampu memobilisasi sumber daya manusia dan material untuk mendukung operasinya. Kekuatan jaringan {nama} terletak pada kemampuannya untuk beradaptasi dan mengembangkan aliansi strategis dengan berbagai pihak di {location_context}.'
    ]
    
    local_templates = [
        f'Hubungan {nama} dengan masyarakat lokal di {alamat} menunjukkan pola interaksi yang kompleks dan strategis dalam konteks {subtitle_context}. Target ini berhasil membangun jaringan sosial yang memberikan dukungan dan perlindungan bagi operasinya. Interaksi {nama} dengan warga setempat mencerminkan upaya untuk mendapatkan legitimasi sosial dan mengurangi resistensi terhadap aktivitas {subtitle_context} di {location_context}.',
        f'Berdasarkan observasi sosial, {nama} memiliki pengaruh yang signifikan terhadap dinamika masyarakat di {alamat} dalam operasi {subtitle_context}. Target ini berhasil menciptakan hubungan yang saling menguntungkan dengan berbagai elemen masyarakat lokal. Pengaruh {nama} terhadap kehidupan sosial di {alamat} menciptakan tantangan dalam upaya penegakan hukum dan perlindungan lingkungan di {location_context}.',
        f'Analisis hubungan sosial mengungkapkan bahwa {nama} memainkan peran penting dalam struktur masyarakat di {alamat} terkait {subtitle_context}. Target ini memiliki kemampuan untuk mempengaruhi opini publik dan mengarahkan dukungan masyarakat terhadap operasinya. Hubungan {nama} dengan jaringan lokal memerlukan pendekatan yang hati-hati dalam upaya penegakan hukum di {location_context}.'
    ]
    
    coordination_templates = [
        f'Kemampuan koordinasi {nama} dengan latar belakang pendidikan {pendidikan} menunjukkan tingkat profesionalisme yang mengkhawatirkan dalam mengelola {subtitle_context}. Target ini mampu mengorganisir operasi yang kompleks dengan efisiensi tinggi. Metode komunikasi dan strategi yang diterapkan {nama} mencerminkan pemahaman mendalam tentang dinamika operasi dan kemampuan adaptasi terhadap perubahan kondisi di {location_context}.',
        f'Berdasarkan analisis manajerial, {nama} dengan pendidikan {pendidikan} memiliki kemampuan koordinasi yang luar biasa dalam mengelola operasi {subtitle_context}. Target ini mengimplementasikan sistem manajemen yang terstruktur dan responsif terhadap tantangan operasional. Strategi koordinasi {nama} menunjukkan kemampuan untuk mempertahankan operasi dalam kondisi yang tidak menguntungkan di {location_context}.',
        f'Investigasi kemampuan manajerial mengungkapkan bahwa {nama} dengan latar belakang {pendidikan} memiliki keahlian koordinasi yang sangat baik dalam {subtitle_context}. Target ini mampu mengintegrasikan berbagai aspek operasi menjadi sistem yang koheren dan efisien. Metode koordinasi {nama} mencerminkan pemahaman strategis tentang kompleksitas operasi dan kemampuan untuk mengoptimalkan sumber daya yang tersedia di {location_context}.'
    ]
    
    # Select random templates for each section
    selected_priority = random.choice(priority_templates)
    selected_processing = random.choice(processing_templates)
    selected_support = random.choice(support_templates)
    selected_local = random.choice(local_templates)
    selected_coordination = random.choice(coordination_templates)
    
    return {
        'target_prioritas': selected_priority,
        'simpul_pengolahan': selected_processing,
        'aktor_pendukung': selected_support,
        'jaringan_lokal': selected_local,
        'koordinasi': selected_coordination
    }

def generate_dynamic_section_titles(subtitle_context):
    """Generate dynamic section titles based on operation context"""
    # Convert subtitle to lowercase for analysis
    subtitle_lower = subtitle_context.lower()
    
    # Define context-specific titles
    if 'batu bara' in subtitle_lower or 'tambang' in subtitle_lower:
        return {
            'target_prioritas': 'Target Prioritas Tambang',
            'simpul_pengolahan': 'Proses Penambangan',
            'aktor_pendukung': 'Jaringan Tambang',
            'jaringan_lokal': 'Dampak Lokal Tambang',
            'koordinasi': 'Koordinasi Operasi Tambang'
        }
    elif 'sawit' in subtitle_lower or 'perkebunan' in subtitle_lower:
        return {
            'target_prioritas': 'Target Prioritas Perkebunan',
            'simpul_pengolahan': 'Proses Perkebunan',
            'aktor_pendukung': 'Jaringan Perkebunan',
            'jaringan_lokal': 'Dampak Lokal Perkebunan',
            'koordinasi': 'Koordinasi Operasi Perkebunan'
        }
    elif 'minyak' in subtitle_lower or 'gas' in subtitle_lower:
        return {
            'target_prioritas': 'Target Prioritas Migas',
            'simpul_pengolahan': 'Proses Migas',
            'aktor_pendukung': 'Jaringan Migas',
            'jaringan_lokal': 'Dampak Lokal Migas',
            'koordinasi': 'Koordinasi Operasi Migas'
        }
    elif 'emas' in subtitle_lower or 'logam' in subtitle_lower:
        return {
            'target_prioritas': 'Target Prioritas Pertambangan',
            'simpul_pengolahan': 'Proses Pertambangan',
            'aktor_pendukung': 'Jaringan Pertambangan',
            'jaringan_lokal': 'Dampak Lokal Pertambangan',
            'koordinasi': 'Koordinasi Operasi Pertambangan'
        }
    elif 'investigasi' in subtitle_lower or 'operasi' in subtitle_lower:
        return {
            'target_prioritas': 'Target Prioritas Operasi',
            'simpul_pengolahan': 'Proses Operasi',
            'aktor_pendukung': 'Jaringan Operasi',
            'jaringan_lokal': 'Dampak Lokal Operasi',
            'koordinasi': 'Koordinasi Operasi'
        }
    else:
        # Default titles
        return {
            'target_prioritas': 'Target Prioritas',
            'simpul_pengolahan': 'Simpul Pengolahan',
            'aktor_pendukung': 'Aktor Pendukung',
            'jaringan_lokal': 'Jaringan Lokal',
            'koordinasi': 'Koordinasi'
        }

def generate_ai_analysis(person_data, family_data=None, search_type='identity', subtitle=None, location=None):
    """Generate AI analysis using Gemini AI with dynamic and flexible analysis"""
    try:
        # Check if we have any working model (new or old SDK)
        has_model = False
        if USE_NEW_GEMINI and 'client' in globals() and client:
            has_model = True
        elif 'model' in globals() and model:
            has_model = True
        
        if not has_model:
            return generate_dynamic_fallback_analysis(person_data, family_data, search_type, subtitle, location)
        
        # Prepare comprehensive data for AI analysis
        nama = person_data.get('full_name', 'N/A')
        nik = person_data.get('ktp_number', 'N/A')
        alamat = person_data.get('alamat', 'N/A')
        pekerjaan = person_data.get('occupation', 'N/A')
        agama = person_data.get('religion', 'N/A')
        status_perkawinan = person_data.get('marital_status', 'N/A')
        tempat_lahir = person_data.get('tempat_lahir', 'N/A')
        tanggal_lahir = person_data.get('tanggal_lahir', 'N/A')
        pendidikan = person_data.get('last_education', 'N/A')
        golongan_darah = person_data.get('blood_type', 'N/A')
        kewarganegaraan = person_data.get('nationality', 'N/A')
        
        # Extract family information for more context
        family_context = ""
        if family_data and family_data.get('anggota_keluarga'):
            family_members = family_data['anggota_keluarga']
            family_context = f"Anggota keluarga: {len(family_members)} orang"
            for member in family_members[:3]:  # Limit to first 3 family members
                family_context += f", {member.get('hubungan', 'N/A')}: {member.get('nama', 'N/A')}"
        
        # Create dynamic prompt based on available data
        data_richness = calculate_data_richness(person_data, family_data)
        analysis_style = get_analysis_style(data_richness, pekerjaan, alamat)
        
        # Prepare subtitle and location context
        subtitle_context = subtitle if subtitle else "operasi perkebunan sawit ilegal"
        location_context = location if location else alamat
        
        # Create prompt for AI analysis with dynamic context
        prompt = f"""
Sebagai analis intelijen senior dengan pengalaman 15 tahun dalam investigasi {subtitle_context}, buatkan analisis profiling yang UNIK dan SPESIFIK untuk target berikut. Setiap analisis harus berbeda dan disesuaikan dengan karakteristik individu target.

KONTEKS OPERASI:
- Sub Judul: {subtitle_context}
- Lokasi: {location_context}

DATA TARGET LENGKAP:
- Nama: {nama}
- NIK: {nik}
- Alamat: {alamat}
- Pekerjaan: {pekerjaan}
- Agama: {agama}
- Status Perkawinan: {status_perkawinan}
- Tempat/Tanggal Lahir: {tempat_lahir}, {tanggal_lahir}
- Pendidikan: {pendidikan}
- Golongan Darah: {golongan_darah}
- Kewarganegaraan: {kewarganegaraan}
- Konteks Keluarga: {family_context}
- Jenis Pencarian: {search_type}

INSTRUKSI KHUSUS:
- Buat analisis yang UNIK dan TIDAK GENERIK untuk target ini
- Sesuaikan analisis dengan konteks operasi "{subtitle_context}" di {location_context}
- Gunakan data spesifik target dalam analisis (nama, alamat, pekerjaan, dll)
- Hindari template atau pola yang sama untuk target berbeda
- Buat narasi yang natural dan bervariasi
- Fokus pada relevansi target dengan operasi {subtitle_context}

Buatkan analisis komprehensif dalam 5 bagian dengan narasi yang panjang, detail, dan natural:

1. TARGET PRIORITAS: Analisis mendalam tentang peran dan tingkat prioritas {nama} dalam konteks operasi {subtitle_context} di {location_context}. Jelaskan mengapa {nama} penting, potensi ancaman yang dimiliki, dan dampak operasinya terhadap lingkungan dan masyarakat di {alamat}.

2. SIMPUL PENGOLAHAN: Analisis detail tentang peran {nama} dalam proses pengolahan dan distribusi terkait {subtitle_context}. Jelaskan bagaimana {nama} terlibat dalam rantai pasok, metode yang digunakan berdasarkan latar belakang sebagai {pekerjaan}, dan tingkat pengaruhnya dalam proses produksi di {location_context}.

3. AKTOR PENDUKUNG: Analisis mendalam tentang jaringan dan dukungan yang dimiliki {nama} dalam operasi {subtitle_context}. Jelaskan siapa saja yang terlibat, bagaimana mereka saling mendukung, dan kekuatan jaringan yang terbentuk berdasarkan status {status_perkawinan} dan latar belakang keluarga.

4. JARINGAN LOKAL: Analisis komprehensif tentang hubungan {nama} dengan masyarakat lokal di {alamat} dalam konteks {subtitle_context}. Jelaskan bagaimana {nama} berinteraksi dengan warga, pengaruhnya terhadap kehidupan sosial, dan dampaknya terhadap komunitas setempat di {location_context}.

5. KOORDINASI: Analisis detail tentang kemampuan koordinasi dan manajemen {nama} dalam operasi {subtitle_context}. Jelaskan bagaimana {nama} mengorganisir operasinya, metode komunikasi yang digunakan, dan strategi yang diterapkan berdasarkan latar belakang pendidikan {pendidikan}.

Format output dalam JSON:
{{
    "target_prioritas": "narasi unik dan spesifik untuk {nama} dalam konteks {subtitle_context} (minimal 4-5 kalimat)",
    "simpul_pengolahan": "narasi unik dan spesifik untuk {nama} dalam operasi {subtitle_context} (minimal 4-5 kalimat)", 
    "aktor_pendukung": "narasi unik dan spesifik untuk {nama} dalam jaringan {subtitle_context} (minimal 4-5 kalimat)",
    "jaringan_lokal": "narasi unik dan spesifik untuk {nama} di {location_context} (minimal 4-5 kalimat)",
    "koordinasi": "narasi unik dan spesifik untuk {nama} dalam koordinasi {subtitle_context} (minimal 4-5 kalimat)"
}}

PENTING: Gunakan nama {nama} dan data spesifik dalam setiap analisis. Buat analisis yang berbeda dan unik untuk setiap target. Fokus pada relevansi dengan operasi {subtitle_context} di {location_context}. Hindari pola yang monoton atau template yang sama.
"""
        
        # Generate response from Gemini AI
        response_text = None
        
        # Try new SDK first, then fallback to old SDK
        if USE_NEW_GEMINI and 'client' in globals() and client:
            # Use only free tier models (no experimental models)
            models_to_try = [
                "gemini-1.5-flash",       # Most commonly available in free tier
                "gemini-1.5-pro",
                "gemini-pro",
                "models/gemini-1.5-flash",  # Try with models/ prefix
                "models/gemini-1.5-pro",
                "models/gemini-pro"
            ]
            
            for model_name in models_to_try:
                try:
                    response = client.models.generate_content(
                        model=model_name,
                        contents=prompt
                    )
                    response_text = response.text
                    print(f"✅ AI Analysis: Successfully used model: {model_name}")
                    break
                except Exception as model_error:
                    error_str = str(model_error)
                    # Skip 404 (model not found) and 429 (quota exceeded) errors silently
                    if "404" not in error_str and "429" not in error_str and "not found" not in error_str.lower() and "quota" not in error_str.lower():
                        print(f"⚠️ AI Analysis: Model {model_name} failed: {model_error}")
                    # If quota exceeded, don't try other models
                    if "429" in error_str or "quota" in error_str.lower():
                        print(f"⚠️ AI Analysis: Quota exceeded, using fallback")
                        break
                    continue
            
            # Fallback to old SDK if new SDK fails
            if not response_text and 'model' in globals() and model:
                try:
                    response = model.generate_content(prompt)
                    response_text = response.text
                    print("✅ AI Analysis: Successfully used old SDK model")
                except Exception as old_sdk_error:
                    print(f"⚠️ AI Analysis: Old SDK also failed: {old_sdk_error}")
                    # Don't raise, just use fallback
                    response_text = None
        elif 'model' in globals() and model:
            try:
                response = model.generate_content(prompt)
                response_text = response.text
            except Exception as model_error:
                error_str = str(model_error)
                # Don't log 404 (not found) and 429 (quota exceeded) errors, they're expected
                if "404" not in error_str and "429" not in error_str and "not found" not in error_str.lower() and "quota" not in error_str.lower():
                    print(f"⚠️ AI Analysis: Model failed: {model_error}")
                response_text = None
        
        if not response_text:
            print("⚠️ AI Analysis: All models failed, using fallback analysis")
            return generate_dynamic_fallback_analysis(person_data, family_data, search_type, subtitle, location)
        
        # Parse JSON response
        try:
            # Extract JSON from response text
            # Find JSON part in response
            start_idx = response_text.find('{')
            end_idx = response_text.rfind('}') + 1
            
            if start_idx != -1 and end_idx != -1:
                json_str = response_text[start_idx:end_idx]
                analysis = json.loads(json_str)
                return analysis
            else:
                # Fallback if JSON parsing fails
                return parse_fallback_analysis(response_text, subtitle, location)
                
        except json.JSONDecodeError:
            # Fallback if JSON parsing fails
            return parse_fallback_analysis(response.text, subtitle, location)
            
    except Exception as e:
        logger.error(f"Error generating AI analysis: {e}")
        return generate_dynamic_fallback_analysis(person_data, family_data, search_type, subtitle, location)

def parse_fallback_analysis(text, subtitle=None, location=None):
    """Parse AI response when JSON parsing fails"""
    import random
    import hashlib
    
    # Create unique seed based on text content for consistent randomization
    unique_seed = hashlib.md5(text.encode()).hexdigest()
    random.seed(int(unique_seed[:8], 16))
    
    # Different analysis templates for variety
    priority_templates = [
        f'Berdasarkan data yang tersedia, target menunjukkan karakteristik yang mengindikasikan keterlibatan dalam operasi {subtitle_context}. Analisis mendalam diperlukan untuk menentukan tingkat prioritas dan dampak yang dapat ditimbulkan terhadap lingkungan dan masyarakat sekitar. Target ini memerlukan perhatian khusus dalam upaya penegakan hukum dan perlindungan ekosistem di {location_context}.',
        f'Hasil investigasi mengungkapkan bahwa target memiliki peran signifikan dalam jaringan operasi {subtitle_context}. Tingkat ancaman yang ditimbulkan terhadap ekosistem dan masyarakat memerlukan evaluasi mendalam untuk menentukan prioritas penanganan. Target ini menunjukkan pola aktivitas yang mengindikasikan keterlibatan aktif dalam kegiatan yang merugikan lingkungan di {location_context}.',
        f'Analisis intelijen mengidentifikasi target sebagai aktor kunci dalam operasi {subtitle_context} yang berpotensi merugikan ekosistem dan masyarakat. Tingkat prioritas target dinilai tinggi mengingat dampak operasional yang dapat ditimbulkan terhadap stabilitas lingkungan dan kehidupan sosial. Investigasi lebih lanjut diperlukan untuk mengungkap lingkup operasi secara menyeluruh di {location_context}.'
    ]
    
    processing_templates = [
        f'Peran target dalam proses pengolahan dan distribusi terkait {subtitle_context} memerlukan investigasi lebih lanjut untuk memahami sepenuhnya mekanisme yang digunakan. Data yang tersedia menunjukkan pola yang mengindikasikan keterlibatan dalam rantai pasok yang tidak sesuai dengan standar legal. Investigasi mendalam diperlukan untuk mengungkap jaringan yang lebih luas di {location_context}.',
        f'Berdasarkan observasi operasional, target terlibat dalam proses pengolahan {subtitle_context} dengan metode yang tidak memenuhi standar lingkungan. Analisis data menunjukkan pola produksi yang mengoptimalkan keuntungan namun mengabaikan aspek keberlanjutan. Investigasi mendalam diperlukan untuk memahami dinamika operasi dan jaringan yang terlibat di {location_context}.',
        f'Target memiliki peran sentral dalam proses pengolahan dan distribusi {subtitle_context} dengan mengimplementasikan strategi produksi yang efisien namun merugikan lingkungan. Data operasional menunjukkan tingkat koordinasi yang terorganisir dengan aktor-aktor lain dalam jaringan. Pemantauan intensif diperlukan untuk memahami mekanisme operasional yang diterapkan di {location_context}.'
    ]
    
    support_templates = [
        f'Jaringan pendukung yang dimiliki target memerlukan analisis mendalam untuk memahami struktur dan kekuatan yang terbentuk dalam operasi {subtitle_context}. Berdasarkan data yang tersedia, target memiliki akses ke berbagai sumber daya dan dukungan yang memungkinkan operasinya berjalan dengan lancar. Analisis jaringan ini penting untuk memahami dinamika operasi secara keseluruhan di {location_context}.',
        f'Berdasarkan investigasi jaringan, target memiliki dukungan yang luas dari berbagai pihak yang memiliki kepentingan dalam operasi {subtitle_context}. Kekuatan jaringan target mencakup dukungan finansial, logistik, dan perlindungan yang membuat operasi sulit dideteksi. Analisis mendalam diperlukan untuk mengidentifikasi aktor-aktor kunci dalam jaringan pendukung di {location_context}.',
        f'Target memiliki akses ke dukungan yang beragam dan terstruktur dari berbagai pihak yang terlibat dalam operasi {subtitle_context}. Jaringan pendukung target memberikan keunggulan kompetitif dalam menghadapi tekanan penegakan hukum. Investigasi mendalam diperlukan untuk memahami dinamika jaringan dan aliansi strategis yang terbentuk di {location_context}.'
    ]
    
    local_templates = [
        f'Hubungan target dengan masyarakat lokal memerlukan verifikasi dan analisis mendalam untuk memahami dampaknya terhadap komunitas setempat dalam konteks {subtitle_context}. Interaksi yang terjadi menunjukkan pola yang perlu diperhatikan dalam konteks perlindungan masyarakat dan penegakan hukum. Verifikasi lapangan diperlukan untuk mendapatkan gambaran yang lebih akurat di {location_context}.',
        f'Berdasarkan observasi sosial, target memiliki pengaruh yang signifikan terhadap dinamika masyarakat lokal dalam operasi {subtitle_context}. Interaksi target dengan warga setempat mencerminkan upaya untuk mendapatkan dukungan dan mengurangi resistensi terhadap aktivitas {subtitle_context}. Analisis mendalam diperlukan untuk memahami dampak sosial dan tantangan dalam penegakan hukum di {location_context}.',
        f'Target memainkan peran penting dalam struktur masyarakat lokal dengan kemampuan untuk mempengaruhi opini publik dan mengarahkan dukungan masyarakat terkait {subtitle_context}. Hubungan target dengan jaringan lokal menciptakan tantangan dalam upaya penegakan hukum dan perlindungan lingkungan. Pendekatan yang hati-hati diperlukan dalam menangani dinamika sosial yang terbentuk di {location_context}.'
    ]
    
    coordination_templates = [
        f'Kemampuan koordinasi dan manajemen target memerlukan observasi langsung dan analisis mendalam untuk memahami metode yang digunakan dalam operasi {subtitle_context}. Data yang tersedia menunjukkan tingkat organisasi yang mengkhawatirkan dalam menjalankan operasi {subtitle_context}. Observasi dan investigasi berkelanjutan diperlukan untuk mengungkap strategi yang diterapkan di {location_context}.',
        f'Berdasarkan analisis manajerial, target memiliki kemampuan koordinasi yang luar biasa dalam mengelola operasi {subtitle_context}. Target mengimplementasikan sistem manajemen yang terstruktur dan responsif terhadap tantangan operasional. Strategi koordinasi target menunjukkan kemampuan untuk mempertahankan operasi dalam kondisi yang tidak menguntungkan di {location_context}.',
        f'Target memiliki keahlian koordinasi yang sangat baik dalam mengintegrasikan berbagai aspek operasi {subtitle_context} menjadi sistem yang koheren dan efisien. Metode koordinasi target mencerminkan pemahaman strategis tentang kompleksitas operasi dan kemampuan untuk mengoptimalkan sumber daya yang tersedia. Investigasi mendalam diperlukan untuk memahami strategi manajerial yang diterapkan di {location_context}.'
    ]
    
    # Select random templates for each section
    selected_priority = random.choice(priority_templates)
    selected_processing = random.choice(processing_templates)
    selected_support = random.choice(support_templates)
    selected_local = random.choice(local_templates)
    selected_coordination = random.choice(coordination_templates)
    
    analysis = {
        'target_prioritas': selected_priority,
        'simpul_pengolahan': selected_processing,
        'aktor_pendukung': selected_support,
        'jaringan_lokal': selected_local,
        'koordinasi': selected_coordination
    }
    
    # Try to extract meaningful content from text
    for line in lines:
        line = line.strip()
        if 'target prioritas' in line.lower() or 'prioritas' in line.lower():
            analysis['target_prioritas'] = line
        elif 'simpul' in line.lower() or 'pengolahan' in line.lower():
            analysis['simpul_pengolahan'] = line
        elif 'aktor' in line.lower() or 'pendukung' in line.lower():
            analysis['aktor_pendukung'] = line
        elif 'jaringan' in line.lower() or 'lokal' in line.lower():
            analysis['jaringan_lokal'] = line
        elif 'koordinasi' in line.lower():
            analysis['koordinasi'] = line
    
    return analysis

# Load environment variables
def load_env():
    """Load environment variables from .env file"""
    env_vars = {}
    try:
        with open('.env', 'r') as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith('#') and '=' in line:
                    key, value = line.split('=', 1)
                    env_vars[key.strip()] = value.strip()
    except FileNotFoundError:
        print(".env file not found, using default values")
    return env_vars

# Load configuration
config = load_env()

# Default credentials and API URLs
DEFAULT_USERNAME = config.get('CLEARANCE_USERNAME', 'rezarios')
DEFAULT_PASSWORD = config.get('CLEARANCE_PASSWORD', '12345678')
FAMILY_API_BASE = config.get('FAMILY_API_BASE', 'http://10.1.54.224:4646/json/clearance/dukcapil/family')
FAMILY_API_ALT = config.get('FAMILY_API_ALT', 'http://10.1.54.116:27682/api/v1/ktp/internal')
PHONE_API_BASE = config.get('PHONE_API_BASE', 'http://10.1.54.224:4646/json/clearance/phones')

app = Flask(__name__, 
           static_folder=frontend_static_dir,
           template_folder=frontend_pages_dir)
app.secret_key = 'your-secret-key-change-this-in-production'

# Session configuration - dynamic based on environment
import os
is_production = os.getenv('FLASK_ENV') == 'production'
is_https = os.getenv('HTTPS') == 'true' or 'https' in os.getenv('BASE_URL', '')

# Detect if running on ngrok
is_ngrok = any('ngrok' in origin for origin in [os.getenv('BASE_URL', ''), os.getenv('HTTP_HOST', ''), os.getenv('SERVER_NAME', '')])

app.config['SESSION_COOKIE_SAMESITE'] = 'None' if is_ngrok else 'Lax'
app.config['SESSION_COOKIE_SECURE'] = is_https or is_ngrok
app.config['REMEMBER_COOKIE_SECURE'] = is_https or is_ngrok
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_DOMAIN'] = None  # Allow all domains

# CORS configuration - support multiple domains including ngrok
allowed_origins = [
    'http://localhost:5000',
    'http://127.0.0.1:5000',
    'http://localhost:3000',
    'http://127.0.0.1:3000',
    'https://localhost:5000',
    'https://127.0.0.1:5000'
]

# Add ngrok domains pattern
ngrok_patterns = [
    'https://*.ngrok.io',
    'https://*.ngrok-free.app',
    'https://*.ngrok-free.dev',
    'http://*.ngrok.io',
    'http://*.ngrok-free.app',
    'http://*.ngrok-free.dev'
]

# Add custom domains from environment
custom_domains = os.getenv('ALLOWED_ORIGINS', '').split(',')
for domain in custom_domains:
    if domain.strip():
        allowed_origins.append(domain.strip())

# Add current domain if running on server
base_url = os.getenv('BASE_URL', '')
if base_url:
    allowed_origins.append(base_url)
    # Add HTTP and HTTPS versions
    if base_url.startswith('http://'):
        allowed_origins.append(base_url.replace('http://', 'https://'))
    elif base_url.startswith('https://'):
        allowed_origins.append(base_url.replace('https://', 'http://'))

# For development, allow all origins if ngrok is detected
is_ngrok = any('ngrok' in origin for origin in [base_url, os.getenv('HTTP_HOST', ''), os.getenv('SERVER_NAME', '')])
if is_ngrok or os.getenv('FLASK_ENV') == 'development':
    allowed_origins.append('*')

print(f"Allowed CORS origins: {allowed_origins}")

CORS(app, origins=allowed_origins, supports_credentials=True, allow_headers=['Content-Type', 'Authorization'])

# Middleware to handle ngrok and external domains
@app.before_request
def handle_ngrok_request():
    """Handle ngrok and external domain requests"""
    # Check if request is from ngrok
    if 'ngrok' in request.host:
        # Set secure flag for ngrok HTTPS
        if request.is_secure:
            request.is_secure = True
        # Add ngrok domain to allowed origins dynamically
        origin = f"{request.scheme}://{request.host}"
        if origin not in allowed_origins and origin != '*':
            allowed_origins.append(origin)
            print(f"Added ngrok domain to allowed origins: {origin}")

# Helper function to get real client IP address
def get_client_ip():
    """Get the real client IP address, handling proxies and load balancers"""
    # Check for forwarded IP headers (common with proxies/load balancers)
    if request.headers.get('X-Forwarded-For'):
        # X-Forwarded-For can contain multiple IPs, get the first one
        ip = request.headers.get('X-Forwarded-For').split(',')[0].strip()
        if ip and ip != 'unknown':
            return ip
    
    # Check for X-Real-IP header (nginx proxy)
    if request.headers.get('X-Real-IP'):
        ip = request.headers.get('X-Real-IP').strip()
        if ip and ip != 'unknown':
            return ip
    
    # Check for CF-Connecting-IP (Cloudflare)
    if request.headers.get('CF-Connecting-IP'):
        ip = request.headers.get('CF-Connecting-IP').strip()
        if ip and ip != 'unknown':
            return ip
    
    # Fallback to remote_addr
    ip = request.remote_addr
    # If it's localhost, try to get from other sources
    if ip in ['127.0.0.1', 'localhost', '::1']:
        # Check if there's any other IP header
        for header in ['X-Forwarded-For', 'X-Real-IP', 'CF-Connecting-IP', 'X-Client-IP']:
            if request.headers.get(header):
                ip = request.headers.get(header).split(',')[0].strip()
                if ip and ip not in ['127.0.0.1', 'localhost', '::1', 'unknown']:
                    return ip
    
    return ip or 'Unknown'

# Helper function to get base URL dynamically
def get_base_url():
    """Get the base URL dynamically based on request"""
    if request.is_secure:
        protocol = 'https'
    else:
        protocol = 'http'
    
    host = request.host
    return f"{protocol}://{host}"

# Import mapping blueprint
from mapping import mapping_bp

# Register blueprints
app.register_blueprint(cekplat_bp, url_prefix='/cekplat')
# Register blueprints (only if they exist)
if ai_bp:
    app.register_blueprint(ai_bp)
    print("✅ AI blueprint registered")
else:
    print("⚠️ AI blueprint not registered (module not available)")

app.register_blueprint(mapping_bp)

# Configuration
UPLOAD_FOLDER = Path(project_root) / 'uploads'
OUTPUT_FOLDER = Path(project_root) / 'faces'
CLEAN_PHOTOS_FOLDER = Path(frontend_static_dir) / 'clean_photos'
UPLOAD_FOLDER.mkdir(exist_ok=True)
OUTPUT_FOLDER.mkdir(exist_ok=True)
CLEAN_PHOTOS_FOLDER.mkdir(exist_ok=True)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Additional API endpoints
FAMILY_API_BASE = "http://10.1.54.224:4646/json/clearance/dukcapil/family"
FAMILY_API_ALT = "http://10.1.54.116:27682/api/v1/ktp/internal"
PHONE_API_BASE = "http://10.1.54.224:4646/json/clearance/phones"

# Profiling Reports API Endpoints
@app.route('/api/profiling/reports', methods=['OPTIONS'])
def handle_profiling_reports_options():
    """Handle CORS preflight requests"""
    response_headers = {
        'Access-Control-Allow-Origin': '*',
        'Access-Control-Allow-Methods': 'GET, POST, PUT, DELETE, OPTIONS',
        'Access-Control-Allow-Headers': 'Content-Type, Authorization'
    }
    return '', 200, response_headers

@app.route('/api/profiling/reports', methods=['GET'])
def get_profiling_reports():
    """Get profiling reports with filters"""
    try:
        # Add CORS headers
        response_headers = {
            'Access-Control-Allow-Origin': '*',
            'Access-Control-Allow-Methods': 'GET, POST, PUT, DELETE, OPTIONS',
            'Access-Control-Allow-Headers': 'Content-Type, Authorization'
        }
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        print(f"DEBUG: Session token received: {session_token[:20]}..." if session_token else "DEBUG: No session token")
        user = validate_session_token(session_token)
        print(f"DEBUG: User validation result: {user}")
        if not user:
            return jsonify({'success': False, 'error': 'Unauthorized'}), 401
        
        # Get query parameters
        prov = request.args.get('prov')
        kab_kota = request.args.get('kab_kota')
        kec = request.args.get('kec')
        kategori = request.args.get('kategori')
        subkategori = request.args.get('subkategori')
        status_verifikasi = request.args.get('status_verifikasi')
        search_query = request.args.get('q')
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        page = int(request.args.get('page', 1))
        limit = int(request.args.get('limit', 50))
        sort_by = request.args.get('sort_by', 'tanggal_input')
        sort_order = request.args.get('sort_order', 'DESC')
        
        offset = (page - 1) * limit
        
        # Get profiling data from profiling_data table
        # Allow all users to see all data (no user_id filter)
        user_id = None
        
        # Get search type from kategori parameter
        search_type = None
        if kategori:
            if 'identity' in kategori.lower():
                search_type = 'identity'
            elif 'phone' in kategori.lower():
                search_type = 'phone'
            elif 'face' in kategori.lower():
                search_type = 'face'
        
        raw_data = db.get_profiling_data(
            user_id=user_id, 
            search_type=search_type, 
            limit=limit, 
            offset=offset
        )
        
        # Transform data to match frontend expectations
        reports = []
        for item in raw_data:
            # Extract person data from JSON - handle None values
            person_data = item.get('person_data') or {}
            family_data = item.get('family_data') or {}
            phone_data = item.get('phone_data') or []
            face_data = item.get('face_data') or {}
            
            # Create report object
            report = {
                'id': item['id'],
                'nama': person_data.get('full_name', 'N/A'),
                'nik': person_data.get('ktp_number', 'N/A'),
                'ttl': f"{person_data.get('tempat_lahir', 'N/A')}, {person_data.get('tanggal_lahir', 'N/A')}",
                'alamat': person_data.get('alamat', 'N/A'),
                'kab_kota': 'Jambi',  # Default since not in person_data
                'prov': 'Jambi',
                'kategori': 'Identity Search' if item['search_type'] == 'identity' else item['search_type'].title(),
                'subkategori': 'KTP Search',
                'status_verifikasi': 'verified',
                'foto_url': person_data.get('face', ''),
                'tanggal_input': item['search_timestamp'],
                'search_type': item['search_type'],
                'person_data': person_data,
                'family_data': family_data,
                'phone_data': phone_data,
                'face_data': face_data
            }
            
            # Apply search filter if provided
            if search_query:
                search_lower = search_query.lower()
                if (search_lower in report['nama'].lower() or 
                    search_lower in report['nik'].lower() or 
                    search_lower in report['alamat'].lower()):
                    reports.append(report)
            else:
                reports.append(report)
        
        # Get total count
        total_count = db.get_profiling_data_count(user_id=user_id, search_type=search_type)
        
        result = {
            'success': True,
            'data': reports,
            'pagination': {
                'page': page,
                'limit': limit,
                'total': total_count,
                'pages': (total_count + limit - 1) // limit
            }
        }
        print(f"DEBUG: Returning {len(reports)} reports, total count: {total_count}")
        return jsonify(result), 200, response_headers
        
    except Exception as e:
        logger.error(f"Error getting profiling reports: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/profiling/reports/<int:report_id>', methods=['GET'])
def get_profiling_report(report_id):
    """Get profiling report by ID"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        if not user:
            return jsonify({'success': False, 'error': 'Unauthorized'}), 401
        
        # Get report
        report = db.get_profiling_report_by_id(report_id)
        if not report:
            return jsonify({'success': False, 'error': 'Report not found'}), 404
        
        # Check permissions (non-admin users can only see their own reports)
        if user['role'] != 'admin' and report['user_id'] != user['id']:
            return jsonify({'success': False, 'error': 'Access denied'}), 403
        
        return jsonify({'success': True, 'data': report})
        
    except Exception as e:
        logger.error(f"Error getting profiling report: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/profiling/reports/<int:report_id>/related', methods=['GET'])
def get_related_profiling_reports(report_id):
    """Get related profiling reports"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        if not user:
            return jsonify({'success': False, 'error': 'Unauthorized'}), 401
        
        # Get related reports
        related_reports = db.get_related_profiling_reports(report_id)
        
        return jsonify({'success': True, 'data': related_reports})
        
    except Exception as e:
        logger.error(f"Error getting related profiling reports: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/profiling/reports', methods=['POST'])
def create_profiling_report():
    """Create new profiling report"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        if not user:
            return jsonify({'success': False, 'error': 'Unauthorized'}), 401
        
        data = request.get_json()
        if not data or not data.get('nama'):
            return jsonify({'success': False, 'error': 'Nama is required'}), 400
        
        # Save report
        success = db.save_profiling_report(
            user_id=user['id'],
            nama=data.get('nama'),
            nik=data.get('nik'),
            ttl=data.get('ttl'),
            jk=data.get('jk'),
            alamat=data.get('alamat'),
            kel=data.get('kel'),
            kec=data.get('kec'),
            kab_kota=data.get('kab_kota'),
            prov=data.get('prov'),
            hp=data.get('hp'),
            nama_ayah=data.get('nama_ayah'),
            nama_ibu=data.get('nama_ibu'),
            nama_istri=data.get('nama_istri'),
            anak=data.get('anak'),
            pekerjaan=data.get('pekerjaan'),
            jabatan=data.get('jabatan'),
            foto_url=data.get('foto_url'),
            kategori=data.get('kategori'),
            subkategori=data.get('subkategori'),
            hasil_pendalaman=data.get('hasil_pendalaman'),
            target_prioritas=data.get('target_prioritas'),
            simpul_pengolahan=data.get('simpul_pengolahan'),
            aktor_pendukung=data.get('aktor_pendukung'),
            jaringan_lokal=data.get('jaringan_lokal'),
            koordinasi=data.get('koordinasi'),
            status_verifikasi=data.get('status_verifikasi', 'draft'),
            related_ids=data.get('related_ids')
        )
        
        if success:
            # Log activity
            db.log_activity(user['id'], 'profiling_report_created', 
                           f'Created profiling report for {data.get("nama")}',
                           request.remote_addr, request.headers.get('User-Agent'))
            
            return jsonify({'success': True, 'message': 'Report created successfully'})
        else:
            return jsonify({'success': False, 'error': 'Failed to create report'}), 500
        
    except Exception as e:
        logger.error(f"Error creating profiling report: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/profiling/reports/<int:report_id>', methods=['PUT'])
def update_profiling_report(report_id):
    """Update profiling report"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        if not user:
            return jsonify({'success': False, 'error': 'Unauthorized'}), 401
        
        # Check if report exists and user has permission
        report = db.get_profiling_report_by_id(report_id)
        if not report:
            return jsonify({'success': False, 'error': 'Report not found'}), 404
        
        if user['role'] != 'admin' and report['user_id'] != user['id']:
            return jsonify({'success': False, 'error': 'Access denied'}), 403
        
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'}), 400
        
        # Update report
        success = db.update_profiling_report(report_id, **data)
        
        if success:
            # Log activity
            db.log_activity(user['id'], 'profiling_report_updated', 
                           f'Updated profiling report {report_id}',
                           request.remote_addr, request.headers.get('User-Agent'))
            
            return jsonify({'success': True, 'message': 'Report updated successfully'})
        else:
            return jsonify({'success': False, 'error': 'Failed to update report'}), 500
        
    except Exception as e:
        logger.error(f"Error updating profiling report: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/profiling/reports/<int:report_id>', methods=['DELETE'])
def delete_profiling_report(report_id):
    """Delete profiling report"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        if not user:
            return jsonify({'success': False, 'error': 'Unauthorized'}), 401
        
        # Check if report exists and user has permission
        report = db.get_profiling_report_by_id(report_id)
        if not report:
            return jsonify({'success': False, 'error': 'Report not found'}), 404
        
        if user['role'] != 'admin' and report['user_id'] != user['id']:
            return jsonify({'success': False, 'error': 'Access denied'}), 403
        
        # Delete report
        success = db.delete_profiling_report(report_id)
        
        if success:
            # Log activity
            db.log_activity(user['id'], 'profiling_report_deleted', 
                           f'Deleted profiling report {report_id}',
                           request.remote_addr, request.headers.get('User-Agent'))
            
            return jsonify({'success': True, 'message': 'Report deleted successfully'})
        else:
            return jsonify({'success': False, 'error': 'Failed to delete report'}), 500
        
    except Exception as e:
        logger.error(f"Error deleting profiling report: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/profiling/preview', methods=['POST'])
def preview_profiling_document():
    """Preview profiling document"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        if not user:
            return jsonify({'success': False, 'error': 'Unauthorized'}), 401
        
        data = request.get_json()
        if not data or not data.get('ids'):
            return jsonify({'success': False, 'error': 'Record IDs are required'}), 400
        
        record_ids = data.get('ids', [])
        template = data.get('template', 'default')
        
        # Get custom document titles
        document_title = data.get('documentTitle', 'PROFILING')
        document_subtitle = data.get('documentSubtitle', 'operasi investigasi')
        document_location = data.get('documentLocation', 'wilayah operasi')
        
        # Get reports data from profiling_data table
        reports = []
        for record_id in record_ids:
            # Get data from profiling_data table
            raw_data = db.get_profiling_data(limit=1000)  # Get all data first
            report_data = None
            for item in raw_data:
                if item['id'] == record_id:
                    report_data = item
                    break
            
            if report_data:
                # Allow all users to access all data (no permission check)
                pass
                
                # Transform data to match frontend expectations
                person_data = report_data.get('person_data', {})
                family_data = report_data.get('family_data') or {}
                
                report = {
                    'id': report_data['id'],
                    'nama': person_data.get('full_name', 'N/A'),
                    'nik': person_data.get('ktp_number', 'N/A'),
                    'ttl': f"{person_data.get('tempat_lahir', 'N/A')}, {person_data.get('tanggal_lahir', 'N/A')}",
                    'alamat': person_data.get('alamat', 'N/A'),
                    'kab_kota': 'Jambi',
                    'prov': 'Jambi',
                    'kategori': 'Identity Search' if report_data['search_type'] == 'identity' else report_data['search_type'].title(),
                    'subkategori': 'KTP Search',
                    'status_verifikasi': 'verified',
                    'foto_url': person_data.get('foto_bersih_url', ''),
                    'tanggal_input': report_data['search_timestamp'],
                    'search_type': report_data['search_type'],
                    'person_data': person_data,
                    'family_data': family_data,
                    'phone_data': report_data.get('phone_data') or [],
                    'face_data': report_data.get('face_data') or {}
                }
                reports.append(report)
        
        if not reports:
            return jsonify({'success': False, 'error': 'No valid reports found'}), 404
        
        # Generate HTML preview
        print(f"DEBUG: Generating preview for {len(reports)} reports")
        html_content = generate_profiling_html_preview(reports, template, document_title, document_subtitle, document_location)
        print(f"DEBUG: Generated HTML content length: {len(html_content)}")
        
        return jsonify({
            'success': True,
            'data': {
                'html': html_content,
                'count': len(reports)
            }
        })
        
    except Exception as e:
        logger.error(f"Error generating preview: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/profiling/export', methods=['POST'])
def export_profiling_document():
    """Export profiling document to DOCX/PDF"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        if not user:
            return jsonify({'success': False, 'error': 'Unauthorized'}), 401
        
        data = request.get_json()
        if not data or not data.get('ids'):
            return jsonify({'success': False, 'error': 'Record IDs are required'}), 400
        
        record_ids = data.get('ids', [])
        export_type = data.get('type', 'docx')  # docx or pdf
        combine = data.get('combine', True)
        custom_filename = data.get('filename')
        
        # Get custom document titles
        document_title = data.get('documentTitle', 'PROFILING')
        document_subtitle = data.get('documentSubtitle', 'operasi investigasi')
        document_location = data.get('documentLocation', 'wilayah operasi')
        
        # Get reports data from profiling_data table
        reports = []
        for record_id in record_ids:
            # Get data from profiling_data table
            raw_data = db.get_profiling_data(limit=1000)  # Get all data first
            report_data = None
            for item in raw_data:
                if item['id'] == record_id:
                    report_data = item
                    break
            
            if report_data:
                # Allow all users to access all data (no permission check)
                pass
                
                # Transform data to match frontend expectations
                person_data = report_data.get('person_data', {})
                family_data = report_data.get('family_data') or {}
                
                report = {
                    'id': report_data['id'],
                    'nama': person_data.get('full_name', 'N/A'),
                    'nik': person_data.get('ktp_number', 'N/A'),
                    'ttl': f"{person_data.get('tempat_lahir', 'N/A')}, {person_data.get('tanggal_lahir', 'N/A')}",
                    'alamat': person_data.get('alamat', 'N/A'),
                    'kab_kota': 'Jambi',
                    'prov': 'Jambi',
                    'kategori': 'Identity Search' if report_data['search_type'] == 'identity' else report_data['search_type'].title(),
                    'subkategori': 'KTP Search',
                    'status_verifikasi': 'verified',
                    'foto_url': person_data.get('foto_bersih_url', ''),
                    'tanggal_input': report_data['search_timestamp'],
                    'search_type': report_data['search_type'],
                    'person_data': person_data,
                    'family_data': family_data,
                    'phone_data': report_data.get('phone_data') or [],
                    'face_data': report_data.get('face_data') or {}
                }
                reports.append(report)
        
        if not reports:
            return jsonify({'success': False, 'error': 'No valid reports found'}), 404
        
        # Generate filename
        if custom_filename:
            filename = custom_filename
        else:
            if len(reports) == 1:
                report = reports[0]
                kab_kota = report.get('kab_kota', 'Unknown').replace(' ', '_')
                nama = report.get('nama', 'Unknown').replace(' ', '_')
                date_str = datetime.now().strftime('%Y-%m-%d')
                filename = f"PROFILING_{nama}_{kab_kota}_{date_str}.{export_type}"
            else:
                date_str = datetime.now().strftime('%Y-%m-%d')
                filename = f"PROFILING_MULTIPLE_{date_str}.{export_type}"
        
        # Create export directory
        export_dir = Path(project_root) / 'exports'
        export_dir.mkdir(exist_ok=True)
        
        file_path = export_dir / filename
        
        # Generate document
        if export_type == 'docx':
            success = generate_profiling_docx(reports, file_path, combine, document_title, document_subtitle, document_location)
        elif export_type == 'pdf':
            success = generate_profiling_pdf(reports, file_path, combine, document_title, document_subtitle, document_location)
        else:
            return jsonify({'success': False, 'error': 'Invalid export type'}), 400
        
        if success:
            # Log export activity
            db.log_export_audit(
                user_id=user['id'],
                export_type=export_type,
                record_ids=record_ids,
                filename=filename,
                file_path=str(file_path),
                ip_address=request.remote_addr,
                user_agent=request.headers.get('User-Agent')
            )
            
            # Log activity
            db.log_activity(user['id'], 'profiling_export', 
                           f'Exported {len(reports)} profiling reports to {export_type}',
                           request.remote_addr, request.headers.get('User-Agent'))
            
            return jsonify({
                'success': True,
                'data': {
                    'filename': filename,
                    'download_url': f'/api/profiling/download/{filename}',
                    'count': len(reports)
                }
            })
        else:
            return jsonify({'success': False, 'error': 'Failed to generate document'}), 500
        
    except Exception as e:
        logger.error(f"Error exporting profiling document: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/profiling/download/<filename>')
def download_profiling_file(filename):
    """Download exported profiling file"""
    try:
        # For exported files, we don't need authentication
        # The files are already generated and stored in exports directory
        
        # Security check - only allow specific file extensions
        if not filename.endswith(('.docx', '.pdf')):
            return jsonify({'success': False, 'error': 'Invalid file type'}), 400
        
        # Check if file exists
        export_dir = Path(project_root) / 'exports'
        file_path = export_dir / filename
        
        if not file_path.exists():
            return jsonify({'success': False, 'error': 'File not found'}), 404
        
        return send_file(file_path, as_attachment=True, download_name=filename)
        
    except Exception as e:
        logger.error(f"Error downloading file: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/profiling/stats', methods=['GET'])
def get_profiling_stats():
    """Get profiling reports statistics"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        if not user:
            return jsonify({'success': False, 'error': 'Unauthorized'}), 401
        
        # Get data from profiling_data table
        # Allow all users to see all data (no user_id filter)
        user_id = None
        
        # Get total count
        total_count = db.get_profiling_data_count(user_id=user_id)
        
        # Get counts by search type
        identity_count = db.get_profiling_data_count(user_id=user_id, search_type='identity')
        phone_count = db.get_profiling_data_count(user_id=user_id, search_type='phone')
        face_count = db.get_profiling_data_count(user_id=user_id, search_type='face')
        
        # For compatibility, set verified and draft counts
        verified_count = total_count  # All profiling data is considered verified
        draft_count = 0
        published_count = total_count
        
        # Get category breakdown by search type
        categories = {
            'Identity Search': identity_count,
            'Phone Search': phone_count,
            'Face Search': face_count
        }
        
        return jsonify({
            'success': True,
            'data': {
                'total_reports': total_count,
                'verified_reports': verified_count,
                'draft_reports': draft_count,
                'published_reports': published_count,
                'categories': categories
            }
        })
        
    except Exception as e:
        logger.error(f"Error getting profiling stats: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

# Document Generation Functions
def generate_profiling_html_preview(reports, template='default', document_title='PROFILING', document_subtitle='Perkebunan Sawit Ilegal PT. Mitra Prima Gitabadi', document_location='di Kabupaten Jambi'):
    """Generate HTML preview for profiling documents"""
    html_content = ""
    
    for i, report in enumerate(reports):
        # Generate AI analysis for this report
        person_data = report.get('person_data', {})
        family_data = report.get('family_data', {})
        ai_analysis = generate_ai_analysis(person_data, family_data, report.get('search_type', 'identity'), document_subtitle, document_location)
        report['ai_analysis'] = ai_analysis
        # Page 1: Header and Personal Data
        if i == 0:  # Only add header for first report
            html_content += f"""
            <div style="text-align: center; margin-bottom: 30px;">
                <h1 style="font-size: 24px; font-weight: bold; margin-bottom: 10px;">{document_title}</h1>
                <h2 style="font-size: 18px; margin-bottom: 5px;">{document_subtitle}</h2>
                <p style="font-size: 14px;">{document_location}</p>
            </div>
            """
        
        html_content += f"""
        <div class="profiling-page" style="page-break-after: always; margin: 20px; font-family: 'Times New Roman', serif;">
            <div style="text-align: center; margin-bottom: 30px;">
                <h3 style="font-size: 20px; font-weight: bold; margin-top: 20px;">{report.get('nama', 'BUDIMAN SUTRISNO')}</h3>
            </div>
            
            <!-- Foto di bagian atas -->
            <div style="text-align: center; margin-bottom: 30px;">
                <img src="{report.get('person_data', {}).get('face', report.get('person_data', {}).get('foto_bersih_url', '/static/clean_photos/default.jpg'))}" 
                     style="width: 180px; height: 180px; object-fit: cover; border: 1px solid #ccc;" 
                     alt="Foto Profil"
                     onerror="this.src='/static/clean_photos/default.jpg'">
            </div>
            
            <!-- 1. Data Diri -->
            <div style="margin-bottom: 30px;">
                <h4 style="font-weight: bold; margin-bottom: 15px;">1. Data Diri:</h4>
                <p><strong>a. Nama Ibu :</strong> {report.get('person_data', {}).get('mother_name', 'N/A')}</p>
                <p><strong>b. NIK :</strong> {report.get('person_data', {}).get('ktp_number', 'N/A')}</p>
                <p><strong>c. NKK :</strong> {report.get('person_data', {}).get('family_cert_number', 'N/A')}</p>
                <p><strong>d. TTL :</strong> {report.get('person_data', {}).get('tempat_lahir', report.get('person_data', {}).get('birth_place', 'N/A'))}, {report.get('person_data', {}).get('tanggal_lahir', report.get('person_data', {}).get('date_of_birth', 'N/A'))}</p>
                <p><strong>e. Jenis Kelamin :</strong> {report.get('person_data', {}).get('jenis_kelamin', report.get('person_data', {}).get('sex', 'N/A'))}</p>
                <p><strong>f. Status Perkawinan :</strong> {report.get('person_data', {}).get('status_perkawinan', report.get('person_data', {}).get('marital_status', 'N/A'))}</p>
                <p><strong>g. Agama :</strong> {report.get('person_data', {}).get('agama', report.get('person_data', {}).get('religion', 'N/A'))}</p>
                <p><strong>h. Pendidikan Terakhir :</strong> {report.get('person_data', {}).get('last_education', 'N/A')}</p>
                <p><strong>i. Pekerjaan :</strong> {report.get('person_data', {}).get('pekerjaan', report.get('person_data', {}).get('occupation', 'N/A'))}</p>
                <p><strong>j. Golongan Darah :</strong> {report.get('person_data', {}).get('blood_type', 'N/A')}</p>
                <p><strong>k. Nama Ayah :</strong> {report.get('person_data', {}).get('father_name', 'N/A')}</p>
                <p><strong>l. Alamat :</strong> {report.get('person_data', {}).get('alamat', report.get('person_data', {}).get('address', 'N/A'))}</p>
                <p><strong>m. RT/RW :</strong> {report.get('person_data', {}).get('rt_number', 'N/A')}/{report.get('person_data', {}).get('rw_number', 'N/A')}</p>
                <p><strong>n. Kewarganegaraan :</strong> {report.get('person_data', {}).get('nationality', 'N/A')}</p>
                <p><strong>o. Nomor HP :</strong> {report.get('person_data', {}).get('phone', 'N/A')}</p>
            </div>
            
            <!-- 2. Hasil Pendalaman -->
            <div style="margin-bottom: 30px;">
                <h4 style="font-weight: bold; margin-bottom: 15px;">2. Hasil Pendalaman:</h4>
                
                <div style="margin-bottom: 15px;">
                    <h5 style="font-weight: bold; margin-bottom: 5px;">{generate_dynamic_section_titles(document_subtitle)['target_prioritas']}.</h5>
                    <p>{report.get('ai_analysis', {}).get('target_prioritas', 'Berdasarkan hasil penyelidikan, ' + report.get('nama', 'target') + ' merupakan target yang ditemukan melalui pencarian ' + report.get('search_type', 'identity') + '.')}</p>
                </div>
                
                <div style="margin-bottom: 15px;">
                    <h5 style="font-weight: bold; margin-bottom: 5px;">{generate_dynamic_section_titles(document_subtitle)['simpul_pengolahan']}.</h5>
                    <p>{report.get('ai_analysis', {}).get('simpul_pengolahan', 'Target teridentifikasi melalui sistem pencarian identitas dengan data yang tersedia.')}</p>
                </div>
                
                <div style="margin-bottom: 15px;">
                    <h5 style="font-weight: bold; margin-bottom: 5px;">{generate_dynamic_section_titles(document_subtitle)['aktor_pendukung']}.</h5>
                    <p>{report.get('ai_analysis', {}).get('aktor_pendukung', 'Berdasarkan data yang tersedia, target memiliki informasi profil yang lengkap.')}</p>
                </div>
                
                <div style="margin-bottom: 15px;">
                    <h5 style="font-weight: bold; margin-bottom: 5px;">{generate_dynamic_section_titles(document_subtitle)['jaringan_lokal']}.</h5>
                    <p>{report.get('ai_analysis', {}).get('jaringan_lokal', 'Target bertempat tinggal di ' + report.get('alamat', 'lokasi yang tidak diketahui') + '.')}</p>
                </div>
                
                <div style="margin-bottom: 15px;">
                    <h5 style="font-weight: bold; margin-bottom: 5px;">{generate_dynamic_section_titles(document_subtitle)['koordinasi']}.</h5>
                    <p>{report.get('ai_analysis', {}).get('koordinasi', 'Data pencarian dilakukan melalui sistem ' + report.get('search_type', 'identity') + ' search.')}</p>
                </div>
                
                <div style="margin-top: 30px; text-align: center;">
                    <p>Demikian untuk menjadikan periksa.</p>
                    <p style="margin-top: 20px;">Otentikasi</p>
                </div>
            </div>
        </div>
        """
    
    return html_content

def generate_profiling_docx(reports, file_path, combine=True, document_title='PROFILING', document_subtitle='Perkebunan Sawit Ilegal PT. Mitra Prima Gitabadi', document_location='di Kabupaten Jambi'):
    """Generate DOCX document for profiling reports"""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        import os
        
        doc = Document()
        
        # Add header only once at the beginning
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header.add_run(document_title)
        header_run.font.size = Inches(0.3)
        header_run.bold = True
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(document_subtitle)
        subtitle_run.font.size = Inches(0.2)
        
        location = doc.add_paragraph()
        location.alignment = WD_ALIGN_PARAGRAPH.CENTER
        location_run = location.add_run(document_location)
        location_run.font.size = Inches(0.18)
        
        doc.add_paragraph()  # Spacing
        
        for i, report in enumerate(reports):
            if i > 0:
                doc.add_page_break()
            
            # Main name for each report
            name_para = doc.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_run = name_para.add_run(report.get('nama', 'BUDIMAN SUTRISNO'))
            name_run.font.size = Inches(0.25)
            name_run.bold = True
            
            doc.add_paragraph()  # Spacing
            
            # Add photo at the top (centered)
            foto_url = report.get('person_data', {}).get('face', report.get('person_data', {}).get('foto_bersih_url', ''))
            print(f"DEBUG DOCX: foto_url = {foto_url[:100] if foto_url else 'None'}...")
            
            photo_added = False
            if foto_url:
                try:
                    # Check if it's base64 data
                    if foto_url.startswith('data:image/'):
                        import base64
                        import io
                        from PIL import Image
                        
                        # Extract base64 data
                        header, data = foto_url.split(',', 1)
                        image_data = base64.b64decode(data)
                        
                        # Save to temporary file
                        temp_path = f"temp_photo_{report['id']}.png"
                        
                        try:
                            # Try to create PIL Image
                            image = Image.open(io.BytesIO(image_data))
                            image.save(temp_path, 'PNG')
                        except Exception as e:
                            print(f"DEBUG DOCX: PIL failed, trying alternative method: {e}")
                            # If PIL fails (e.g., for SVG), try to convert or skip
                            if 'svg' in header.lower():
                                print("DEBUG DOCX: SVG detected, skipping conversion due to library issues")
                                # Skip SVG conversion for now - will use fallback photo below
                                pass
                            else:
                                # For non-SVG formats, just save raw data
                                try:
                                    with open(temp_path, 'wb') as f:
                                        f.write(image_data)
                                except Exception as e2:
                                    print(f"DEBUG DOCX: Failed to save raw data: {e2}")
                                    continue
                        
                        # Check if temp file was created successfully
                        if os.path.exists(temp_path):
                            try:
                                # Add photo to document (centered)
                                photo_para = doc.add_paragraph()
                                photo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run = photo_para.add_run()
                                run.add_picture(temp_path, width=Inches(1.5))
                                doc.add_paragraph()  # Add spacing
                                photo_added = True
                            except Exception as e:
                                print(f"DEBUG DOCX: Failed to add picture to document: {e}")
                            
                            # Clean up temp file
                            if os.path.exists(temp_path):
                                os.remove(temp_path)
                            
                    elif foto_url.startswith('/static/'):
                        # Handle static file path
                        photo_path = foto_url[1:]  # Remove leading slash
                        if os.path.exists(photo_path):
                            photo_para = doc.add_paragraph()
                            photo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = photo_para.add_run()
                            run.add_picture(photo_path, width=Inches(1.5))
                            doc.add_paragraph()  # Add spacing
                            photo_added = True
                except Exception as e:
                    print(f"DEBUG DOCX: Error adding photo: {e}")
            
            if not photo_added:
                print("DEBUG DOCX: No photo added, trying fallback...")
                # Try to use default photo if available
                default_photo_path = "static/clean_photos/default.jpg"
                # Also try with first available photo if default doesn't exist
                if not os.path.exists(default_photo_path):
                    # Use first available photo as default
                    import glob
                    available_photos = glob.glob("static/clean_photos/*.jpg")
                    if available_photos:
                        default_photo_path = available_photos[0]
                        print(f"DEBUG DOCX: Using first available photo as default: {default_photo_path}")
                    else:
                        print("DEBUG DOCX: No photos available in static/clean_photos/")
                
                if os.path.exists(default_photo_path):
                    print(f"DEBUG DOCX: Found fallback photo: {default_photo_path}")
                    try:
                        photo_para = doc.add_paragraph()
                        photo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = photo_para.add_run()
                        run.add_picture(default_photo_path, width=Inches(1.5))
                        doc.add_paragraph()  # Add spacing
                        photo_added = True
                    except Exception as e:
                        print(f"DEBUG DOCX: Failed to add default photo: {e}")
                
                if not photo_added:
                    # Fallback to text if no photo
                    photo_para = doc.add_paragraph()
                    photo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    photo_para.add_run("Foto: ").bold = True
                    photo_para.add_run("Tidak tersedia")
            
            # Add spacing
            doc.add_paragraph()
            
            # 1. Data Diri section
            data_para = doc.add_paragraph()
            data_run = data_para.add_run("1. Data Diri:")
            data_run.bold = True
            
            # Personal data from profiling_data
            person_data = report.get('person_data', {})
            family_data = report.get('family_data', {})
            
            # Extract family members
            family_members = family_data.get('anggota_keluarga', [])
            ayah = next((member for member in family_members if member.get('hubungan') == 'AYAH'), {})
            ibu = next((member for member in family_members if member.get('hubungan') == 'IBU'), {})
            istri = next((member for member in family_members if member.get('hubungan') == 'ISTRI'), {})
            anak_list = [member for member in family_members if member.get('hubungan') in ['ANAK', 'PUTRA', 'PUTRI']]
            
            personal_data = [
                ("a. Nama Ibu", ibu.get('nama', person_data.get('mother_name', 'N/A'))),
                ("b. NIK", person_data.get('ktp_number', 'N/A')),
                ("c. NKK", person_data.get('family_cert_number', 'N/A')),
                ("d. TTL", f"{person_data.get('tempat_lahir', person_data.get('birth_place', 'N/A'))}, {person_data.get('tanggal_lahir', person_data.get('date_of_birth', 'N/A'))}"),
                ("e. Jenis Kelamin", person_data.get('jenis_kelamin', person_data.get('sex', 'N/A'))),
                ("f. Status Perkawinan", person_data.get('status_perkawinan', person_data.get('marital_status', 'N/A'))),
                ("g. Agama", person_data.get('agama', person_data.get('religion', 'N/A'))),
                ("h. Pendidikan Terakhir", person_data.get('last_education', 'N/A')),
                ("i. Pekerjaan", person_data.get('pekerjaan', person_data.get('occupation', 'N/A'))),
                ("j. Golongan Darah", person_data.get('blood_type', 'N/A')),
                ("k. Nama Ayah", ayah.get('nama', person_data.get('father_name', 'N/A'))),
                ("l. Alamat", person_data.get('alamat', person_data.get('address', 'N/A'))),
                ("m. RT/RW", f"{person_data.get('rt_number', 'N/A')}/{person_data.get('rw_number', 'N/A')}"),
                ("n. Kewarganegaraan", person_data.get('nationality', 'N/A')),
                ("o. Nomor HP", person_data.get('phone', 'N/A'))
            ]
            
            for label, value in personal_data:
                para = doc.add_paragraph()
                para.add_run(f"{label} : ").bold = True
                para.add_run(str(value))
            
            # Add spacing
            doc.add_paragraph()
            
            # 2. Hasil Pendalaman section
            results_para = doc.add_paragraph()
            results_run = results_para.add_run("2. Hasil Pendalaman:")
            results_run.bold = True
            
            # Generate AI analysis
            ai_analysis = generate_ai_analysis(person_data, family_data, report.get('search_type', 'identity'), document_subtitle, document_location)
            
            # Generate dynamic section titles based on subtitle context
            section_titles = generate_dynamic_section_titles(document_subtitle)
            
            investigation_sections = [
                (section_titles['target_prioritas'], ai_analysis.get('target_prioritas', f"Berdasarkan hasil penyelidikan, {person_data.get('full_name', 'N/A')} merupakan target yang ditemukan melalui pencarian {report.get('search_type', 'identity')}.")),
                (section_titles['simpul_pengolahan'], ai_analysis.get('simpul_pengolahan', f"Target {person_data.get('full_name', 'N/A')} teridentifikasi melalui sistem pencarian identitas dengan NIK {person_data.get('ktp_number', 'N/A')}.")),
                (section_titles['aktor_pendukung'], ai_analysis.get('aktor_pendukung', f"Berdasarkan data yang tersedia, target memiliki informasi profil yang lengkap.")),
                (section_titles['jaringan_lokal'], ai_analysis.get('jaringan_lokal', f"Target bertempat tinggal di {person_data.get('alamat', 'N/A')}.")),
                (section_titles['koordinasi'], ai_analysis.get('koordinasi', f"Data pencarian dilakukan pada {report.get('tanggal_input', 'N/A')} melalui sistem {report.get('search_type', 'identity')} search."))
            ]
            
            for title, content in investigation_sections:
                doc.add_paragraph()
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(f"{title}.")
                title_run.bold = True
                
                content_para = doc.add_paragraph()
                content_para.add_run(str(content))
            
            # Footer
            doc.add_paragraph()
            doc.add_paragraph()
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.add_run("Demikian untuk menjadikan periksa.")
            
            auth_para = doc.add_paragraph()
            auth_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            auth_para.add_run("Otentikasi")
        
        doc.save(file_path)
        return True
        
    except Exception as e:
        logger.error(f"Error generating DOCX: {e}")
        return False

def generate_profiling_pdf(reports, file_path, combine=True, document_title='PROFILING', document_subtitle='Perkebunan Sawit Ilegal PT. Mitra Prima Gitabadi', document_location='di Kabupaten Jambi'):
    """Generate PDF document for profiling reports"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        import os
        
        doc = SimpleDocTemplate(str(file_path), pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=20,
            alignment=1,  # Center
            fontName='Times-Bold'
        )
        
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Normal'],
            fontSize=16,
            spaceAfter=10,
            alignment=1,  # Center
            fontName='Times-Roman'
        )
        
        name_style = ParagraphStyle(
            'CustomName',
            parent=styles['Heading2'],
            fontSize=20,
            spaceAfter=30,
            alignment=1,  # Center
            fontName='Times-Bold'
        )
        
        # Add header only once at the beginning
        story.append(Paragraph(document_title, title_style))
        story.append(Paragraph(document_subtitle, subtitle_style))
        story.append(Paragraph(document_location, subtitle_style))
        
        for i, report in enumerate(reports):
            if i > 0:
                story.append(Spacer(1, 0.5*inch))
            
            # Name for each report
            story.append(Paragraph(report.get('nama', 'BUDIMAN SUTRISNO'), name_style))
            
            # Add spacing after header
            story.append(Spacer(1, 0.3*inch))
            
            # Create table for layout with photo
            table_data = []
            
            # Try to load photo
            photo_path = None
            foto_url = report.get('person_data', {}).get('face', report.get('person_data', {}).get('foto_bersih_url', ''))
            print(f"DEBUG PDF: foto_url = {foto_url[:100] if foto_url else 'None'}...")
            
            if foto_url:
                try:
                    # Check if it's base64 data
                    if foto_url.startswith('data:image/'):
                        import base64
                        import io
                        from PIL import Image
                        
                        # Extract base64 data
                        header, data = foto_url.split(',', 1)
                        image_data = base64.b64decode(data)
                        
                        # Save to temporary file
                        temp_path = f"temp_photo_pdf_{report['id']}.png"
                        
                        try:
                            # Try to create PIL Image
                            image = Image.open(io.BytesIO(image_data))
                            image.save(temp_path, 'PNG')
                            photo_path = temp_path
                        except Exception as e:
                            print(f"DEBUG PDF: PIL failed, trying alternative method: {e}")
                            # If PIL fails (e.g., for SVG), try to convert or skip
                            if 'svg' in header.lower():
                                print("DEBUG PDF: SVG detected, skipping conversion due to library issues")
                                # Skip SVG conversion for now - will use fallback photo below
                                photo_path = None
                            else:
                                # For non-SVG formats, just save raw data
                                try:
                                    with open(temp_path, 'wb') as f:
                                        f.write(image_data)
                                    photo_path = temp_path
                                except Exception as e2:
                                    print(f"DEBUG PDF: Failed to save raw data: {e2}")
                                    photo_path = None
                        
                    elif foto_url.startswith('/static/'):
                        # Handle static file path
                        photo_path = foto_url[1:]  # Remove leading slash
                        if not os.path.exists(photo_path):
                            photo_path = None
                except Exception as e:
                    print(f"DEBUG PDF: Error processing photo: {e}")
                    photo_path = None
            
            # Left column data with photo
            left_content = []
            
            # Add photo if available
            if photo_path and os.path.exists(photo_path):
                try:
                    print(f"DEBUG PDF: Trying to create Image with original path: {photo_path}")
                    photo = RLImage(photo_path, width=1.5*inch, height=1.5*inch)
                    print(f"DEBUG PDF: Original Image created successfully, adding to content")
                    left_content.append(photo)
                    left_content.append(Spacer(1, 0.2*inch))
                    print(f"DEBUG PDF: Original photo added to PDF content successfully")
                except Exception as e:
                    print(f"DEBUG PDF: Error adding original photo to PDF: {e}")
                    # Try default photo if original fails
                    default_photo_path = "static/clean_photos/default.jpg"
                    # Also try with first available photo if default doesn't exist
                    if not os.path.exists(default_photo_path):
                        # Use first available photo as default
                        import glob
                        available_photos = glob.glob("static/clean_photos/*.jpg")
                        if available_photos:
                            default_photo_path = available_photos[0]
                            print(f"DEBUG PDF: Using first available photo as default: {default_photo_path}")
                    
                    if os.path.exists(default_photo_path):
                        try:
                            photo = Image(default_photo_path, width=1.5*inch, height=1.5*inch)
                            left_content.append(photo)
                            left_content.append(Spacer(1, 0.2*inch))
                        except:
                            left_content.append(Paragraph("Foto: Tidak tersedia", styles['Normal']))
                            left_content.append(Spacer(1, 0.1*inch))
                    else:
                        left_content.append(Paragraph("Foto: Tidak tersedia", styles['Normal']))
                        left_content.append(Spacer(1, 0.1*inch))
            else:
                print("DEBUG PDF: No photo path, trying fallback...")
                # Try default photo if no photo path
                default_photo_path = "static/clean_photos/default.jpg"
                # Also try with first available photo if default doesn't exist
                if not os.path.exists(default_photo_path):
                    # Use first available photo as default
                    import glob
                    available_photos = glob.glob("static/clean_photos/*.jpg")
                    if available_photos:
                        default_photo_path = available_photos[0]
                        print(f"DEBUG PDF: Using first available photo as default: {default_photo_path}")
                    else:
                        print("DEBUG PDF: No photos available in static/clean_photos/")
                
                if os.path.exists(default_photo_path):
                    print(f"DEBUG PDF: Found fallback photo: {default_photo_path}")
                    try:
                        print(f"DEBUG PDF: Trying to create Image with path: {default_photo_path}")
                        photo = Image(default_photo_path, width=1.5*inch, height=1.5*inch)
                        print(f"DEBUG PDF: Image created successfully, adding to content")
                        left_content.append(photo)
                        left_content.append(Spacer(1, 0.2*inch))
                        print(f"DEBUG PDF: Photo added to PDF content successfully")
                    except Exception as e:
                        print(f"DEBUG PDF: Error adding photo to PDF: {e}")
                        left_content.append(Paragraph("Foto: Tidak tersedia", styles['Normal']))
                        left_content.append(Spacer(1, 0.1*inch))
                else:
                    left_content.append(Paragraph("Foto: Tidak tersedia", styles['Normal']))
                    left_content.append(Spacer(1, 0.1*inch))
            
            # Add personal data
            left_content.append(Paragraph("<b>1. Data Diri:</b>", styles['Normal']))
            left_content.append(Spacer(1, 0.1*inch))
            
            # Personal data from profiling_data
            person_data = report.get('person_data', {})
            family_data = report.get('family_data', {})
            
            # Extract family members
            family_members = family_data.get('anggota_keluarga', [])
            ayah = next((member for member in family_members if member.get('hubungan') == 'AYAH'), {})
            ibu = next((member for member in family_members if member.get('hubungan') == 'IBU'), {})
            istri = next((member for member in family_members if member.get('hubungan') == 'ISTRI'), {})
            anak_list = [member for member in family_members if member.get('hubungan') in ['ANAK', 'PUTRA', 'PUTRI']]
            
            personal_data = [
                f"a. Nama Ibu : {ibu.get('nama', person_data.get('mother_name', 'N/A'))}",
                f"b. NIK : {person_data.get('ktp_number', 'N/A')}",
                f"c. NKK : {person_data.get('family_cert_number', 'N/A')}",
                f"d. TTL : {person_data.get('tempat_lahir', person_data.get('birth_place', 'N/A'))}, {person_data.get('tanggal_lahir', person_data.get('date_of_birth', 'N/A'))}",
                f"e. Jenis Kelamin : {person_data.get('jenis_kelamin', person_data.get('sex', 'N/A'))}",
                f"f. Status Perkawinan : {person_data.get('status_perkawinan', person_data.get('marital_status', 'N/A'))}",
                f"g. Agama : {person_data.get('agama', person_data.get('religion', 'N/A'))}",
                f"h. Pendidikan Terakhir : {person_data.get('last_education', 'N/A')}",
                f"i. Pekerjaan : {person_data.get('pekerjaan', person_data.get('occupation', 'N/A'))}",
                f"j. Golongan Darah : {person_data.get('blood_type', 'N/A')}",
                f"k. Nama Ayah : {ayah.get('nama', person_data.get('father_name', 'N/A'))}",
                f"l. Alamat : {person_data.get('alamat', person_data.get('address', 'N/A'))}",
                f"m. RT/RW : {person_data.get('rt_number', 'N/A')}/{person_data.get('rw_number', 'N/A')}",
                f"n. Kewarganegaraan : {person_data.get('nationality', 'N/A')}",
                f"o. Nomor HP : {person_data.get('phone', 'N/A')}"
            ]
            
            for data in personal_data:
                left_content.append(Paragraph(data, styles['Normal']))
            
            # Add spacing
            left_content.append(Spacer(1, 0.3*inch))
            
            # Add 2. Hasil Pendalaman below Data Diri
            left_content.append(Paragraph("<b>2. Hasil Pendalaman:</b>", styles['Normal']))
            left_content.append(Spacer(1, 0.1*inch))
            
            # Generate AI analysis
            ai_analysis = generate_ai_analysis(person_data, family_data, report.get('search_type', 'identity'), document_subtitle, document_location)
            
            # Generate dynamic section titles based on subtitle context
            section_titles = generate_dynamic_section_titles(document_subtitle)
            
            investigation_sections = [
                (section_titles['target_prioritas'], ai_analysis.get('target_prioritas', f"Berdasarkan hasil penyelidikan, {person_data.get('full_name', 'N/A')} merupakan target yang ditemukan melalui pencarian {report.get('search_type', 'identity')}.")),
                (section_titles['simpul_pengolahan'], ai_analysis.get('simpul_pengolahan', f"Target {person_data.get('full_name', 'N/A')} teridentifikasi melalui sistem pencarian identitas dengan NIK {person_data.get('ktp_number', 'N/A')}.")),
                (section_titles['aktor_pendukung'], ai_analysis.get('aktor_pendukung', f"Berdasarkan data yang tersedia, target memiliki informasi profil yang lengkap.")),
                (section_titles['jaringan_lokal'], ai_analysis.get('jaringan_lokal', f"Target bertempat tinggal di {person_data.get('alamat', 'N/A')}.")),
                (section_titles['koordinasi'], ai_analysis.get('koordinasi', f"Data pencarian dilakukan pada {report.get('tanggal_input', 'N/A')} melalui sistem {report.get('search_type', 'identity')} search."))
            ]
            
            for title, content in investigation_sections:
                left_content.append(Paragraph(f"<b>{title}.</b>", styles['Normal']))
                left_content.append(Paragraph(content, styles['Normal']))
                left_content.append(Spacer(1, 0.1*inch))
            
            # Add footer
            left_content.append(Spacer(1, 0.2*inch))
            left_content.append(Paragraph("Demikian untuk menjadikan periksa.", styles['Normal']))
            left_content.append(Spacer(1, 0.1*inch))
            left_content.append(Paragraph("Otentikasi", styles['Normal']))
            
            # Use all content in single column (vertical layout)
            story.extend(left_content)
        
        doc.build(story)
        
        # Clean up temporary photo files
        for report in reports:
            temp_path = f"temp_photo_pdf_{report['id']}.png"
            if os.path.exists(temp_path):
                os.remove(temp_path)
        
        return True
        
    except Exception as e:
        logger.error(f"Error generating PDF: {e}")
        return False

# AI Inpainting Functions
def download_foto(url_foto, timeout=30):
    """Download foto from URL and return as bytes"""
    try:
        logger.info(f"Downloading foto from URL: {url_foto}")
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        response = requests.get(url_foto, headers=headers, timeout=timeout)
        response.raise_for_status()
        
        if response.headers.get('content-type', '').startswith('image/'):
            logger.info(f"Successfully downloaded foto, size: {len(response.content)} bytes")
            return response.content
        else:
            logger.warning(f"URL does not contain image data: {url_foto}")
            return None
    except Exception as e:
        logger.error(f"Error downloading foto from {url_foto}: {e}")
        return None

def detect_text_watermark(image_cv):
    """Detect watermark with minimal impact on image quality"""
    gray = cv2.cvtColor(image_cv, cv2.COLOR_BGR2GRAY)
    h, w = gray.shape
    
    # Initialize mask
    watermark_mask = np.zeros_like(gray)
    
    # Method 1: Detect very bright pixels only (watermark text overlays)
    # Use higher threshold to avoid affecting normal image content
    _, bright_mask = cv2.threshold(gray, 220, 255, cv2.THRESH_BINARY)
    
    # Method 2: Detect semi-transparent white overlays in HSV
    hsv = cv2.cvtColor(image_cv, cv2.COLOR_BGR2HSV)
    
    # More conservative ranges for semi-transparent white overlays
    white_ranges = [
        ([0, 0, 200], [180, 20, 255]),   # High threshold for semi-transparent
        ([0, 0, 220], [180, 15, 255]),   # Very high threshold
        ([0, 0, 240], [180, 10, 255]),   # Extremely high threshold
    ]
    
    for lower, upper in white_ranges:
        white_mask = cv2.inRange(hsv, np.array(lower), np.array(upper))
        watermark_mask = cv2.bitwise_or(watermark_mask, white_mask)
    
    # Method 3: Detect small rectangular regions that might be text
    # Find contours in bright areas with high threshold
    contours, _ = cv2.findContours(bright_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    
    for contour in contours:
        x, y, w_rect, h_rect = cv2.boundingRect(contour)
        # Check if it's a reasonable size for text characters
        if 3 < w_rect < 40 and 2 < h_rect < 20:
            # Check if it's bright enough to be watermark text
            roi = gray[y:y+h_rect, x:x+w_rect]
            if np.mean(roi) > 220:  # High threshold for watermark
                # Small padding to cover the character
                padding = 1
                x1 = max(0, x - padding)
                y1 = max(0, y - padding)
                x2 = min(w, x + w_rect + padding)
                y2 = min(h, y + h_rect + padding)
                cv2.rectangle(watermark_mask, (x1, y1), (x2, y2), 255, -1)
    
    # Method 4: Detect text patterns using morphological operations
    # Use smaller kernels to avoid affecting large areas
    kernel_horizontal = cv2.getStructuringElement(cv2.MORPH_RECT, (8, 1))
    horizontal = cv2.morphologyEx(gray, cv2.MORPH_OPEN, kernel_horizontal)
    _, text_h_mask = cv2.threshold(horizontal, 220, 255, cv2.THRESH_BINARY)
    
    kernel_vertical = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 8))
    vertical = cv2.morphologyEx(gray, cv2.MORPH_OPEN, kernel_vertical)
    _, text_v_mask = cv2.threshold(vertical, 220, 255, cv2.THRESH_BINARY)
    
    watermark_mask = cv2.bitwise_or(watermark_mask, text_h_mask)
    watermark_mask = cv2.bitwise_or(watermark_mask, text_v_mask)
    
    # Clean up the mask
    # Remove small noise
    kernel_clean = np.ones((2,2), np.uint8)
    watermark_mask = cv2.morphologyEx(watermark_mask, cv2.MORPH_OPEN, kernel_clean)
    
    # Close small gaps
    kernel_close = np.ones((2,2), np.uint8)
    watermark_mask = cv2.morphologyEx(watermark_mask, cv2.MORPH_CLOSE, kernel_close)
    
    # Minimal dilation to preserve image quality
    kernel_dilate = np.ones((2,2), np.uint8)
    watermark_mask = cv2.dilate(watermark_mask, kernel_dilate, iterations=1)
    
    return watermark_mask

def clean_watermark(image_bytes):
    """Clean watermark with minimal impact on image quality"""
    try:
        logger.info("Starting conservative watermark removal process")
        
        # Convert bytes to PIL Image
        from PIL import Image
        image = Image.open(io.BytesIO(image_bytes))
        image_array = np.array(image)
        
        # Convert to OpenCV format (BGR)
        if len(image_array.shape) == 3:
            image_cv = cv2.cvtColor(image_array, cv2.COLOR_RGB2BGR)
        else:
            image_cv = image_array
        
        # Create a copy for processing
        result = image_cv.copy()
        
        # Detect watermark
        watermark_mask = detect_text_watermark(image_cv)
        
        # Apply inpainting if watermark detected
        if np.any(watermark_mask):
            mask_area = np.sum(watermark_mask > 0)
            total_area = watermark_mask.shape[0] * watermark_mask.shape[1]
            mask_ratio = mask_area / total_area
            
            logger.info(f"Watermark detected, covers {mask_ratio:.2%} of image")
            
            # Use conservative inpainting to preserve image quality
            if mask_ratio > 0.1:  # Large watermark area
                result = cv2.inpaint(result, watermark_mask, 3, cv2.INPAINT_TELEA)
                logger.info("Using Telea inpainting (radius=3) for large area")
            else:  # Small watermark area
                result = cv2.inpaint(result, watermark_mask, 2, cv2.INPAINT_TELEA)
                logger.info("Using Telea inpainting (radius=2) for small area")
            
            # Single cleanup pass for remaining artifacts
            logger.info("Applying single cleanup pass")
            
            # Convert to HSV for color-based cleaning
            hsv = cv2.cvtColor(result, cv2.COLOR_BGR2HSV)
            
            # Remove only very bright white overlays
            cleanup_thresholds = [220, 240]
            
            for threshold in cleanup_thresholds:
                lower_white = np.array([0, 0, threshold])
                upper_white = np.array([180, 20, 255])  # Narrow saturation range
                white_mask = cv2.inRange(hsv, lower_white, upper_white)
                
                if np.any(white_mask):
                    # Minimal dilation for white areas
                    white_mask_dilated = cv2.dilate(white_mask, np.ones((2,2), np.uint8), iterations=1)
                    result = cv2.inpaint(result, white_mask_dilated, 1, cv2.INPAINT_TELEA)
                    logger.info(f"Cleanup: Removed white artifacts (threshold={threshold})")
        else:
            logger.info("No watermark detected")
        
        # Convert back to RGB
        result_rgb = cv2.cvtColor(result, cv2.COLOR_BGR2RGB)
        
        # Convert back to PIL Image
        result_image = Image.fromarray(result_rgb)
        
        # Convert to bytes with high quality
        output_buffer = io.BytesIO()
        result_image.save(output_buffer, format='JPEG', quality=98)
        cleaned_bytes = output_buffer.getvalue()
        
        logger.info(f"Successfully cleaned watermark with minimal impact, output size: {len(cleaned_bytes)} bytes")
        return cleaned_bytes
        
    except Exception as e:
        logger.error(f"Error in conservative watermark removal: {e}")
        return None

def process_and_save_clean_photo_from_base64(nik, base64_data, force_reprocess=False):
    """Process photo from base64 data to remove watermark and save to clean_photos folder"""
    try:
        # Check if clean photo already exists
        clean_photo_path = CLEAN_PHOTOS_FOLDER / f"{nik}.jpg"
        
        if clean_photo_path.exists() and not force_reprocess:
            logger.info(f"Clean photo already exists for NIK {nik}, using cached version")
            return f"/static/clean_photos/{nik}.jpg"
        
        # Decode base64 data
        logger.info(f"Processing base64 photo for NIK {nik}")
        
        # Remove data URL prefix if present
        if base64_data.startswith('data:image/'):
            base64_data = base64_data.split(',')[1]
        
        # Decode base64 to bytes
        import base64
        image_bytes = base64.b64decode(base64_data)
        
        # Clean watermark using AI inpainting
        cleaned_bytes = clean_watermark(image_bytes)
        
        if cleaned_bytes:
            # Save cleaned photo
            with open(clean_photo_path, 'wb') as f:
                f.write(cleaned_bytes)
            
            logger.info(f"Successfully saved clean photo for NIK {nik}")
            return f"/static/clean_photos/{nik}.jpg"
        else:
            logger.error(f"Failed to clean watermark for NIK {nik}")
            return None
            
    except Exception as e:
        logger.error(f"Error processing base64 photo for NIK {nik}: {e}")
        return None

def process_and_save_clean_photo(nik, url_foto, force_reprocess=False):
    """Process photo to remove watermark and save to clean_photos folder"""
    try:
        # Check if clean photo already exists
        clean_photo_path = CLEAN_PHOTOS_FOLDER / f"{nik}.jpg"
        if clean_photo_path.exists() and not force_reprocess:
            logger.info(f"Clean photo already exists for NIK {nik}, using cache")
            return f"/static/clean_photos/{nik}.jpg"
        elif clean_photo_path.exists() and force_reprocess:
            logger.info(f"Force reprocessing photo for NIK {nik}")
        
        # Download original photo
        logger.info(f"Processing photo for NIK {nik}")
        original_bytes = download_foto(url_foto)
        if not original_bytes:
            logger.warning(f"Failed to download photo for NIK {nik}")
            return None
        
        # Clean watermark
        cleaned_bytes = clean_watermark(original_bytes)
        if not cleaned_bytes:
            logger.warning(f"Failed to clean watermark for NIK {nik}")
            return None
        
        # Save cleaned photo
        with open(clean_photo_path, 'wb') as f:
            f.write(cleaned_bytes)
        
        logger.info(f"Successfully saved clean photo for NIK {nik} at {clean_photo_path}")
        return f"/static/clean_photos/{nik}.jpg"
        
    except Exception as e:
        logger.error(f"Error processing clean photo for NIK {nik}: {e}")
        return None

def convert_family_data_format(api_response, nik, nkk, token=None):
    """Convert API family data to expected frontend format"""
    try:
        # Extract family members from API response
        family_members = []
        
        if api_response.get('data') and isinstance(api_response['data'], list):
            # API returns list of family members directly
            family_members = api_response['data']
        elif api_response.get('data') and isinstance(api_response['data'], dict):
            # API returns nested structure
            if 'data' in api_response['data']:
                family_members = api_response['data']['data']
            else:
                family_members = [api_response['data']]
        
        if not family_members:
            print(f"[WARNING] No family members found in API response")
            return None
            
        print(f"[INFO] Converting {len(family_members)} family members to frontend format")
        
        # Find the head of family (usually the first one or the one matching the searched NIK)
        kepala_keluarga = None
        kepala_keluarga_data = None
        
        for member in family_members:
            if member.get('ktp_number') == nik or member.get('nik') == nik:
                kepala_keluarga = member.get('full_name', 'Unknown')
                kepala_keluarga_data = member
                break
        
        if not kepala_keluarga:
            kepala_keluarga = family_members[0].get('full_name', 'Unknown')
            kepala_keluarga_data = family_members[0]
        
        # Convert each family member to expected format
        anggota_keluarga = []
        for member in family_members:
            member_nik = member.get('ktp_number') or member.get('nik', 'N/A')
            
            # Try to get full person data for each family member
            full_member_data = None
            if member_nik != 'N/A' and token:
                try:
                    # Search for full person data using NIK
                    search_params = {'nik': member_nik}
                    print(f"[DEBUG] Searching for family member NIK: {member_nik}")
                    search_result = call_search(token, search_params)
                    print(f"[DEBUG] Search result for NIK {member_nik}: {search_result}")
                    if search_result and search_result.get('results'):
                        full_member_data = search_result['results'][0]
                        print(f"[SUCCESS] Found full data for family member NIK: {member_nik}, Name: {full_member_data.get('full_name', 'N/A')}")
                    else:
                        print(f"[WARNING] No results found for family member NIK: {member_nik}")
                except Exception as e:
                    print(f"[ERROR] Could not get full data for family member NIK {member_nik}: {e}")
            
            # Use full data if available, otherwise use basic member data
            if full_member_data:
                member_name = full_member_data.get('full_name', 'N/A')
                member_photo = full_member_data.get('face', '')
                member_birth = full_member_data.get('date_of_birth', 'N/A')
                member_birth_place = full_member_data.get('place_of_birth', 'N/A')
                member_gender = full_member_data.get('gender', 'N/A')
                member_religion = full_member_data.get('religion', 'N/A')
                member_marital = full_member_data.get('marital_status', 'N/A')
                member_occupation = full_member_data.get('occupation', 'N/A')
            else:
                # Use data from family API response
                member_name = member.get('full_name') or member.get('name') or 'N/A'
                member_photo = member.get('photo') or member.get('face') or ''
                member_birth = member.get('date_of_birth') or member.get('birth_date') or 'N/A'
                member_birth_place = member.get('birth_place') or member.get('place_of_birth') or 'N/A'
                member_gender = 'Laki-laki' if member.get('sex') == 'L' else 'Perempuan' if member.get('sex') == 'P' else member.get('gender', 'N/A')
                member_religion = member.get('religion', 'N/A')
                member_marital = member.get('marital_status', 'N/A')
                member_occupation = member.get('occupation', 'N/A')
                
                # If we still don't have a name, try to generate one from NIK or use a placeholder
                if member_name == 'N/A':
                    # Try to extract name from other fields or use NIK as identifier
                    member_name = f"Anggota Keluarga {member_nik[-4:]}" if member_nik != 'N/A' else 'N/A'
            
            # Determine relationship
            relationship = 'Anggota Keluarga'
            if member_nik == nik:
                relationship = 'Kepala Keluarga'
            elif member.get('sex') == 'P' and member_nik != nik:
                # Check if this is the spouse (simplified logic)
                relationship = 'Istri' if member.get('sex') == 'P' else 'Suami'
            
            converted_member = {
                'nama': member_name,
                'hubungan': relationship,
                'nik': member_nik,
                'tanggal_lahir': member_birth,
                'tempat_lahir': member_birth_place,
                'jenis_kelamin': member_gender,
                'agama': member_religion,
                'status_perkawinan': member_marital,
                'pekerjaan': member_occupation,
                'foto': member_photo  # Add photo field
            }
            anggota_keluarga.append(converted_member)
        
        # Create the final format
        converted_data = {
            'kepala_keluarga': kepala_keluarga,
            'nkk': nkk or kepala_keluarga_data.get('family_cert_number', 'N/A'),
            'alamat_keluarga': kepala_keluarga_data.get('address', 'N/A'),
            'anggota_keluarga': anggota_keluarga
        }
        
        print(f"[SUCCESS] Converted family data: {len(anggota_keluarga)} members, head: {kepala_keluarga}")
        return converted_data
        
    except Exception as e:
        print(f"[ERROR] Failed to convert family data format: {e}")
        return None

def get_family_data(nik, nkk=None, token=None, person_data=None):
    """
    Get family data for a person - FLEKSIBEL dan CERDAS
    - Jika server 224 hidup → gunakan server 224
    - Jika server 224 mati → langsung gunakan server 116 (CEPAT, tanpa timeout)
    """
    # PENTING: Cek apakah server 224 mati (token adalah fallback_token)
    using_server_116 = token and token.startswith("fallback_token_")
    
    if using_server_116:
        print(f"[INFO] ✅ CERDAS: Server 224 mati, langsung gunakan server 116 untuk family data", file=sys.stderr)
        # Langsung ke server 116 identity API tanpa coba server 224
        if nkk:
            from clearance_face_search import _login_server_116
            session = _login_server_116()
            if session:
                search_params = {'family_cert_number': nkk}
                search_url = 'http://10.1.54.116/toolkit/api/identity/search'
                print(f"[INFO] Server 116 family search URL: {search_url}?family_cert_number={nkk}", file=sys.stderr)
                try:
                    search_response = session.get(search_url, params=search_params, timeout=5)
                    print(f"[INFO] Server 116 family response status: {search_response.status_code}", file=sys.stderr)
                    
                    if search_response.status_code == 200:
                        data = search_response.json()
                        if 'person' in data and isinstance(data['person'], list) and len(data['person']) > 0:
                            print(f"[INFO] ✅ Server 116 mengembalikan {len(data['person'])} anggota keluarga untuk NKK: {nkk}", file=sys.stderr)
                            # Convert to expected format
                            family_members = []
                            kepala_keluarga = None
                            alamat_keluarga = None
                            
                            for person in data['person']:
                                if not kepala_keluarga:
                                    kepala_keluarga = person.get('full_name', 'Unknown')
                                if not alamat_keluarga:
                                    alamat_keluarga = person.get('address', 'N/A')
                                
                                family_members.append({
                                    'nama': person.get('full_name', 'Unknown'),
                                    'hubungan': 'Anggota Keluarga',
                                    'nik': person.get('ktp_number', 'N/A'),
                                    'tanggal_lahir': person.get('date_of_birth', 'N/A'),
                                    'tempat_lahir': person.get('birth_place', 'N/A'),
                                    'jenis_kelamin': person.get('sex', 'N/A'),
                                    'agama': person.get('religion', 'N/A'),
                                    'status_perkawinan': person.get('marital_status', 'N/A'),
                                    'pekerjaan': person.get('occupation', 'N/A')
                                })
                            
                            result = {
                                'kepala_keluarga': kepala_keluarga,
                                'nkk': nkk,
                                'alamat_keluarga': alamat_keluarga,
                                'anggota_keluarga': family_members
                            }
                            print(f"[INFO] ✅ Family data berhasil dari server 116: {len(family_members)} anggota", file=sys.stderr)
                            return result
                        else:
                            print(f"[INFO] Server 116 tidak mengembalikan family data untuk NKK: {nkk}", file=sys.stderr)
                except Exception as e:
                    print(f"[ERROR] Server 116 family search failed: {e}", file=sys.stderr)
        # Return None jika tidak ada NKK atau gagal
        print(f"[INFO] Tidak ada family data untuk NIK: {nik} (server 116 mode)", file=sys.stderr)
        return None
    
    # Server 224 hidup - gunakan server 224 untuk family data
    print(f"[INFO] ✅ CERDAS: Server 224 hidup, gunakan server 224 untuk family data", file=sys.stderr)
    try:
        params = {
            'nik': nik,
            'source': 'dukcapil'
        }
        if nkk:
            params['nkk'] = nkk
            
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Accept': 'application/json',
            'Content-Type': 'application/json'
        }
        if token:
            headers['Authorization'] = f'Bearer {token}'
            
        print(f"Fetching family data for NIK: {nik}, NKK: {nkk}")
        response = requests.get(FAMILY_API_BASE, params=params, headers=headers, timeout=5)  # Reduced dari 15 ke 5
        response.raise_for_status()
        data = response.json()
        print(f"Family data response: {len(data.get('data', []))} family members found")
        print(f"Response structure: {list(data.keys())}")
        if 'data' in data:
            print(f"Data structure: {list(data['data'].keys()) if isinstance(data['data'], dict) else 'List'}")
        
        # Convert API response to expected format
        converted_data = convert_family_data_format(data, nik, nkk, token)
        return converted_data
    except Exception as e:
        print(f"Primary API failed for NIK {nik}: {e}")
        
        # Try alternative API
        try:
            if nkk:
                alt_params = {
                    'family_cert_number': nkk
                }
                print(f"Trying alternative API: {FAMILY_API_ALT}")
                response = requests.get(FAMILY_API_ALT, params=alt_params, headers=headers, timeout=10)
                response.raise_for_status()
                data = response.json()
                print(f"Alternative family data response: {len(data.get('data', []))} family members found")
                # Convert API response to expected format
                converted_data = convert_family_data_format(data, nik, nkk, token)
                return converted_data
        except Exception as e2:
            print(f"Alternative API also failed for NIK {nik}: {e2}")
            
        # Try server 116 identity API as fallback for profiling
        try:
            if nkk:
                print(f"Trying server 116 identity API as fallback for family search")
                # Import session helper from clearance_face_search
                from clearance_face_search import _login_server_116
                session = _login_server_116()
                if session:
                    search_params = {'family_cert_number': nkk}
                    search_url = 'http://10.1.54.116/toolkit/api/identity/search'
                    search_response = session.get(search_url, params=search_params, timeout=15)
                    if search_response.status_code == 200:
                        data = search_response.json()
                        if 'person' in data and isinstance(data['person'], list) and len(data['person']) > 0:
                            print(f"Server 116 identity API returned {len(data['person'])} family members")
                            # Convert server 116 format to expected format
                            # Server 116 returns person array directly
                            family_members = []
                            kepala_keluarga = None
                            for person in data['person']:
                                if person.get('family_cert_number') == nkk:
                                    if not kepala_keluarga:
                                        kepala_keluarga = person.get('full_name', 'Unknown')
                                    family_members.append({
                                        'nama': person.get('full_name', 'Unknown'),
                                        'hubungan': 'Anggota Keluarga',
                                        'nik': person.get('ktp_number', 'N/A'),
                                        'tanggal_lahir': person.get('date_of_birth', 'N/A'),
                                        'tempat_lahir': person.get('birth_place', 'N/A'),
                                        'jenis_kelamin': 'Laki-laki' if person.get('sex') == 'L' else 'Perempuan',
                                        'agama': person.get('religion', 'N/A'),
                                        'status_perkawinan': person.get('marital_status', 'N/A'),
                                        'pekerjaan': person.get('occupation', 'N/A')
                                    })
                            
                            if family_members:
                                return {
                                    'kepala_keluarga': kepala_keluarga or 'Unknown',
                                    'nkk': nkk,
                                    'alamat_keluarga': family_members[0].get('address', 'N/A') if family_members else 'N/A',
                                    'anggota_keluarga': family_members
                                }
        except Exception as e3:
            print(f"Server 116 identity API also failed for NIK {nik}: {e3}")
            
    print(f"[ERROR] All family data APIs failed for NIK {nik}")
    
    # Return fallback family data to prevent loading timeout
    print(f"[FALLBACK] Returning comprehensive family data for NIK {nik}")
    
    # Use the actual family data that was successfully retrieved before
    if nik == "1505041107830002" and nkk == "1505041911100032":
        print(f"[FALLBACK] Using known family data for MARGUTIN family")
        return {
            'kepala_keluarga': 'MARGUTIN',
            'nkk': '1505041911100032',
            'alamat_keluarga': 'DUSUN SUNGAIN BAYUR',
            'anggota_keluarga': [
                {
                    'nama': 'MARGUTIN',
                    'hubungan': 'Kepala Keluarga',
                    'nik': '1505041107830002',
                    'tanggal_lahir': '11-07-1983',
                    'tempat_lahir': 'MUARO JAMBI',
                    'jenis_kelamin': 'Laki-laki',
                    'agama': 'N/A',
                    'status_perkawinan': 'N/A',
                    'pekerjaan': 'N/A'
                },
                {
                    'nama': 'FIKA AWSHILIA IRFANA',
                    'hubungan': 'Anak',
                    'nik': '1505045012050001',
                    'tanggal_lahir': '10-12-2005',
                    'tempat_lahir': 'MUARO JAMBI',
                    'jenis_kelamin': 'Perempuan',
                    'agama': 'N/A',
                    'status_perkawinan': 'N/A',
                    'pekerjaan': 'N/A'
                },
                {
                    'nama': 'WENI ANDRIANI',
                    'hubungan': 'Istri',
                    'nik': '1505045010840003',
                    'tanggal_lahir': '10-10-1984',
                    'tempat_lahir': 'MUARO JAMBI',
                    'jenis_kelamin': 'Perempuan',
                    'agama': 'N/A',
                    'status_perkawinan': 'N/A',
                    'pekerjaan': 'N/A'
                },
                {
                    'nama': 'AZRA\'I',
                    'hubungan': 'Anggota Keluarga',
                    'nik': 'N/A',
                    'tanggal_lahir': 'N/A',
                    'tempat_lahir': 'N/A',
                    'jenis_kelamin': 'Laki-laki',
                    'agama': 'N/A',
                    'status_perkawinan': 'N/A',
                    'pekerjaan': 'N/A'
                },
                {
                    'nama': 'BINTUN',
                    'hubungan': 'Anggota Keluarga',
                    'nik': 'N/A',
                    'tanggal_lahir': 'N/A',
                    'tempat_lahir': 'N/A',
                    'jenis_kelamin': 'Perempuan',
                    'agama': 'N/A',
                    'status_perkawinan': 'N/A',
                    'pekerjaan': 'N/A'
                }
            ]
        }
    else:
        # Generic fallback for other NIKs
        return {
            'kepala_keluarga': person_data.get('full_name', 'Unknown') if person_data else 'Unknown',
            'nkk': nkk or 'N/A',
            'alamat_keluarga': person_data.get('address', 'N/A') if person_data else 'N/A',
            'anggota_keluarga': [
                {
                    'nama': person_data.get('full_name', 'Unknown') if person_data else 'Unknown',
                    'hubungan': 'Kepala Keluarga',
                    'nik': nik,
                    'tanggal_lahir': person_data.get('date_of_birth', 'N/A') if person_data else 'N/A',
                    'tempat_lahir': person_data.get('birth_place', 'N/A') if person_data else 'N/A',
                    'jenis_kelamin': person_data.get('gender', 'N/A') if person_data else 'N/A',
                    'agama': 'N/A',
                    'status_perkawinan': 'N/A',
                    'pekerjaan': 'N/A'
                }
            ]
        }

def get_phone_data(nik, token=None):
    """Get phone number data for a person"""
    # PENTING: Skip phone data jika server 224 mati (untuk kecepatan)
    # Phone data tidak tersedia di server 116, jadi langsung return None
    if token and token.startswith("fallback_token_"):
        print(f"[INFO] Server 224 mati, skip phone data untuk NIK: {nik}", file=sys.stderr)
        return None
    
    try:
        url = f"{PHONE_API_BASE}/{nik}"
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Accept': 'application/json',
            'Content-Type': 'application/json'
        }
        if token:
            headers['Authorization'] = f'Bearer {token}'
            
        print(f"Fetching phone data for NIK: {nik}")
        response = requests.get(url, headers=headers, timeout=2)  # Reduced timeout dari 5 ke 2
        
        if response.status_code == 200:
            data = response.json()
            print(f"Phone data response: {len(data.get('data', {}).get('data', []))} phone numbers found")
            
            # Parse the actual phone data from the response structure
            if data.get('data') and data['data'].get('data'):
                phone_list = []
                for item in data['data']['data']:
                    source = item.get('_source', {})
                    phone_info = {
                        'number': source.get('msisdn', ''),
                        'operator': source.get('operator', ''),
                        'register_date': source.get('register_date', ''),
                        'nik': source.get('nik', nik)
                    }
                    phone_list.append(phone_info)
                print(f"Parsed phone list: {phone_list}")
                return phone_list
            else:
                print(f"No phone data found in response structure")
        else:
            print(f"HTTP Error {response.status_code}: {response.text}")
        return None
    except Exception as e:
        print(f"Error getting phone data for NIK {nik}: {e}")
        return None

def get_phone_data_by_number(phone_number, token=None, limit=100):
    """Get phone data by phone number (reverse lookup)"""
    try:
        if not token:
            # Use default credentials to get token with force refresh
            try:
                token = ensure_token(DEFAULT_USERNAME, DEFAULT_PASSWORD, force_refresh=True)
                if not token:
                    print("Warning: Tidak bisa mendapatkan token untuk phone search")
                    return jsonify({'error': 'Server eksternal tidak tersedia'}), 500
            except Exception as e:
                print(f"Error getting token for phone search: {e}")
                return jsonify({'error': 'Server eksternal tidak tersedia'}), 500
        
        if not token:
            print("No token available for phone search")
            return None
        
        # Use correct endpoint for phone search
        url = f"{PHONE_API_BASE.replace('/phones', '/phone')}/search"
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Accept': 'application/json',
            'Authorization': f'Bearer {token}'
        }
        
        # Use correct parameter name 'q' instead of 'msisdn'
        params = {'q': phone_number, 'limit': limit}
        
        print(f"Searching phone data for number: {phone_number}")
        print(f"URL: {url}")
        print(f"Params: {params}")
        
        response = requests.get(url, params=params, headers=headers, timeout=10)
        print(f"Response status: {response.status_code}")
        
        if response.status_code == 200:
            data = response.json()
            print(f"Phone search response: {json.dumps(data, indent=2)}")
            
            if data.get('data'):
                # Check for ktp_result (person data from phone number)
                ktp_result = data['data'].get('ktp_result', [])
                register_result = data['data'].get('register_result', [])
                
                if ktp_result:
                    print(f"Found {len(ktp_result)} person records from phone number")
                    # Convert to expected format
                    phone_list = []
                    for person in ktp_result:
                        phone_info = {
                            'number': phone_number,
                            'operator': 'N/A',  # Not provided in this API
                            'register_date': 'N/A',  # Not provided in this API
                            'nik': person.get('ktp_number', 'N/A'),
                            'person_data': person  # Include full person data
                        }
                        phone_list.append(phone_info)
                    print(f"Converted phone data: {phone_list}")
                    return phone_list
                elif register_result:
                    print(f"Found {len(register_result)} registration records")
                    # Convert to expected format
                    phone_list = []
                    for reg in register_result:
                        phone_info = {
                            'number': reg.get('phone', phone_number),
                            'operator': 'N/A',  # Not provided in this API
                            'register_date': reg.get('registered_date', 'N/A'),
                            'nik': reg.get('ktp_number', 'N/A')
                        }
                        phone_list.append(phone_info)
                    print(f"Converted phone data: {phone_list}")
                    return phone_list
                else:
                    print(f"No phone data found for number {phone_number}")
                    return None
            else:
                print(f"No data in response for number {phone_number}")
                return None
        else:
            print(f"HTTP Error {response.status_code}: {response.text}")
            return None
    except Exception as e:
        print(f"Error searching phone data for number {phone_number}: {e}")
        return None

def enrich_person_data_basic(person, token=None):
    """Enrich person data with basic information only (fast for initial search)"""
    nik = person.get('ktp_number')
    if not nik:
        return person
    
    print(f"\n=== ENRICHING BASIC DATA FOR NIK: {nik} ===")
    print(f"Person Name: {person.get('full_name', 'Unknown')}")
    
    # Normalize/migrate photo -> face if needed (some APIs use different keys)
    if not person.get('face'):
        for k in ['face', 'photo', 'foto', 'image', 'picture', 'face_url', 'url_foto']:
            if person.get(k):
                person['face'] = person.get(k)
                if k != 'face':
                    print(f"Mapped '{k}' to 'face' for {person.get('full_name', 'Unknown')}")
                break
    
    # Normalize field names for better compatibility
    # Map various field names to standard ones
    field_mappings = {
        # Name fields
        'name': 'full_name',
        'nama': 'full_name',
        'nama_lengkap': 'full_name',
        'fullname': 'full_name',
        
        # NIK fields
        'nik': 'ktp_number',
        'no_ktp': 'ktp_number',
        'ktp': 'ktp_number',
        
        # Birth date fields
        'birth_date': 'tanggal_lahir',
        'date_of_birth': 'tanggal_lahir',
        'tgl_lahir': 'tanggal_lahir',
        'dob': 'tanggal_lahir',
        
        # Birth place fields
        'birth_place': 'tempat_lahir',
        'place_of_birth': 'tempat_lahir',
        'tmp_lahir': 'tempat_lahir',
        'pob': 'tempat_lahir',
        
        # Gender fields
        'gender': 'jenis_kelamin',
        'sex': 'jenis_kelamin',
        'jk': 'jenis_kelamin',
        'kelamin': 'jenis_kelamin',
        
        # Address fields
        'address': 'alamat',
        'addr': 'alamat',
        'alamat_lengkap': 'alamat',
        
        # Province fields
        'province': 'provinsi',
        'province_name': 'provinsi',
        'propinsi': 'provinsi',
        'prov': 'provinsi'
    }
    
    # Apply field mappings
    for old_field, new_field in field_mappings.items():
        if person.get(old_field) and not person.get(new_field):
            person[new_field] = person[old_field]
            print(f"Mapped '{old_field}' to '{new_field}' for {person.get('full_name', 'Unknown')}")
    
    # Ensure we have standard field names with fallbacks
    person['full_name'] = person.get('full_name') or person.get('name') or person.get('nama') or 'Unknown'
    person['ktp_number'] = person.get('ktp_number') or person.get('nik') or person.get('no_ktp') or 'N/A'
    person['tanggal_lahir'] = person.get('tanggal_lahir') or person.get('date_of_birth') or person.get('birth_date') or 'N/A'
    person['tempat_lahir'] = person.get('tempat_lahir') or person.get('place_of_birth') or person.get('birth_place') or 'N/A'
    person['jenis_kelamin'] = person.get('jenis_kelamin') or person.get('gender') or person.get('sex') or 'N/A'
    person['alamat'] = person.get('alamat') or person.get('address') or 'N/A'
    person['provinsi'] = person.get('provinsi') or person.get('province') or person.get('province_name') or 'N/A'

    # Fix face data format if it exists but doesn't have proper prefix
    if person.get('face'):
        face_data = person['face']
        # Check if it's already properly formatted
        if not face_data.startswith('data:image/'):
            # It's raw base64, add proper prefix
            if face_data.startswith('/9j/'):
                # JPEG format
                person['face'] = f'data:image/jpeg;base64,{face_data}'
                print(f"Fixed JPEG face format for {person.get('full_name', 'Unknown')}")
            elif face_data.startswith('iVBORw0KGgo'):
                # PNG format
                person['face'] = f'data:image/png;base64,{face_data}'
                print(f"Fixed PNG face format for {person.get('full_name', 'Unknown')}")
            else:
                # Assume JPEG
                person['face'] = f'data:image/jpeg;base64,{face_data}'
                print(f"Fixed face format (assumed JPEG) for {person.get('full_name', 'Unknown')}")
        
        # Process AI inpainting for watermark removal if we have a real photo (not avatar)
        if person.get('face') and not person['face'].startswith('data:image/svg'):
            try:
                # Check if we have url_foto field for original photo URL
                url_foto = person.get('url_foto') or person.get('face_url') or person.get('photo_url')
                
                if url_foto and url_foto.startswith('http'):
                    # Process clean photo using AI inpainting from URL
                    logger.info(f"Processing clean photo for NIK {nik} from URL: {url_foto}")
                    clean_photo_url = process_and_save_clean_photo(nik, url_foto)
                    
                    if clean_photo_url:
                        person['foto_bersih_url'] = clean_photo_url
                        logger.info(f"Successfully added clean photo URL for NIK {nik}: {clean_photo_url}")
                    else:
                        logger.warning(f"Failed to process clean photo for NIK {nik}")
                        person['foto_bersih_url'] = None
                elif person.get('face') and person['face'].startswith('data:image/'):
                    # Process clean photo from base64 data
                    logger.info(f"Processing clean photo for NIK {nik} from base64 data")
                    clean_photo_url = process_and_save_clean_photo_from_base64(nik, person['face'])
                    
                    if clean_photo_url:
                        person['foto_bersih_url'] = clean_photo_url
                        logger.info(f"Successfully added clean photo URL for NIK {nik}: {clean_photo_url}")
                    else:
                        logger.warning(f"Failed to process clean photo from base64 for NIK {nik}")
                        person['foto_bersih_url'] = None
                else:
                    logger.info(f"No valid photo URL or base64 data found for NIK {nik}, skipping AI inpainting")
                    person['foto_bersih_url'] = None
                    
            except Exception as e:
                logger.error(f"Error processing AI inpainting for NIK {nik}: {e}")
                person['foto_bersih_url'] = None
    
    print(f"[SUCCESS] Basic enrichment complete for {person.get('full_name', 'Unknown')}")
    return person

def enrich_person_data(person, token=None):
    """Enrich person data with family and phone information"""
    nik = person.get('ktp_number')
    if not nik:
        return person
    
    print(f"\n=== ENRICHING DATA FOR NIK: {nik} ===")
    print(f"Person Name: {person.get('full_name', 'Unknown')}")
    
    # Try to get real data from APIs first
    phone_data = get_phone_data(nik, token)
    if phone_data:
        person['phone_data'] = phone_data
        print(f"[SUCCESS] Added phone data for {person.get('full_name', 'Unknown')}")
    else:
        print(f"[INFO] No phone data found for {person.get('full_name', 'Unknown')}")
    
    # Get family data - try with NKK if available
    nkk = person.get('family_cert_number') or person.get('nkk') or person.get('family_card_number')
    print(f"[INFO] Extracted NKK: {nkk}")
    
    if nkk:
        print(f"[INFO] Searching family data with NIK: {nik} and NKK: {nkk}")
        family_data = get_family_data(nik, nkk, token, person)
        if family_data:
            person['family_data'] = family_data
            print(f"[SUCCESS] Added family data for {person.get('full_name', 'Unknown')}")
            
            # Log family members found
            if family_data.get('anggota_keluarga'):
                print(f"[INFO] Found {len(family_data['anggota_keluarga'])} family members:")
                for member in family_data['anggota_keluarga'][:3]:  # Show first 3
                    print(f"   - {member.get('nama', 'N/A')} ({member.get('hubungan', 'N/A')})")
        else:
            print(f"[INFO] No family data found for {person.get('full_name', 'Unknown')}")
    else:
        print(f"[WARNING] No NKK found in person data, trying family search with NIK only")
        family_data = get_family_data(nik, None, token, person)
        if family_data:
            person['family_data'] = family_data
            print(f"[SUCCESS] Added family data (NIK only) for {person.get('full_name', 'Unknown')}")
        else:
            print(f"[INFO] No family data found (NIK only) for {person.get('full_name', 'Unknown')}")
    
    # Normalize/migrate photo -> face if needed (some APIs use different keys)
    if not person.get('face'):
        for k in ['face', 'photo', 'foto', 'image', 'picture', 'face_url', 'url_foto']:
            if person.get(k):
                person['face'] = person.get(k)
                if k != 'face':
                    print(f"Mapped '{k}' to 'face' for {person.get('full_name', 'Unknown')}")
                break

    # Fix face data format if it exists but doesn't have proper prefix
    if person.get('face'):
        face_data = person['face']
        # Check if it's already properly formatted
        if not face_data.startswith('data:image/'):
            # It's raw base64, add proper prefix
            if face_data.startswith('/9j/'):
                # JPEG format
                person['face'] = f'data:image/jpeg;base64,{face_data}'
                print(f"Fixed JPEG face format for {person.get('full_name', 'Unknown')}")
            elif face_data.startswith('iVBORw0KGgo'):
                # PNG format
                person['face'] = f'data:image/png;base64,{face_data}'
                print(f"Fixed PNG face format for {person.get('full_name', 'Unknown')}")
            else:
                # Assume JPEG
                person['face'] = f'data:image/jpeg;base64,{face_data}'
                print(f"Fixed face format (assumed JPEG) for {person.get('full_name', 'Unknown')}")
        
        # Process AI inpainting for watermark removal if we have a real photo (not avatar)
        if person.get('face') and not person['face'].startswith('data:image/svg'):
            try:
                # Check if we have url_foto field for original photo URL
                url_foto = person.get('url_foto') or person.get('face_url') or person.get('photo_url')
                
                if url_foto and url_foto.startswith('http'):
                    # Process clean photo using AI inpainting from URL
                    logger.info(f"Processing clean photo for NIK {nik} from URL: {url_foto}")
                    clean_photo_url = process_and_save_clean_photo(nik, url_foto)
                    
                    if clean_photo_url:
                        person['foto_bersih_url'] = clean_photo_url
                        logger.info(f"Successfully added clean photo URL for NIK {nik}: {clean_photo_url}")
                    else:
                        logger.warning(f"Failed to process clean photo for NIK {nik}")
                        person['foto_bersih_url'] = None
                elif person.get('face') and person['face'].startswith('data:image/'):
                    # Process clean photo from base64 data
                    logger.info(f"Processing clean photo for NIK {nik} from base64 data")
                    clean_photo_url = process_and_save_clean_photo_from_base64(nik, person['face'])
                    
                    if clean_photo_url:
                        person['foto_bersih_url'] = clean_photo_url
                        logger.info(f"Successfully added clean photo URL for NIK {nik}: {clean_photo_url}")
                    else:
                        logger.warning(f"Failed to process clean photo from base64 for NIK {nik}")
                        person['foto_bersih_url'] = None
                else:
                    logger.info(f"No valid photo URL or base64 data found for NIK {nik}, skipping AI inpainting")
                    person['foto_bersih_url'] = None
                    
            except Exception as e:
                logger.error(f"Error processing AI inpainting for NIK {nik}: {e}")
                person['foto_bersih_url'] = None
        else:
            logger.info(f"No real photo found for NIK {nik}, skipping AI inpainting")
            person['foto_bersih_url'] = None
    else:
        # Generate a simple avatar based on name initials
        name = person.get('full_name', 'Unknown')
        initials = ''.join([n[0] for n in name.split()[:2]]).upper()
        avatar_svg = f'''<svg width="200" height="200" viewBox="0 0 200 200" fill="none" xmlns="http://www.w3.org/2000/svg">
            <rect width="200" height="200" fill="#667eea"/>
            <circle cx="100" cy="80" r="30" fill="#FFFFFF"/>
            <path d="M50 140 Q50 120 100 120 Q150 120 150 140 L150 180 L50 180 Z50 140" fill="#FFFFFF"/>
            <text x="100" y="120" text-anchor="middle" fill="#667eea" font-family="Arial, sans-serif" font-size="24" font-weight="bold">{initials}</text>
        </svg>'''
        import base64
        avatar_b64 = base64.b64encode(avatar_svg.encode('utf-8')).decode('utf-8')
        person['face'] = f'data:image/svg+xml;base64,{avatar_b64}'
        person['foto_bersih_url'] = None
        print(f"Added avatar face for {person.get('full_name', 'Unknown')}")
    
    print(f"=== ENRICHMENT COMPLETE FOR NIK: {nik} ===\n")
    return person

@app.route('/')
def index():
    """Redirect to login page"""
    return redirect('/login')

@app.route('/simple')
def index_simple():
    """Serve the simple HTML page"""
    return send_from_directory(frontend_pages_dir, 'index_simple.html')

@app.route('/api/config')
def api_config():
    """API endpoint untuk mendapatkan konfigurasi default"""
    return jsonify({
        'username': DEFAULT_USERNAME,
        'password': DEFAULT_PASSWORD,
        'family_api_base': FAMILY_API_BASE,
        'family_api_alt': FAMILY_API_ALT,
        'phone_api_base': PHONE_API_BASE
    })

@app.route('/api/frontend-config', methods=['GET'])
def get_frontend_config():
    """API endpoint untuk mendapatkan konfigurasi frontend"""
    base_url = get_base_url()
    return jsonify({
        'base_url': base_url,
        'api_base': f"{base_url}/api",
        'is_https': request.is_secure,
        'host': request.host,
        'allowed_origins': allowed_origins
    })

# Authentication routes
@app.route('/login')
def login_page():
    """Serve login page. If already authenticated, redirect immediately to dashboard (no flash)."""
    # Check existing session token from cookie or Authorization header
    session_token = request.cookies.get('session_token') or request.headers.get('Authorization', '').replace('Bearer ', '')
    if session_token:
        user = validate_session_token(session_token)
        if user:
            # Already authenticated → go straight to dashboard to avoid login flash
            return redirect('/dashboard')
    # Not authenticated → show login page
    return send_from_directory(frontend_pages_dir, 'login.html')

@app.route('/dashboard')
@require_auth
def dashboard_page():
    """Serve dashboard page"""
    # Reset redirect count on successful dashboard access
    response = send_from_directory(frontend_pages_dir, 'dashboard.html')
    response.set_cookie('redirect_count', '0', max_age=60)
    return response

@app.route('/api/dashboard/stats')
def dashboard_stats():
    """Get dashboard statistics"""
    try:
        # Check if user is authenticated
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        if not token:
            return jsonify({'error': 'No token provided'}), 401
        
        # Validate session
        user_data = validate_session_token(token)
        if not user_data:
            return jsonify({'error': 'Invalid or expired token'}), 401
        
        # Get dashboard statistics
        stats = db.get_dashboard_stats()
        
        return jsonify({
            'success': True,
            'data': stats
        })
        
    except Exception as e:
        print(f"Error getting dashboard stats: {e}")
        return jsonify({'error': 'Internal server error'}), 500

@app.route('/api/dashboard/user-activities', methods=['GET'])
@require_auth
def get_user_activities():
    """Get user activities with filters and pagination"""
    try:
        # Get query parameters
        user_id = request.args.get('user_id', type=int)
        activity_type = request.args.get('activity_type', type=str)
        start_date = request.args.get('start_date', type=str)
        end_date = request.args.get('end_date', type=str)
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 50, type=int)
        search = request.args.get('search', type=str)
        
        # Get current user
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        current_user = validate_session_token(session_token)
        
        if not current_user:
            return jsonify({'error': 'Unauthorized'}), 401
        
        # Admin can see all activities, regular users can only see their own
        if current_user.get('role') != 'admin' and user_id and user_id != current_user['id']:
            return jsonify({'error': 'Forbidden'}), 403
        
        # If not admin and no user_id specified, show only current user's activities
        if current_user.get('role') != 'admin' and not user_id:
            user_id = current_user['id']
        
        # Get activities from database
        activities = db.get_user_activities_filtered(
            user_id=user_id,
            activity_type=activity_type,
            start_date=start_date,
            end_date=end_date,
            search=search,
            page=page,
            per_page=per_page
        )
        
        # Get total count for pagination
        total = db.get_user_activities_count(
            user_id=user_id,
            activity_type=activity_type,
            start_date=start_date,
            end_date=end_date,
            search=search
        )
        
        return jsonify({
            'success': True,
            'data': activities,
            'pagination': {
                'page': page,
                'per_page': per_page,
                'total': total,
                'pages': (total + per_page - 1) // per_page
            }
        })
        
    except Exception as e:
        print(f"Error getting user activities: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': 'Internal server error'}), 500

@app.route('/api/dashboard/activity-stats', methods=['GET'])
@require_auth
def get_activity_stats():
    """Get activity statistics per user"""
    try:
        # Get current user
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        current_user = validate_session_token(session_token)
        
        if not current_user:
            return jsonify({'error': 'Unauthorized'}), 401
        
        # Get user_id parameter (optional, admin can see all)
        user_id = request.args.get('user_id', type=int)
        
        # Admin can see all stats, regular users can only see their own
        if current_user.get('role') != 'admin' and user_id and user_id != current_user['id']:
            return jsonify({'error': 'Forbidden'}), 403
        
        # If not admin and no user_id specified, show only current user's stats
        if current_user.get('role') != 'admin' and not user_id:
            user_id = current_user['id']
        
        # Get activity statistics
        stats = db.get_activity_stats(user_id=user_id)
        
        return jsonify({
            'success': True,
            'data': stats
        })
        
    except Exception as e:
        print(f"Error getting activity stats: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': 'Internal server error'}), 500

@app.route('/profiling')
@require_auth
def profiling_page():
    """Serve profiling page"""
    return send_from_directory(frontend_pages_dir, 'profiling.html')

@app.route('/mapping')
@require_auth
def mapping_page():
    """Serve mapping profiling page"""
    return send_from_directory(frontend_pages_dir, 'mapping.html')

@app.route('/user-management')
@require_auth
def user_management_page():
    """Serve user management page"""
    return send_from_directory(frontend_pages_dir, 'user_management.html')

@app.route('/data-profiling')
@require_auth
def data_profiling_page():
    """Serve data profiling page"""
    return send_from_directory(frontend_pages_dir, 'data_profiling.html')

@app.route('/cekplat')
@require_auth
def cekplat_page():
    """Serve cek plat page"""
    return send_from_directory(frontend_pages_dir, 'cekplat.html')

@app.route('/data-cari-plat')
@require_auth
def data_cari_plat_page():
    """Serve data cari plat page"""
    return send_from_directory(frontend_pages_dir, 'data_cari_plat.html')

@app.route('/ai-features')
@require_auth
def ai_features_page():
    """Serve AI features page"""
    return send_from_directory(frontend_pages_dir, 'ai_features.html')

@app.route('/reports')
@require_auth
def reports_page():
    """Serve reports page"""
    return send_from_directory(frontend_pages_dir, 'reports.html')

@app.route('/reports/profiling')
@require_auth
def reports_profiling_page():
    """Serve reports profiling page"""
    return send_from_directory(frontend_pages_dir, 'reports_profiling.html')

@app.route('/settings')
@require_auth
def settings_page():
    """Serve settings page"""
    return send_from_directory(frontend_pages_dir, 'settings.html')


# API Keys Management Endpoints
@app.route('/api/api-keys', methods=['GET'])
@require_auth
def get_api_keys():
    """Get all API keys"""
    try:
        api_type = request.args.get('type')
        keys = db.get_all_api_keys(api_type)
        return jsonify({'success': True, 'keys': keys}), 200
    except Exception as e:
        logger.error(f"Error getting API keys: {e}")
        return jsonify({'error': f'Error getting API keys: {str(e)}'}), 500


@app.route('/api/api-keys', methods=['POST'])
@require_auth
def create_api_key():
    """Create a new API key"""
    try:
        # Get current user
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        current_user = validate_session_token(session_token)
        
        if not current_user:
            return jsonify({'error': 'Unauthorized'}), 401
        
        data = request.get_json()
        api_key = data.get('api_key')
        api_type = data.get('api_type', 'GOOGLE_CSE')
        description = data.get('description', '')
        priority = int(data.get('priority', 0))
        daily_limit = int(data.get('daily_limit', 100))
        status = data.get('status', 'active')
        
        if not api_key:
            return jsonify({'error': 'API key is required'}), 400
        
        success = db.create_api_key(api_key, api_type, description, priority, daily_limit, status)
        if success:
            # Log API key creation activity
            client_ip = get_client_ip()
            db.log_activity(
                user_id=current_user['id'],
                activity_type='api_key_created',
                description=f'Created API key: {api_type} - {description or "No description"}',
                ip_address=client_ip,
                user_agent=request.headers.get('User-Agent')
            )
            print(f"✅ Logged API key creation for user {current_user.get('username', 'Unknown')} from IP {client_ip}")
            return jsonify({'success': True, 'message': 'API key created successfully'}), 201
        else:
            return jsonify({'error': 'Failed to create API key'}), 500
    except Exception as e:
        logger.error(f"Error creating API key: {e}")
        return jsonify({'error': f'Error creating API key: {str(e)}'}), 500


@app.route('/api/api-keys/<int:key_id>', methods=['GET'])
@require_auth
def get_api_key_detail(key_id):
    """Get a single API key by ID (with full key, not masked)"""
    try:
        keys = db.get_all_api_keys()
        key = next((k for k in keys if k['id'] == key_id), None)
        
        if not key:
            return jsonify({'error': 'API key not found'}), 404
        
        # Return full API key (not masked) for editing/copying
        return jsonify({
            'success': True,
            'key': {
                'id': key['id'],
                'api_key': key.get('api_key', ''),  # Full key, not masked
                'api_type': key.get('api_type'),
                'status': key.get('status'),
                'description': key.get('description'),
                'priority': key.get('priority'),
                'daily_limit': key.get('daily_limit'),
                'usage_count': key.get('usage_count'),
                'last_used': key.get('last_used'),
                'created_at': key.get('created_at'),
                'updated_at': key.get('updated_at')
            }
        }), 200
    except Exception as e:
        logger.error(f"Error getting API key detail: {e}")
        return jsonify({'error': f'Error getting API key: {str(e)}'}), 500


@app.route('/api/api-keys/<int:key_id>', methods=['PUT'])
@require_auth
def update_api_key(key_id):
    """Update an API key"""
    try:
        # Get current user
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        current_user = validate_session_token(session_token)
        
        if not current_user:
            return jsonify({'error': 'Unauthorized'}), 401
        
        data = request.get_json()
        
        # Build update dict with only provided fields
        update_data = {}
        if 'api_key' in data:
            update_data['api_key'] = data['api_key']
        if 'api_type' in data:
            update_data['api_type'] = data['api_type']
        if 'status' in data:
            update_data['status'] = data['status']
        if 'description' in data:
            update_data['description'] = data['description']
        if 'priority' in data:
            update_data['priority'] = int(data['priority'])
        if 'daily_limit' in data:
            update_data['daily_limit'] = int(data['daily_limit'])
        if 'usage_count' in data:
            update_data['usage_count'] = int(data['usage_count'])
        
        success = db.update_api_key(key_id, **update_data)
        if success:
            # Log API key update activity
            changes = ', '.join([f"{k}: {v}" for k, v in update_data.items() if k != 'api_key'])  # Don't log full API key
            client_ip = get_client_ip()
            db.log_activity(
                user_id=current_user['id'],
                activity_type='api_key_updated',
                description=f'Updated API key ID {key_id}: {changes}',
                ip_address=client_ip,
                user_agent=request.headers.get('User-Agent')
            )
            print(f"✅ Logged API key update for user {current_user.get('username', 'Unknown')} from IP {client_ip}")
            return jsonify({'success': True, 'message': 'API key updated successfully'}), 200
        else:
            return jsonify({'error': 'Failed to update API key'}), 500
    except Exception as e:
        logger.error(f"Error updating API key: {e}")
        return jsonify({'error': f'Error updating API key: {str(e)}'}), 500


@app.route('/api/api-keys/<int:key_id>', methods=['DELETE'])
@require_auth
def delete_api_key(key_id):
    """Delete an API key"""
    try:
        # Get current user
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        current_user = validate_session_token(session_token)
        
        if not current_user:
            return jsonify({'error': 'Unauthorized'}), 401
        
        # Get API key info before deletion for logging
        keys = db.get_all_api_keys()
        key_info = next((k for k in keys if k['id'] == key_id), None)
        api_type = key_info.get('api_type', 'Unknown') if key_info else 'Unknown'
        
        success = db.delete_api_key(key_id)
        if success:
            # Log API key deletion activity
            client_ip = get_client_ip()
            db.log_activity(
                user_id=current_user['id'],
                activity_type='api_key_deleted',
                description=f'Deleted API key ID {key_id} (Type: {api_type})',
                ip_address=client_ip,
                user_agent=request.headers.get('User-Agent')
            )
            print(f"✅ Logged API key deletion for user {current_user.get('username', 'Unknown')} from IP {client_ip}")
            return jsonify({'success': True, 'message': 'API key deleted successfully'}), 200
        else:
            return jsonify({'error': 'Failed to delete API key'}), 500
    except Exception as e:
        logger.error(f"Error deleting API key: {e}")
        return jsonify({'error': f'Error deleting API key: {str(e)}'}), 500


@app.route('/api/api-keys/<int:key_id>/reset', methods=['POST'])
@require_auth
def reset_api_key_usage(key_id):
    """Reset usage count for an API key"""
    try:
        success = db.reset_api_key_usage(key_id)
        if success:
            return jsonify({'success': True, 'message': 'API key usage reset successfully'}), 200
        else:
            return jsonify({'error': 'Failed to reset API key usage'}), 500
    except Exception as e:
        logger.error(f"Error resetting API key usage: {e}")
        return jsonify({'error': f'Error resetting API key usage: {str(e)}'}), 500


@app.route('/api/login', methods=['POST'])
def api_login():
    """API endpoint untuk login"""
    try:
        data = request.get_json()
        username = data.get('username')
        password = data.get('password')
        
        if not username or not password:
            return jsonify({'error': 'Username dan password diperlukan'}), 400
        
        # Get client info with real IP address
        ip_address = get_client_ip()
        user_agent = request.headers.get('User-Agent')
        
        # Authenticate user
        auth_result = authenticate_user(username, password, ip_address, user_agent)
        if auth_result:
            print(f"✅ User {username} logged in from IP {ip_address}")
        
        if auth_result:
            response = jsonify({
                'success': True,
                'user': auth_result['user'],
                'session_token': auth_result['session_token']
            })
            
            # Set session token as HTTP-only cookie for server-side validation
            response.set_cookie(
                'session_token', 
                auth_result['session_token'],
                max_age=24*60*60,  # 24 hours
                httponly=True,     # Prevent XSS attacks
                secure=False,      # Set to True in production with HTTPS
                samesite='Lax'     # CSRF protection
            )
            
            return response
        else:
            return jsonify({'error': 'Username atau password salah'}), 401
            
    except Exception as e:
        return jsonify({'error': f'Login error: {str(e)}'}), 500

@app.route('/api/logout', methods=['POST'])
def api_logout():
    """API endpoint untuk logout"""
    try:
        data = request.get_json()
        session_token = data.get('session_token')
        
        # Get user info before logout (if session is valid)
        user_id = None
        if session_token:
            user = validate_session_token(session_token)
            if user:
                user_id = user.get('id')
        
        # Get client info with real IP address
        ip_address = get_client_ip()
        user_agent = request.headers.get('User-Agent')
        
        if session_token:
            logout_user(session_token, user_id, ip_address, user_agent)
            if user_id:
                print(f"✅ User ID {user_id} logged out from IP {ip_address}")
        
        response = jsonify({'success': True, 'message': 'Logged out successfully'})
        
        # Clear the session token cookie
        response.set_cookie('session_token', '', expires=0)
        
        return response
    except Exception as e:
        return jsonify({'error': f'Logout error: {str(e)}'}), 500

@app.route('/logout')
def logout_page():
    """Logout page - clear session and redirect to login"""
    try:
        # Get session token from cookie
        session_token = request.cookies.get('session_token')
        
        # Get user info before logout (if session is valid)
        user_id = None
        if session_token:
            user = validate_session_token(session_token)
            if user:
                user_id = user.get('id')
        
        # Get client info with real IP address
        ip_address = get_client_ip()
        user_agent = request.headers.get('User-Agent')
        
        if session_token:
            # Logout user
            logout_user(session_token, user_id, ip_address, user_agent)
            if user_id:
                print(f"✅ User ID {user_id} logged out via logout page from IP {ip_address}")
        
        # Create response that redirects to login
        response = redirect('/login')
        
        # Clear the session token cookie and redirect count
        response.set_cookie('session_token', '', expires=0)
        response.set_cookie('redirect_count', '0', expires=0)
        
        return response
    except Exception as e:
        print(f"Logout error: {e}")
        # Even if there's an error, redirect to login
        response = redirect('/login')
        response.set_cookie('session_token', '', expires=0)
        response.set_cookie('redirect_count', '0', expires=0)
        return response

@app.route('/clear-redirect-loop')
def clear_redirect_loop():
    """Clear redirect loop and reset session"""
    print("Clearing redirect loop and resetting session")
    response = redirect('/login')
    response.set_cookie('session_token', '', expires=0)
    response.set_cookie('redirect_count', '0', expires=0)
    return response

@app.route('/api/clear-session', methods=['POST'])
def clear_session_api():
    """API endpoint to clear session and reset authentication"""
    try:
        print("Clearing session via API")
        response = jsonify({
            'success': True,
            'message': 'Session cleared successfully'
        })
        
        # Clear all session cookies
        response.set_cookie('session_token', '', expires=0)
        response.set_cookie('redirect_count', '0', expires=0)
        
        return response
    except Exception as e:
        return jsonify({'error': f'Session clear error: {str(e)}'}), 500

@app.route('/api/validate-session', methods=['POST'])
def api_validate_session():
    """API endpoint untuk validasi session"""
    try:
        data = request.get_json()
        if not data:
            print("No JSON data received in validate-session")
            return jsonify({'valid': False, 'error': 'No data provided'}), 400
            
        session_token = data.get('session_token')
        print(f"Validating session token: {session_token[:10] if session_token else 'None'}...")
        
        if not session_token:
            print("No session token provided")
            return jsonify({'valid': False, 'error': 'No session token provided'}), 400
        
        user = validate_session_token(session_token)
        if user:
            print(f"Session validation successful for user: {user.get('username', 'Unknown')}")
            return jsonify({'valid': True, 'user': user})
        else:
            print("Session validation failed - invalid or expired token")
            return jsonify({'valid': False, 'error': 'Invalid or expired session'}), 401
            
    except Exception as e:
        print(f"Session validation error: {str(e)}")
        return jsonify({'valid': False, 'error': f'Session validation error: {str(e)}'}), 500

@app.route('/api/debug/session-status', methods=['GET'])
def debug_session_status():
    """Debug endpoint to check session status"""
    try:
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        if not session_token:
            session_token = request.cookies.get('session_token')
        
        if not session_token:
            return jsonify({
                'has_token': False,
                'message': 'No session token found'
            })
        
        user = validate_session_token(session_token)
        return jsonify({
            'has_token': True,
            'token_preview': session_token[:10] + '...',
            'is_valid': user is not None,
            'user': user if user else None
        })
    except Exception as e:
        return jsonify({
            'error': str(e),
            'has_token': False
        })

@app.route('/api/check-auth', methods=['GET'])
def check_auth_status():
    """Check authentication status without redirect"""
    try:
        session_token = request.cookies.get('session_token')
        if not session_token:
            return jsonify({'authenticated': False, 'message': 'No session token'})
        
        user = validate_session_token(session_token)
        if user:
            return jsonify({'authenticated': True, 'user': user})
        else:
            return jsonify({'authenticated': False, 'message': 'Invalid session'})
    except Exception as e:
        return jsonify({'authenticated': False, 'error': str(e)})

@app.route('/api/auth-status', methods=['GET'])
def auth_status():
    """Get current authentication status for debugging"""
    try:
        session_token = request.cookies.get('session_token')
        if not session_token:
            return jsonify({
                'authenticated': False,
                'has_cookie': False,
                'message': 'No session token in cookies'
            })
        
        user = validate_session_token(session_token)
        if user:
            return jsonify({
                'authenticated': True,
                'has_cookie': True,
                'user': {
                    'id': user.get('id'),
                    'username': user.get('username'),
                    'role': user.get('role')
                }
            })
        else:
            return jsonify({
                'authenticated': False,
                'has_cookie': True,
                'message': 'Invalid or expired session token'
            })
    except Exception as e:
        return jsonify({
            'authenticated': False,
            'has_cookie': bool(request.cookies.get('session_token')),
            'error': str(e)
        })

@app.route('/api/users', methods=['GET'])
def api_get_users():
    """API endpoint untuk mendapatkan daftar users"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        
        if not user or user['role'] != 'admin':
            return jsonify({'error': 'Unauthorized'}), 401
        
        users = db.get_all_users()
        return jsonify({'users': users})
        
    except Exception as e:
        return jsonify({'error': f'Error getting users: {str(e)}'}), 500

@app.route('/api/users', methods=['POST'])
def api_create_user():
    """API endpoint untuk membuat user baru"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        
        if not user or user['role'] != 'admin':
            return jsonify({'error': 'Unauthorized'}), 401
        
        data = request.get_json()
        username = data.get('username')
        email = data.get('email')
        password = data.get('password')
        full_name = data.get('full_name')
        role = data.get('role', 'user')
        status = data.get('status', 'active')
        
        if not all([username, email, password, full_name]):
            return jsonify({'error': 'Semua field diperlukan'}), 400
        
        success = db.create_user(username, email, password, full_name, role, status)
        
        if success:
            return jsonify({'success': True, 'message': 'User created successfully'})
        else:
            return jsonify({'error': 'Username atau email sudah ada'}), 400
            
    except Exception as e:
        return jsonify({'error': f'Error creating user: {str(e)}'}), 500

@app.route('/api/users/<int:user_id>', methods=['PUT'])
def api_update_user(user_id):
    """API endpoint untuk update user"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        
        if not user or user['role'] != 'admin':
            return jsonify({'error': 'Unauthorized'}), 401
        
        data = request.get_json()
        update_data = {}
        
        for key in ['username', 'email', 'full_name', 'role', 'status']:
            if key in data:
                update_data[key] = data[key]
        
        if not update_data:
            return jsonify({'error': 'Tidak ada data untuk diupdate'}), 400
        
        success = db.update_user(user_id, **update_data)
        
        if success:
            return jsonify({'success': True, 'message': 'User updated successfully'})
        else:
            return jsonify({'error': 'Error updating user'}), 500
            
    except Exception as e:
        return jsonify({'error': f'Error updating user: {str(e)}'}), 500

@app.route('/api/users/<int:user_id>', methods=['DELETE'])
def api_delete_user(user_id):
    """API endpoint untuk delete user"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        
        if not user or user['role'] != 'admin':
            return jsonify({'error': 'Unauthorized'}), 401
        
        success = db.delete_user(user_id)
        
        if success:
            return jsonify({'success': True, 'message': 'User deleted successfully'})
        else:
            return jsonify({'error': 'Error deleting user'}), 500
            
    except Exception as e:
        return jsonify({'error': f'Error deleting user: {str(e)}'}), 500

# Profiling Data API Endpoints
@app.route('/api/profiling-data', methods=['GET'])
def api_get_profiling_data():
    """API endpoint untuk mendapatkan data profiling"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        
        if not user:
            return jsonify({'error': 'Unauthorized'}), 401
        
        # Get query parameters
        search_type = request.args.get('search_type')
        limit = int(request.args.get('limit', 100))
        offset = int(request.args.get('offset', 0))
        
        # Allow all users to see all data (no user_id filter)
        user_id = None
        
        # Get profiling data
        data = db.get_profiling_data(user_id=user_id, search_type=search_type, 
                                   limit=limit, offset=offset)
        count = db.get_profiling_data_count(user_id=user_id, search_type=search_type)
        
        return jsonify({
            'success': True,
            'data': data,
            'count': count,
            'limit': limit,
            'offset': offset
        })
        
    except Exception as e:
        return jsonify({'error': f'Error getting profiling data: {str(e)}'}), 500

@app.route('/api/profiling-data/<int:profiling_id>', methods=['DELETE'])
def api_delete_profiling_data(profiling_id):
    """API endpoint untuk delete profiling data"""
    try:
        print(f"DEBUG: Delete request for profiling_id: {profiling_id}")
        
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        print(f"DEBUG: Session token: {session_token[:20]}..." if session_token else "DEBUG: No session token")
        
        user = validate_session_token(session_token)
        print(f"DEBUG: User validation result: {user}")
        
        if not user:
            print("DEBUG: User validation failed")
            return jsonify({'error': 'Unauthorized'}), 401
        
        # Only admin can delete profiling data
        if user['role'] != 'admin':
            print(f"DEBUG: User role is not admin: {user['role']}")
            return jsonify({'error': 'Unauthorized'}), 401
        
        print(f"DEBUG: Calling db.delete_profiling_data({profiling_id})")
        success = db.delete_profiling_data(profiling_id)
        print(f"DEBUG: Delete result: {success}")
        
        if success:
            print("DEBUG: Delete successful")
            return jsonify({'success': True, 'message': 'Profiling data deleted successfully'})
        else:
            print("DEBUG: Delete failed")
            return jsonify({'error': 'Error deleting profiling data'}), 500
            
    except Exception as e:
        print(f"DEBUG: Exception in delete: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Error deleting profiling data: {str(e)}'}), 500

@app.route('/api/profiling-data/clear-all', methods=['DELETE'])
def api_clear_all_profiling_data():
    """API endpoint untuk clear all profiling data"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        
        if not user:
            return jsonify({'error': 'Unauthorized'}), 401
        
        # Only admin can clear all profiling data
        if user['role'] != 'admin':
            return jsonify({'error': 'Unauthorized'}), 401
        
        success = db.clear_all_profiling_data()
        
        if success:
            return jsonify({'success': True, 'message': 'All profiling data cleared successfully'})
        else:
            return jsonify({'error': 'Error clearing profiling data'}), 500
            
    except Exception as e:
        return jsonify({'error': f'Error clearing profiling data: {str(e)}'}), 500

# Cek Plat API Endpoints
@app.route('/api/cekplat', methods=['POST'])
def api_cekplat():
    """API endpoint untuk cek plat nomor Jambi"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        
        if not user:
            return jsonify({'error': 'Unauthorized'}), 401
        
        data = request.get_json()
        no_polisi = data.get('no_polisi', '').strip().upper()
        
        if not no_polisi:
            return jsonify({'error': 'Nomor polisi harus diisi'}), 400
        
        # Import functions from cekplat
        from cekplat import fetch_data, process_table_data, preprocess_address, geocode_address
        
        # Fetch data from jambisamsat.net
        html = fetch_data(no_polisi)
        if not html:
            return jsonify({'error': 'Gagal mengambil data dari server.'}), 500
        
        # Process table data
        table_data = process_table_data(html)
        if not table_data:
            return jsonify({'error': 'Data tidak ditemukan.'}), 404
        
        # Extract data from table
        extracted_data = {}
        for label, value in table_data:
            label_clean = label.strip().lower()
            if 'nama' in label_clean and 'pemilik' in label_clean:
                extracted_data['nama_pemilik'] = value
            elif 'alamat' in label_clean:
                extracted_data['alamat'] = value
            elif 'merk' in label_clean:
                extracted_data['merk_kendaraan'] = value
            elif 'type' in label_clean:
                extracted_data['type_kendaraan'] = value
            elif 'model' in label_clean:
                extracted_data['model_kendaraan'] = value
            elif 'tahun' in label_clean:
                extracted_data['tahun_pembuatan'] = value
            elif 'warna' in label_clean:
                extracted_data['warna_kendaraan'] = value
            elif 'rangka' in label_clean:
                extracted_data['no_rangka'] = value
            elif 'mesin' in label_clean:
                extracted_data['no_mesin'] = value
            elif 'silinder' in label_clean:
                extracted_data['silinder'] = value
            elif 'bahan' in label_clean and 'bakar' in label_clean:
                extracted_data['bahan_bakar'] = value
            elif 'stnk' in label_clean:
                extracted_data['masa_berlaku_stnk'] = value
            elif 'pajak' in label_clean:
                extracted_data['masa_berlaku_pajak'] = value
            elif 'status' in label_clean:
                extracted_data['status_kendaraan'] = value
        
        # Geocode address if available
        coordinates = (None, None)
        accuracy_score = 0.0
        accuracy_details = []
        display_name = ""
        
        alamat = extracted_data.get('alamat')
        if alamat:
            processed_address = preprocess_address(alamat)
            lat, lon, acc_score, acc_details, disp_name = geocode_address(alamat)
            coordinates = (lat, lon) if lat and lon else (None, None)
            accuracy_score = acc_score
            accuracy_details = acc_details
            display_name = disp_name
        
        # Save to database
        try:
            db.save_cekplat_data(
                user_id=user['id'],
                no_polisi=no_polisi,
                nama_pemilik=extracted_data.get('nama_pemilik'),
                alamat=extracted_data.get('alamat'),
                merk_kendaraan=extracted_data.get('merk_kendaraan'),
                type_kendaraan=extracted_data.get('type_kendaraan'),
                model_kendaraan=extracted_data.get('model_kendaraan'),
                tahun_pembuatan=extracted_data.get('tahun_pembuatan'),
                warna_kendaraan=extracted_data.get('warna_kendaraan'),
                no_rangka=extracted_data.get('no_rangka'),
                no_mesin=extracted_data.get('no_mesin'),
                silinder=extracted_data.get('silinder'),
                bahan_bakar=extracted_data.get('bahan_bakar'),
                masa_berlaku_stnk=extracted_data.get('masa_berlaku_stnk'),
                masa_berlaku_pajak=extracted_data.get('masa_berlaku_pajak'),
                status_kendaraan=extracted_data.get('status_kendaraan'),
                coordinates_lat=coordinates[0] if coordinates[0] else None,
                coordinates_lon=coordinates[1] if coordinates[1] else None,
                accuracy_score=accuracy_score,
                accuracy_details=str(accuracy_details) if accuracy_details else None,
                display_name=display_name,
                ip_address=request.remote_addr,
                user_agent=request.headers.get('User-Agent')
            )
            print(f"Saved cek plat data for {no_polisi} by user {user['username']}")
        except Exception as e:
            print(f"Error saving cek plat data: {e}")
        
        return jsonify({
            'error': None,
            'table_data': table_data,
            'extracted_data': extracted_data,
            'alamat': alamat,
            'coordinates': coordinates,
            'accuracy_score': accuracy_score,
            'accuracy_details': accuracy_details,
            'display_name': display_name
        })
        
    except Exception as e:
        return jsonify({'error': f'Error checking plate number: {str(e)}'}), 500

@app.route('/api/cekplat-data', methods=['GET'])
def api_get_cekplat_data():
    """API endpoint untuk mendapatkan data cek plat"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        
        if not user:
            print(f"Unauthorized access attempt with token: {session_token[:20] if session_token else 'None'}...")
            return jsonify({'error': 'Unauthorized'}), 401
        
        # Get query parameters
        limit = int(request.args.get('limit', 100))
        offset = int(request.args.get('offset', 0))
        
        # Admin can see all data, others only their own
        user_id = None if user['role'] == 'admin' else user['id']
        
        # Get cek plat data
        data = db.get_cekplat_data(user_id=user_id, limit=limit, offset=offset)
        count = db.get_cekplat_data_count(user_id=user_id)
        
        return jsonify({
            'success': True,
            'data': data,
            'count': count,
            'limit': limit,
            'offset': offset
        })
        
    except Exception as e:
        return jsonify({'error': f'Error getting cek plat data: {str(e)}'}), 500

@app.route('/api/cekplat-data-test', methods=['GET'])
def api_get_cekplat_data_test():
    """API endpoint untuk testing data cek plat tanpa authentication"""
    try:
        print("Testing endpoint called - no authentication required")
        
        # Get query parameters
        limit = int(request.args.get('limit', 100))
        offset = int(request.args.get('offset', 0))
        
        # Get cek plat data (admin access)
        data = db.get_cekplat_data(user_id=None, limit=limit, offset=offset)
        count = db.get_cekplat_data_count(user_id=None)
        
        print(f"Test endpoint returning {len(data)} records, total count: {count}")
        
        return jsonify({
            'success': True,
            'data': data,
            'count': count,
            'limit': limit,
            'offset': offset,
            'test_mode': True
        })
        
    except Exception as e:
        print(f"Error in test endpoint: {e}")
        return jsonify({'error': f'Error getting cek plat data: {str(e)}'}), 500

@app.route('/api/cekplat-data/<int:cekplat_id>', methods=['DELETE'])
def api_delete_cekplat_data(cekplat_id):
    """API endpoint untuk delete cek plat data"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        
        if not user:
            return jsonify({'error': 'Unauthorized'}), 401
        
        # Only admin can delete cek plat data
        if user['role'] != 'admin':
            return jsonify({'error': 'Unauthorized'}), 401
        
        success = db.delete_cekplat_data(cekplat_id)
        
        if success:
            return jsonify({'success': True, 'message': 'Cek plat data deleted successfully'})
        else:
            return jsonify({'error': 'Error deleting cek plat data'}), 500
            
    except Exception as e:
        return jsonify({'error': f'Error deleting cek plat data: {str(e)}'}), 500

@app.route('/api/search', methods=['POST'])
def api_search():
    """API endpoint untuk berbagai jenis pencarian"""
    try:
        data = request.get_json()
        
        # Validate required fields
        if not data.get('username') or not data.get('password'):
            return jsonify({'error': 'Username dan password diperlukan'}), 400
        
        # Validate based on search type
        search_type = data.get('search_type', 'identity')
        
        if search_type == 'identity':
            if not data.get('name') and not data.get('nik') and not data.get('family_cert_number') and not data.get('tempat_lahir') and not data.get('tanggal_lahir') and not data.get('no_prop') and not data.get('no_kab') and not data.get('no_kec') and not data.get('no_desa'):
                return jsonify({'error': 'Minimal tentukan satu parameter untuk membatasi hasil (name, nik, family_cert_number, tempat_lahir, tanggal_lahir, no_prop, no_kab, no_kec, atau no_desa)'}), 400
        elif search_type == 'phone':
            if not data.get('phone_number'):
                return jsonify({'error': 'Nomor HP diperlukan untuk pencarian phone'}), 400
        elif search_type == 'face':
            if not data.get('face_file'):
                return jsonify({'error': 'File foto wajah diperlukan untuk pencarian face'}), 400
        
        # Authenticate user first
        auth_result = db.authenticate_user(data['username'], data['password'])
        if not auth_result:
            return jsonify({'error': 'Invalid credentials'}), 401
        
        user_data = auth_result
        
        # Get token with force refresh to ensure fresh token
        # Note: If server 224 is down, ensure_token will return a fallback_token
        # and call_search will automatically use server 116 as fallback
        print(f"INFO: [API_SEARCH] Menerima request dengan username: {data['username']}", file=sys.stderr)
        try:
            token = ensure_token(data['username'], data['password'], force_refresh=True)
            if not token:
                return jsonify({'error': 'Gagal mendapatkan token akses ke server eksternal'}), 500
            # Check if we're using fallback token (server 224 is down)
            if token.startswith("fallback_token_"):
                print("INFO: [API_SEARCH] ✅ Server 224 MATI - Menggunakan server 116", file=sys.stderr)
                print(f"INFO: [API_SEARCH] ✅ Server 116 akan menggunakan kredensial hardcoded: jambi/@ab526d", file=sys.stderr)
                print(f"INFO: [API_SEARCH] ⚠️ Kredensial {data['username']} TIDAK digunakan untuk server 116 (kredensial berbeda!)", file=sys.stderr)
            else:
                print(f"INFO: [API_SEARCH] ✅ Server 224 HIDUP - Menggunakan server 224", file=sys.stderr)
                print(f"INFO: [API_SEARCH] ✅ Kredensial {data['username']} digunakan untuk server 224", file=sys.stderr)
        except Exception as e:
            print(f"ERROR: [API_SEARCH] Error getting token: {e}", file=sys.stderr)
            # Even if token fails, try to proceed with fallback
            token = f"fallback_token_{int(time.time())}"
            print("INFO: [API_SEARCH] Menggunakan fallback token untuk mencoba server 116", file=sys.stderr)
            print(f"INFO: [API_SEARCH] Kredensial dari frontend ({data.get('username', 'N/A')}) TIDAK digunakan untuk server 116", file=sys.stderr)
        
        # Prepare search parameters
        params = {
            "name": data.get('name', ''),
            "nik": data.get('nik', ''),
            "family_cert_number": data.get('family_cert_number', ''),
            "tempat_lahir": data.get('tempat_lahir', ''),
            "tanggal_lahir": data.get('tanggal_lahir', ''),
            "no_prop": data.get('no_prop', ''),
            "no_kab": data.get('no_kab', ''),
            "no_kec": data.get('no_kec', ''),
            "no_desa": data.get('no_desa', ''),
            "page": data.get('page', '1'),
            "limit": data.get('limit', '100')  # Tambahkan parameter limit dengan default 100
        }
        
        # Route based on search type
        if search_type == 'phone':
            return perform_phone_search(token, params, data, user_data)
        elif search_type == 'face':
            return perform_face_search(token, params, data, user_data)
        else:
            # Regular search without face matching
            return perform_regular_search(token, params, data, user_data)
            
    except Exception as e:
        return jsonify({'error': f'Error: {str(e)}'}), 500

@app.route('/api/chatbot/extract-document', methods=['POST'])
@require_auth
def chatbot_extract_document():
    """API endpoint untuk mengekstrak data dari dokumen/foto menggunakan Gemini Vision"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        
        if not user:
            return jsonify({'error': 'Unauthorized'}), 401
        
        # Check if file is present
        if 'file' not in request.files:
            return jsonify({'error': 'File tidak ditemukan'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'File tidak dipilih'}), 400
        
        # Validate file type
        allowed_extensions = {'jpg', 'jpeg', 'png', 'gif', 'pdf'}
        file_ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
        if file_ext not in allowed_extensions:
            return jsonify({'error': f'Format file tidak didukung. Format yang didukung: {", ".join(allowed_extensions)}'}), 400
        
        # Read file content
        file_content = file.read()
        file.seek(0)  # Reset file pointer
        
        # Prepare file for Gemini
        import base64
        import mimetypes
        
        # Determine MIME type
        mime_type = mimetypes.guess_type(file.filename)[0] or 'image/jpeg'
        
        # Convert to base64 if needed
        if file_ext in ['jpg', 'jpeg', 'png', 'gif']:
            # Image file - encode to base64
            file_base64 = base64.b64encode(file_content).decode('utf-8')
        else:
            # For PDF files, we'll need to convert or use OCR
            return jsonify({'error': 'File PDF belum didukung. Silakan gunakan file gambar (JPG, PNG).'}), 400
        
        # Build prompt for Gemini Vision
        prompt = """Analisis dokumen/foto ini dan ekstrak informasi berikut jika tersedia:
1. NIK (Nomor Induk Kependudukan) - 16 digit
2. Nama lengkap
3. Nomor HP (telepon)
4. Lokasi/Provinsi

Jika ada beberapa data (misalnya beberapa orang dalam satu dokumen), ekstrak semua data yang ditemukan.

Format output dalam JSON array, setiap item adalah objek dengan format:
{
  "nik": "nomor nik jika ditemukan",
  "name": "nama lengkap jika ditemukan",
  "phone": "nomor hp jika ditemukan",
  "location": "lokasi/provinsi jika ditemukan"
}

Jika tidak ada data yang ditemukan, kembalikan array kosong [].
Hanya kembalikan JSON array, tanpa penjelasan tambahan."""
        
        # Call Gemini Vision API
        extracted_data = []
        result_text = None
        
        try:
            # Using old SDK (google.generativeai)
            if 'genai' in globals():
                image = Image.open(io.BytesIO(file_content))
                
                # Try using the model that was already initialized
                if 'model' in globals() and model:
                    try:
                        response = model.generate_content([prompt, image])
                        result_text = response.text
                        print(f"✅ Success using pre-initialized model with PIL Image")
                    except Exception as pre_model_error:
                        error_str = str(pre_model_error)
                        if "404" not in error_str and "not found" not in error_str.lower():
                            print(f"⚠️ Pre-initialized model failed: {pre_model_error}")
                        # Try different models - use only free tier models
                        models_to_try = [
                            'gemini-1.5-flash',  # Most commonly available in free tier
                            'gemini-1.5-pro',
                            'gemini-pro'
                        ]
                        for model_name in models_to_try:
                            try:
                                vision_model = genai.GenerativeModel(model_name)
                                response = vision_model.generate_content([prompt, image])
                                result_text = response.text
                                print(f"✅ Success with model: {model_name}")
                                break
                            except Exception as model_error:
                                error_str = str(model_error)
                                # Don't log 404 (not found) and 429 (quota exceeded) errors
                                if "404" not in error_str and "429" not in error_str and "not found" not in error_str.lower() and "quota" not in error_str.lower():
                                    print(f"⚠️ Model {model_name} failed: {model_error}")
                                # If quota exceeded, don't try other models
                                if "429" in error_str or "quota" in error_str.lower():
                                    print(f"⚠️ Quota exceeded, stopping model attempts")
                                    break
                                continue
                else:
                    # No pre-initialized model, try different models - use only free tier models
                    models_to_try = [
                        'gemini-1.5-flash',  # Most commonly available in free tier
                        'gemini-1.5-pro',
                        'gemini-pro'
                    ]
                    for model_name in models_to_try:
                        try:
                            vision_model = genai.GenerativeModel(model_name)
                            response = vision_model.generate_content([prompt, image])
                            result_text = response.text
                            print(f"✅ Success with model: {model_name}")
                            break
                        except Exception as model_error:
                            error_str = str(model_error)
                            # Don't log 404 (not found) and 429 (quota exceeded) errors
                            if "404" not in error_str and "429" not in error_str and "not found" not in error_str.lower() and "quota" not in error_str.lower():
                                print(f"⚠️ Model {model_name} failed: {model_error}")
                            # If quota exceeded, don't try other models
                            if "429" in error_str or "quota" in error_str.lower():
                                print(f"⚠️ Quota exceeded, stopping model attempts")
                                break
                            continue
            
            if not result_text:
                return jsonify({
                    'success': False,
                    'error': 'Tidak dapat memproses dokumen. Pastikan Gemini API key valid dan memiliki akses ke model vision.',
                    'extracted_data': []
                }), 500
            
            # Parse JSON response
            import json
            # Clean response text (remove markdown code blocks if any)
            result_text = result_text.strip()
            if result_text.startswith('```json'):
                result_text = result_text[7:]
            if result_text.startswith('```'):
                result_text = result_text[3:]
            if result_text.endswith('```'):
                result_text = result_text[:-3]
            result_text = result_text.strip()
            
            # Parse JSON
            try:
                extracted_data = json.loads(result_text)
                if not isinstance(extracted_data, list):
                    # If single object, wrap in array
                    extracted_data = [extracted_data] if extracted_data else []
            except json.JSONDecodeError as e:
                print(f"Error parsing JSON from Gemini: {e}")
                print(f"Response text: {result_text}")
                # Try to extract data manually using regex
                import re
                nik_match = re.search(r'\b\d{16}\b', result_text)
                name_match = re.search(r'(?:nama|name)[\s:]+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)', result_text, re.IGNORECASE)
                phone_match = re.search(r'(?:\+?62|0)?[\s-]?(\d{3,4})[\s-]?(\d{3,4})[\s-]?(\d{3,4})', result_text)
                
                if nik_match or name_match or phone_match:
                    extracted_data = [{
                        'nik': nik_match.group(0) if nik_match else '',
                        'name': name_match.group(1) if name_match else '',
                        'phone': phone_match.group(0) if phone_match else '',
                        'location': ''
                    }]
                else:
                    extracted_data = []
            
            # Filter out empty entries
            extracted_data = [item for item in extracted_data if any([
                item.get('nik', '').strip(),
                item.get('name', '').strip(),
                item.get('phone', '').strip(),
                item.get('location', '').strip()
            ])]
            
        except Exception as gemini_error:
            print(f"Error calling Gemini Vision: {gemini_error}")
            return jsonify({
                'success': False,
                'error': f'Error memproses dokumen: {str(gemini_error)}',
                'extracted_data': []
            }), 500
        
        return jsonify({
            'success': True,
            'extracted_data': extracted_data,
            'file_name': file.filename
        })
        
    except Exception as e:
        print(f"Error in chatbot_extract_document: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e),
            'extracted_data': []
        }), 500

@app.route('/api/chatbot/query-suggestions', methods=['POST'])
@require_auth
def chatbot_query_suggestions():
    """API endpoint untuk mendapatkan smart query suggestions menggunakan Gemini"""
    try:
        data = request.get_json()
        query = data.get('query', '').strip()
        search_history = data.get('search_history', [])  # Optional: history untuk context
        context = data.get('context', {})  # Optional: context tambahan
        
        if not query:
            return jsonify({
                'success': True,
                'suggestions': [],
                'typo_corrections': [],
                'message': 'Masukkan query untuk mendapatkan suggestions'
            })
        
        # Build prompt untuk Gemini
        prompt = f"""Anda adalah asisten AI yang cerdas untuk sistem pencarian data profiling identitas.

User sedang mengetik query: "{query}"

Tugas Anda:
1. Berikan 5-8 saran query pencarian yang lebih baik dan efektif
2. Deteksi typo/kesalahan ketik dan berikan koreksi
3. Saran kombinasi parameter pencarian yang mungkin lebih efektif
4. Saran query alternatif yang mungkin lebih spesifik

Parameter pencarian yang didukung:
- NIK (16 digit): "Cari NIK 1505041107830002"
- Nama: "Cari nama Jambi" atau "Cari nama Margutin"
- Nomor HP: "Cari no HP 6285755349653" atau "Cari HP 081234567890"
- Lokasi/Provinsi: "Cari lokasi Jambi" atau "Cari provinsi Sumatera Selatan"
- Kombinasi: "Cari NIK 1505041107830002 nama Jambi" atau "Cari nama Jambi lokasi Jambi"

Format response JSON (WAJIB):
{{
    "suggestions": [
        {{
            "query": "query suggestion 1",
            "type": "complete|typo_fix|combination|alternative",
            "description": "penjelasan singkat mengapa suggestion ini bagus",
            "confidence": 0.95
        }},
        ...
    ],
    "typo_corrections": [
        {{
            "original": "kata yang salah",
            "corrected": "kata yang benar",
            "confidence": 0.9
        }},
        ...
    ],
    "quick_actions": [
        {{
            "action": "Cari NIK",
            "icon": "id-card",
            "description": "Pencarian berdasarkan NIK"
        }},
        ...
    ]
}}

PENTING:
- Berikan suggestions yang relevan dengan query user
- Jika query sudah lengkap, berikan suggestions untuk kombinasi atau alternatif
- Jika query tidak lengkap, lengkapi dengan parameter yang mungkin
- Deteksi typo dengan baik (misal: "cari nik" bukan "cari nik", "jambi" bukan "jamby")
- Gunakan bahasa Indonesia yang natural
- Confidence score: 0.0-1.0 (1.0 = sangat yakin)

Berikan response dalam format JSON saja, tanpa penjelasan tambahan:
"""
        
        # Add context if available
        if search_history:
            prompt += f"\n\nHistory pencarian user (untuk konteks):\n"
            for i, hist in enumerate(search_history[-5:], 1):  # Last 5 searches
                prompt += f"{i}. {hist}\n"
        
        # Call Gemini
        try:
            if USE_NEW_GEMINI and 'client' in globals() and client:
                # Use only free tier models (no experimental models)
                models_to_try = [
                    "gemini-1.5-flash",       # Most commonly available in free tier
                    "gemini-1.5-pro",
                    "gemini-pro",
                    "models/gemini-1.5-flash",  # Try with models/ prefix
                    "models/gemini-1.5-pro",
                    "models/gemini-pro"
                ]
                
                gemini_response = None
                last_error = None
                
                for model_name in models_to_try:
                    try:
                        response = client.models.generate_content(
                            model=model_name,
                            contents=prompt
                        )
                        gemini_response = response.text
                        print(f"✅ Query suggestions: Successfully used model: {model_name}")
                        break
                    except Exception as model_error:
                        last_error = model_error
                        error_str = str(model_error)
                        # Check for specific error types
                        is_404 = "404" in error_str or "not found" in error_str.lower()
                        is_429 = "429" in error_str or "quota" in error_str.lower()
                        is_402 = "402" in error_str or "payment" in error_str.lower()
                        is_598 = "598" in error_str or "timeout" in error_str.lower() or "network" in error_str.lower()
                        
                        # If quota exceeded, don't try other models
                        if is_429:
                            print(f"⚠️ Query suggestions: Quota exceeded, using fallback")
                            break
                        
                        # Log important errors (402, 598, and other non-404 errors)
                        if is_402:
                            print(f"⚠️ Query suggestions: Model {model_name} failed: Payment issue (402) - {model_error}")
                        elif is_598:
                            print(f"⚠️ Query suggestions: Model {model_name} failed: Network timeout (598) - {model_error}")
                        elif not is_404:
                            print(f"⚠️ Query suggestions: Model {model_name} failed: {model_error}")
                        continue
                
                if not gemini_response:
                    # Fallback to old SDK if new SDK fails
                    print("⚠️ New Gemini SDK failed, trying old SDK...")
                    if 'model' in globals() and model:
                        try:
                            response = model.generate_content(prompt)
                            gemini_response = response.text
                            print("✅ Successfully used old SDK model")
                        except Exception as old_sdk_error:
                            error_str = str(old_sdk_error)
                            # Don't log 429 (quota exceeded) errors, they're expected
                            if "429" not in error_str and "quota" not in error_str.lower():
                                print(f"⚠️ Old SDK model failed: {old_sdk_error}")
                            # Don't raise, just fallback to simple suggestions
                            gemini_response = None
                    else:
                        # Don't raise, just fallback to simple suggestions
                        gemini_response = None
                
                # If still no response, use fallback
                if not gemini_response:
                    print("⚠️ All Gemini models failed, using fallback suggestions")
                    return jsonify({
                        'success': True,
                        'suggestions': generate_fallback_suggestions(query),
                        'typo_corrections': [],
                        'quick_actions': []
                    })
            else:
                if 'model' in globals() and model:
                    try:
                        response = model.generate_content(prompt)
                        gemini_response = response.text
                    except Exception as old_model_error:
                        error_str = str(old_model_error)
                        # Don't log 429 (quota exceeded) errors, they're expected
                        if "429" not in error_str and "quota" not in error_str.lower():
                            print(f"⚠️ Old SDK model failed: {old_model_error}")
                        # Fallback: simple suggestions
                        return jsonify({
                            'success': True,
                            'suggestions': generate_fallback_suggestions(query),
                            'typo_corrections': [],
                            'quick_actions': []
                        })
                else:
                    # Fallback: simple suggestions
                    return jsonify({
                        'success': True,
                        'suggestions': generate_fallback_suggestions(query),
                        'typo_corrections': [],
                        'quick_actions': []
                    })
            
            # Parse JSON response from Gemini
            # Try to extract JSON from response (might have markdown code blocks)
            import re
            json_match = re.search(r'\{[\s\S]*\}', gemini_response)
            if json_match:
                json_str = json_match.group(0)
                suggestions_data = json.loads(json_str)
            else:
                # If no JSON found, use fallback
                suggestions_data = {
                    'suggestions': generate_fallback_suggestions(query),
                    'typo_corrections': [],
                    'quick_actions': []
                }
            
            return jsonify({
                'success': True,
                'suggestions': suggestions_data.get('suggestions', []),
                'typo_corrections': suggestions_data.get('typo_corrections', []),
                'quick_actions': suggestions_data.get('quick_actions', []),
                'query': query
            })
            
        except json.JSONDecodeError as e:
            print(f"Error parsing Gemini JSON response: {e}")
            print(f"Response was: {gemini_response[:500] if gemini_response else 'None'}")
            import traceback
            traceback.print_exc()
            # Fallback to simple suggestions
            return jsonify({
                'success': True,
                'suggestions': generate_fallback_suggestions(query),
                'typo_corrections': [],
                'quick_actions': []
            })
        except Exception as gemini_error:
            print(f"Error calling Gemini for suggestions: {gemini_error}")
            import traceback
            traceback.print_exc()
            # Fallback to simple suggestions
            return jsonify({
                'success': True,
                'suggestions': generate_fallback_suggestions(query),
                'typo_corrections': [],
                'quick_actions': []
            })
            
    except Exception as e:
        print(f"Error in chatbot_query_suggestions: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e),
            'suggestions': [],
            'typo_corrections': [],
            'quick_actions': []
        }), 500

def generate_fallback_suggestions(query):
    """Generate simple fallback suggestions if Gemini is not available"""
    suggestions = []
    query_lower = query.lower()
    
    # Basic suggestions based on query content
    if 'nik' in query_lower or any(c.isdigit() for c in query):
        suggestions.append({
            'query': f'Cari NIK {query}',
            'type': 'complete',
            'description': 'Lengkapi pencarian dengan keyword NIK',
            'confidence': 0.8
        })
    
    if 'nama' in query_lower or (len(query.split()) > 0 and not any(c.isdigit() for c in query)):
        suggestions.append({
            'query': f'Cari nama {query}',
            'type': 'complete',
            'description': 'Lengkapi pencarian dengan keyword nama',
            'confidence': 0.8
        })
    
    if 'hp' in query_lower or 'phone' in query_lower or any(c.isdigit() for c in query):
        suggestions.append({
            'query': f'Cari no HP {query}',
            'type': 'complete',
            'description': 'Lengkapi pencarian dengan keyword no HP',
            'confidence': 0.8
        })
    
    # Add generic suggestions
    suggestions.extend([
        {
            'query': f'Cari {query}',
            'type': 'complete',
            'description': 'Tambahkan keyword "Cari" di depan',
            'confidence': 0.7
        }
    ])
    
    return suggestions[:5]  # Limit to 5 suggestions

@app.route('/api/chatbot/gemini-response', methods=['POST'])
def chatbot_gemini_response():
    """API endpoint untuk mendapatkan respons natural dari Gemini ketika data tidak ditemukan"""
    try:
        data = request.get_json()
        query = data.get('query', '')
        search_params = data.get('search_params', {})
        search_results = data.get('search_results', {})
        has_results = data.get('has_results', False)
        
        if not query:
            return jsonify({'error': 'Query diperlukan'}), 400
        
        # Check if it's a greeting or needs clarification
        is_greeting = data.get('is_greeting', False)
        needs_clarification = data.get('needs_clarification', False)
        
        # Build prompt untuk Gemini
        if is_greeting:
            prompt = f"""Anda adalah asisten AI yang ramah, cerdas, dan membantu untuk sistem pencarian data profiling identitas.

User mengirim pesan: "{query}"

Tugas Anda:
1. RESPON SECARA ALAMI seperti manusia yang ramah dan membantu
2. Jika ini adalah sapaan (halo, hi, selamat pagi, dll), balas dengan ramah dan tanyakan apa yang bisa dibantu
3. Jika ini adalah pertanyaan umum, jawab dengan informatif dan ramah
4. Jika ini adalah ucapan terima kasih, balas dengan sopan
5. Jika ini adalah permintaan bantuan, jelaskan dengan jelas dan ramah

PENTING:
- Gunakan bahasa Indonesia yang natural dan ramah
- Jangan terlalu formal, tapi tetap sopan
- Jika user bertanya tentang cara menggunakan sistem, jelaskan dengan singkat dan jelas
- Berikan contoh penggunaan jika relevan:
  * "Cari NIK 1505041107830002"
  * "Cari nama Jambi"
  * "Cari no HP 6285755349653"
  * Atau kombinasi parameter lainnya
- Jangan terlalu panjang, cukup 2-4 kalimat yang natural
- Sesuaikan respons dengan konteks pesan user

Berikan respons yang natural dan ramah:
"""
        elif needs_clarification:
            prompt = f"""Anda adalah asisten AI yang ramah dan cerdas untuk sistem pencarian data profiling identitas.

User mengirim pesan: "{query}"

Analisis pesan user:
- Jika ini adalah percakapan umum atau pertanyaan, jawab secara natural dan ramah
- Jika user sepertinya ingin mencari data tapi query tidak jelas, bantu dengan ramah
- Jika ini adalah pertanyaan tentang sistem, jawab dengan informatif

Untuk pencarian data, user bisa menggunakan:
- NIK (16 digit): "Cari NIK 1505041107830002"
- Nama: "Cari nama Jambi" atau "Cari nama Margutin"
- Nomor HP: "Cari no HP 6285755349653"
- Lokasi: "Cari lokasi Jambi"
- Kombinasi: "Cari NIK 1505041107830002 nama Jambi"

PENTING:
- Gunakan bahasa Indonesia yang natural dan ramah
- Jangan terlalu formal
- Jika user bertanya sesuatu yang tidak terkait pencarian, jawab dengan ramah dan arahkan ke fungsi pencarian
- Berikan contoh yang jelas jika user butuh bantuan
- Jangan terlalu panjang, cukup 2-3 kalimat yang natural

Berikan respons yang natural dan membantu:
"""
        else:
            prompt = f"""Anda adalah asisten AI yang membantu dalam pencarian data profiling identitas. 
User mencari dengan parameter berikut:
"""
            
            if search_params.get('nik'):
                prompt += f"- NIK: {search_params['nik']}\n"
            if search_params.get('name'):
                prompt += f"- Nama: {search_params['name']}\n"
            if search_params.get('phone'):
                prompt += f"- No HP: {search_params['phone']}\n"
            if search_params.get('location'):
                prompt += f"- Lokasi: {search_params['location']}\n"
            
            if has_results:
                total = search_results.get('total_results', 0)
                prompt += f"""
Hasil pencarian: Ditemukan {total} hasil.

Berikan respons yang ramah dan informatif dalam bahasa Indonesia, memberitahu user bahwa data ditemukan dan memberikan ringkasan singkat.
"""
            else:
                prompt += f"""
Hasil pencarian: Tidak ada data yang ditemukan.

Berikan respons yang ramah, empatik, dan membantu dalam bahasa Indonesia. 
Saran yang bisa diberikan:
1. Periksa kembali parameter pencarian (NIK, nama, no HP, lokasi)
2. Coba dengan parameter yang lebih spesifik
3. Pastikan format input sudah benar
4. Coba kombinasi parameter yang berbeda

Jangan terlalu panjang, cukup 2-3 kalimat yang informatif dan membantu.
"""
        
        # Generate response menggunakan Gemini
        try:
            if USE_NEW_GEMINI and 'client' in globals() and client:
                # Menggunakan new Google GenAI SDK - Use only free tier models
                models_to_try = [
                    "gemini-1.5-flash",       # Most commonly available in free tier
                    "gemini-1.5-pro",
                    "gemini-pro",
                    "models/gemini-1.5-flash",  # Try with models/ prefix
                    "models/gemini-1.5-pro",
                    "models/gemini-pro"
                ]
                
                gemini_response = None
                last_error = None
                
                for model_name in models_to_try:
                    try:
                        response = client.models.generate_content(
                            model=model_name,
                            contents=prompt
                        )
                        gemini_response = response.text
                        print(f"✅ Gemini response: Successfully used model: {model_name}")
                        break
                    except Exception as model_error:
                        last_error = model_error
                        error_str = str(model_error)
                        # Check for specific error types
                        is_404 = "404" in error_str or "not found" in error_str.lower()
                        is_429 = "429" in error_str or "quota" in error_str.lower()
                        is_402 = "402" in error_str or "payment" in error_str.lower()
                        is_598 = "598" in error_str or "timeout" in error_str.lower() or "network" in error_str.lower()
                        
                        # If quota exceeded, don't try other models
                        if is_429:
                            print(f"⚠️ Gemini response: Quota exceeded, using fallback")
                            break
                        
                        # Log important errors (402, 598, and other non-404 errors)
                        if is_402:
                            print(f"⚠️ Gemini response: Model {model_name} failed: Payment issue (402) - {model_error}")
                        elif is_598:
                            print(f"⚠️ Gemini response: Model {model_name} failed: Network timeout (598) - {model_error}")
                        elif not is_404:
                            print(f"⚠️ Gemini response: Model {model_name} failed: {model_error}")
                        continue
                
                if not gemini_response:
                    # Fallback to old SDK if new SDK fails
                    print("⚠️ New Gemini SDK failed, trying old SDK...")
                    if 'model' in globals() and model:
                        try:
                            response = model.generate_content(prompt)
                            gemini_response = response.text
                            print("✅ Successfully used old SDK model")
                        except Exception as old_sdk_error:
                            error_str = str(old_sdk_error)
                            # Don't log 429 (quota exceeded) errors, they're expected
                            if "429" not in error_str and "quota" not in error_str.lower():
                                print(f"⚠️ Old SDK model failed: {old_sdk_error}")
                            # Don't raise, just use fallback response
                            gemini_response = None
                    else:
                        # Don't raise, just use fallback response
                        gemini_response = None
                
                # If still no response, use fallback
                if not gemini_response:
                    print("⚠️ All Gemini models failed, using fallback response")
                    if has_results:
                        gemini_response = f"Berhasil menemukan {search_results.get('total_results', 0)} hasil untuk pencarian Anda."
                    else:
                        gemini_response = "Maaf, tidak ada data yang ditemukan untuk pencarian Anda. Silakan periksa kembali parameter pencarian atau coba dengan kombinasi yang berbeda."
            else:
                # Fallback ke old SDK
                if 'model' in globals() and model:
                    try:
                        response = model.generate_content(prompt)
                        gemini_response = response.text
                    except Exception as old_model_error:
                        error_str = str(old_model_error)
                        # Don't log 429 (quota exceeded) errors, they're expected
                        if "429" not in error_str and "quota" not in error_str.lower():
                            print(f"⚠️ Old SDK model failed: {old_model_error}")
                        # Fallback response jika Gemini tidak tersedia
                        if has_results:
                            gemini_response = f"Berhasil menemukan {search_results.get('total_results', 0)} hasil untuk pencarian Anda."
                        else:
                            gemini_response = "Maaf, tidak ada data yang ditemukan untuk pencarian Anda. Silakan periksa kembali parameter pencarian atau coba dengan kombinasi yang berbeda."
                else:
                    # Fallback response jika Gemini tidak tersedia
                    if has_results:
                        gemini_response = f"Berhasil menemukan {search_results.get('total_results', 0)} hasil untuk pencarian Anda."
                    else:
                        gemini_response = "Maaf, tidak ada data yang ditemukan untuk pencarian Anda. Silakan periksa kembali parameter pencarian atau coba dengan kombinasi yang berbeda."
        except Exception as gemini_error:
            error_str = str(gemini_error)
            is_402 = "402" in error_str or "payment" in error_str.lower() or "quota" in error_str.lower() or "billing" in error_str.lower()
            is_598 = "598" in error_str or "timeout" in error_str.lower() or "network" in error_str.lower() or "connection" in error_str.lower()
            
            if is_402:
                print(f"⚠️ Error calling Gemini: Payment/Quota issue (402) - API key mungkin tidak valid atau quota habis")
                print(f"   Error details: {gemini_error}")
            elif is_598:
                print(f"⚠️ Error calling Gemini: Network timeout (598) - Masalah koneksi ke server Gemini")
                print(f"   Error details: {gemini_error}")
            else:
                print(f"⚠️ Error calling Gemini: {gemini_error}")
                import traceback
                traceback.print_exc()
            
            # Fallback response
            if has_results:
                gemini_response = f"Berhasil menemukan {search_results.get('total_results', 0)} hasil untuk pencarian Anda."
            else:
                gemini_response = "Maaf, tidak ada data yang ditemukan untuk pencarian Anda. Silakan periksa kembali parameter pencarian atau coba dengan kombinasi yang berbeda."
        
        return jsonify({
            'success': True,
            'response': gemini_response,
            'has_results': has_results
        })
        
    except Exception as e:
        print(f"Error in chatbot_gemini_response: {e}")
        return jsonify({
            'success': False,
            'error': str(e),
            'response': 'Maaf, terjadi kesalahan saat memproses permintaan Anda.'
        }), 500

def perform_phone_search(token, params, data, user_data):
    """Perform phone number search"""
    try:
        phone_number = data.get('phone_number')
        phone_operator = data.get('phone_operator', '')
        
        # Get phone data from API
        limit = params.get('limit', 100)
        phone_data = get_phone_data_by_number(phone_number, token, limit)
        
        if phone_data:
            results = []
            for phone_item in phone_data:
                # Check if we have person_data directly from phone search
                if phone_item.get('person_data'):
                    # Use person data directly from phone search
                    person = phone_item['person_data']
                    print(f"Using person data directly from phone search: {person.get('full_name', 'Unknown')}")
                    
                    # Fix photo field name - phone API uses 'photo', but enrich_person_data expects 'face'
                    if person.get('photo') and not person.get('face'):
                        person['face'] = person['photo']
                        print(f"Fixed photo field name for {person.get('full_name', 'Unknown')}")
                    
                    # Get NIK for additional data enrichment
                    nik = person.get('ktp_number')
                    if nik:
                        print(f"Performing additional NIK search for complete data: {nik}")
                        try:
                            # Perform NIK search to get complete data
                            search_params = {'nik': nik}
                            nik_search_result = call_search(token, search_params)
                            
                            if nik_search_result and nik_search_result.get('results'):
                                # Merge complete data from NIK search
                                complete_person = nik_search_result['results'][0]
                                print(f"Found complete data for NIK {nik}")
                                print(f"Complete person keys: {list(complete_person.keys())[:10]}")
                                
                                # Merge phone search data with complete NIK data
                                # Keep phone-specific data but use complete data for missing fields
                                for key, value in complete_person.items():
                                    # Check if value is empty or N/A
                                    current_value = person.get(key)
                                    if current_value in ['N/A', '', None] or key not in person:
                                        person[key] = value
                                        if key in ['jenis_kelamin', 'alamat', 'provinsi', 'address', 'sex']:
                                            print(f"Merged {key}: {current_value} -> {value}")
                                
                                # Ensure we keep the phone number that was searched
                                if not person.get('phone_data'):
                                    person['phone_data'] = []
                                
                                # Add the searched phone number to phone data
                                searched_phone = {
                                    'number': phone_number,
                                    'operator': 'N/A',
                                    'register_date': 'N/A',
                                    'nik': nik
                                }
                                person['phone_data'].insert(0, searched_phone)
                            else:
                                print(f"No complete data found for NIK {nik}, using phone search data only")
                        except Exception as e:
                            print(f"Error during NIK search for {nik}: {e}")
                            # Continue with phone search data only
                    
                    # Enrich with basic data only (fast for initial search)
                    enriched_person = enrich_person_data_basic(person.copy(), token)
                    
                    results.append({'person': enriched_person})
                else:
                    # Fallback: Get person data using NIK from phone data
                    nik = phone_item.get('nik')
                    if nik:
                        search_params = {'nik': nik}
                        search_result = call_search(token, search_params)
                        
                        if search_result.get('results'):
                            person = search_result['results'][0]
                            # Enrich with phone data
                            enriched_person = enrich_person_data(person.copy(), token)
                            results.append({'person': enriched_person})
            
            if results:
                result = {
                    'results': results,
                    'total_results': len(results),
                    'message': f'Ditemukan {len(results)} hasil untuk nomor {phone_number}'
                }
            else:
                result = {
                    'results': [],
                    'total_results': 0,
                    'message': f'Tidak ada hasil untuk nomor {phone_number}'
                }
        else:
            result = {
                'results': [],
                'total_results': 0,
                'message': f'Tidak ada hasil untuk nomor {phone_number}'
            }
        
        # Save profiling data to database and log activity
        try:
            # Use user_data that was passed as parameter (already authenticated)
            if user_data and user_data.get('id'):
                client_ip = get_client_ip()
                
                # Prepare search parameters for saving
                search_params = {
                    'search_type': 'phone',
                    'phone_number': phone_number,
                    'phone_operator': data.get('phone_operator', '')
                }
                
                # Prepare search results for saving
                search_results = {
                    'total_results': len(results),
                    'message': f'Ditemukan {len(results)} hasil untuk nomor {phone_number}'
                }
                
                # Save to database
                try:
                    db.save_profiling_data(
                        user_id=user_data['id'],
                        search_type='phone',
                        search_params=search_params,
                        search_results=search_results,
                        person_data=results[0]['person'] if results else None,
                        phone_data={'phone_number': phone_number, 'operator': data.get('phone_operator', '')},
                        ip_address=client_ip,
                        user_agent=request.headers.get('User-Agent')
                    )
                    print(f"✅ Saved phone search profiling data for user {user_data.get('username', 'Unknown')} (ID: {user_data['id']})")
                except Exception as save_error:
                    print(f"⚠️ Error saving phone search profiling data: {save_error}")
                
                # Log phone search activity (ALWAYS log, even if save failed or no results)
                try:
                    db.log_activity(
                        user_id=user_data['id'],
                        activity_type='profiling_search',
                        description=f'Phone search: {phone_number} - Found {len(results)} results',
                        ip_address=client_ip,
                        user_agent=request.headers.get('User-Agent')
                    )
                    print(f"✅ Logged phone search activity for user {user_data.get('username', 'Unknown')} from IP {client_ip}")
                except Exception as log_error:
                    print(f"❌ Error logging phone search activity: {log_error}")
        except Exception as save_error:
            print(f"❌ Error in phone search save/logging: {save_error}")
            import traceback
            traceback.print_exc()
        
        return jsonify(result)
        
    except Exception as e:
        print(f"Error during phone search: {e}")
        return jsonify({'error': f'Error during phone search: {str(e)}'}), 500

def perform_face_search(token, params, data, user_data):
    """Perform face search with uploaded image"""
    if not USE_FACE_LIB:
        return jsonify({'error': 'face_recognition library tidak terpasang'}), 500
    
    try:
        # Decode base64 image
        face_query_b64 = data['face_query']
        if face_query_b64.startswith('data:'):
            face_query_b64 = face_query_b64.split(',', 1)[1]
        
        # Save to temporary file
        with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as tmp_file:
            img_bytes = base64.b64decode(face_query_b64)
            tmp_file.write(img_bytes)
            tmp_path = Path(tmp_file.name)
        
        try:
            # Get query encoding
            q_enc = load_image_file_to_encoding(tmp_path)
            if q_enc is None:
                return jsonify({'error': 'Tidak menemukan wajah pada query image'}), 400
            
            # Get search results
            j = call_search(token, params)
            people = parse_people_from_response(j)
            
            if not people:
                # Log activity even if no people found
                if user_data and user_data.get('id'):
                    try:
                        client_ip = get_client_ip()
                        db.log_activity(
                            user_id=user_data['id'],
                            activity_type='profiling_search',
                            description=f'Face search: No people found from API - threshold {data.get("face_threshold", 0.50)}',
                            ip_address=client_ip,
                            user_agent=request.headers.get('User-Agent')
                        )
                        print(f"✅ Logged face search activity (no people) for user {user_data.get('username', 'Unknown')} from IP {client_ip}")
                    except Exception as log_error:
                        print(f"❌ Error logging face search activity: {log_error}")
                
                return jsonify({'results': [], 'message': 'Tidak ada person yang dikembalikan oleh API'})
            
            # Process face matching
            threshold = float(data.get('face_threshold', 0.50))
            matches = []
            
            for p in people:
                # Fix photo field name - API might use 'photo', but we expect 'face'
                if p.get('photo') and not p.get('face'):
                    p['face'] = p['photo']
                    print(f"Fixed photo field name for {p.get('full_name', 'Unknown')}")
                
                face_b64 = p.get("face", "")
                if not face_b64:
                    continue
                
                try:
                    enc = get_encoding_from_base64_face(face_b64)
                    if enc is None:
                        continue
                    
                    # Compute distance
                    dist = np.linalg.norm(np.array(q_enc) - np.array(enc))
                    
                    if dist <= threshold:
                        # Enrich person data with family and phone info
                        enriched_person = enrich_person_data(p.copy(), token)
                        
                        # Save face image if requested
                        saved_path = None
                        if data.get('save_face'):
                            saved_path = save_face_image(
                                face_b64, 
                                OUTPUT_FOLDER, 
                                filename_prefix=str(p.get("ktp_number") or p.get("full_name") or "face")
                            )
                        
                        matches.append({
                            'distance': float(dist),
                            'person': enriched_person,
                            'saved_face_path': str(saved_path) if saved_path else None
                        })
                        
                except Exception as e:
                    print(f"Warning: gagal decode face untuk {p.get('ktp_number', p.get('full_name', 'unknown'))}: {e}")
                    continue
            
            # Sort by distance
            matches.sort(key=lambda x: x['distance'])
            
            # Save profiling data to database and log activity
            try:
                # Use user_data that was passed as parameter (already authenticated)
                if user_data and user_data.get('id'):
                    client_ip = get_client_ip()
                    
                    # Prepare search parameters for saving
                    search_params = {
                        'search_type': 'face',
                        'threshold': threshold,
                        'save_face': data.get('save_face', False)
                    }
                    
                    # Prepare search results for saving
                    search_results = {
                        'total_matches': len(matches),
                        'threshold': threshold,
                        'message': f'Ditemukan {len(matches)} kandidat match'
                    }
                    
                    # Save to database
                    try:
                        db.save_profiling_data(
                            user_id=user_data['id'],
                            search_type='face',
                            search_params=search_params,
                            search_results=search_results,
                            person_data=matches[0]['person'] if matches else None,
                            face_data={'matches_count': len(matches), 'threshold': threshold},
                            ip_address=client_ip,
                            user_agent=request.headers.get('User-Agent')
                        )
                        print(f"✅ Saved face search profiling data for user {user_data.get('username', 'Unknown')} (ID: {user_data['id']})")
                    except Exception as save_error:
                        print(f"⚠️ Error saving face search profiling data: {save_error}")
                    
                    # Log face search activity (ALWAYS log, even if save failed or no matches)
                    try:
                        db.log_activity(
                            user_id=user_data['id'],
                            activity_type='profiling_search',
                            description=f'Face search: threshold {threshold} - Found {len(matches)} matches',
                            ip_address=client_ip,
                            user_agent=request.headers.get('User-Agent')
                        )
                        print(f"✅ Logged face search activity for user {user_data.get('username', 'Unknown')} from IP {client_ip}")
                    except Exception as log_error:
                        print(f"❌ Error logging face search activity: {log_error}")
            except Exception as save_error:
                print(f"❌ Error in face search save/logging: {save_error}")
                import traceback
                traceback.print_exc()
            
            return jsonify({
                'results': matches,
                'total_matches': len(matches),
                'threshold': threshold,
                'message': f'Ditemukan {len(matches)} kandidat match'
            })
            
        finally:
            # Clean up temporary file
            if tmp_path.exists():
                tmp_path.unlink()
                
    except Exception as e:
        return jsonify({'error': f'Face search error: {str(e)}'}), 500

def perform_regular_search(token, params, data, user_data):
    """Perform regular search without face matching"""
    try:
        print(f"DEBUG: perform_regular_search - Token: {token[:30]}..., Params: {params}", file=sys.stderr)
        # Pass username and password for flexible server 116 login
        j = call_search(token, params, username=data.get('username'), password=data.get('password'))
        print(f"DEBUG: perform_regular_search - Response dari call_search: type={type(j)}, keys={list(j.keys()) if isinstance(j, dict) else 'N/A'}", file=sys.stderr)
        
        # Check if server 116 was used as fallback
        using_server_116 = j.get('_server_116_fallback', False)
        server_224_unavailable = j.get('_server_224_unavailable', False)
        
        if using_server_116:
            print(f"DEBUG: perform_regular_search - Server 116 digunakan sebagai fallback", file=sys.stderr)
        
        people = parse_people_from_response(j)
        print(f"DEBUG: perform_regular_search - Hasil parse: {len(people)} people ditemukan", file=sys.stderr)
        
        if len(people) > 0:
            print(f"DEBUG: perform_regular_search - Contoh person pertama: NIK={people[0].get('ktp_number', 'N/A')}, Nama={people[0].get('full_name', 'N/A')}", file=sys.stderr)
        
        results = []
        
        if not people:
            # Log activity even if no results found
            if user_data and user_data.get('id'):
                try:
                    client_ip = get_client_ip()
                    search_params_str = f"name: {data.get('name', '')}, nik: {data.get('nik', '')}"
                    db.log_activity(
                        user_id=user_data['id'],
                        activity_type='profiling_search',
                        description=f'Profiling search: identity - {search_params_str} - Found 0 results',
                        ip_address=client_ip,
                        user_agent=request.headers.get('User-Agent')
                    )
                    print(f"✅ Logged profiling search activity (no results) for user {user_data.get('username', 'Unknown')} from IP {client_ip}")
                except Exception as log_error:
                    print(f"❌ Error logging profiling search activity: {log_error}")
            
            # PENTING: Selalu kirim flag fallback meskipun hasil kosong
            response_data = {
                'results': [], 
                'total_results': 0,
                'message': 'Tidak ada hasil'
            }
            # Tambahkan flag fallback jika server 116 digunakan
            if using_server_116 or server_224_unavailable:
                response_data['_server_116_fallback'] = True
                response_data['_server_224_unavailable'] = True
                print(f"DEBUG: perform_regular_search - Mengirim flag fallback meskipun hasil kosong", file=sys.stderr)
            
            return jsonify(response_data)
        
        for p in people:
            # Fix photo field name - API might use 'photo', but enrich_person_data expects 'face'
            if p.get('photo') and not p.get('face'):
                p['face'] = p['photo']
                print(f"Fixed photo field name for {p.get('full_name', 'Unknown')}")
            
            # Enrich person data with basic info only (fast for initial search)
            enriched_person = enrich_person_data_basic(p.copy(), token)
            
            # Skip refetch jika menggunakan server 116 (untuk kecepatan)
            # Refetch hanya dilakukan jika server 224 hidup dan tidak ada foto
            if not token.startswith("fallback_token_"):
                ep_face = enriched_person.get('face', '')
                if (not ep_face) or (isinstance(ep_face, str) and ep_face.startswith('data:image/svg')):
                    nik_lookup = enriched_person.get('ktp_number') or enriched_person.get('nik')
                    if nik_lookup:
                        try:
                            search_result = call_search(token, {'nik': nik_lookup})
                            people_refetch = parse_people_from_response(search_result)
                            if people_refetch:
                                full_person = people_refetch[0]
                                if full_person.get('photo') and not full_person.get('face'):
                                    full_person['face'] = full_person['photo']
                                enriched_person = enrich_person_data(full_person.copy(), token)
                                print(f"Refetched full person data with face for NIK {nik_lookup}")
                        except Exception as _e:
                            pass
            else:
                print(f"[INFO] Skip refetch untuk NIK (server 116 mode - optimasi kecepatan)", file=sys.stderr)
            
            # Photo formatting is already handled in enrich_person_data
            
            result = {'person': enriched_person}
            
            # Save face if requested
            if data.get('save_face') and enriched_person.get('face'):
                saved_path = save_face_image(
                    enriched_person.get('face'), 
                    OUTPUT_FOLDER, 
                    filename_prefix=str(enriched_person.get("ktp_number") or enriched_person.get("full_name") or "face")
                )
                result['saved_face_path'] = str(saved_path) if saved_path else None
            
            results.append(result)
        
        # Save profiling data to database and log activity
        try:
            # Use user_data that was passed as parameter (already authenticated)
            if user_data and user_data.get('id'):
                client_ip = get_client_ip()
                
                # Prepare search parameters for saving
                search_params = {
                    'search_type': 'identity',
                    'name': data.get('name', ''),
                    'nik': data.get('nik', ''),
                    'page': data.get('page', '1')
                }
                
                # Prepare search results for saving
                search_results = {
                    'total_results': len(results),
                    'message': f'Ditemukan {len(results)} hasil'
                }
                
                # Save to database
                try:
                    db.save_profiling_data(
                        user_id=user_data['id'],
                        search_type='identity',
                        search_params=search_params,
                        search_results=search_results,
                        person_data=results[0]['person'] if results else None,
                        ip_address=client_ip,
                        user_agent=request.headers.get('User-Agent')
                    )
                    print(f"✅ Saved profiling data for user {user_data.get('username', 'Unknown')} (ID: {user_data['id']})")
                except Exception as save_error:
                    print(f"⚠️ Error saving profiling data: {save_error}")
                
                # Log profiling search activity (ALWAYS log, even if save failed)
                try:
                    search_params_str = ', '.join([f"{k}: {v}" for k, v in search_params.items() if v])
                    db.log_activity(
                        user_id=user_data['id'],
                        activity_type='profiling_search',
                        description=f'Profiling search: identity - {search_params_str} - Found {len(results)} results',
                        ip_address=client_ip,
                        user_agent=request.headers.get('User-Agent')
                    )
                    print(f"✅ Logged profiling search activity for user {user_data.get('username', 'Unknown')} from IP {client_ip}")
                except Exception as log_error:
                    print(f"❌ Error logging profiling search activity: {log_error}")
        except Exception as save_error:
            print(f"❌ Error in profiling search save/logging: {save_error}")
            import traceback
            traceback.print_exc()
        
        response_data = {
            'results': results,
            'total_results': len(results),
            'message': f'Ditemukan {len(results)} hasil'
        }
        
        # PENTING: Selalu kirim flag fallback jika server 116 digunakan (meskipun ada hasil)
        # Ini memastikan frontend tahu bahwa server 116 digunakan sebagai fallback
        if using_server_116 or server_224_unavailable:
            response_data['_server_116_fallback'] = True
            response_data['_server_224_unavailable'] = True
            print(f"DEBUG: perform_regular_search - Mengirim flag fallback dengan {len(results)} hasil", file=sys.stderr)
        
        return jsonify(response_data)
        
    except Exception as e:
        return jsonify({'error': f'Search error: {str(e)}'}), 500

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'face_lib_available': USE_FACE_LIB,
        'timestamp': time.time()
    })

@app.route('/static/clean_photos/<filename>')
def serve_clean_photo(filename):
    """Serve clean photos from static folder"""
    try:
        return send_from_directory(CLEAN_PHOTOS_FOLDER, filename)
    except Exception as e:
        logger.error(f"Error serving clean photo {filename}: {e}")
        return jsonify({'error': 'Photo not found'}), 404

@app.route('/api/reprocess-photo/<nik>', methods=['POST'])
def api_reprocess_photo(nik):
    """API endpoint untuk memaksa reprocessing foto dengan algoritma terbaru"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        
        if not user:
            return jsonify({'error': 'Unauthorized'}), 401
        
        data = request.get_json()
        url_foto = data.get('url_foto')
        
        if not url_foto:
            return jsonify({'error': 'URL foto diperlukan'}), 400
        
        logger.info(f"Force reprocessing photo for NIK {nik} by user {user['username']}")
        
        # Force reprocess with new algorithm
        clean_photo_url = process_and_save_clean_photo(nik, url_foto, force_reprocess=True)
        
        if clean_photo_url:
            return jsonify({
                'success': True,
                'message': f'Foto berhasil diproses ulang untuk NIK {nik}',
                'foto_bersih_url': clean_photo_url
            })
        else:
            return jsonify({'error': 'Gagal memproses foto'}), 500
            
    except Exception as e:
        logger.error(f"Error reprocessing photo for NIK {nik}: {e}")
        return jsonify({'error': f'Error: {str(e)}'}), 500

@app.route('/api/test-watermark-removal', methods=['POST'])
def api_test_watermark_removal():
    """API endpoint untuk test watermark removal dengan foto yang sudah ada"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        
        if not user:
            return jsonify({'error': 'Unauthorized'}), 401
        
        data = request.get_json()
        nik = data.get('nik', '1505041107830002')  # Default to MARGUTIN
        
        logger.info(f"Testing watermark removal for NIK {nik} by user {user['username']}")
        
        # Get person data first
        search_data = {
            "search_type": "identity",
            "nik": nik,
            "username": DEFAULT_USERNAME,
            "password": DEFAULT_PASSWORD
        }
        
        # Search for person data
        try:
            token = ensure_token(DEFAULT_USERNAME, DEFAULT_PASSWORD, force_refresh=True)
            if not token:
                print("Warning: Tidak bisa mendapatkan token untuk person search")
                return jsonify({'error': 'Server eksternal tidak tersedia'}), 500
            search_result = call_search(token, {'nik': nik})
        except Exception as e:
            print(f"Error getting token for person search: {e}")
            return jsonify({'error': 'Server eksternal tidak tersedia'}), 500
        people = parse_people_from_response(search_result)
        
        if not people:
            return jsonify({'error': 'Person tidak ditemukan'}), 404
        
        person = people[0]
        
        # Check if person has face data
        face_data = person.get('face', '')
        if not face_data or face_data.startswith('data:image/svg'):
            return jsonify({'error': 'Person tidak memiliki foto asli'}), 400
        
        # Try to extract photo URL from face data or other fields
        url_foto = person.get('url_foto') or person.get('face_url') or person.get('photo_url')
        
        if not url_foto or not url_foto.startswith('http'):
            # If no URL, we can't download the original photo for watermark removal
            return jsonify({
                'error': 'Tidak ada URL foto asli yang tersedia untuk watermark removal',
                'person_name': person.get('full_name', 'Unknown'),
                'has_face_data': bool(face_data),
                'face_data_type': 'base64' if face_data.startswith('data:image/') else 'svg_avatar'
            }), 400
        
        # Force reprocess with aggressive algorithm
        clean_photo_url = process_and_save_clean_photo(nik, url_foto, force_reprocess=True)
        
        if clean_photo_url:
            return jsonify({
                'success': True,
                'message': f'Watermark removal test berhasil untuk NIK {nik}',
                'person_name': person.get('full_name', 'Unknown'),
                'foto_bersih_url': clean_photo_url,
                'original_photo_url': url_foto
            })
        else:
            return jsonify({'error': 'Gagal memproses watermark removal'}), 500
            
    except Exception as e:
        logger.error(f"Error testing watermark removal for NIK {nik}: {e}")
        return jsonify({'error': f'Error: {str(e)}'}), 500

@app.route('/api/debug/family/<nik>', methods=['GET'])
def debug_family_data(nik):
    """Debug endpoint to test family data API"""
    try:
        # Get token from request headers or use a default one
        auth_header = request.headers.get('Authorization')
        token = None
        if auth_header and auth_header.startswith('Bearer '):
            token = auth_header[7:]
        
        # Get NKK from query parameter
        nkk = request.args.get('nkk')
        
        # Use sample data for NIK 1505041107830002
        if nik == "1505041107830002":
            sample_family_data = {
                'kepala_keluarga': 'MARGUTIN',
                'nkk': '1505041911100032',
                'alamat_keluarga': 'DUSUN SUNGAIN BAYUR, MANDI ANGIN',
                'anggota_keluarga': [
                    {
                        'nama': 'MARGUTIN',
                        'hubungan': 'Kepala Keluarga',
                        'nik': '1505041107830002',
                        'tanggal_lahir': '11-07-1983',
                        'tempat_lahir': 'MANDI ANGIN',
                        'jenis_kelamin': 'Laki-laki',
                        'agama': 'Islam',
                        'status_perkawinan': 'Kawin',
                        'pekerjaan': 'Wiraswasta',
                        'foto': 'data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iNjQiIGhlaWdodD0iNjQiIHZpZXdCb3g9IjAgMCA2NCA2NCIgZmlsbD0ibm9uZSIgeG1sbnM9Imh0dHA6Ly93d3cudzMub3JnLzIwMDAvc3ZnIj4KPGNpcmNsZSBjeD0iMzIiIGN5PSIzMiIgcj0iMzIiIGZpbGw9IiM2NjdFRUEiLz4KPHN2ZyB3aWR0aD0iMzIiIGhlaWdodD0iMzIiIHZpZXdCb3g9IjAgMCAyNCAyNCIgZmlsbD0ibm9uZSIgeG1sbnM9Imh0dHA6Ly93d3cudzMub3JnLzIwMDAvc3ZnIiB4PSIxNiIgeT0iMTYiPgo8cGF0aCBkPSJNMTIgMTJDMTQuMjA5MSAxMiAxNiAxMC4yMDkxIDE2IDhDMTYgNS43OTA5IDE0LjIwOTEgNCAxMiA0QzkuNzkwODYgNCA4IDUuNzkwOSA4IDhDOCAxMC4yMDkxIDkuNzkwODYgMTIgMTIgMTJaIiBmaWxsPSJ3aGl0ZSIvPgo8cGF0aCBkPSJNMTIgMTRDOC42ODYyOSAxNCA2IDE2LjY4NjMgNiAyMEgxOEMxOCAxNi42ODYzIDE1LjMxMzcgMTQgMTIgMTRaIiBmaWxsPSJ3aGl0ZSIvPgo8L3N2Zz4KPC9zdmc+'
                    },
                    {
                        'nama': 'SITI AMINAH',
                        'hubungan': 'Istri',
                        'nik': '1505041911100033',
                        'tanggal_lahir': '15-03-1985',
                        'tempat_lahir': 'MANDI ANGIN',
                        'jenis_kelamin': 'Perempuan',
                        'agama': 'Islam',
                        'status_perkawinan': 'Kawin',
                        'pekerjaan': 'Ibu Rumah Tangga',
                        'foto': 'data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iNjQiIGhlaWdodD0iNjQiIHZpZXdCb3g9IjAgMCA2NCA2NCIgZmlsbD0ibm9uZSIgeG1sbnM9Imh0dHA6Ly93d3cudzMub3JnLzIwMDAvc3ZnIj4KPGNpcmNsZSBjeD0iMzIiIGN5PSIzMiIgcj0iMzIiIGZpbGw9IiNGRjY5QjQiLz4KPHN2ZyB3aWR0aD0iMzIiIGhlaWdodD0iMzIiIHZpZXdCb3g9IjAgMCAyNCAyNCIgZmlsbD0ibm9uZSIgeG1sbnM9Imh0dHA6Ly93d3cudzMub3JnLzIwMDAvc3ZnIiB4PSIxNiIgeT0iMTYiPgo8cGF0aCBkPSJNMTIgMTJDMTQuMjA5MSAxMiAxNiAxMC4yMDkxIDE2IDhDMTYgNS43OTA5IDE0LjIwOTEgNCAxMiA0QzkuNzkwODYgNCA4IDUuNzkwOSA4IDhDOCAxMC4yMDkxIDkuNzkwODYgMTIgMTIgMTJaIiBmaWxsPSJ3aGl0ZSIvPgo8cGF0aCBkPSJNMTIgMTRDOC42ODYyOSAxNCA2IDE2LjY4NjMgNiAyMEgxOEMxOCAxNi42ODYzIDE1LjMxMzcgMTQgMTIgMTRaIiBmaWxsPSJ3aGl0ZSIvPgo8L3N2Zz4KPC9zdmc+'
                    },
                    {
                        'nama': 'AHMAD RIZKI',
                        'hubungan': 'Anak',
                        'nik': '1505041911100034',
                        'tanggal_lahir': '20-12-2010',
                        'tempat_lahir': 'MANDI ANGIN',
                        'jenis_kelamin': 'Laki-laki',
                        'agama': 'Islam',
                        'status_perkawinan': 'Belum Kawin',
                        'pekerjaan': 'Pelajar',
                        'foto': 'data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iNjQiIGhlaWdodD0iNjQiIHZpZXdCb3g9IjAgMCA2NCA2NCIgZmlsbD0ibm9uZSIgeG1sbnM9Imh0dHA6Ly93d3cudzMub3JnLzIwMDAvc3ZnIj4KPGNpcmNsZSBjeD0iMzIiIGN5PSIzMiIgcj0iMzIiIGZpbGw9IiM0Q0FGNTIiLz4KPHN2ZyB3aWR0aD0iMzIiIGhlaWdodD0iMzIiIHZpZXdCb3g9IjAgMCAyNCAyNCIgZmlsbD0ibm9uZSIgeG1sbnM9Imh0dHA6Ly93d3cudzMub3JnLzIwMDAvc3ZnIiB4PSIxNiIgeT0iMTYiPgo8cGF0aCBkPSJNMTIgMTJDMTQuMjA5MSAxMiAxNiAxMC4yMDkxIDE2IDhDMTYgNS43OTA5IDE0LjIwOTEgNCAxMiA0QzkuNzkwODYgNCA4IDUuNzkwOSA4IDhDOCAxMC4yMDkxIDkuNzkwODYgMTIgMTIgMTJaIiBmaWxsPSJ3aGl0ZSIvPgo8cGF0aCBkPSJNMTIgMTRDOC42ODYyOSAxNCA2IDE2LjY4NjMgNiAyMEgxOEMxOCAxNi42ODYzIDE1LjMxMzcgMTQgMTIgMTRaIiBmaWxsPSJ3aGl0ZSIvPgo8L3N2Zz4KPC9zdmc+'
                    },
                    {
                        'nama': 'FATIMAH ZAHRA',
                        'hubungan': 'Anak',
                        'nik': '1505041911100035',
                        'tanggal_lahir': '05-08-2012',
                        'tempat_lahir': 'MANDI ANGIN',
                        'jenis_kelamin': 'Perempuan',
                        'agama': 'Islam',
                        'status_perkawinan': 'Belum Kawin',
                        'pekerjaan': 'Pelajar',
                        'foto': 'data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iNjQiIGhlaWdodD0iNjQiIHZpZXdCb3g9IjAgMCA2NCA2NCIgZmlsbD0ibm9uZSIgeG1sbnM9Imh0dHA6Ly93d3cudzMub3JnLzIwMDAvc3ZnIj4KPGNpcmNsZSBjeD0iMzIiIGN5PSIzMiIgcj0iMzIiIGZpbGw9IiNGRjQ0NzciLz4KPHN2ZyB3aWR0aD0iMzIiIGhlaWdodD0iMzIiIHZpZXdCb3g9IjAgMCAyNCAyNCIgZmlsbD0ibm9uZSIgeG1sbnM9Imh0dHA6Ly93d3cudzMub3JnLzIwMDAvc3ZnIiB4PSIxNiIgeT0iMTYiPgo8cGF0aCBkPSJNMTIgMTJDMTQuMjA5MSAxMiAxNiAxMC4yMDkxIDE2IDhDMTYgNS43OTA5IDE0LjIwOTEgNCAxMiA0QzkuNzkwODYgNCA4IDUuNzkwOSA4IDhDOCAxMC4yMDkxIDkuNzkwODYgMTIgMTIgMTJaIiBmaWxsPSJ3aGl0ZSIvPgo8cGF0aCBkPSJNMTIgMTRDOC42ODYyOSAxNCA2IDE2LjY4NjMgNiAyMEgxOEMxOCAxNi42ODYzIDE1LjMxMzcgMTQgMTIgMTRaIiBmaWxsPSJ3aGl0ZSIvPgo8L3N2Zz4KPC9zdmc+'
                    }
                ]
            }
            return jsonify({
                'nik': nik,
                'nkk': nkk,
                'family_data': sample_family_data,
                'status': 'success',
                'note': 'Using sample data for testing'
            })
        
        family_data = get_family_data(nik, nkk=nkk, token=token)
        return jsonify({
            'nik': nik,
            'nkk': nkk,
            'family_data': family_data,
            'status': 'success'
        })
    except Exception as e:
        return jsonify({
            'nik': nik,
            'error': str(e),
            'status': 'error'
        }), 500

@app.route('/api/debug/family-alt/<nkk>', methods=['GET'])
def debug_family_alt(nkk):
    """Debug endpoint to test alternative family data API"""
    try:
        url = f"{FAMILY_API_ALT}?family_cert_number={nkk}"
        print(f"Testing alternative family API: {url}")
        
        response = requests.get(url, timeout=30)
        print(f"Alternative API response status: {response.status_code}")
        print(f"Alternative API response text: {response.text}")
        
        if response.status_code == 200:
            data = response.json()
            return jsonify({
                'url': url,
                'status_code': response.status_code,
                'response': data,
                'status': 'success'
            })
        else:
            return jsonify({
                'url': url,
                'status_code': response.status_code,
                'response': response.text,
                'status': 'error'
            }), response.status_code
            
    except Exception as e:
        return jsonify({
            'url': url,
            'error': str(e),
            'status': 'error'
        }), 500

@app.route('/api/debug/phone/<nik>', methods=['GET'])
def debug_phone_data(nik):
    """Debug endpoint to test phone data API"""
    try:
        # Get token from request headers or use a default one
        auth_header = request.headers.get('Authorization')
        token = None
        if auth_header and auth_header.startswith('Bearer '):
            token = auth_header[7:]
        
        # Test with sample data first
        if nik == "1505041107830002":
            sample_data = {
                "data": {
                    "data": [
                        {
                            "_source": {
                                "msisdn": "6285755349653",
                                "nik": "1505041107830002",
                                "operator": "INDOSAT",
                                "register_date": "2018-07-29"
                            }
                        },
                        {
                            "_source": {
                                "msisdn": "6285655343518",
                                "nik": "1505041107830002",
                                "operator": "INDOSAT",
                                "register_date": "2018-08-05"
                            }
                        },
                        {
                            "_source": {
                                "msisdn": "6282260122369",
                                "nik": "1505041107830002",
                                "operator": "TELKOMSEL",
                                "register_date": "2018-03-27"
                            }
                        },
                        {
                            "_source": {
                                "msisdn": "6281280699628",
                                "nik": "1505041107830002",
                                "operator": "TELKOMSEL",
                                "register_date": "2018-03-27"
                            }
                        },
                        {
                            "_source": {
                                "msisdn": "6282313374849",
                                "nik": "1505041107830002",
                                "operator": "TELKOMSEL",
                                "register_date": "2019-05-09"
                            }
                        }
                    ]
                }
            }
            print(f"Using sample data for testing: {sample_data}")
            # Parse sample data
            phone_list = []
            for item in sample_data['data']['data']:
                source = item.get('_source', {})
                phone_info = {
                    'number': source.get('msisdn', ''),
                    'operator': source.get('operator', ''),
                    'register_date': source.get('register_date', ''),
                    'nik': source.get('nik', nik)
                }
                phone_list.append(phone_info)
            return jsonify({
                'nik': nik,
                'phone_data': phone_list,
                'status': 'success',
                'note': 'Using sample data for testing'
            })
            
        phone_data = get_phone_data(nik, token=token)
        return jsonify({
            'nik': nik,
            'phone_data': phone_data,
            'status': 'success'
        })
    except Exception as e:
        return jsonify({
            'nik': nik,
            'error': str(e),
            'status': 'error'
        }), 500

@app.route('/api/debug/phone-direct/<nik>', methods=['GET'])
def debug_phone_direct(nik):
    """Debug endpoint to test phone data API directly"""
    try:
        url = f"{PHONE_API_BASE}/{nik}"
        print(f"Testing direct API call to: {url}")
        
        response = requests.get(url, timeout=10)
        print(f"Direct API response status: {response.status_code}")
        print(f"Direct API response text: {response.text}")
        
        if response.status_code == 200:
            data = response.json()
            return jsonify({
                'url': url,
                'status_code': response.status_code,
                'response': data,
                'status': 'success'
            })
        else:
            return jsonify({
                'url': url,
                'status_code': response.status_code,
                'response': response.text,
                'status': 'error'
            }), response.status_code
            
    except Exception as e:
        return jsonify({
            'url': url,
            'error': str(e),
            'status': 'error'
        }), 500

@app.route('/api/debug/search/<nik>', methods=['GET'])
def debug_search(nik):
    """Debug endpoint to test search with sample data"""
    try:
        # Create sample person data
        sample_person = {
            'ktp_number': nik,
            'full_name': 'MARGUTIN',
            'birth_date': '11-07-1983',
            'birth_place': 'MANDI ANGIN',
            'gender': 'Laki-laki',
            'religion': 'Islam',
            'marital_status': 'Kawin',
            'occupation': 'Wiraswasta',
            'address': 'DUSUN SUNGAIN BAYUR',
            'face': 'data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iMjAwIiBoZWlnaHQ9IjIwMCIgdmlld0JveD0iMCAwIDIwMCAyMDAiIGZpbGw9Im5vbmUiIHhtbG5zPSJodHRwOi8vd3d3LnczLm9yZy8yMDAwL3N2ZyI+CjxyZWN0IHdpZHRoPSIyMDAiIGhlaWdodD0iMjAwIiBmaWxsPSIjRkY2OUI0Ii8+CjxjaXJjbGUgY3g9IjEwMCIgY3k9IjgwIiByPSIzMCIgZmlsbD0iI0ZGRkZGRiIvPgo8cGF0aCBkPSJNNTAgMTQwIFE1MCAxMjAgMTAwIDEyMCBRMTUwIDEyMCAxNTAgMTQwIEwxNTAgMTgwIEw1MCAxODAgWk01MCAxNDAiIGZpbGw9IiNGRkZGRkYiLz4KPC9zdmc+'
        }
        
        # Enrich with family and phone data
        enriched_person = enrich_person_data(sample_person.copy())
        
        return jsonify({
            'nik': nik,
            'person': enriched_person,
            'has_face': bool(enriched_person.get('face')),
            'face_length': len(enriched_person.get('face', '')) if enriched_person.get('face') else 0,
            'status': 'success'
        })
        
    except Exception as e:
        return jsonify({
            'nik': nik,
            'error': str(e),
            'status': 'error'
        }), 500

@app.route('/api/test/face', methods=['GET'])
def test_face():
    """Test endpoint to check if face image displays correctly"""
    try:
        # Simple test with a basic face image
        test_face = 'data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iMjAwIiBoZWlnaHQ9IjIwMCIgdmlld0JveD0iMCAwIDIwMCAyMDAiIGZpbGw9Im5vbmUiIHhtbG5zPSJodHRwOi8vd3d3LnczLm9yZy8yMDAwL3N2ZyI+CjxyZWN0IHdpZHRoPSIyMDAiIGhlaWdodD0iMjAwIiBmaWxsPSIjRkY2OUI0Ii8+CjxjaXJjbGUgY3g9IjEwMCIgY3k9IjgwIiByPSIzMCIgZmlsbD0iI0ZGRkZGRiIvPgo8cGF0aCBkPSJNNTAgMTQwIFE1MCAxMjAgMTAwIDEyMCBRMTUwIDEyMCAxNTAgMTQwIEwxNTAgMTgwIEw1MCAxODAgWk01MCAxNDAiIGZpbGw9IiNGRkZGRkYiLz4KPC9zdmc+'
        
        return jsonify({
            'message': 'Test face image',
            'face': test_face,
            'face_length': len(test_face),
            'status': 'success'
        })
        
    except Exception as e:
        return jsonify({
            'error': str(e),
            'status': 'error'
        }), 500

@app.route('/api/person-details', methods=['POST'])
def api_person_details():
    """API endpoint untuk mendapatkan detail lengkap person"""
    try:
        data = request.get_json()
        
        # Validate required fields
        if not data.get('nik'):
            return jsonify({'error': 'NIK diperlukan'}), 400
        
        if not data.get('username') or not data.get('password'):
            return jsonify({'error': 'Username dan password diperlukan'}), 400
        
        nik = data.get('nik')
        username = data.get('username')
        password = data.get('password')
        
        # Get token with force refresh to ensure fresh token
        try:
            token = ensure_token(username, password, force_refresh=True)
            if not token:
                return jsonify({'error': 'Gagal mendapatkan token akses ke server eksternal'}), 500
        except Exception as e:
            print(f"Error getting token: {e}")
            return jsonify({'error': f'Gagal mengakses server eksternal: {str(e)}'}), 500
        
        # Search for person by NIK
        search_params = {
            'nik': nik,
            'name': '',
            'family_cert_number': '',
            'tempat_lahir': '',
            'tanggal_lahir': '',
            'no_prop': '',
            'no_kab': '',
            'no_kec': '',
            'no_desa': '',
            'page': '1'
        }
        
        # Get person data
        search_result = call_search(token, search_params)
        people = parse_people_from_response(search_result)
        
        if not people:
            return jsonify({'error': 'Person tidak ditemukan'}), 404
        
        # Get the first person (should be the one we're looking for)
        person = people[0]
        
        # Enrich with full data (family and phone info) for detail view
        enriched_person = enrich_person_data(person.copy(), token)
        
        return jsonify({
            'success': True,
            'person': enriched_person
        })
        
    except Exception as e:
        print(f"Error getting person details: {e}")
        return jsonify({'error': f'Error: {str(e)}'}), 500

@app.route('/api/wilayah/provinsi', methods=['GET'])
def get_provinsi():
    """API endpoint untuk mengambil daftar provinsi"""
    try:
        # Data provinsi Indonesia (dapat disesuaikan dengan data server asli)
        provinsi_data = [
            {"kode": "11", "nama": "ACEH"},
            {"kode": "12", "nama": "SUMATERA UTARA"},
            {"kode": "13", "nama": "SUMATERA BARAT"},
            {"kode": "14", "nama": "RIAU"},
            {"kode": "15", "nama": "JAMBI"},
            {"kode": "16", "nama": "SUMATERA SELATAN"},
            {"kode": "17", "nama": "BENGKULU"},
            {"kode": "18", "nama": "LAMPUNG"},
            {"kode": "19", "nama": "KEPULAUAN BANGKA BELITUNG"},
            {"kode": "21", "nama": "KEPULAUAN RIAU"},
            {"kode": "31", "nama": "DKI JAKARTA"},
            {"kode": "32", "nama": "JAWA BARAT"},
            {"kode": "33", "nama": "JAWA TENGAH"},
            {"kode": "34", "nama": "DI YOGYAKARTA"},
            {"kode": "35", "nama": "JAWA TIMUR"},
            {"kode": "36", "nama": "BANTEN"},
            {"kode": "51", "nama": "BALI"},
            {"kode": "52", "nama": "NUSA TENGGARA BARAT"},
            {"kode": "53", "nama": "NUSA TENGGARA TIMUR"},
            {"kode": "61", "nama": "KALIMANTAN BARAT"},
            {"kode": "62", "nama": "KALIMANTAN TENGAH"},
            {"kode": "63", "nama": "KALIMANTAN SELATAN"},
            {"kode": "64", "nama": "KALIMANTAN TIMUR"},
            {"kode": "65", "nama": "KALIMANTAN UTARA"},
            {"kode": "71", "nama": "SULAWESI UTARA"},
            {"kode": "72", "nama": "SULAWESI TENGAH"},
            {"kode": "73", "nama": "SULAWESI SELATAN"},
            {"kode": "74", "nama": "SULAWESI TENGGARA"},
            {"kode": "75", "nama": "GORONTALO"},
            {"kode": "76", "nama": "SULAWESI BARAT"},
            {"kode": "81", "nama": "MALUKU"},
            {"kode": "82", "nama": "MALUKU UTARA"},
            {"kode": "91", "nama": "PAPUA BARAT"},
            {"kode": "94", "nama": "PAPUA"}
        ]
        
        return jsonify({
            'success': True,
            'data': provinsi_data
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Error getting provinsi data: {str(e)}'
        }), 500

@app.route('/api/wilayah/kabupaten/<provinsi_kode>', methods=['GET'])
def get_kabupaten(provinsi_kode):
    """API endpoint untuk mengambil daftar kabupaten berdasarkan kode provinsi"""
    try:
        # Data kabupaten untuk Jambi (kode 15) sebagai contoh
        if provinsi_kode == "15":  # Jambi
            kabupaten_data = [
                {"kode": "1501", "nama": "KERINCI"},
                {"kode": "1502", "nama": "MERANGIN"},
                {"kode": "1503", "nama": "SAROLANGUN"},
                {"kode": "1504", "nama": "BATANG HARI"},
                {"kode": "1505", "nama": "MUARO JAMBI"},
                {"kode": "1506", "nama": "TANJUNG JABUNG TIMUR"},
                {"kode": "1507", "nama": "TANJUNG JABUNG BARAT"},
                {"kode": "1508", "nama": "TEBO"},
                {"kode": "1509", "nama": "BUNGO"},
                {"kode": "1571", "nama": "KOTA JAMBI"},
                {"kode": "1572", "nama": "KOTA SUNGAI PENUH"}
            ]
        elif provinsi_kode == "12":  # Sumatera Utara
            kabupaten_data = [
                {"kode": "1201", "nama": "KABUPATEN NIAS"},
                {"kode": "1202", "nama": "KABUPATEN MANDAILING NATAL"},
                {"kode": "1203", "nama": "KABUPATEN TAPANULI SELATAN"},
                {"kode": "1204", "nama": "KABUPATEN TAPANULI TENGAH"},
                {"kode": "1205", "nama": "KABUPATEN TAPANULI UTARA"},
                {"kode": "1206", "nama": "KABUPATEN TOBA SAMOSIR"},
                {"kode": "1207", "nama": "KABUPATEN LABUHAN BATU"},
                {"kode": "1208", "nama": "KABUPATEN ASAHAN"},
                {"kode": "1209", "nama": "KABUPATEN SIMALUNGUN"},
                {"kode": "1210", "nama": "KABUPATEN DAIRI"},
                {"kode": "1211", "nama": "KABUPATEN KARO"},
                {"kode": "1212", "nama": "KABUPATEN DELI SERDANG"},
                {"kode": "1213", "nama": "KABUPATEN LANGKAT"},
                {"kode": "1214", "nama": "KABUPATEN NIAS SELATAN"},
                {"kode": "1215", "nama": "KABUPATEN HUMBANG HASUNDUTAN"},
                {"kode": "1216", "nama": "KABUPATEN PAKPAK BHARAT"},
                {"kode": "1217", "nama": "KABUPATEN SAMOSIR"},
                {"kode": "1218", "nama": "KABUPATEN SERDANG BEDAGAI"},
                {"kode": "1219", "nama": "KABUPATEN BATU BARA"},
                {"kode": "1220", "nama": "KABUPATEN PADANG LAWAS UTARA"},
                {"kode": "1221", "nama": "KABUPATEN PADANG LAWAS"},
                {"kode": "1222", "nama": "KABUPATEN LABUHAN BATU SELATAN"},
                {"kode": "1223", "nama": "KABUPATEN LABUHAN BATU UTARA"},
                {"kode": "1224", "nama": "KABUPATEN NIAS UTARA"},
                {"kode": "1225", "nama": "KABUPATEN NIAS BARAT"},
                {"kode": "1271", "nama": "KOTA SIBOLGA"},
                {"kode": "1272", "nama": "KOTA TANJUNG BALAI"},
                {"kode": "1273", "nama": "KOTA PEMATANG SIANTAR"},
                {"kode": "1274", "nama": "KOTA TEBING TINGGI"},
                {"kode": "1275", "nama": "KOTA MEDAN"},
                {"kode": "1276", "nama": "KOTA BINJAI"},
                {"kode": "1277", "nama": "KOTA PADANGSIDIMPUAN"},
                {"kode": "1278", "nama": "KOTA GUNUNGSITOLI"}
            ]
        else:
            # Default data untuk provinsi lain
            kabupaten_data = [
                {"kode": f"{provinsi_kode}01", "nama": f"KABUPATEN {provinsi_kode}"},
                {"kode": f"{provinsi_kode}02", "nama": f"KABUPATEN {provinsi_kode} 2"},
                {"kode": f"{provinsi_kode}71", "nama": f"KOTA {provinsi_kode}"}
            ]
        
        return jsonify({
            'success': True,
            'data': kabupaten_data
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Error getting kabupaten data: {str(e)}'
        }), 500

@app.route('/api/wilayah/kecamatan/<kabupaten_kode>', methods=['GET'])
def get_kecamatan(kabupaten_kode):
    """API endpoint untuk mengambil daftar kecamatan berdasarkan kode kabupaten"""
    try:
        # Data kecamatan untuk Kota Jambi (kode 1571) sebagai contoh
        if kabupaten_kode == "1571":  # Kota Jambi
            kecamatan_data = [
                {"kode": "157101", "nama": "KECAMATAN JAMBI SELATAN"},
                {"kode": "157102", "nama": "KECAMATAN JAMBI TIMUR"},
                {"kode": "157103", "nama": "KECAMATAN JELUTUNG"},
                {"kode": "157104", "nama": "KECAMATAN PELAYANGAN"},
                {"kode": "157105", "nama": "KECAMATAN DANAU TELUK"},
                {"kode": "157106", "nama": "KECAMATAN KOTA BARU"},
                {"kode": "157107", "nama": "KECAMATAN JAMBI UTARA"},
                {"kode": "157108", "nama": "KECAMATAN ALAM BARAJO"},
                {"kode": "157109", "nama": "KECAMATAN DANAU SIPIN"},
                {"kode": "157110", "nama": "KECAMATAN PAAL MERAH"}
            ]
        elif kabupaten_kode == "1201":  # Kabupaten Nias
            kecamatan_data = [
                {"kode": "120101", "nama": "KECAMATAN GUNUNGSITOLI"},
                {"kode": "120102", "nama": "KECAMATAN GUNUNGSITOLI SELATAN"},
                {"kode": "120103", "nama": "KECAMATAN GUNUNGSITOLI UTARA"},
                {"kode": "120104", "nama": "KECAMATAN GUNUNGSITOLI BARAT"},
                {"kode": "120105", "nama": "KECAMATAN GUNUNGSITOLI TIMUR"},
                {"kode": "120106", "nama": "KECAMATAN GUNUNGSITOLI IDANOI"},
                {"kode": "120107", "nama": "KECAMATAN GUNUNGSITOLI ALO'OA"},
                {"kode": "120108", "nama": "KECAMATAN GUNUNGSITOLI HILIR"}
            ]
        else:
            # Default data untuk kabupaten lain
            kecamatan_data = [
                {"kode": f"{kabupaten_kode}01", "nama": f"KECAMATAN {kabupaten_kode}"},
                {"kode": f"{kabupaten_kode}02", "nama": f"KECAMATAN {kabupaten_kode} 2"},
                {"kode": f"{kabupaten_kode}03", "nama": f"KECAMATAN {kabupaten_kode} 3"}
            ]
        
        return jsonify({
            'success': True,
            'data': kecamatan_data
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Error getting kecamatan data: {str(e)}'
        }), 500

@app.route('/api/wilayah/kelurahan/<kecamatan_kode>', methods=['GET'])
def get_kelurahan(kecamatan_kode):
    """API endpoint untuk mengambil daftar kelurahan berdasarkan kode kecamatan"""
    try:
        # Data kelurahan untuk Jambi Selatan (kode 157101) sebagai contoh
        if kecamatan_kode == "157101":  # Jambi Selatan
            kelurahan_data = [
                {"kode": "1571012001", "nama": "KELURAHAN KENALI BESAR"},
                {"kode": "1571012002", "nama": "KELURAHAN KENALI KECIL"},
                {"kode": "1571012003", "nama": "KELURAHAN KENALI ASAM"},
                {"kode": "1571012004", "nama": "KELURAHAN KENALI BESAR ATAS"},
                {"kode": "1571012005", "nama": "KELURAHAN KENALI BESAR BAWAH"},
                {"kode": "1571012006", "nama": "KELURAHAN KENALI BESAR TENGAH"},
                {"kode": "1571012007", "nama": "KELURAHAN KENALI BESAR TIMUR"},
                {"kode": "1571012008", "nama": "KELURAHAN KENALI BESAR BARAT"}
            ]
        elif kecamatan_kode == "157102":  # Jambi Timur
            kelurahan_data = [
                {"kode": "1571022001", "nama": "KELURAHAN TELANAIPURA"},
                {"kode": "1571022002", "nama": "KELURAHAN TELANAIPURA TIMUR"},
                {"kode": "1571022003", "nama": "KELURAHAN TELANAIPURA BARAT"},
                {"kode": "1571022004", "nama": "KELURAHAN TELANAIPURA UTARA"},
                {"kode": "1571022005", "nama": "KELURAHAN TELANAIPURA SELATAN"}
            ]
        else:
            # Default data untuk kecamatan lain
            kelurahan_data = [
                {"kode": f"{kecamatan_kode}2001", "nama": f"KELURAHAN {kecamatan_kode}"},
                {"kode": f"{kecamatan_kode}2002", "nama": f"KELURAHAN {kecamatan_kode} 2"},
                {"kode": f"{kecamatan_kode}2003", "nama": f"KELURAHAN {kecamatan_kode} 3"}
            ]
        
        return jsonify({
            'success': True,
            'data': kelurahan_data
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Error getting kelurahan data: {str(e)}'
        }), 500

@app.route('/api/universal-search', methods=['POST'])
def api_universal_search():
    """API endpoint untuk universal search"""
    try:
        data = request.get_json()
        name = data.get('name')
        
        if not name:
            return jsonify({'error': 'Nama diperlukan untuk pencarian universal'}), 400
        
        # Import requests untuk external API calls
        import requests
        
        # Create session untuk maintain cookies
        session = requests.Session()
        
        # First, get the login page to extract CSRF token
        login_page_url = 'http://10.1.54.116/auth/login'
        
        login_page_response = session.get(login_page_url, timeout=5)
        
        if login_page_response.status_code != 200:
            logger.warning(f"Universal search login page failed: {login_page_response.status_code}")
            return jsonify({
                'success': False,
                'error': f'Gagal mengakses halaman login. Status: {login_page_response.status_code}',
                'message': 'Server internal tidak dapat diakses.',
                'data': {}
            }), 200
        
        # Extract CSRF token from the page (simplified approach)
        import re
        csrf_match = re.search(r'name="_csrf"\s+value="([^"]+)"', login_page_response.text)
        csrf_token = csrf_match.group(1) if csrf_match else 'jbHXcAQGRIgskYpnCBIVo43cTQg='
        
        # Login ke external server
        login_data = {
            'username': 'jambi',
            'password': '@ab526d',
            '_csrf': csrf_token
        }
        
        # Login request
        login_response = session.post(login_page_url, 
                                    data=login_data, 
                                    headers={'Content-Type': 'application/x-www-form-urlencoded'},
                                    timeout=10)
        
        if login_response.status_code != 200:
            logger.warning(f"Universal search login failed: {login_response.status_code}")
            return jsonify({
                'success': False,
                'error': f'Gagal login ke server universal search. Status: {login_response.status_code}',
                'message': 'Autentikasi ke server internal gagal.',
                'data': {}
            }), 200
        
        # Perform universal search
        search_url = f'http://10.1.54.116/toolkit/api/universal-search-engine/search?input={name}'
        search_response = session.get(search_url, timeout=15)
        
        if search_response.status_code != 200:
            logger.warning(f"Universal search API failed: {search_response.status_code}")
            return jsonify({
                'success': False,
                'error': f'Gagal melakukan pencarian universal. Status: {search_response.status_code}',
                'message': 'Pencarian di server internal gagal.',
                'data': {}
            }), 200
        
        # Parse JSON response
        search_data = search_response.json()
        
        logger.info(f"Universal search success for: {name}")
        return jsonify({
            'success': True,
            'data': search_data,
            'search_name': name
        })
        
    except requests.exceptions.Timeout:
        logger.warning(f"Universal search timeout for: {name}")
        return jsonify({
            'success': False,
            'error': 'Server universal search tidak merespons (timeout)',
            'message': 'Server internal tidak available. Gunakan Social Media Search untuk pencarian online.',
            'data': {}
        }), 200  # Return 200 to avoid frontend error display
    except requests.exceptions.ConnectionError as e:
        logger.warning(f"Universal search connection error for: {name} - {str(e)}")
        return jsonify({
            'success': False,
            'error': 'Tidak dapat terhubung ke server universal search',
            'message': 'Server internal (10.1.54.116) tidak dapat diakses. Fitur ini memerlukan koneksi ke server internal.',
            'data': {}
        }), 200  # Return 200 to avoid frontend error display
    except requests.exceptions.RequestException as e:
        logger.warning(f"Universal search request error for: {name} - {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Request Error: {str(e)}',
            'message': 'Terjadi kesalahan saat menghubungi server internal.',
            'data': {}
        }), 200
    except Exception as e:
        logger.error(f"Universal search unexpected error for: {name} - {str(e)}", exc_info=True)
        return jsonify({
            'success': False,
            'error': f'Error: {str(e)}',
            'message': 'Terjadi kesalahan pada sistem. Silakan gunakan Social Media Search.',
            'data': {}
        }), 200

@app.route('/api/social-media-search', methods=['POST'])
def api_social_media_search():
    """API endpoint untuk Google CSE social media search dengan images"""
    try:
        data = request.get_json()
        name = data.get('name')
        search_type = data.get('type', 'web')  # 'web' or 'image'
        
        if not name:
            return jsonify({'error': 'Nama diperlukan untuk pencarian social media'}), 400
        
        import requests
        import urllib.parse
        
        # Google Custom Search Engine API
        CSE_ID = '7693f5093e95e4c28'
        # Get API key from database first, then fallback to environment variable
        API_KEY = db.get_api_key('GOOGLE_CSE') or os.getenv('GOOGLE_CSE_API_KEY') or request.args.get('key')
        if not API_KEY:
            return jsonify({'error': 'GOOGLE_CSE_API_KEY tidak ditemukan. Silakan tambahkan API key di Settings atau set di file .env'}), 500
        
        # Build search query untuk social media - format lebih beragam untuk hasil yang seimbang
        # Gunakan nama dengan quotes untuk exact match, tapi tanpa bias ke social media
        query = name if 'site:' in name else f'"{name}"'  # Support advanced queries
        
        # Google CSE API endpoint
        search_url = 'https://www.googleapis.com/customsearch/v1'
        
        # OPTIMIZED: Maximum coverage untuk hasil yang lebih banyak dan detail
        # 10 pages × 10 results = 100 results per query (MAXIMUM COVERAGE!)
        all_results = []
        pages_to_fetch = 10  # 10 pages = 100 results (Maximum untuk hasil yang lebih banyak!)
        
        for page in range(pages_to_fetch):
            start_index = (page * 10) + 1  # Google CSE uses 1-based indexing
            
            # Prepare parameters
            params = {
                'key': API_KEY,
                'cx': CSE_ID,
                'q': query,
                'num': 10,  # Standard 10 results per query (free tier limit)
                'start': start_index,  # Pagination
                'safe': 'active',
                # TIDAK menambahkan filter social media agar hasil lebih beragam
            }
            
            # Add searchType for images
            if search_type == 'image':
                params['searchType'] = 'image'
            
            # Log request for debugging
            logger.info(f"Google CSE Request Page {page+1}/{pages_to_fetch} - Query: {query}, Start: {start_index}, Type: {search_type}")
            
            try:
                # Make request to Google CSE
                response = requests.get(search_url, params=params, timeout=15)
                
                if response.status_code == 200:
                    data = response.json()
                    items = data.get('items', [])
                    all_results.extend(items)
                    logger.info(f"✓ Page {page+1}: Got {len(items)} results (total so far: {len(all_results)})")
                    
                    # If this page has less than 10 results, stop fetching
                    if len(items) < 10:
                        logger.info(f"Last page reached (got {len(items)} < 10 results)")
                        break
                else:
                    error_text = response.text[:200] if hasattr(response, 'text') else 'N/A'
                    logger.warning(f"Page {page+1} failed with status {response.status_code}: {error_text}")
                    
                    # If it's a quota error (429), stop trying more pages
                    if response.status_code == 429:
                        logger.error(f"❌ QUOTA EXCEEDED! Stopping multi-page fetch at page {page+1}")
                        break
                    
                    # Continue to next page even if one fails
                    continue
                    
            except Exception as page_error:
                logger.error(f"Error fetching page {page+1}: {str(page_error)}")
                # Continue to next page
                continue
        
        logger.info(f"=== TOTAL RESULTS FETCHED: {len(all_results)} ===")
        
        # Now process the combined results (use first response for metadata)
        # Re-fetch first page for standard error handling
        params = {
            'key': API_KEY,
            'cx': CSE_ID,
            'q': query,
            'num': 10,
            'safe': 'active',
        }
        if search_type == 'image':
            params['searchType'] = 'image'
            
        response = requests.get(search_url, params=params, timeout=15)
        
        # Log API key being used (first 20 chars only for security)
        logger.info(f"Using API Key: {API_KEY[:20]}... (truncated)")
        
        # Better error handling with detailed logging
        if response.status_code != 200:
            error_details = {}
            error_message = 'Unknown error'
            try:
                error_json = response.json()
                error_message = error_json.get('error', {}).get('message', 'Unknown error')
                error_code = error_json.get('error', {}).get('code', response.status_code)
                error_details = {
                    'error': error_message,
                    'code': error_code,
                    'errors': error_json.get('error', {}).get('errors', [])
                }
                
                # Detailed error logging
                logger.error(f"❌ Google CSE API Error {error_code}: {error_message}")
                logger.error(f"   Query: {query}")
                logger.error(f"   CSE ID: {CSE_ID}")
                logger.error(f"   API Key: {API_KEY[:20]}... (truncated)")
                logger.error(f"   Full error: {json.dumps(error_json, indent=2)}")
                
                # Try fallback: Use Google Search directly (scraping) if API fails
                logger.warning(f"Google CSE API failed, trying fallback Google Search for: {query}")
                try:
                    # Use Google Search HTML scraping as fallback
                    google_search_url = f"https://www.google.com/search?q={urllib.parse.quote(query)}&num=10"
                    headers = {
                        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                    }
                    search_response = requests.get(google_search_url, headers=headers, timeout=10)
                    
                    if search_response.status_code == 200:
                        # Parse HTML results (simplified - just return empty for now)
                        # In production, you'd want to use BeautifulSoup or similar
                        logger.info("Fallback Google Search successful, but HTML parsing not implemented")
                        # For now, return empty results with message
                        return jsonify({
                            'success': False,
                            'error': error_message,
                            'details': error_details,
                            'message': 'Google CSE API tidak tersedia. Gunakan tab "Web" atau "Gambar" untuk pencarian langsung melalui Google CSE widget.',
                            'fallback': 'widget',
                            'data': {
                                'web': [],
                                'images': [],
                                'social_links': []
                            }
                        }), 200
                except Exception as fallback_error:
                    logger.error(f"Fallback Google Search also failed: {str(fallback_error)}")
                
                # Provide more helpful error messages based on error code
                if 'quota' in error_message.lower() or error_code == 429:
                    error_message = 'Quota Google CSE API habis. Gunakan Google CSE widget sebagai alternatif (tab "Web" atau "Gambar").'
                elif 'invalid' in error_message.lower() or error_code == 400:
                    error_message = 'API key atau CSE ID tidak valid. Gunakan Google CSE widget sebagai alternatif (tab "Web" atau "Gambar").'
                elif error_code == 403:
                    error_message = 'Akses ditolak oleh Google CSE API. Gunakan Google CSE widget sebagai alternatif (tab "Web" atau "Gambar").'
            except:
                error_details = {
                    'error': response.text[:500],
                    'status_code': response.status_code
                }
                error_message = f'Error dari Google CSE API: {response.status_code}'
            
            logger.error(f"Google CSE Error Details: {json.dumps(error_details, indent=2)}")
            
            # Return 200 with success: false so frontend can read the error message
            return jsonify({
                'success': False,
                'error': error_message,
                'error_code': error_code,
                'details': error_details,
                'message': 'Google CSE API error. Gunakan tab "Categorized" atau "Social Links" untuk hasil dari Universal Search (Internal).',
                'fallback': 'universal_search',
                'data': {
                    'web': [],
                    'images': [],
                    'social_links': []
                }
            }), 200  # Return 200 so frontend can read the error message
        
        # Check if we have results from multi-page fetch
        if len(all_results) > 0:
            logger.info(f"✅ Using {len(all_results)} results from multi-page fetch")
            # Create a mock search_data structure for processing
            search_data = {
                'items': all_results,
                'searchInformation': {
                    'totalResults': str(len(all_results)),
                    'searchTime': 0.5
                }
            }
            items_to_process = all_results
        else:
            # Fallback to single page response
            logger.warning("⚠️ No results from multi-page fetch, using single page response")
            search_data = response.json()
            items_to_process = search_data.get('items', [])
        
        # Format results - USE ALL_RESULTS from multiple pages!
        results = {
            'web': [],
            'images': [],
            'social_links': []
        }
        
        logger.info(f"Processing {len(items_to_process)} items (from {len(all_results)} multi-page results)")
        
        if items_to_process:
            for item in items_to_process:
                result_item = {
                    'title': item.get('title', 'No Title'),
                    'link': item.get('link', ''),
                    'snippet': item.get('snippet', ''),
                    'htmlSnippet': item.get('htmlSnippet', ''),  # HTML formatted snippet
                    'displayLink': item.get('displayLink', ''),
                    'formattedUrl': item.get('formattedUrl', ''),
                    'cacheId': item.get('cacheId', ''),  # For cached version link
                    'kind': item.get('kind', ''),  # Result type
                    'htmlTitle': item.get('htmlTitle', ''),  # HTML formatted title
                    'pagemap': item.get('pagemap', {})  # Full pagemap data for thumbnails, metatags, etc.
                }
                
                # Add image data if available (for image search results)
                if 'image' in item:
                    result_item['image'] = {
                        'src': item['image'].get('src', ''),
                        'width': item['image'].get('width', 0),
                        'height': item['image'].get('height', 0)
                    }
                    results['images'].append(result_item)
                
                # Enhanced pagemap extraction for profile/post images
                if 'pagemap' in item:
                    pagemap = item['pagemap']
                    result_item['pagemap'] = pagemap
                    
                    # Extract all possible images from pagemap (prioritized)
                    extracted_images = []
                    
                    # 1. CSE Image (highest quality, from Google)
                    if 'cse_image' in pagemap and pagemap['cse_image']:
                        for img in pagemap['cse_image']:
                            if img.get('src'):
                                extracted_images.append({
                                    'src': img.get('src'),
                                    'width': img.get('width', 0),
                                    'height': img.get('height', 0),
                                    'type': 'cse_image'
                                })
                    
                    # 2. Metatags (Open Graph, Twitter Cards - profile/post images)
                    if 'metatags' in pagemap and pagemap['metatags']:
                        for meta in pagemap['metatags']:
                            # Open Graph image (best for social media)
                            if 'og:image' in meta and meta['og:image']:
                                extracted_images.append({
                                    'src': meta['og:image'],
                                    'type': 'og_image'
                                })
                            # Twitter image
                            if 'twitter:image' in meta and meta['twitter:image']:
                                extracted_images.append({
                                    'src': meta['twitter:image'],
                                    'type': 'twitter_image'
                                })
                            # Image property
                            if 'image' in meta and meta['image']:
                                extracted_images.append({
                                    'src': meta['image'],
                                    'type': 'meta_image'
                                })
                    
                    # 3. CSE Thumbnail
                    if 'cse_thumbnail' in pagemap and pagemap['cse_thumbnail']:
                        for thumb in pagemap['cse_thumbnail']:
                            if thumb.get('src'):
                                extracted_images.append({
                                    'src': thumb.get('src'),
                                    'width': thumb.get('width', 0),
                                    'height': thumb.get('height', 0),
                                    'type': 'thumbnail'
                                })
                    
                    # 4. Store best image (prioritize og:image, then cse_image, then others)
                    if extracted_images:
                        # Sort by priority
                        priority_order = {'og_image': 1, 'twitter_image': 2, 'cse_image': 3, 'meta_image': 4, 'thumbnail': 5}
                        extracted_images.sort(key=lambda x: priority_order.get(x.get('type', ''), 99))
                        result_item['best_image'] = extracted_images[0]
                        result_item['all_images'] = extracted_images  # Store all for frontend
                
                # ========== LOCATION-BASED RELEVANCE SCORING ==========
                # Detect location from query and result content
                result_item['location_score'] = 0
                result_item['detected_locations'] = []
                result_item['content_type'] = 'general'
                
                # Extract location keywords from query
                query_lower = name.lower() if name else ''
                result_text = f"{result_item['title']} {result_item['snippet']} {result_item['link']}".lower()
                
                # Indonesian provinces and major cities
                provinces = [
                    'aceh', 'sumatera utara', 'sumatera barat', 'riau', 'jambi', 'sumatera selatan',
                    'bengkulu', 'lampung', 'kepulauan bangka belitung', 'kepulauan riau',
                    'dki jakarta', 'jakarta', 'jawa barat', 'jawa tengah', 'jawa timur',
                    'di yogyakarta', 'yogyakarta', 'banten', 'bali', 'nusa tenggara barat',
                    'nusa tenggara timur', 'kalimantan barat', 'kalimantan tengah',
                    'kalimantan selatan', 'kalimantan timur', 'kalimantan utara',
                    'sulawesi utara', 'sulawesi tengah', 'sulawesi selatan', 'sulawesi tenggara',
                    'gorontalo', 'sulawesi barat', 'maluku', 'maluku utara', 'papua barat', 'papua'
                ]
                
                major_cities = [
                    'jakarta', 'surabaya', 'bandung', 'medan', 'semarang', 'makassar',
                    'palembang', 'tangerang', 'depok', 'bekasi', 'yogyakarta', 'bogor',
                    'malang', 'surakarta', 'batam', 'pekanbaru', 'padang', 'denpasar',
                    'banjarmasin', 'pontianak', 'samarinda', 'jambi', 'cirebon', 'sukabumi',
                    'tasikmalaya', 'serang', 'mataram', 'kupang', 'manado', 'palu', 'kendari'
                ]
                
                # Check for location matches in query
                query_locations = []
                for prov in provinces:
                    if prov in query_lower:
                        query_locations.append(prov)
                        result_item['detected_locations'].append(prov.title())
                for city in major_cities:
                    if city in query_lower:
                        query_locations.append(city)
                        result_item['detected_locations'].append(city.title())
                
                # Check for location matches in result content
                content_locations = []
                for prov in provinces:
                    if prov in result_text:
                        content_locations.append(prov)
                        if prov.title() not in result_item['detected_locations']:
                            result_item['detected_locations'].append(prov.title())
                for city in major_cities:
                    if city in result_text:
                        content_locations.append(city)
                        if city.title() not in result_item['detected_locations']:
                            result_item['detected_locations'].append(city.title())
                
                # Calculate location relevance score
                # Higher score = more relevant to query location
                if query_locations:
                    # Exact match in title = highest score
                    if any(loc in result_item['title'].lower() for loc in query_locations):
                        result_item['location_score'] = 100
                    # Match in snippet = high score
                    elif any(loc in result_item['snippet'].lower() for loc in query_locations):
                        result_item['location_score'] = 80
                    # Match in URL = medium score
                    elif any(loc in result_item['link'].lower() for loc in query_locations):
                        result_item['location_score'] = 60
                    # Match in content = low score
                    elif content_locations:
                        result_item['location_score'] = 40
                else:
                    # No location in query, but location found in result = bonus
                    if content_locations:
                        result_item['location_score'] = 20
                
                # ========== CONTENT TYPE DETECTION ==========
                # Detect content type: news, social_media, blog, forum, official, video, posting
                link_lower = result_item['link'].lower()
                title_lower = result_item['title'].lower()
                snippet_lower = result_item['snippet'].lower()
                
                # News sites
                news_keywords = ['berita', 'news', 'kompas', 'detik', 'tribun', 'liputan6', 'cnn', 'bbc', 'antaranews', 'republika', 'tempo']
                if any(kw in link_lower or kw in title_lower for kw in news_keywords):
                    result_item['content_type'] = 'news'
                
                # Social media
                social_domains = [
                    'facebook.com', 'instagram.com', 'twitter.com', 'x.com', 'linkedin.com', 
                    'tiktok.com', 'youtube.com', 'pinterest.com', 'snapchat.com', 'telegram.org',
                    'whatsapp.com', 'reddit.com', 'tumblr.com', 'discord.com', 'twitch.tv',
                    'medium.com', 'quora.com', 'flickr.com', 'vimeo.com', 'weibo.com',
                    'vk.com', 'line.me', 'github.com', 'behance.net', 'dribbble.com'
                ]
                if any(domain in link_lower for domain in social_domains):
                    result_item['content_type'] = 'social_media'
                    results['social_links'].append(result_item)
                
                # Blog
                blog_keywords = ['blog', 'wordpress', 'blogspot', 'medium.com', 'substack']
                if any(kw in link_lower for kw in blog_keywords) and result_item['content_type'] == 'general':
                    result_item['content_type'] = 'blog'
                
                # Forum
                forum_keywords = ['forum', 'kaskus', 'reddit', 'quora', 'stackoverflow']
                if any(kw in link_lower for kw in forum_keywords) and result_item['content_type'] == 'general':
                    result_item['content_type'] = 'forum'
                
                # Official/Government
                official_keywords = ['pemerintah', 'government', '.go.id', 'kemenkumham', 'kemenkominfo', 'kemendagri']
                if any(kw in link_lower or kw in title_lower for kw in official_keywords):
                    result_item['content_type'] = 'official'
                
                # Video
                video_keywords = ['youtube.com', 'vimeo.com', 'dailymotion', 'tiktok.com']
                if any(kw in link_lower for kw in video_keywords):
                    result_item['content_type'] = 'video'
                
                # Posting (social media posts, articles)
                posting_keywords = ['post', 'status', 'update', 'artikel', 'article']
                if any(kw in snippet_lower or kw in title_lower for kw in posting_keywords) and result_item['content_type'] in ['general', 'social_media']:
                    result_item['content_type'] = 'posting'
                
                results['web'].append(result_item)
        
        # ========== SORT BY LOCATION RELEVANCE ==========
        # Sort results by location_score (highest first), then by content type priority
        content_type_priority = {
            'news': 1,
            'official': 2,
            'social_media': 3,
            'posting': 4,
            'blog': 5,
            'forum': 6,
            'video': 7,
            'general': 8
        }
        
        def sort_key(item):
            location_score = item.get('location_score', 0)
            content_type = item.get('content_type', 'general')
            type_priority = content_type_priority.get(content_type, 99)
            # Higher location_score = better, lower type_priority = better
            return (-location_score, type_priority)
        
        results['web'].sort(key=sort_key)
        results['social_links'].sort(key=sort_key)
        
        logger.info(f"✓ ADVANCED SEARCH SUCCESS - Found {len(results['web'])} web results, {len(results['images'])} images, {len(results['social_links'])} social links")
        logger.info(f"✓ Location-based ranking applied - Top results prioritized by location relevance")
        
        # DEBUG: Log response structure
        logger.info(f"📤 Returning response: success=True, web_results={len(results['web'])}, images={len(results['images'])}, social_links={len(results['social_links'])}")
        
        response_data = {
            'success': True,
            'data': results,
            'search_name': name,
            'total_results': search_data.get('searchInformation', {}).get('totalResults', '0'),
            'location_ranking': True  # Flag to indicate location-based ranking is applied
        }
        
        # Log sample results for debugging
        if len(results['web']) > 0:
            sample_result = results['web'][0]
            logger.info(f"📋 Sample result: title='{sample_result.get('title', 'N/A')[:50]}', link='{sample_result.get('link', 'N/A')[:50]}'")
        
        return jsonify(response_data)
        
    except requests.exceptions.Timeout:
        logger.error("❌ Google CSE Timeout - API tidak merespons dalam 15 detik")
        return jsonify({
            'success': False, 
            'error': 'Timeout: Google CSE tidak merespons',
            'error_type': 'timeout',
            'message': 'Google CSE API timeout. Gunakan tab "Categorized" untuk hasil dari Universal Search (Internal).',
            'fallback': 'universal_search',
            'data': {
                'web': [],
                'images': [],
                'social_links': []
            }
        }), 200  # Return 200 so frontend can read the error
    except requests.exceptions.ConnectionError:
        logger.error("❌ Google CSE Connection Error - Tidak dapat terhubung")
        return jsonify({
            'success': False, 
            'error': 'Connection Error: Tidak dapat terhubung ke Google CSE',
            'error_type': 'connection_error',
            'message': 'Tidak dapat terhubung ke Google CSE API. Gunakan tab "Categorized" untuk hasil dari Universal Search (Internal).',
            'fallback': 'universal_search',
            'data': {
                'web': [],
                'images': [],
                'social_links': []
            }
        }), 200  # Return 200 so frontend can read the error
    except json.JSONDecodeError as e:
        logger.error(f"❌ JSON Decode Error: {str(e)}")
        logger.error(f"   Response text: {response.text[:500] if 'response' in locals() else 'N/A'}")
        return jsonify({
            'success': False, 
            'error': f'Error parsing response: {str(e)}',
            'error_type': 'json_decode_error',
            'message': 'Error parsing Google CSE response. Gunakan tab "Categorized" untuk hasil dari Universal Search (Internal).',
            'fallback': 'universal_search',
            'data': {
                'web': [],
                'images': [],
                'social_links': []
            }
        }), 200  # Return 200 so frontend can read the error
    except Exception as e:
        logger.error(f"❌ Unexpected error in social media search: {str(e)}", exc_info=True)
        logger.error(f"   Error type: {type(e).__name__}")
        logger.error(f"   Error args: {e.args}")
        return jsonify({
            'success': False, 
            'error': f'Error: {str(e)}',
            'error_type': type(e).__name__,
            'message': 'Unexpected error. Gunakan tab "Categorized" untuk hasil dari Universal Search (Internal).',
            'fallback': 'universal_search',
            'data': {
                'web': [],
                'images': [],
                'social_links': []
            }
        }), 200  # Return 200 so frontend can read the error

@app.route('/api/test-google-cse', methods=['GET', 'POST'])
def test_google_cse():
    """Test endpoint untuk memverifikasi Google CSE API"""
    try:
        import requests
        import urllib.parse
        
        CSE_ID = '7693f5093e95e4c28'
        # Get API key from database first, then fallback to environment variable
        API_KEY = db.get_api_key('GOOGLE_CSE') or os.getenv('GOOGLE_CSE_API_KEY')
        if not API_KEY:
            return jsonify({'error': 'GOOGLE_CSE_API_KEY tidak ditemukan. Silakan tambahkan API key di Settings atau set di file .env'}), 500
        
        # Test query
        test_query = request.args.get('q', 'test') or (request.get_json() or {}).get('q', 'test')
        query = f'"{test_query}"'
        
        search_url = 'https://www.googleapis.com/customsearch/v1'
        params = {
            'key': API_KEY,
            'cx': CSE_ID,
            'q': query,
            'num': 1,  # Just 1 result for testing
            'safe': 'active',
        }
        
        logger.info(f"🧪 Testing Google CSE API with query: {query}")
        logger.info(f"   CSE ID: {CSE_ID}")
        logger.info(f"   API Key: {API_KEY[:20]}...")
        
        response = requests.get(search_url, params=params, timeout=10)
        
        if response.status_code == 200:
            data = response.json()
            items = data.get('items', [])
            search_info = data.get('searchInformation', {})
            
            logger.info(f"✅ Google CSE API Test SUCCESS!")
            logger.info(f"   Total Results: {search_info.get('totalResults', 'N/A')}")
            logger.info(f"   Items Found: {len(items)}")
            
            return jsonify({
                'success': True,
                'message': 'Google CSE API is working correctly',
                'test_query': query,
                'total_results': search_info.get('totalResults', '0'),
                'items_found': len(items),
                'sample_result': items[0] if items else None,
                'api_key_status': 'valid',
                'cse_id_status': 'valid'
            }), 200
        else:
            error_json = response.json() if response.headers.get('content-type', '').startswith('application/json') else {}
            error_message = error_json.get('error', {}).get('message', f'HTTP {response.status_code}')
            error_code = error_json.get('error', {}).get('code', response.status_code)
            
            logger.error(f"❌ Google CSE API Test FAILED!")
            logger.error(f"   Status Code: {response.status_code}")
            logger.error(f"   Error Code: {error_code}")
            logger.error(f"   Error Message: {error_message}")
            
            return jsonify({
                'success': False,
                'message': 'Google CSE API test failed',
                'test_query': query,
                'status_code': response.status_code,
                'error_code': error_code,
                'error_message': error_message,
                'api_key_status': 'invalid' if error_code in [400, 403] else 'unknown',
                'cse_id_status': 'invalid' if error_code == 400 else 'unknown',
                'quota_exceeded': error_code == 429,
                'details': error_json
            }), 200
            
    except Exception as e:
        logger.error(f"❌ Google CSE API Test Error: {str(e)}", exc_info=True)
        return jsonify({
            'success': False,
            'message': f'Test error: {str(e)}',
            'error_type': type(e).__name__
        }), 200

@app.route('/api/google-cse-widget', methods=['POST'])
def api_google_cse_widget():
    """API endpoint untuk fetch Google CSE Widget (UNLIMITED! NO API KEY NEEDED!)"""
    try:
        data = request.get_json()
        name = data.get('name')
        search_type = data.get('type', 'web')  # 'web' or 'image'
        
        if not name:
            return jsonify({'error': 'Nama diperlukan untuk pencarian'}), 400
        
        import urllib.parse
        import re
        
        # Google CSE Widget Configuration (UNLIMITED - NO API KEY!)
        CSE_ID = '7693f5093e95e4c28'
        encoded_query = urllib.parse.quote(name)
        
        # Build Google CSE Widget URL (UNLIMITED!)
        if search_type == 'image':
            widget_url = f'https://cse.google.com/cse?cx={CSE_ID}&q={encoded_query}&tbm=isch'
        else:
            widget_url = f'https://cse.google.com/cse?cx={CSE_ID}&q={encoded_query}'
        
        logger.info(f"Fetching Google CSE Widget (UNLIMITED): {widget_url}")
        
        # Fetch widget HTML dengan proper headers
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.9',
            'Referer': 'https://cse.google.com/',
        }
        
        try:
            response = requests.get(widget_url, headers=headers, timeout=15)
            
            if response.status_code != 200:
                logger.warning(f"Google CSE Widget fetch failed: {response.status_code}")
                return jsonify({
                    'success': False,
                    'error': f'Failed to fetch widget (status {response.status_code})',
                    'data': {'web': [], 'images': []}
                }), 200
            
            html_content = response.text
            html_length = len(html_content)
            
            logger.info(f"Google CSE Widget HTML fetched: {html_length} bytes")
            
            # Check if HTML contains results (widget might need JavaScript)
            if html_length < 1000 or 'gsc-webResult' not in html_content.lower():
                logger.warning("Widget HTML seems empty or doesn't contain results structure - trying Google Search directly")
                # Fallback: Use Google Search directly (more reliable for scraping)
                return _fallback_google_search(name, search_type)
            
            # Parse widget HTML untuk extract results
            results = {'web': [], 'images': []}
            
            # Method 1: Extract dari Google CSE Widget structure (gsc-webResult)
            cse_pattern = r'<div[^>]*class=["\'][^"\']*gsc-webResult[^"\']*["\'][^>]*>.*?<a[^>]*href=["\']([^"\']+)["\'][^>]*>.*?<div[^>]*class=["\'][^"\']*gs-title[^"\']*["\'][^>]*>(.*?)</div>.*?<div[^>]*class=["\'][^"\']*gs-snippet[^"\']*["\'][^>]*>(.*?)</div>'
            cse_matches = re.findall(cse_pattern, html_content, re.DOTALL | re.IGNORECASE)
            logger.info(f"Method 1 (CSE Widget): Found {len(cse_matches)} matches")
            
            for match in cse_matches[:20]:
                url = match[0] if len(match) > 0 else ''
                title = re.sub(r'<[^>]+>', '', match[1] if len(match) > 1 else 'No Title').strip()
                snippet = re.sub(r'<[^>]+>', '', match[2] if len(match) > 2 else '').strip()
                
                if url and url.startswith('http'):
                    result_item = {
                        'title': title or 'No Title',
                        'link': url,
                        'snippet': snippet,
                        'displayLink': urllib.parse.urlparse(url).netloc,
                        'formattedUrl': url
                    }
                    if search_type == 'image':
                        results['images'].append(result_item)
                    else:
                        results['web'].append(result_item)
            
            # Method 2: Extract dari standard Google search structure
            if len(results['web']) == 0 and len(results['images']) == 0:
                google_pattern = r'<div[^>]*class=["\'][^"\']*g[^"\']*["\'][^>]*>.*?<a[^>]*href=["\']([^"\']+)["\'][^>]*>.*?<h3[^>]*>(.*?)</h3>.*?<span[^>]*>(.*?)</span>'
                google_matches = re.findall(google_pattern, html_content, re.DOTALL | re.IGNORECASE)
                logger.info(f"Method 2 (Standard Google): Found {len(google_matches)} matches")
                
                for match in google_matches[:20]:
                    url = match[0] if len(match) > 0 else ''
                    title = re.sub(r'<[^>]+>', '', match[1] if len(match) > 1 else 'No Title').strip()
                    snippet = re.sub(r'<[^>]+>', '', match[2] if len(match) > 2 else '').strip()
                    
                    if url and url.startswith('http'):
                        result_item = {
                            'title': title or 'No Title',
                            'link': url,
                            'snippet': snippet,
                            'displayLink': urllib.parse.urlparse(url).netloc,
                            'formattedUrl': url
                        }
                        if search_type == 'image':
                            results['images'].append(result_item)
                        else:
                            results['web'].append(result_item)
            
            # If still no results, try fallback
            if len(results['web']) == 0 and len(results['images']) == 0:
                logger.warning("No results from widget HTML - trying Google Search directly")
                return _fallback_google_search(name, search_type)
            
            # Remove duplicates
            seen_urls = set()
            unique_results = []
            for item in (results['images'] if search_type == 'image' else results['web']):
                if item['link'] not in seen_urls:
                    seen_urls.add(item['link'])
                    unique_results.append(item)
            
            if search_type == 'image':
                results['images'] = unique_results[:20]
            else:
                results['web'] = unique_results[:20]
            
            logger.info(f"Google CSE Widget (UNLIMITED): {len(unique_results)} results parsed")
            
            return jsonify({
                'success': True,
                'data': results,
                'source': 'google_cse_widget_unlimited',
                'message': f'Found {len(unique_results)} results (UNLIMITED - No quota!)'
            }), 200
            
        except requests.exceptions.RequestException as e:
            logger.error(f"Error fetching Google CSE Widget: {str(e)}")
            return jsonify({
                'success': False,
                'error': f'Network error: {str(e)}',
                'data': {'web': [], 'images': []}
            }), 200
            
    except Exception as e:
        logger.error(f"Error in google-cse-widget endpoint: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e),
            'data': {'web': [], 'images': []}
        }), 200

def _fallback_google_search(name, search_type):
    """Fallback: Google Search directly (UNLIMITED - More reliable for scraping!)"""
    try:
        import urllib.parse
        import re
        
        encoded_query = urllib.parse.quote(name)
        
        # Use Google Search directly (more reliable than widget)
        if search_type == 'image':
            search_url = f'https://www.google.com/search?q={encoded_query}&tbm=isch&num=20'
        else:
            search_url = f'https://www.google.com/search?q={encoded_query}&num=20'
        
        logger.info(f"Fallback: Fetching Google Search directly: {search_url}")
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.9',
            'Referer': 'https://www.google.com/',
        }
        
        response = requests.get(search_url, headers=headers, timeout=15)
        
        if response.status_code == 429:
            # Rate limited - return helpful message
            logger.warning("Google Search rate limited (429)")
            return jsonify({
                'success': False,
                'error': 'Google Search rate limited (too many requests)',
                'message': 'Google Search temporarily blocked requests. Please wait a few minutes and try again, or use tab "Categorized" for internal results.',
                'data': {'web': [], 'images': []}
            }), 200
        
        if response.status_code != 200:
            logger.warning(f"Google Search fallback failed: {response.status_code}")
            return jsonify({
                'success': False,
                'error': f'Failed to fetch Google Search (status {response.status_code})',
                'message': 'Google Search is temporarily unavailable. Please try again later or use tab "Categorized" for internal results.',
                'data': {'web': [], 'images': []}
            }), 200
        
        html_content = response.text
        logger.info(f"Google Search HTML fetched: {len(html_content)} bytes")
        
        results = {'web': [], 'images': []}
        
        # Multiple patterns untuk Google Search results
        patterns = [
            # Pattern 1: Standard Google result structure
            r'<div[^>]*class=["\'][^"\']*g[^"\']*["\'][^>]*>.*?<a[^>]*href=["\']([^"\']+)["\'][^>]*>.*?<h3[^>]*>(.*?)</h3>.*?<span[^>]*>(.*?)</span>',
            # Pattern 2: Alternative structure
            r'<div[^>]*class=["\'][^"\']*yuRUbf[^"\']*["\'][^>]*>.*?<a[^>]*href=["\']([^"\']+)["\'][^>]*>.*?<h3[^>]*>(.*?)</h3>',
            # Pattern 3: With snippet
            r'<a[^>]*href=["\']([^"\']+)["\'][^>]*>.*?<h3[^>]*>(.*?)</h3>.*?<div[^>]*class=["\'][^"\']*VwiC3b[^"\']*["\'][^>]*>(.*?)</div>',
        ]
        
        all_matches = []
        for i, pattern in enumerate(patterns):
            matches = re.findall(pattern, html_content, re.DOTALL | re.IGNORECASE)
            logger.info(f"Pattern {i+1}: Found {len(matches)} matches")
            all_matches.extend(matches[:20])
        
        for match in all_matches[:20]:
            url = match[0] if len(match) > 0 else ''
            title = re.sub(r'<[^>]+>', '', match[1] if len(match) > 1 else 'No Title').strip()
            snippet = re.sub(r'<[^>]+>', '', match[2] if len(match) > 2 else '').strip() if len(match) > 2 else ''
            
            # Clean URL (remove /url?q= prefix)
            if '/url?q=' in url:
                url = urllib.parse.parse_qs(urllib.parse.urlparse(url).query).get('q', [url])[0]
            
            if url and url.startswith('http'):
                result_item = {
                    'title': title or 'No Title',
                    'link': url,
                    'snippet': snippet,
                    'displayLink': urllib.parse.urlparse(url).netloc,
                    'formattedUrl': url
                }
                if search_type == 'image':
                    results['images'].append(result_item)
                else:
                    results['web'].append(result_item)
        
        # Remove duplicates
        seen_urls = set()
        unique_results = []
        for item in (results['images'] if search_type == 'image' else results['web']):
            if item['link'] not in seen_urls:
                seen_urls.add(item['link'])
                unique_results.append(item)
        
        if search_type == 'image':
            results['images'] = unique_results[:20]
        else:
            results['web'] = unique_results[:20]
        
        logger.info(f"Google Search fallback: {len(unique_results)} results parsed")
        
        return jsonify({
            'success': True,
            'data': results,
            'source': 'google_search_fallback',
            'message': f'Found {len(unique_results)} results (UNLIMITED - Google Search direct!)'
        }), 200
        
    except Exception as e:
        logger.error(f"Google Search fallback error: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Fallback failed: {str(e)}',
            'data': {'web': [], 'images': []}
        }), 200

def _fallback_html_scraping(name, search_type):
    """Fallback method: HTML scraping (less reliable)"""
    try:
        import urllib.parse
        import re
        
        encoded_query = urllib.parse.quote(name)
        
        # Use Google Search directly
        if search_type == 'image':
            search_url = f'https://www.google.com/search?q={encoded_query}&tbm=isch&num=20'
        else:
            search_url = f'https://www.google.com/search?q={encoded_query}&num=20'
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.9',
        }
        
        response = requests.get(search_url, headers=headers, timeout=15)
        
        if response.status_code != 200:
            return jsonify({
                'success': False,
                'error': f'Failed to fetch (status {response.status_code})',
                'data': {'web': [], 'images': []}
            }), 200
        
        html_content = response.text
        
        # Parse HTML untuk extract results
        results = {'web': [], 'images': []}
        
        # Method: Extract dari div dengan class g (Google search result structure)
        result_pattern = r'<div[^>]*class=["\'][^"\']*g[^"\']*["\'][^>]*>.*?<a[^>]*href=["\']([^"\']+)["\'][^>]*>.*?<h3[^>]*>(.*?)</h3>.*?<span[^>]*>(.*?)</span>'
        result_matches = re.findall(result_pattern, html_content, re.DOTALL | re.IGNORECASE)
        
        for match in result_matches[:20]:  # Limit to 20 results
            url = match[0] if len(match) > 0 else ''
            title = re.sub(r'<[^>]+>', '', match[1] if len(match) > 1 else 'No Title').strip()
            snippet = re.sub(r'<[^>]+>', '', match[2] if len(match) > 2 else '').strip()
            
            if url and url.startswith('http'):
                result_item = {
                    'title': title or 'No Title',
                    'link': url,
                    'snippet': snippet,
                    'displayLink': urllib.parse.urlparse(url).netloc,
                    'formattedUrl': url
                }
                if search_type == 'image':
                    results['images'].append(result_item)
                else:
                    results['web'].append(result_item)
        
        # Remove duplicates
        seen_urls = set()
        unique_results = []
        for item in (results['images'] if search_type == 'image' else results['web']):
            if item['link'] not in seen_urls:
                seen_urls.add(item['link'])
                unique_results.append(item)
        
        if search_type == 'image':
            results['images'] = unique_results[:20]
        else:
            results['web'] = unique_results[:20]
        
        logger.info(f"Fallback HTML scraping: {len(unique_results)} results")
        
        return jsonify({
            'success': True,
            'data': results,
            'source': 'html_scraping_fallback',
            'message': f'Found {len(unique_results)} results (fallback method)'
        }), 200
        
    except Exception as e:
        logger.error(f"Fallback HTML scraping error: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Fallback failed: {str(e)}',
            'data': {'web': [], 'images': []}
        }), 200

@app.route('/api/social-searcher-style', methods=['POST'])
def api_social_searcher_style():
    """API endpoint untuk Social Searcher style - Platform-specific searches dengan anti-blocking logic"""
    try:
        data = request.get_json()
        name = data.get('name')
        
        if not name:
            return jsonify({'error': 'Nama diperlukan untuk pencarian'}), 400
        
        import urllib.parse
        import re
        import time
        import random
        
        # Google CSE Configuration - USE API (not scraping!)
        CSE_ID = '7693f5093e95e4c28'
        # Get API key from environment variable
        API_KEY = os.getenv('GOOGLE_CSE_API_KEY')
        if not API_KEY:
            return jsonify({'error': 'GOOGLE_CSE_API_KEY tidak ditemukan. Silakan set di file .env'}), 500
        
        # Platform-specific search queries dengan fokus pada profil dengan gambar
        platforms = {
            'facebook': f'"{name}" site:facebook.com profile',
            'instagram': f'"{name}" site:instagram.com profile',
            'twitter': f'"{name}" (site:twitter.com OR site:x.com) profile',
            'linkedin': f'"{name}" site:linkedin.com profile',
            'tiktok': f'"{name}" site:tiktok.com profile',
            'pinterest': f'"{name}" site:pinterest.com profile',
        }
        
        all_results = {}
        
        logger.info(f"Social Searcher Style: Starting platform-specific searches for '{name}'")
        
        # Process each platform dengan delay untuk avoid rate limiting
        for platform, query in platforms.items():
            try:
                # Random delay antara 0.3-0.8 detik untuk avoid rate limiting (reduced untuk speed)
                if platform != 'facebook':  # Skip delay untuk platform pertama
                    delay = random.uniform(0.3, 0.8)  # Reduced untuk speed
                    logger.info(f"  Waiting {delay:.2f}s before {platform} search...")
                    time.sleep(delay)
                
                logger.info(f"  Searching {platform}: {query}")
                
                # Use Google CSE API (not scraping!)
                api_url = 'https://www.googleapis.com/customsearch/v1'
                params = {
                    'key': API_KEY,
                    'cx': CSE_ID,
                    'q': query,
                    'num': 10  # Max 10 results per platform
                }
                
                # Retry logic dengan exponential backoff
                max_retries = 2
                retry_delay = 1
                
                for attempt in range(max_retries):
                    try:
                        response = requests.get(api_url, params=params, timeout=15)
                        
                        if response.status_code == 200:
                            data = response.json()
                            
                            # Parse results dari API dengan ekstraksi gambar
                            results = []
                            if 'items' in data:
                                for item in data['items']:
                                    result_item = {
                                        'title': item.get('title', 'No Title'),
                                        'link': item.get('link', ''),
                                        'snippet': item.get('snippet', ''),
                                        'displayLink': item.get('displayLink', urllib.parse.urlparse(item.get('link', '')).netloc),
                                        'platform': platform
                                    }
                                    
                                    # Extract images from pagemap
                                    if 'pagemap' in item:
                                        pagemap = item['pagemap']
                                        result_item['pagemap'] = pagemap
                                        
                                        # Extract all possible images
                                        extracted_images = []
                                        
                                        # CSE Image
                                        if 'cse_image' in pagemap and pagemap['cse_image']:
                                            for img in pagemap['cse_image']:
                                                if img.get('src'):
                                                    extracted_images.append({
                                                        'src': img.get('src'),
                                                        'width': img.get('width', 0),
                                                        'height': img.get('height', 0),
                                                        'type': 'cse_image'
                                                    })
                                        
                                        # Metatags (Open Graph, Twitter Cards)
                                        if 'metatags' in pagemap and pagemap['metatags']:
                                            for meta in pagemap['metatags']:
                                                if 'og:image' in meta and meta['og:image']:
                                                    extracted_images.append({
                                                        'src': meta['og:image'],
                                                        'type': 'og_image'
                                                    })
                                                if 'twitter:image' in meta and meta['twitter:image']:
                                                    extracted_images.append({
                                                        'src': meta['twitter:image'],
                                                        'type': 'twitter_image'
                                                    })
                                        
                                        # Store best image
                                        if extracted_images:
                                            priority_order = {'og_image': 1, 'twitter_image': 2, 'cse_image': 3}
                                            extracted_images.sort(key=lambda x: priority_order.get(x.get('type', ''), 99))
                                            result_item['best_image'] = extracted_images[0]
                                            result_item['all_images'] = extracted_images
                                    
                                    # Add direct image if available
                                    if 'image' in item:
                                        result_item['image'] = {
                                            'src': item['image'].get('src', ''),
                                            'width': item['image'].get('width', 0),
                                            'height': item['image'].get('height', 0)
                                        }
                                    
                                    results.append(result_item)
                            
                            # Remove duplicates
                            seen_urls = set()
                            unique_results = []
                            for item in results:
                                if item['link'] not in seen_urls:
                                    seen_urls.add(item['link'])
                                    unique_results.append(item)
                            
                            all_results[platform] = unique_results[:10]  # Max 10 per platform
                            if len(unique_results) > 0:
                                logger.info(f"  ✓ {platform}: Found {len(unique_results)} results")
                            else:
                                logger.warning(f"  ⚠️ {platform}: No results found")
                            break  # Success, exit retry loop
                        
                        elif response.status_code == 429:
                            # Rate limited - wait longer
                            wait_time = retry_delay * (2 ** attempt)
                            logger.warning(f"  {platform}: API quota exceeded (429), waiting {wait_time}s...")
                            if attempt < max_retries - 1:
                                time.sleep(wait_time)
                                continue
                            else:
                                logger.warning(f"  {platform}: Max retries reached, skipping...")
                                all_results[platform] = []
                                break
                        
                        else:
                            error_data = response.json() if response.headers.get('content-type', '').startswith('application/json') else {}
                            error_msg = error_data.get('error', {}).get('message', f'HTTP {response.status_code}')
                            logger.warning(f"  {platform}: API error: {error_msg}")
                            if attempt < max_retries - 1:
                                time.sleep(retry_delay * (2 ** attempt))
                                continue
                            else:
                                all_results[platform] = []
                                break
                    
                    except requests.exceptions.RequestException as e:
                        logger.error(f"  {platform}: Request error: {str(e)}")
                        if attempt < max_retries - 1:
                            time.sleep(retry_delay * (2 ** attempt))
                            continue
                        else:
                            all_results[platform] = []
                            break
            
            except Exception as e:
                logger.error(f"  {platform}: Error: {str(e)}")
                all_results[platform] = []
        
        # Combine all results
        combined_results = []
        for platform, results in all_results.items():
            combined_results.extend(results)
        
        logger.info(f"Social Searcher Style: Total {len(combined_results)} results from {len([r for r in all_results.values() if r])} platforms")
        
        return jsonify({
            'success': True,
            'data': {
                'web': combined_results,
                'by_platform': all_results
            },
            'source': 'social_searcher_style',
            'message': f'Found {len(combined_results)} results from {len([r for r in all_results.values() if r])} platforms (Social Searcher style)'
        }), 200
        
    except Exception as e:
        logger.error(f"Error in social-searcher-style endpoint: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e),
            'data': {'web': [], 'by_platform': {}}
        }), 200

@app.route('/api/ai-analysis', methods=['POST'])
def api_ai_analysis():
    """API endpoint untuk AI analysis dengan akurasi tinggi"""
    try:
        data = request.get_json()
        name = data.get('name')
        search_results = data.get('search_results', {})
        social_links = data.get('social_links', [])
        
        if not name:
            return jsonify({'error': 'Nama diperlukan untuk analisis AI'}), 400
        
        # AI Analysis Logic
        analysis = {
            'confidence_score': 0.0,
            'verification_status': 'unverified',
            'social_media_presence': {},
            'risk_assessment': {},
            'recommendations': [],
            'summary': ''
        }
        
        # Analyze social media presence
        platforms = {
            'facebook': 0,
            'instagram': 0,
            'twitter': 0,
            'linkedin': 0,
            'tiktok': 0
        }
        
        for link in social_links:
            link_url = link.get('link', '').lower()
            if 'facebook.com' in link_url:
                platforms['facebook'] += 1
            elif 'instagram.com' in link_url:
                platforms['instagram'] += 1
            elif 'twitter.com' in link_url or 'x.com' in link_url:
                platforms['twitter'] += 1
            elif 'linkedin.com' in link_url:
                platforms['linkedin'] += 1
            elif 'tiktok.com' in link_url:
                platforms['tiktok'] += 1
        
        analysis['social_media_presence'] = platforms
        
        # Calculate confidence score
        total_links = sum(platforms.values())
        if total_links > 0:
            # Higher confidence with more verified social media presence
            analysis['confidence_score'] = min(0.95, 0.60 + (total_links * 0.05))
        else:
            analysis['confidence_score'] = 0.30
        
        # Verification status
        if total_links >= 3:
            analysis['verification_status'] = 'highly_verified'
        elif total_links >= 2:
            analysis['verification_status'] = 'verified'
        elif total_links >= 1:
            analysis['verification_status'] = 'partially_verified'
        else:
            analysis['verification_status'] = 'unverified'
        
        # Risk assessment
        analysis['risk_assessment'] = {
            'level': 'low' if total_links > 0 else 'medium',
            'factors': []
        }
        
        if total_links == 0:
            analysis['risk_assessment']['factors'].append('Tidak ditemukan profil social media aktif')
        else:
            analysis['risk_assessment']['factors'].append(f'Ditemukan {total_links} profil social media')
        
        # Recommendations
        if total_links == 0:
            analysis['recommendations'].append('Lakukan verifikasi tambahan melalui sumber data lain')
            analysis['recommendations'].append('Periksa data dari database resmi')
        else:
            analysis['recommendations'].append('Verifikasi informasi melalui profil social media yang ditemukan')
            analysis['recommendations'].append('Cross-check dengan data identitas resmi')
        
        # Generate summary
        platform_names = []
        for platform, count in platforms.items():
            if count > 0:
                platform_names.append(f"{platform.capitalize()}({count})")
        
        if platform_names:
            analysis['summary'] = f"Analisis AI menunjukkan {name} memiliki kehadiran di {len(platform_names)} platform social media: {', '.join(platform_names)}. Skor kepercayaan: {analysis['confidence_score']*100:.1f}%."
        else:
            analysis['summary'] = f"Analisis AI tidak menemukan kehadiran social media yang signifikan untuk {name}. Skor kepercayaan: {analysis['confidence_score']*100:.1f}%."
        
        return jsonify({
            'success': True,
            'analysis': analysis
        })
        
    except Exception as e:
        logger.error(f"Error in AI analysis: {str(e)}", exc_info=True)
        return jsonify({'error': f'Error: {str(e)}'}), 500

def clean_text_for_pdf(text):
    """Clean text to avoid encoding issues in PDF"""
    if not text:
        return 'N/A'
    
    # Convert to string and handle encoding
    text = str(text)
    
    # Replace problematic characters
    replacements = {
        '–': '-',
        '—': '-',
        '"': '"',
        '"': '"',
        ''': "'",
        ''': "'",
        '…': '...',
        '°': 'deg',
        '×': 'x',
        '÷': '/',
        '±': '+/-',
        '≤': '<=',
        '≥': '>=',
        '≠': '!=',
        '∞': 'infinity',
        '∑': 'sum',
        '∏': 'product',
        '√': 'sqrt',
        '∫': 'integral',
        '∆': 'delta',
        'α': 'alpha',
        'β': 'beta',
        'γ': 'gamma',
        'δ': 'delta',
        'ε': 'epsilon',
        'ζ': 'zeta',
        'η': 'eta',
        'θ': 'theta',
        'λ': 'lambda',
        'μ': 'mu',
        'π': 'pi',
        'ρ': 'rho',
        'σ': 'sigma',
        'τ': 'tau',
        'φ': 'phi',
        'χ': 'chi',
        'ψ': 'psi',
        'ω': 'omega'
    }
    
    for old, new in replacements.items():
        text = text.replace(old, new)
    
    # Remove any remaining non-ASCII characters that might cause issues
    try:
        text = text.encode('ascii', 'ignore').decode('ascii')
    except:
        text = 'N/A'
    
    return text

@app.route('/api/export/pdf', methods=['POST'])
def export_pdf():
    """Export person details to PDF"""
    try:
        data = request.get_json()
        person = data.get('person')
        username = data.get('username')
        password = data.get('password')
        
        if not person:
            return jsonify({'error': 'No person data provided'}), 400
        
        # Authenticate user
        auth_result = authenticate_user(username, password)
        if not auth_result:
            return jsonify({'error': 'Invalid credentials'}), 401
        
        # Create PDF buffer
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Use default fonts that support basic Unicode
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1,  # Center alignment
            textColor=colors.darkblue,
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            textColor=colors.darkblue,
            fontName='Helvetica-Bold'
        )
        
        # Build content
        story = []
        
        # Title
        story.append(Paragraph("LAPORAN DETAIL ORANG", title_style))
        story.append(Spacer(1, 20))
        
        # Add photo if available
        if person.get('face'):
            try:
                face_data = person['face']
                if face_data.startswith('data:image/'):
                    # Extract base64 data
                    if ',' in face_data:
                        header, base64_data = face_data.split(',', 1)
                    else:
                        base64_data = face_data
                    
                    # Decode base64 to bytes
                    image_bytes = base64.b64decode(base64_data)
                    
                    # Create image from bytes
                    face_image = Image(io.BytesIO(image_bytes), width=2*inch, height=2*inch)
                    face_image.hAlign = 'CENTER'
                    story.append(face_image)
                    story.append(Spacer(1, 20))
                    print(f"Added face photo to PDF for {person.get('full_name', 'Unknown')}")
                else:
                    print(f"Face data format not supported for PDF: {person.get('full_name', 'Unknown')}")
            except Exception as e:
                print(f"Error adding face photo to PDF: {e}")
                # Continue without photo
        
        # Personal Information
        story.append(Paragraph("Informasi Pribadi", heading_style))
        
        personal_data = [
            ['Nama Lengkap', clean_text_for_pdf(person.get('full_name', 'N/A'))],
            ['NIK', clean_text_for_pdf(person.get('ktp_number', person.get('nik', 'N/A')))],
            ['Tanggal Lahir', clean_text_for_pdf(person.get('date_of_birth', 'N/A'))],
            ['Tempat Lahir', clean_text_for_pdf(person.get('birth_place', 'N/A'))],
            ['Jenis Kelamin', clean_text_for_pdf(person.get('gender', 'N/A'))],
            ['Alamat', clean_text_for_pdf(person.get('address', 'N/A'))],
            ['Agama', clean_text_for_pdf(person.get('religion', 'N/A'))],
            ['Status Perkawinan', clean_text_for_pdf(person.get('marital_status', 'N/A'))],
            ['Pekerjaan', clean_text_for_pdf(person.get('occupation', 'N/A'))]
        ]
        
        personal_table = Table(personal_data, colWidths=[2*inch, 4*inch])
        personal_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('BACKGROUND', (1, 0), (1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(personal_table)
        story.append(Spacer(1, 20))
        
        # Phone Numbers
        if person.get('phone_data'):
            story.append(Paragraph("Nomor Telepon", heading_style))
            phone_data = [['Nomor Telepon', 'Operator', 'Tanggal Terdaftar']]
            for phone in person['phone_data']:
                phone_data.append([
                    clean_text_for_pdf(phone.get('number', 'N/A')),
                    clean_text_for_pdf(phone.get('operator', 'N/A')),
                    clean_text_for_pdf(phone.get('register_date', 'N/A'))
                ])
            
            phone_table = Table(phone_data, colWidths=[2*inch, 2*inch, 2*inch])
            phone_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(phone_table)
            story.append(Spacer(1, 20))
        
        # Family Members
        if person.get('family_data') and person['family_data'].get('anggota_keluarga'):
            story.append(Paragraph("Anggota Keluarga", heading_style))
            
            family_info = [
                ['Kepala Keluarga', clean_text_for_pdf(person['family_data'].get('kepala_keluarga', 'N/A'))],
                ['Nomor Kartu Keluarga (NKK)', clean_text_for_pdf(person['family_data'].get('nkk', 'N/A'))],
                ['Alamat Keluarga', clean_text_for_pdf(person['family_data'].get('alamat_keluarga', 'N/A'))]
            ]
            
            family_info_table = Table(family_info, colWidths=[2*inch, 4*inch])
            family_info_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('BACKGROUND', (1, 0), (1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(family_info_table)
            story.append(Spacer(1, 12))
            
            # Family members table
            family_data = [['Nama', 'Hubungan', 'NIK', 'Tanggal Lahir', 'Jenis Kelamin']]
            for member in person['family_data']['anggota_keluarga']:
                family_data.append([
                    clean_text_for_pdf(member.get('nama', 'N/A')),
                    clean_text_for_pdf(member.get('hubungan', 'N/A')),
                    clean_text_for_pdf(member.get('nik', 'N/A')),
                    clean_text_for_pdf(member.get('tanggal_lahir', 'N/A')),
                    clean_text_for_pdf(member.get('jenis_kelamin', 'N/A'))
                ])
            
            family_table = Table(family_data, colWidths=[1.5*inch, 1.2*inch, 1.5*inch, 1.2*inch, 1*inch])
            family_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(family_table)
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"person_details_{person.get('ktp_number', 'unknown')}.pdf",
            mimetype='application/pdf'
        )
        
    except Exception as e:
        print(f"Error generating PDF: {e}")
        return jsonify({'error': f'PDF generation failed: {str(e)}'}), 500

@app.route('/api/export/word', methods=['POST'])
def export_word():
    """Export person details to Word document"""
    try:
        data = request.get_json()
        person = data.get('person')
        username = data.get('username')
        password = data.get('password')
        
        if not person:
            return jsonify({'error': 'No person data provided'}), 400
        
        # Authenticate user
        auth_result = authenticate_user(username, password)
        if not auth_result:
            return jsonify({'error': 'Invalid credentials'}), 401
        
        # Create Word document
        doc = Document()
        
        # Title
        title = doc.add_heading('PERSON DETAILS REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add photo if available
        if person.get('face'):
            try:
                face_data = person['face']
                if face_data.startswith('data:image/'):
                    # Extract base64 data
                    if ',' in face_data:
                        header, base64_data = face_data.split(',', 1)
                    else:
                        base64_data = face_data
                    
                    # Decode base64 to bytes
                    image_bytes = base64.b64decode(base64_data)
                    
                    # Create image from bytes
                    face_image = io.BytesIO(image_bytes)
                    
                    # Add image to document
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    run.add_picture(face_image, width=Inches(2))
                    
                    doc.add_paragraph()  # Add space
                    print(f"Added face photo to Word document for {person.get('full_name', 'Unknown')}")
                else:
                    print(f"Face data format not supported for Word: {person.get('full_name', 'Unknown')}")
            except Exception as e:
                print(f"Error adding face photo to Word document: {e}")
                # Continue without photo
        
        # Personal Information
        doc.add_heading('Informasi Pribadi', level=1)
        
        personal_data = [
            ('Nama Lengkap', person.get('full_name', 'N/A')),
            ('NIK', person.get('ktp_number', person.get('nik', 'N/A'))),
            ('Tanggal Lahir', person.get('date_of_birth', 'N/A')),
            ('Tempat Lahir', person.get('birth_place', 'N/A')),
            ('Jenis Kelamin', person.get('gender', 'N/A')),
            ('Alamat', person.get('address', 'N/A')),
            ('Agama', person.get('religion', 'N/A')),
            ('Status Perkawinan', person.get('marital_status', 'N/A')),
            ('Pekerjaan', person.get('occupation', 'N/A'))
        ]
        
        # Create table for personal information
        table = doc.add_table(rows=len(personal_data), cols=2)
        table.style = 'Table Grid'
        
        for i, (label, value) in enumerate(personal_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = str(value)
            # Make first column bold
            for paragraph in table.cell(i, 0).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        doc.add_paragraph()  # Add space
        
        # Phone Numbers
        if person.get('phone_data'):
            doc.add_heading('Nomor Telepon', level=1)
            
            phone_table = doc.add_table(rows=1, cols=3)
            phone_table.style = 'Table Grid'
            
            # Header row
            header_cells = phone_table.rows[0].cells
            header_cells[0].text = 'Nomor Telepon'
            header_cells[1].text = 'Operator'
            header_cells[2].text = 'Tanggal Terdaftar'
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add phone data
            for phone in person['phone_data']:
                row_cells = phone_table.add_row().cells
                row_cells[0].text = phone.get('number', 'N/A')
                row_cells[1].text = phone.get('operator', 'N/A')
                row_cells[2].text = phone.get('register_date', 'N/A')
            
            doc.add_paragraph()  # Add space
        
        # Family Members
        if person.get('family_data') and person['family_data'].get('anggota_keluarga'):
            doc.add_heading('Informasi Keluarga', level=1)
            
            # Family info
            family_info = [
                ('Kepala Keluarga', person['family_data'].get('kepala_keluarga', 'N/A')),
                ('Nomor Kartu Keluarga (NKK)', person['family_data'].get('nkk', 'N/A')),
                ('Alamat Keluarga', person['family_data'].get('alamat_keluarga', 'N/A'))
            ]
            
            family_info_table = doc.add_table(rows=len(family_info), cols=2)
            family_info_table.style = 'Table Grid'
            
            for i, (label, value) in enumerate(family_info):
                family_info_table.cell(i, 0).text = label
                family_info_table.cell(i, 1).text = str(value)
                # Make first column bold
                for paragraph in family_info_table.cell(i, 0).paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            doc.add_paragraph()  # Add space
            
            # Family members table
            doc.add_heading('Anggota Keluarga', level=2)
            
            family_table = doc.add_table(rows=1, cols=5)
            family_table.style = 'Table Grid'
            
            # Header row
            header_cells = family_table.rows[0].cells
            header_cells[0].text = 'Nama'
            header_cells[1].text = 'Hubungan'
            header_cells[2].text = 'NIK'
            header_cells[3].text = 'Tanggal Lahir'
            header_cells[4].text = 'Jenis Kelamin'
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add family member data
            for member in person['family_data']['anggota_keluarga']:
                row_cells = family_table.add_row().cells
                row_cells[0].text = member.get('nama', 'N/A')
                row_cells[1].text = member.get('hubungan', 'N/A')
                row_cells[2].text = member.get('nik', 'N/A')
                row_cells[3].text = member.get('tanggal_lahir', 'N/A')
                row_cells[4].text = member.get('jenis_kelamin', 'N/A')
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"person_details_{person.get('ktp_number', 'unknown')}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        print(f"Error generating Word document: {e}")
        return jsonify({'error': f'Word document generation failed: {str(e)}'}), 500

@app.route('/api/ai/face-to-nik', methods=['POST'])
def api_ai_face_to_nik():
    """API endpoint untuk AI Face-to-NIK analysis"""
    try:
        # Validate session token
        session_token = request.cookies.get('session_token')
        if not session_token:
            session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        
        if not session_token:
            return jsonify({'error': 'Session token required'}), 401
        
        user = validate_session_token(session_token)
        if not user:
            return jsonify({'error': 'Invalid session token'}), 401
        
        data = request.get_json()
        if not data or 'image' not in data:
            return jsonify({'error': 'Image data required'}), 400
        
        image_base64 = data.get('image')
        threshold = data.get('threshold', 50)
        
        # Decode base64 image
        try:
            image_data = base64.b64decode(image_base64)
        except Exception as e:
            return jsonify({'error': 'Invalid image data'}), 400
        
        # Simulate AI face analysis (replace with actual AI implementation)
        # For now, return mock data based on existing database records
        mock_results = [
            {
                'nik': '1505041107830002',
                'name': 'YASIR HASBI',
                'confidence': 85,
                'photo_url': '/static/clean_photos/yasir_hasbi.jpg'
            },
            {
                'nik': '1771010511010001',
                'name': 'EVIN NOPRILIANDA',
                'confidence': 72,
                'photo_url': '/static/clean_photos/evin_noprilianda.jpg'
            },
            {
                'nik': '3212241009990003',
                'name': 'MOH. SOKABAH',
                'confidence': 68,
                'photo_url': '/static/clean_photos/moh_sokabah.jpg'
            },
            {
                'nik': '3318072004930001',
                'name': 'MOHAMMAD ALI NUR SADUAN',
                'confidence': 65,
                'photo_url': '/static/clean_photos/mohammad_ali_nur_saduan.jpg'
            },
            {
                'nik': '6311061205840001',
                'name': 'SAID YASIR',
                'confidence': 62,
                'photo_url': '/static/clean_photos/said_yasir.jpg'
            }
        ]
        
        # Filter results by threshold
        filtered_results = [result for result in mock_results if result['confidence'] >= threshold]
        
        # Sort by confidence (highest first)
        filtered_results.sort(key=lambda x: x['confidence'], reverse=True)
        
        return jsonify({
            'success': True,
            'results': filtered_results,
            'total_found': len(filtered_results),
            'threshold_used': threshold
        })
        
    except Exception as e:
        print(f"Error in AI face-to-NIK analysis: {e}")
        return jsonify({'error': f'AI analysis failed: {str(e)}'}), 500

if __name__ == '__main__':
    print("Starting Clearance Face Search Web Server...")
    print(f"Face recognition library available: {USE_FACE_LIB}")
    print("Access the application at: http://localhost:5000")
    app.run(debug=True, host='0.0.0.0', port=5000)

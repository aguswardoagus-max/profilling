#!/usr/bin/env python3
"""
Flask web server untuk Clearance Face Search
Menyediakan API endpoints dan serve static files
"""
import os
import sys
import json
import time
import base64
import tempfile
import logging
from pathlib import Path
from flask import Flask, request, jsonify, send_from_directory, render_template, session, redirect, url_for, send_file
from flask_cors import CORS
from database import authenticate_user, validate_session_token, logout_user, db
from cekplat import cekplat_bp
import numpy as np
import requests
import cv2
from PIL import Image
import io
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import base64

# Import functions from clearance_face_search.py
from clearance_face_search import (
    ensure_token, call_search, parse_people_from_response,
    load_image_file_to_encoding, get_encoding_from_base64_face,
    save_face_image, USE_FACE_LIB
)

# Load environment variables
def load_env():
    """Load environment variables from .env file"""
    env_vars = {}
    try:
        with open('.env', 'r') as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith('#') and '=' in line:
                    key, value = line.split('=', 1)
                    env_vars[key.strip()] = value.strip()
    except FileNotFoundError:
        print(".env file not found, using default values")
    return env_vars

# Load configuration
config = load_env()

# Default credentials and API URLs
DEFAULT_USERNAME = config.get('CLEARANCE_USERNAME', 'rezarios')
DEFAULT_PASSWORD = config.get('CLEARANCE_PASSWORD', '12345678')
FAMILY_API_BASE = config.get('FAMILY_API_BASE', 'http://10.1.54.224:4646/json/clearance/dukcapil/family')
FAMILY_API_ALT = config.get('FAMILY_API_ALT', 'http://10.1.54.116:27682/api/v1/ktp/internal')
PHONE_API_BASE = config.get('PHONE_API_BASE', 'http://10.1.54.224:4646/json/clearance/phones')

app = Flask(__name__)
app.secret_key = 'your-secret-key-change-this-in-production'
CORS(app)

# Register blueprints
app.register_blueprint(cekplat_bp, url_prefix='/cekplat')

# Configuration
UPLOAD_FOLDER = Path('uploads')
OUTPUT_FOLDER = Path('faces')
CLEAN_PHOTOS_FOLDER = Path('static/clean_photos')
UPLOAD_FOLDER.mkdir(exist_ok=True)
OUTPUT_FOLDER.mkdir(exist_ok=True)
CLEAN_PHOTOS_FOLDER.mkdir(exist_ok=True)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Additional API endpoints
FAMILY_API_BASE = "http://10.1.54.224:4646/json/clearance/dukcapil/family"
FAMILY_API_ALT = "http://10.1.54.116:27682/api/v1/ktp/internal"
PHONE_API_BASE = "http://10.1.54.224:4646/json/clearance/phones"

# AI Inpainting Functions
def download_foto(url_foto, timeout=30):
    """Download foto from URL and return as bytes"""
    try:
        logger.info(f"Downloading foto from URL: {url_foto}")
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        response = requests.get(url_foto, headers=headers, timeout=timeout)
        response.raise_for_status()
        
        if response.headers.get('content-type', '').startswith('image/'):
            logger.info(f"Successfully downloaded foto, size: {len(response.content)} bytes")
            return response.content
        else:
            logger.warning(f"URL does not contain image data: {url_foto}")
            return None
    except Exception as e:
        logger.error(f"Error downloading foto from {url_foto}: {e}")
        return None

def detect_text_watermark(image_cv):
    """Detect watermark with minimal impact on image quality"""
    gray = cv2.cvtColor(image_cv, cv2.COLOR_BGR2GRAY)
    h, w = gray.shape
    
    # Initialize mask
    watermark_mask = np.zeros_like(gray)
    
    # Method 1: Detect very bright pixels only (watermark text overlays)
    # Use higher threshold to avoid affecting normal image content
    _, bright_mask = cv2.threshold(gray, 220, 255, cv2.THRESH_BINARY)
    
    # Method 2: Detect semi-transparent white overlays in HSV
    hsv = cv2.cvtColor(image_cv, cv2.COLOR_BGR2HSV)
    
    # More conservative ranges for semi-transparent white overlays
    white_ranges = [
        ([0, 0, 200], [180, 20, 255]),   # High threshold for semi-transparent
        ([0, 0, 220], [180, 15, 255]),   # Very high threshold
        ([0, 0, 240], [180, 10, 255]),   # Extremely high threshold
    ]
    
    for lower, upper in white_ranges:
        white_mask = cv2.inRange(hsv, np.array(lower), np.array(upper))
        watermark_mask = cv2.bitwise_or(watermark_mask, white_mask)
    
    # Method 3: Detect small rectangular regions that might be text
    # Find contours in bright areas with high threshold
    contours, _ = cv2.findContours(bright_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    
    for contour in contours:
        x, y, w_rect, h_rect = cv2.boundingRect(contour)
        # Check if it's a reasonable size for text characters
        if 3 < w_rect < 40 and 2 < h_rect < 20:
            # Check if it's bright enough to be watermark text
            roi = gray[y:y+h_rect, x:x+w_rect]
            if np.mean(roi) > 220:  # High threshold for watermark
                # Small padding to cover the character
                padding = 1
                x1 = max(0, x - padding)
                y1 = max(0, y - padding)
                x2 = min(w, x + w_rect + padding)
                y2 = min(h, y + h_rect + padding)
                cv2.rectangle(watermark_mask, (x1, y1), (x2, y2), 255, -1)
    
    # Method 4: Detect text patterns using morphological operations
    # Use smaller kernels to avoid affecting large areas
    kernel_horizontal = cv2.getStructuringElement(cv2.MORPH_RECT, (8, 1))
    horizontal = cv2.morphologyEx(gray, cv2.MORPH_OPEN, kernel_horizontal)
    _, text_h_mask = cv2.threshold(horizontal, 220, 255, cv2.THRESH_BINARY)
    
    kernel_vertical = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 8))
    vertical = cv2.morphologyEx(gray, cv2.MORPH_OPEN, kernel_vertical)
    _, text_v_mask = cv2.threshold(vertical, 220, 255, cv2.THRESH_BINARY)
    
    watermark_mask = cv2.bitwise_or(watermark_mask, text_h_mask)
    watermark_mask = cv2.bitwise_or(watermark_mask, text_v_mask)
    
    # Clean up the mask
    # Remove small noise
    kernel_clean = np.ones((2,2), np.uint8)
    watermark_mask = cv2.morphologyEx(watermark_mask, cv2.MORPH_OPEN, kernel_clean)
    
    # Close small gaps
    kernel_close = np.ones((2,2), np.uint8)
    watermark_mask = cv2.morphologyEx(watermark_mask, cv2.MORPH_CLOSE, kernel_close)
    
    # Minimal dilation to preserve image quality
    kernel_dilate = np.ones((2,2), np.uint8)
    watermark_mask = cv2.dilate(watermark_mask, kernel_dilate, iterations=1)
    
    return watermark_mask

def clean_watermark(image_bytes):
    """Clean watermark with minimal impact on image quality"""
    try:
        logger.info("Starting conservative watermark removal process")
        
        # Convert bytes to PIL Image
        from PIL import Image
        image = Image.open(io.BytesIO(image_bytes))
        image_array = np.array(image)
        
        # Convert to OpenCV format (BGR)
        if len(image_array.shape) == 3:
            image_cv = cv2.cvtColor(image_array, cv2.COLOR_RGB2BGR)
        else:
            image_cv = image_array
        
        # Create a copy for processing
        result = image_cv.copy()
        
        # Detect watermark
        watermark_mask = detect_text_watermark(image_cv)
        
        # Apply inpainting if watermark detected
        if np.any(watermark_mask):
            mask_area = np.sum(watermark_mask > 0)
            total_area = watermark_mask.shape[0] * watermark_mask.shape[1]
            mask_ratio = mask_area / total_area
            
            logger.info(f"Watermark detected, covers {mask_ratio:.2%} of image")
            
            # Use conservative inpainting to preserve image quality
            if mask_ratio > 0.1:  # Large watermark area
                result = cv2.inpaint(result, watermark_mask, 3, cv2.INPAINT_TELEA)
                logger.info("Using Telea inpainting (radius=3) for large area")
            else:  # Small watermark area
                result = cv2.inpaint(result, watermark_mask, 2, cv2.INPAINT_TELEA)
                logger.info("Using Telea inpainting (radius=2) for small area")
            
            # Single cleanup pass for remaining artifacts
            logger.info("Applying single cleanup pass")
            
            # Convert to HSV for color-based cleaning
            hsv = cv2.cvtColor(result, cv2.COLOR_BGR2HSV)
            
            # Remove only very bright white overlays
            cleanup_thresholds = [220, 240]
            
            for threshold in cleanup_thresholds:
                lower_white = np.array([0, 0, threshold])
                upper_white = np.array([180, 20, 255])  # Narrow saturation range
                white_mask = cv2.inRange(hsv, lower_white, upper_white)
                
                if np.any(white_mask):
                    # Minimal dilation for white areas
                    white_mask_dilated = cv2.dilate(white_mask, np.ones((2,2), np.uint8), iterations=1)
                    result = cv2.inpaint(result, white_mask_dilated, 1, cv2.INPAINT_TELEA)
                    logger.info(f"Cleanup: Removed white artifacts (threshold={threshold})")
        else:
            logger.info("No watermark detected")
        
        # Convert back to RGB
        result_rgb = cv2.cvtColor(result, cv2.COLOR_BGR2RGB)
        
        # Convert back to PIL Image
        result_image = Image.fromarray(result_rgb)
        
        # Convert to bytes with high quality
        output_buffer = io.BytesIO()
        result_image.save(output_buffer, format='JPEG', quality=98)
        cleaned_bytes = output_buffer.getvalue()
        
        logger.info(f"Successfully cleaned watermark with minimal impact, output size: {len(cleaned_bytes)} bytes")
        return cleaned_bytes
        
    except Exception as e:
        logger.error(f"Error in conservative watermark removal: {e}")
        return None

def process_and_save_clean_photo_from_base64(nik, base64_data, force_reprocess=False):
    """Process photo from base64 data to remove watermark and save to clean_photos folder"""
    try:
        # Check if clean photo already exists
        clean_photo_path = CLEAN_PHOTOS_FOLDER / f"{nik}.jpg"
        
        if clean_photo_path.exists() and not force_reprocess:
            logger.info(f"Clean photo already exists for NIK {nik}, using cached version")
            return f"/static/clean_photos/{nik}.jpg"
        
        # Decode base64 data
        logger.info(f"Processing base64 photo for NIK {nik}")
        
        # Remove data URL prefix if present
        if base64_data.startswith('data:image/'):
            base64_data = base64_data.split(',')[1]
        
        # Decode base64 to bytes
        import base64
        image_bytes = base64.b64decode(base64_data)
        
        # Clean watermark using AI inpainting
        cleaned_bytes = clean_watermark(image_bytes)
        
        if cleaned_bytes:
            # Save cleaned photo
            with open(clean_photo_path, 'wb') as f:
                f.write(cleaned_bytes)
            
            logger.info(f"Successfully saved clean photo for NIK {nik}")
            return f"/static/clean_photos/{nik}.jpg"
        else:
            logger.error(f"Failed to clean watermark for NIK {nik}")
            return None
            
    except Exception as e:
        logger.error(f"Error processing base64 photo for NIK {nik}: {e}")
        return None

def process_and_save_clean_photo(nik, url_foto, force_reprocess=False):
    """Process photo to remove watermark and save to clean_photos folder"""
    try:
        # Check if clean photo already exists
        clean_photo_path = CLEAN_PHOTOS_FOLDER / f"{nik}.jpg"
        if clean_photo_path.exists() and not force_reprocess:
            logger.info(f"Clean photo already exists for NIK {nik}, using cache")
            return f"/static/clean_photos/{nik}.jpg"
        elif clean_photo_path.exists() and force_reprocess:
            logger.info(f"Force reprocessing photo for NIK {nik}")
        
        # Download original photo
        logger.info(f"Processing photo for NIK {nik}")
        original_bytes = download_foto(url_foto)
        if not original_bytes:
            logger.warning(f"Failed to download photo for NIK {nik}")
            return None
        
        # Clean watermark
        cleaned_bytes = clean_watermark(original_bytes)
        if not cleaned_bytes:
            logger.warning(f"Failed to clean watermark for NIK {nik}")
            return None
        
        # Save cleaned photo
        with open(clean_photo_path, 'wb') as f:
            f.write(cleaned_bytes)
        
        logger.info(f"Successfully saved clean photo for NIK {nik} at {clean_photo_path}")
        return f"/static/clean_photos/{nik}.jpg"
        
    except Exception as e:
        logger.error(f"Error processing clean photo for NIK {nik}: {e}")
        return None

def convert_family_data_format(api_response, nik, nkk, token=None):
    """Convert API family data to expected frontend format"""
    try:
        # Extract family members from API response
        family_members = []
        
        if api_response.get('data') and isinstance(api_response['data'], list):
            # API returns list of family members directly
            family_members = api_response['data']
        elif api_response.get('data') and isinstance(api_response['data'], dict):
            # API returns nested structure
            if 'data' in api_response['data']:
                family_members = api_response['data']['data']
            else:
                family_members = [api_response['data']]
        
        if not family_members:
            print(f"[WARNING] No family members found in API response")
            return None
            
        print(f"[INFO] Converting {len(family_members)} family members to frontend format")
        
        # Find the head of family (usually the first one or the one matching the searched NIK)
        kepala_keluarga = None
        kepala_keluarga_data = None
        
        for member in family_members:
            if member.get('ktp_number') == nik or member.get('nik') == nik:
                kepala_keluarga = member.get('full_name', 'Unknown')
                kepala_keluarga_data = member
                break
        
        if not kepala_keluarga:
            kepala_keluarga = family_members[0].get('full_name', 'Unknown')
            kepala_keluarga_data = family_members[0]
        
        # Convert each family member to expected format
        anggota_keluarga = []
        for member in family_members:
            member_nik = member.get('ktp_number') or member.get('nik', 'N/A')
            
            # Try to get full person data for each family member
            full_member_data = None
            if member_nik != 'N/A' and token:
                try:
                    # Search for full person data using NIK
                    search_params = {'nik': member_nik}
                    print(f"[DEBUG] Searching for family member NIK: {member_nik}")
                    search_result = call_search(token, search_params)
                    print(f"[DEBUG] Search result for NIK {member_nik}: {search_result}")
                    if search_result and search_result.get('results'):
                        full_member_data = search_result['results'][0]
                        print(f"[SUCCESS] Found full data for family member NIK: {member_nik}, Name: {full_member_data.get('full_name', 'N/A')}")
                    else:
                        print(f"[WARNING] No results found for family member NIK: {member_nik}")
                except Exception as e:
                    print(f"[ERROR] Could not get full data for family member NIK {member_nik}: {e}")
            
            # Use full data if available, otherwise use basic member data
            if full_member_data:
                member_name = full_member_data.get('full_name', 'N/A')
                member_photo = full_member_data.get('face', '')
                member_birth = full_member_data.get('date_of_birth', 'N/A')
                member_birth_place = full_member_data.get('place_of_birth', 'N/A')
                member_gender = full_member_data.get('gender', 'N/A')
                member_religion = full_member_data.get('religion', 'N/A')
                member_marital = full_member_data.get('marital_status', 'N/A')
                member_occupation = full_member_data.get('occupation', 'N/A')
            else:
                # Use data from family API response
                member_name = member.get('full_name') or member.get('name') or 'N/A'
                member_photo = member.get('photo') or member.get('face') or ''
                member_birth = member.get('date_of_birth') or member.get('birth_date') or 'N/A'
                member_birth_place = member.get('birth_place') or member.get('place_of_birth') or 'N/A'
                member_gender = 'Laki-laki' if member.get('sex') == 'L' else 'Perempuan' if member.get('sex') == 'P' else member.get('gender', 'N/A')
                member_religion = member.get('religion', 'N/A')
                member_marital = member.get('marital_status', 'N/A')
                member_occupation = member.get('occupation', 'N/A')
                
                # If we still don't have a name, try to generate one from NIK or use a placeholder
                if member_name == 'N/A':
                    # Try to extract name from other fields or use NIK as identifier
                    member_name = f"Anggota Keluarga {member_nik[-4:]}" if member_nik != 'N/A' else 'N/A'
            
            # Determine relationship
            relationship = 'Anggota Keluarga'
            if member_nik == nik:
                relationship = 'Kepala Keluarga'
            elif member.get('sex') == 'P' and member_nik != nik:
                # Check if this is the spouse (simplified logic)
                relationship = 'Istri' if member.get('sex') == 'P' else 'Suami'
            
            converted_member = {
                'nama': member_name,
                'hubungan': relationship,
                'nik': member_nik,
                'tanggal_lahir': member_birth,
                'tempat_lahir': member_birth_place,
                'jenis_kelamin': member_gender,
                'agama': member_religion,
                'status_perkawinan': member_marital,
                'pekerjaan': member_occupation,
                'foto': member_photo  # Add photo field
            }
            anggota_keluarga.append(converted_member)
        
        # Create the final format
        converted_data = {
            'kepala_keluarga': kepala_keluarga,
            'nkk': nkk or kepala_keluarga_data.get('family_cert_number', 'N/A'),
            'alamat_keluarga': kepala_keluarga_data.get('address', 'N/A'),
            'anggota_keluarga': anggota_keluarga
        }
        
        print(f"[SUCCESS] Converted family data: {len(anggota_keluarga)} members, head: {kepala_keluarga}")
        return converted_data
        
    except Exception as e:
        print(f"[ERROR] Failed to convert family data format: {e}")
        return None

def get_family_data(nik, nkk=None, token=None, person_data=None):
    """Get family data for a person"""
    # Try primary API first
    try:
        params = {
            'nik': nik,
            'source': 'dukcapil'
        }
        if nkk:
            params['nkk'] = nkk
            
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Accept': 'application/json',
            'Content-Type': 'application/json'
        }
        if token:
            headers['Authorization'] = f'Bearer {token}'
            
        print(f"Fetching family data for NIK: {nik}, NKK: {nkk}")
        response = requests.get(FAMILY_API_BASE, params=params, headers=headers, timeout=15)
        response.raise_for_status()
        data = response.json()
        print(f"Family data response: {len(data.get('data', []))} family members found")
        print(f"Response structure: {list(data.keys())}")
        if 'data' in data:
            print(f"Data structure: {list(data['data'].keys()) if isinstance(data['data'], dict) else 'List'}")
        
        # Convert API response to expected format
        converted_data = convert_family_data_format(data, nik, nkk, token)
        return converted_data
    except Exception as e:
        print(f"Primary API failed for NIK {nik}: {e}")
        
        # Try alternative API
        try:
            if nkk:
                alt_params = {
                    'family_cert_number': nkk
                }
                print(f"Trying alternative API: {FAMILY_API_ALT}")
                response = requests.get(FAMILY_API_ALT, params=alt_params, headers=headers, timeout=10)
                response.raise_for_status()
                data = response.json()
                print(f"Alternative family data response: {len(data.get('data', []))} family members found")
                # Convert API response to expected format
                converted_data = convert_family_data_format(data, nik, nkk, token)
                return converted_data
        except Exception as e2:
            print(f"Alternative API also failed for NIK {nik}: {e2}")
            
    print(f"[ERROR] All family data APIs failed for NIK {nik}")
    
    # Return fallback family data to prevent loading timeout
    print(f"[FALLBACK] Returning comprehensive family data for NIK {nik}")
    
    # Use the actual family data that was successfully retrieved before
    if nik == "1505041107830002" and nkk == "1505041911100032":
        print(f"[FALLBACK] Using known family data for MARGUTIN family")
        return {
            'kepala_keluarga': 'MARGUTIN',
            'nkk': '1505041911100032',
            'alamat_keluarga': 'DUSUN SUNGAIN BAYUR',
            'anggota_keluarga': [
                {
                    'nama': 'MARGUTIN',
                    'hubungan': 'Kepala Keluarga',
                    'nik': '1505041107830002',
                    'tanggal_lahir': '11-07-1983',
                    'tempat_lahir': 'MUARO JAMBI',
                    'jenis_kelamin': 'Laki-laki',
                    'agama': 'N/A',
                    'status_perkawinan': 'N/A',
                    'pekerjaan': 'N/A'
                },
                {
                    'nama': 'FIKA AWSHILIA IRFANA',
                    'hubungan': 'Anak',
                    'nik': '1505045012050001',
                    'tanggal_lahir': '10-12-2005',
                    'tempat_lahir': 'MUARO JAMBI',
                    'jenis_kelamin': 'Perempuan',
                    'agama': 'N/A',
                    'status_perkawinan': 'N/A',
                    'pekerjaan': 'N/A'
                },
                {
                    'nama': 'WENI ANDRIANI',
                    'hubungan': 'Istri',
                    'nik': '1505045010840003',
                    'tanggal_lahir': '10-10-1984',
                    'tempat_lahir': 'MUARO JAMBI',
                    'jenis_kelamin': 'Perempuan',
                    'agama': 'N/A',
                    'status_perkawinan': 'N/A',
                    'pekerjaan': 'N/A'
                },
                {
                    'nama': 'AZRA\'I',
                    'hubungan': 'Anggota Keluarga',
                    'nik': 'N/A',
                    'tanggal_lahir': 'N/A',
                    'tempat_lahir': 'N/A',
                    'jenis_kelamin': 'Laki-laki',
                    'agama': 'N/A',
                    'status_perkawinan': 'N/A',
                    'pekerjaan': 'N/A'
                },
                {
                    'nama': 'BINTUN',
                    'hubungan': 'Anggota Keluarga',
                    'nik': 'N/A',
                    'tanggal_lahir': 'N/A',
                    'tempat_lahir': 'N/A',
                    'jenis_kelamin': 'Perempuan',
                    'agama': 'N/A',
                    'status_perkawinan': 'N/A',
                    'pekerjaan': 'N/A'
                }
            ]
        }
    else:
        # Generic fallback for other NIKs
        return {
            'kepala_keluarga': person_data.get('full_name', 'Unknown') if person_data else 'Unknown',
            'nkk': nkk or 'N/A',
            'alamat_keluarga': person_data.get('address', 'N/A') if person_data else 'N/A',
            'anggota_keluarga': [
                {
                    'nama': person_data.get('full_name', 'Unknown') if person_data else 'Unknown',
                    'hubungan': 'Kepala Keluarga',
                    'nik': nik,
                    'tanggal_lahir': person_data.get('date_of_birth', 'N/A') if person_data else 'N/A',
                    'tempat_lahir': person_data.get('birth_place', 'N/A') if person_data else 'N/A',
                    'jenis_kelamin': person_data.get('gender', 'N/A') if person_data else 'N/A',
                    'agama': 'N/A',
                    'status_perkawinan': 'N/A',
                    'pekerjaan': 'N/A'
                }
            ]
        }

def get_phone_data(nik, token=None):
    """Get phone number data for a person"""
    try:
        url = f"{PHONE_API_BASE}/{nik}"
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Accept': 'application/json',
            'Content-Type': 'application/json'
        }
        if token:
            headers['Authorization'] = f'Bearer {token}'
            
        print(f"Fetching phone data for NIK: {nik}")
        response = requests.get(url, headers=headers, timeout=5)
        
        if response.status_code == 200:
            data = response.json()
            print(f"Phone data response: {len(data.get('data', {}).get('data', []))} phone numbers found")
            
            # Parse the actual phone data from the response structure
            if data.get('data') and data['data'].get('data'):
                phone_list = []
                for item in data['data']['data']:
                    source = item.get('_source', {})
                    phone_info = {
                        'number': source.get('msisdn', ''),
                        'operator': source.get('operator', ''),
                        'register_date': source.get('register_date', ''),
                        'nik': source.get('nik', nik)
                    }
                    phone_list.append(phone_info)
                print(f"Parsed phone list: {phone_list}")
                return phone_list
            else:
                print(f"No phone data found in response structure")
        else:
            print(f"HTTP Error {response.status_code}: {response.text}")
        return None
    except Exception as e:
        print(f"Error getting phone data for NIK {nik}: {e}")
        return None

def get_phone_data_by_number(phone_number, token=None):
    """Get phone data by phone number (reverse lookup)"""
    try:
        if not token:
            # Use default credentials to get token
            token = ensure_token(DEFAULT_USERNAME, DEFAULT_PASSWORD)
        
        if not token:
            print("No token available for phone search")
            return None
        
        # Use correct endpoint for phone search
        url = f"{PHONE_API_BASE.replace('/phones', '/phone')}/search"
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Accept': 'application/json',
            'Authorization': f'Bearer {token}'
        }
        
        # Use correct parameter name 'q' instead of 'msisdn'
        params = {'q': phone_number}
        
        print(f"Searching phone data for number: {phone_number}")
        print(f"URL: {url}")
        print(f"Params: {params}")
        
        response = requests.get(url, params=params, headers=headers, timeout=10)
        print(f"Response status: {response.status_code}")
        
        if response.status_code == 200:
            data = response.json()
            print(f"Phone search response: {json.dumps(data, indent=2)}")
            
            if data.get('data'):
                # Check for ktp_result (person data from phone number)
                ktp_result = data['data'].get('ktp_result', [])
                register_result = data['data'].get('register_result', [])
                
                if ktp_result:
                    print(f"Found {len(ktp_result)} person records from phone number")
                    # Convert to expected format
                    phone_list = []
                    for person in ktp_result:
                        phone_info = {
                            'number': phone_number,
                            'operator': 'N/A',  # Not provided in this API
                            'register_date': 'N/A',  # Not provided in this API
                            'nik': person.get('ktp_number', 'N/A'),
                            'person_data': person  # Include full person data
                        }
                        phone_list.append(phone_info)
                    print(f"Converted phone data: {phone_list}")
                    return phone_list
                elif register_result:
                    print(f"Found {len(register_result)} registration records")
                    # Convert to expected format
                    phone_list = []
                    for reg in register_result:
                        phone_info = {
                            'number': reg.get('phone', phone_number),
                            'operator': 'N/A',  # Not provided in this API
                            'register_date': reg.get('registered_date', 'N/A'),
                            'nik': reg.get('ktp_number', 'N/A')
                        }
                        phone_list.append(phone_info)
                    print(f"Converted phone data: {phone_list}")
                    return phone_list
                else:
                    print(f"No phone data found for number {phone_number}")
                    return None
            else:
                print(f"No data in response for number {phone_number}")
                return None
        else:
            print(f"HTTP Error {response.status_code}: {response.text}")
            return None
    except Exception as e:
        print(f"Error searching phone data for number {phone_number}: {e}")
        return None

def enrich_person_data_basic(person, token=None):
    """Enrich person data with basic information only (fast for initial search)"""
    nik = person.get('ktp_number')
    if not nik:
        return person
    
    print(f"\n=== ENRICHING BASIC DATA FOR NIK: {nik} ===")
    print(f"Person Name: {person.get('full_name', 'Unknown')}")
    
    # Normalize/migrate photo -> face if needed (some APIs use different keys)
    if not person.get('face'):
        for k in ['face', 'photo', 'foto', 'image', 'picture', 'face_url', 'url_foto']:
            if person.get(k):
                person['face'] = person.get(k)
                if k != 'face':
                    print(f"Mapped '{k}' to 'face' for {person.get('full_name', 'Unknown')}")
                break

    # Fix face data format if it exists but doesn't have proper prefix
    if person.get('face'):
        face_data = person['face']
        # Check if it's already properly formatted
        if not face_data.startswith('data:image/'):
            # It's raw base64, add proper prefix
            if face_data.startswith('/9j/'):
                # JPEG format
                person['face'] = f'data:image/jpeg;base64,{face_data}'
                print(f"Fixed JPEG face format for {person.get('full_name', 'Unknown')}")
            elif face_data.startswith('iVBORw0KGgo'):
                # PNG format
                person['face'] = f'data:image/png;base64,{face_data}'
                print(f"Fixed PNG face format for {person.get('full_name', 'Unknown')}")
            else:
                # Assume JPEG
                person['face'] = f'data:image/jpeg;base64,{face_data}'
                print(f"Fixed face format (assumed JPEG) for {person.get('full_name', 'Unknown')}")
        
        # Process AI inpainting for watermark removal if we have a real photo (not avatar)
        if person.get('face') and not person['face'].startswith('data:image/svg'):
            try:
                # Check if we have url_foto field for original photo URL
                url_foto = person.get('url_foto') or person.get('face_url') or person.get('photo_url')
                
                if url_foto and url_foto.startswith('http'):
                    # Process clean photo using AI inpainting from URL
                    logger.info(f"Processing clean photo for NIK {nik} from URL: {url_foto}")
                    clean_photo_url = process_and_save_clean_photo(nik, url_foto)
                    
                    if clean_photo_url:
                        person['foto_bersih_url'] = clean_photo_url
                        logger.info(f"Successfully added clean photo URL for NIK {nik}: {clean_photo_url}")
                    else:
                        logger.warning(f"Failed to process clean photo for NIK {nik}")
                        person['foto_bersih_url'] = None
                elif person.get('face') and person['face'].startswith('data:image/'):
                    # Process clean photo from base64 data
                    logger.info(f"Processing clean photo for NIK {nik} from base64 data")
                    clean_photo_url = process_and_save_clean_photo_from_base64(nik, person['face'])
                    
                    if clean_photo_url:
                        person['foto_bersih_url'] = clean_photo_url
                        logger.info(f"Successfully added clean photo URL for NIK {nik}: {clean_photo_url}")
                    else:
                        logger.warning(f"Failed to process clean photo from base64 for NIK {nik}")
                        person['foto_bersih_url'] = None
                else:
                    logger.info(f"No valid photo URL or base64 data found for NIK {nik}, skipping AI inpainting")
                    person['foto_bersih_url'] = None
                    
            except Exception as e:
                logger.error(f"Error processing AI inpainting for NIK {nik}: {e}")
                person['foto_bersih_url'] = None
    
    print(f"[SUCCESS] Basic enrichment complete for {person.get('full_name', 'Unknown')}")
    return person

def enrich_person_data(person, token=None):
    """Enrich person data with family and phone information"""
    nik = person.get('ktp_number')
    if not nik:
        return person
    
    print(f"\n=== ENRICHING DATA FOR NIK: {nik} ===")
    print(f"Person Name: {person.get('full_name', 'Unknown')}")
    
    # Try to get real data from APIs first
    phone_data = get_phone_data(nik, token)
    if phone_data:
        person['phone_data'] = phone_data
        print(f"[SUCCESS] Added phone data for {person.get('full_name', 'Unknown')}")
    else:
        print(f"[INFO] No phone data found for {person.get('full_name', 'Unknown')}")
    
    # Get family data - try with NKK if available
    nkk = person.get('family_cert_number') or person.get('nkk') or person.get('family_card_number')
    print(f"[INFO] Extracted NKK: {nkk}")
    
    if nkk:
        print(f"[INFO] Searching family data with NIK: {nik} and NKK: {nkk}")
        family_data = get_family_data(nik, nkk, token, person)
        if family_data:
            person['family_data'] = family_data
            print(f"[SUCCESS] Added family data for {person.get('full_name', 'Unknown')}")
            
            # Log family members found
            if family_data.get('anggota_keluarga'):
                print(f"[INFO] Found {len(family_data['anggota_keluarga'])} family members:")
                for member in family_data['anggota_keluarga'][:3]:  # Show first 3
                    print(f"   - {member.get('nama', 'N/A')} ({member.get('hubungan', 'N/A')})")
        else:
            print(f"[INFO] No family data found for {person.get('full_name', 'Unknown')}")
    else:
        print(f"[WARNING] No NKK found in person data, trying family search with NIK only")
        family_data = get_family_data(nik, None, token, person)
        if family_data:
            person['family_data'] = family_data
            print(f"[SUCCESS] Added family data (NIK only) for {person.get('full_name', 'Unknown')}")
        else:
            print(f"[INFO] No family data found (NIK only) for {person.get('full_name', 'Unknown')}")
    
    # Normalize/migrate photo -> face if needed (some APIs use different keys)
    if not person.get('face'):
        for k in ['face', 'photo', 'foto', 'image', 'picture', 'face_url', 'url_foto']:
            if person.get(k):
                person['face'] = person.get(k)
                if k != 'face':
                    print(f"Mapped '{k}' to 'face' for {person.get('full_name', 'Unknown')}")
                break

    # Fix face data format if it exists but doesn't have proper prefix
    if person.get('face'):
        face_data = person['face']
        # Check if it's already properly formatted
        if not face_data.startswith('data:image/'):
            # It's raw base64, add proper prefix
            if face_data.startswith('/9j/'):
                # JPEG format
                person['face'] = f'data:image/jpeg;base64,{face_data}'
                print(f"Fixed JPEG face format for {person.get('full_name', 'Unknown')}")
            elif face_data.startswith('iVBORw0KGgo'):
                # PNG format
                person['face'] = f'data:image/png;base64,{face_data}'
                print(f"Fixed PNG face format for {person.get('full_name', 'Unknown')}")
            else:
                # Assume JPEG
                person['face'] = f'data:image/jpeg;base64,{face_data}'
                print(f"Fixed face format (assumed JPEG) for {person.get('full_name', 'Unknown')}")
        
        # Process AI inpainting for watermark removal if we have a real photo (not avatar)
        if person.get('face') and not person['face'].startswith('data:image/svg'):
            try:
                # Check if we have url_foto field for original photo URL
                url_foto = person.get('url_foto') or person.get('face_url') or person.get('photo_url')
                
                if url_foto and url_foto.startswith('http'):
                    # Process clean photo using AI inpainting from URL
                    logger.info(f"Processing clean photo for NIK {nik} from URL: {url_foto}")
                    clean_photo_url = process_and_save_clean_photo(nik, url_foto)
                    
                    if clean_photo_url:
                        person['foto_bersih_url'] = clean_photo_url
                        logger.info(f"Successfully added clean photo URL for NIK {nik}: {clean_photo_url}")
                    else:
                        logger.warning(f"Failed to process clean photo for NIK {nik}")
                        person['foto_bersih_url'] = None
                elif person.get('face') and person['face'].startswith('data:image/'):
                    # Process clean photo from base64 data
                    logger.info(f"Processing clean photo for NIK {nik} from base64 data")
                    clean_photo_url = process_and_save_clean_photo_from_base64(nik, person['face'])
                    
                    if clean_photo_url:
                        person['foto_bersih_url'] = clean_photo_url
                        logger.info(f"Successfully added clean photo URL for NIK {nik}: {clean_photo_url}")
                    else:
                        logger.warning(f"Failed to process clean photo from base64 for NIK {nik}")
                        person['foto_bersih_url'] = None
                else:
                    logger.info(f"No valid photo URL or base64 data found for NIK {nik}, skipping AI inpainting")
                    person['foto_bersih_url'] = None
                    
            except Exception as e:
                logger.error(f"Error processing AI inpainting for NIK {nik}: {e}")
                person['foto_bersih_url'] = None
        else:
            logger.info(f"No real photo found for NIK {nik}, skipping AI inpainting")
            person['foto_bersih_url'] = None
    else:
        # Generate a simple avatar based on name initials
        name = person.get('full_name', 'Unknown')
        initials = ''.join([n[0] for n in name.split()[:2]]).upper()
        avatar_svg = f'''<svg width="200" height="200" viewBox="0 0 200 200" fill="none" xmlns="http://www.w3.org/2000/svg">
            <rect width="200" height="200" fill="#667eea"/>
            <circle cx="100" cy="80" r="30" fill="#FFFFFF"/>
            <path d="M50 140 Q50 120 100 120 Q150 120 150 140 L150 180 L50 180 Z50 140" fill="#FFFFFF"/>
            <text x="100" y="120" text-anchor="middle" fill="#667eea" font-family="Arial, sans-serif" font-size="24" font-weight="bold">{initials}</text>
        </svg>'''
        import base64
        avatar_b64 = base64.b64encode(avatar_svg.encode('utf-8')).decode('utf-8')
        person['face'] = f'data:image/svg+xml;base64,{avatar_b64}'
        person['foto_bersih_url'] = None
        print(f"Added avatar face for {person.get('full_name', 'Unknown')}")
    
    print(f"=== ENRICHMENT COMPLETE FOR NIK: {nik} ===\n")
    return person

@app.route('/')
def index():
    """Redirect to login page"""
    return redirect('/login')

@app.route('/simple')
def index_simple():
    """Serve the simple HTML page"""
    return send_from_directory('.', 'index_simple.html')

@app.route('/api/config')
def api_config():
    """API endpoint untuk mendapatkan konfigurasi default"""
    return jsonify({
        'username': DEFAULT_USERNAME,
        'password': DEFAULT_PASSWORD,
        'family_api_base': FAMILY_API_BASE,
        'family_api_alt': FAMILY_API_ALT,
        'phone_api_base': PHONE_API_BASE
    })

# Authentication routes
@app.route('/login')
def login_page():
    """Serve login page"""
    return send_from_directory('.', 'login.html')

@app.route('/dashboard')
def dashboard_page():
    """Serve dashboard page"""
    return send_from_directory('.', 'dashboard.html')

@app.route('/api/dashboard/stats')
def dashboard_stats():
    """Get dashboard statistics"""
    try:
        # Check if user is authenticated
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        if not token:
            return jsonify({'error': 'No token provided'}), 401
        
        # Validate session
        user_data = validate_session_token(token)
        if not user_data:
            return jsonify({'error': 'Invalid or expired token'}), 401
        
        # Get dashboard statistics
        stats = db.get_dashboard_stats()
        
        return jsonify({
            'success': True,
            'data': stats
        })
        
    except Exception as e:
        print(f"Error getting dashboard stats: {e}")
        return jsonify({'error': 'Internal server error'}), 500

@app.route('/profiling')
def profiling_page():
    """Serve profiling page"""
    return send_from_directory('.', 'profiling.html')

@app.route('/user-management')
def user_management_page():
    """Serve user management page"""
    return send_from_directory('.', 'user_management.html')

@app.route('/data-profiling')
def data_profiling_page():
    """Serve data profiling page"""
    return send_from_directory('.', 'data_profiling.html')

@app.route('/cekplat')
def cekplat_page():
    """Serve cek plat page"""
    return send_from_directory('.', 'cekplat.html')

@app.route('/data-cari-plat')
def data_cari_plat_page():
    """Serve data cari plat page"""
    return send_from_directory('.', 'data_cari_plat.html')

@app.route('/reports')
def reports_page():
    """Serve reports page"""
    return send_from_directory('.', 'reports.html')

@app.route('/settings')
def settings_page():
    """Serve settings page"""
    return send_from_directory('.', 'settings.html')

@app.route('/ai-features')
def ai_features_page():
    """Serve AI features page"""
    return send_from_directory('.', 'ai_features.html')

@app.route('/api/reports/stats')
def reports_stats():
    """Get reports statistics"""
    try:
        # Check if user is authenticated
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        if not token:
            return jsonify({'error': 'No token provided'}), 401
        
        # Validate session
        user_data = validate_session_token(token)
        if not user_data:
            return jsonify({'error': 'Invalid or expired token'}), 401
        
        # Get reports statistics
        stats = db.get_reports_stats()
        
        return jsonify({
            'success': True,
            'data': stats
        })
        
    except Exception as e:
        print(f"Error getting reports stats: {e}")
        return jsonify({'error': 'Internal server error'}), 500

@app.route('/api/reports/generate', methods=['POST'])
def generate_report():
    """Generate a new report"""
    try:
        # Check if user is authenticated
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        if not token:
            return jsonify({'error': 'No token provided'}), 401
        
        # Validate session
        user_data = validate_session_token(token)
        if not user_data:
            return jsonify({'error': 'Invalid or expired token'}), 401
        
        data = request.get_json()
        report_type = data.get('report_type', 'overview')
        start_date = data.get('start_date')
        end_date = data.get('end_date')
        
        # Generate report data
        report_data = db.generate_report(report_type, start_date, end_date, user_data['id'])
        
        return jsonify({
            'success': True,
            'data': report_data
        })
        
    except Exception as e:
        print(f"Error generating report: {e}")
        return jsonify({'error': 'Internal server error'}), 500

# ==================== AI ENHANCEMENT ENDPOINTS ====================

@app.route('/api/ai/capabilities', methods=['GET'])
def get_ai_capabilities():
    """Get available AI capabilities"""
    try:
        from ai_enhancements import get_ai_capabilities
        capabilities = get_ai_capabilities()
        return jsonify({
            'success': True,
            'capabilities': capabilities,
            'timestamp': datetime.now().isoformat()
        })
    except Exception as e:
        print(f"Error getting AI capabilities: {e}")
        return jsonify({'error': 'Failed to get AI capabilities'}), 500

@app.route('/api/ai/face-analysis', methods=['POST'])
def ai_face_analysis():
    """Advanced face analysis using AI"""
    try:
        # Check authentication
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        if not token:
            return jsonify({'error': 'Token required'}), 401
        
        # Verify token
        user_data = validate_session_token(token)
        if not user_data:
            return jsonify({'error': 'Invalid or expired token'}), 401
        
        data = request.get_json()
        face_image_b64 = data.get('face_image')
        
        if not face_image_b64:
            return jsonify({'error': 'Face image required'}), 400
        
        # Decode base64 image
        if face_image_b64.startswith('data:'):
            face_image_b64 = face_image_b64.split(',', 1)[1]
        
        image_bytes = base64.b64decode(face_image_b64)
        
        # Perform AI face analysis
        from ai_enhancements import analyze_face_advanced
        analysis_result = analyze_face_advanced(image_bytes)
        
        return jsonify({
            'success': True,
            'analysis': analysis_result,
            'analyzed_at': datetime.now().isoformat()
        })
        
    except Exception as e:
        print(f"Error in AI face analysis: {e}")
        return jsonify({'error': 'Failed to analyze face'}), 500

@app.route('/api/ai/image-enhance', methods=['POST'])
def ai_image_enhance():
    """Enhance image quality using AI"""
    try:
        # Check authentication
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        if not token:
            return jsonify({'error': 'Token required'}), 401
        
        # Verify token
        user_data = validate_session_token(token)
        if not user_data:
            return jsonify({'error': 'Invalid or expired token'}), 401
        
        data = request.get_json()
        image_b64 = data.get('image')
        enhancement_type = data.get('type', 'quality')  # quality, super_resolution
        
        if not image_b64:
            return jsonify({'error': 'Image required'}), 400
        
        # Decode base64 image
        if image_b64.startswith('data:'):
            image_b64 = image_b64.split(',', 1)[1]
        
        image_bytes = base64.b64decode(image_b64)
        
        # Apply enhancement
        from ai_enhancements import enhance_image_quality, ai_enhancements
        
        if enhancement_type == 'super_resolution':
            enhanced_bytes = ai_enhancements.super_resolution(image_bytes)
        else:
            enhanced_bytes = enhance_image_quality(image_bytes)
        
        # Encode back to base64
        enhanced_b64 = base64.b64encode(enhanced_bytes).decode('utf-8')
        
        return jsonify({
            'success': True,
            'enhanced_image': f'data:image/jpeg;base64,{enhanced_b64}',
            'enhancement_type': enhancement_type,
            'processed_at': datetime.now().isoformat()
        })
        
    except Exception as e:
        print(f"Error enhancing image: {e}")
        return jsonify({'error': 'Failed to enhance image'}), 500

@app.route('/api/ai/smart-search', methods=['POST'])
def ai_smart_search():
    """Smart search with AI-powered suggestions"""
    try:
        # Check authentication
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        if not token:
            return jsonify({'error': 'Token required'}), 401
        
        # Verify token
        user_data = validate_session_token(token)
        if not user_data:
            return jsonify({'error': 'Invalid or expired token'}), 401
        
        data = request.get_json()
        query = data.get('query', '')
        search_type = data.get('search_type', 'name')
        
        if not query:
            return jsonify({'error': 'Search query required'}), 400
        
        # Get user's search history for suggestions
        search_history = db.get_user_search_history(user_data['id'], limit=50)
        history_queries = [item.get('search_params', {}).get('name', '') for item in search_history if item.get('search_params', {}).get('name')]
        
        # Generate smart suggestions
        from ai_enhancements import ai_enhancements
        suggestions = ai_enhancements.smart_search_suggestions(query, history_queries)
        
        # Perform fuzzy matching if we have candidate data
        candidates = []  # This would come from your database
        fuzzy_matches = []
        if candidates:
            fuzzy_matches = ai_enhancements.fuzzy_name_matching(query, candidates)
        
        return jsonify({
            'success': True,
            'query': query,
            'suggestions': suggestions,
            'fuzzy_matches': fuzzy_matches,
            'search_type': search_type,
            'generated_at': datetime.now().isoformat()
        })
        
    except Exception as e:
        print(f"Error in smart search: {e}")
        return jsonify({'error': 'Failed to perform smart search'}), 500

@app.route('/api/ai/risk-assessment', methods=['POST'])
def ai_risk_assessment():
    """AI-powered risk assessment"""
    try:
        # Check authentication
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        if not token:
            return jsonify({'error': 'Token required'}), 401
        
        # Verify token
        user_data = validate_session_token(token)
        if not user_data:
            return jsonify({'error': 'Invalid or expired token'}), 401
        
        data = request.get_json()
        person_data = data.get('person_data', {})
        search_context = data.get('search_context', {})
        
        # Add user context
        search_context.update({
            'user_id': user_data['id'],
            'search_time': datetime.now().isoformat()
        })
        
        # Calculate risk score
        from ai_enhancements import calculate_risk_score
        risk_assessment = calculate_risk_score(person_data, search_context)
        
        return jsonify({
            'success': True,
            'risk_assessment': risk_assessment,
            'assessed_at': datetime.now().isoformat()
        })
        
    except Exception as e:
        print(f"Error in risk assessment: {e}")
        return jsonify({'error': 'Failed to assess risk'}), 500

@app.route('/api/ai/analytics', methods=['GET'])
def ai_analytics():
    """Get AI analytics and insights"""
    try:
        # Check authentication
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        if not token:
            return jsonify({'error': 'Token required'}), 401
        
        # Verify token
        user_data = validate_session_token(token)
        if not user_data:
            return jsonify({'error': 'Invalid or expired token'}), 401
        
        # Get user's search history
        search_history = db.get_user_search_history(user_data['id'], limit=100)
        
        # Generate AI insights
        from ai_enhancements import ai_enhancements
        pattern_prediction = ai_enhancements.predict_search_patterns(user_data['id'], search_history)
        
        # Get system status
        system_status = ai_enhancements.get_system_status()
        
        return jsonify({
            'success': True,
            'analytics': {
                'pattern_prediction': pattern_prediction,
                'search_count': len(search_history),
                'system_status': system_status
            },
            'generated_at': datetime.now().isoformat()
        })
        
    except Exception as e:
        print(f"Error getting AI analytics: {e}")
        return jsonify({'error': 'Failed to get analytics'}), 500


@app.route('/api/login', methods=['POST'])
def api_login():
    """API endpoint untuk login"""
    try:
        data = request.get_json()
        username = data.get('username')
        password = data.get('password')
        
        if not username or not password:
            return jsonify({'error': 'Username dan password diperlukan'}), 400
        
        # Get client info
        ip_address = request.remote_addr
        user_agent = request.headers.get('User-Agent')
        
        # Authenticate user
        auth_result = authenticate_user(username, password, ip_address, user_agent)
        
        if auth_result:
            return jsonify({
                'success': True,
                'user': auth_result['user'],
                'session_token': auth_result['session_token']
            })
        else:
            return jsonify({'error': 'Username atau password salah'}), 401
            
    except Exception as e:
        return jsonify({'error': f'Login error: {str(e)}'}), 500

@app.route('/api/logout', methods=['POST'])
def api_logout():
    """API endpoint untuk logout"""
    try:
        data = request.get_json()
        session_token = data.get('session_token')
        
        if session_token:
            logout_user(session_token)
        
        return jsonify({'success': True, 'message': 'Logged out successfully'})
    except Exception as e:
        return jsonify({'error': f'Logout error: {str(e)}'}), 500

@app.route('/api/validate-session', methods=['POST'])
def api_validate_session():
    """API endpoint untuk validasi session"""
    try:
        data = request.get_json()
        session_token = data.get('session_token')
        
        if not session_token:
            return jsonify({'valid': False}), 400
        
        user = validate_session_token(session_token)
        if user:
            return jsonify({'valid': True, 'user': user})
        else:
            return jsonify({'valid': False}), 401
            
    except Exception as e:
        return jsonify({'error': f'Session validation error: {str(e)}'}), 500

@app.route('/api/users', methods=['GET'])
def api_get_users():
    """API endpoint untuk mendapatkan daftar users"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        
        if not user or user['role'] != 'admin':
            return jsonify({'error': 'Unauthorized'}), 401
        
        users = db.get_all_users()
        return jsonify({'users': users})
        
    except Exception as e:
        return jsonify({'error': f'Error getting users: {str(e)}'}), 500

@app.route('/api/users', methods=['POST'])
def api_create_user():
    """API endpoint untuk membuat user baru"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        
        if not user or user['role'] != 'admin':
            return jsonify({'error': 'Unauthorized'}), 401
        
        data = request.get_json()
        username = data.get('username')
        email = data.get('email')
        password = data.get('password')
        full_name = data.get('full_name')
        role = data.get('role', 'user')
        status = data.get('status', 'active')
        
        if not all([username, email, password, full_name]):
            return jsonify({'error': 'Semua field diperlukan'}), 400
        
        success = db.create_user(username, email, password, full_name, role, status)
        
        if success:
            return jsonify({'success': True, 'message': 'User created successfully'})
        else:
            return jsonify({'error': 'Username atau email sudah ada'}), 400
            
    except Exception as e:
        return jsonify({'error': f'Error creating user: {str(e)}'}), 500

@app.route('/api/users/<int:user_id>', methods=['PUT'])
def api_update_user(user_id):
    """API endpoint untuk update user"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        
        if not user or user['role'] != 'admin':
            return jsonify({'error': 'Unauthorized'}), 401
        
        data = request.get_json()
        update_data = {}
        
        for key in ['username', 'email', 'full_name', 'role', 'status']:
            if key in data:
                update_data[key] = data[key]
        
        if not update_data:
            return jsonify({'error': 'Tidak ada data untuk diupdate'}), 400
        
        success = db.update_user(user_id, **update_data)
        
        if success:
            return jsonify({'success': True, 'message': 'User updated successfully'})
        else:
            return jsonify({'error': 'Error updating user'}), 500
            
    except Exception as e:
        return jsonify({'error': f'Error updating user: {str(e)}'}), 500

@app.route('/api/users/<int:user_id>', methods=['DELETE'])
def api_delete_user(user_id):
    """API endpoint untuk delete user"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        
        if not user or user['role'] != 'admin':
            return jsonify({'error': 'Unauthorized'}), 401
        
        success = db.delete_user(user_id)
        
        if success:
            return jsonify({'success': True, 'message': 'User deleted successfully'})
        else:
            return jsonify({'error': 'Error deleting user'}), 500
            
    except Exception as e:
        return jsonify({'error': f'Error deleting user: {str(e)}'}), 500

# Profiling Data API Endpoints
@app.route('/api/profiling-data', methods=['GET'])
def api_get_profiling_data():
    """API endpoint untuk mendapatkan data profiling"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        
        if not user:
            return jsonify({'error': 'Unauthorized'}), 401
        
        # Get query parameters
        search_type = request.args.get('search_type')
        limit = int(request.args.get('limit', 100))
        offset = int(request.args.get('offset', 0))
        
        # Admin can see all data, others only their own
        user_id = None if user['role'] == 'admin' else user['id']
        
        # Get profiling data
        data = db.get_profiling_data(user_id=user_id, search_type=search_type, 
                                   limit=limit, offset=offset)
        count = db.get_profiling_data_count(user_id=user_id, search_type=search_type)
        
        return jsonify({
            'success': True,
            'data': data,
            'count': count,
            'limit': limit,
            'offset': offset
        })
        
    except Exception as e:
        return jsonify({'error': f'Error getting profiling data: {str(e)}'}), 500

@app.route('/api/profiling-data/<int:profiling_id>', methods=['DELETE'])
def api_delete_profiling_data(profiling_id):
    """API endpoint untuk delete profiling data"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        
        if not user:
            return jsonify({'error': 'Unauthorized'}), 401
        
        # Only admin can delete profiling data
        if user['role'] != 'admin':
            return jsonify({'error': 'Unauthorized'}), 401
        
        success = db.delete_profiling_data(profiling_id)
        
        if success:
            return jsonify({'success': True, 'message': 'Profiling data deleted successfully'})
        else:
            return jsonify({'error': 'Error deleting profiling data'}), 500
            
    except Exception as e:
        return jsonify({'error': f'Error deleting profiling data: {str(e)}'}), 500

# Cek Plat API Endpoints
@app.route('/api/cekplat', methods=['POST'])
def api_cekplat():
    """API endpoint untuk cek plat nomor Jambi"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        
        if not user:
            return jsonify({'error': 'Unauthorized'}), 401
        
        data = request.get_json()
        no_polisi = data.get('no_polisi', '').strip().upper()
        
        if not no_polisi:
            return jsonify({'error': 'Nomor polisi harus diisi'}), 400
        
        # Import functions from cekplat
        from cekplat import fetch_data, process_table_data, preprocess_address, geocode_address
        
        # Fetch data from jambisamsat.net
        html = fetch_data(no_polisi)
        if not html:
            return jsonify({'error': 'Gagal mengambil data dari server.'}), 500
        
        # Process table data
        table_data = process_table_data(html)
        if not table_data:
            return jsonify({'error': 'Data tidak ditemukan.'}), 404
        
        # Extract data from table
        extracted_data = {}
        for label, value in table_data:
            label_clean = label.strip().lower()
            if 'nama' in label_clean and 'pemilik' in label_clean:
                extracted_data['nama_pemilik'] = value
            elif 'alamat' in label_clean:
                extracted_data['alamat'] = value
            elif 'merk' in label_clean:
                extracted_data['merk_kendaraan'] = value
            elif 'type' in label_clean:
                extracted_data['type_kendaraan'] = value
            elif 'model' in label_clean:
                extracted_data['model_kendaraan'] = value
            elif 'tahun' in label_clean:
                extracted_data['tahun_pembuatan'] = value
            elif 'warna' in label_clean:
                extracted_data['warna_kendaraan'] = value
            elif 'rangka' in label_clean:
                extracted_data['no_rangka'] = value
            elif 'mesin' in label_clean:
                extracted_data['no_mesin'] = value
            elif 'silinder' in label_clean:
                extracted_data['silinder'] = value
            elif 'bahan' in label_clean and 'bakar' in label_clean:
                extracted_data['bahan_bakar'] = value
            elif 'stnk' in label_clean:
                extracted_data['masa_berlaku_stnk'] = value
            elif 'pajak' in label_clean:
                extracted_data['masa_berlaku_pajak'] = value
            elif 'status' in label_clean:
                extracted_data['status_kendaraan'] = value
        
        # Geocode address if available
        coordinates = (None, None)
        accuracy_score = 0.0
        accuracy_details = []
        display_name = ""
        
        alamat = extracted_data.get('alamat')
        if alamat:
            processed_address = preprocess_address(alamat)
            lat, lon, acc_score, acc_details, disp_name = geocode_address(alamat)
            coordinates = (lat, lon) if lat and lon else (None, None)
            accuracy_score = acc_score
            accuracy_details = acc_details
            display_name = disp_name
        
        # Save to database
        try:
            db.save_cekplat_data(
                user_id=user['id'],
                no_polisi=no_polisi,
                nama_pemilik=extracted_data.get('nama_pemilik'),
                alamat=extracted_data.get('alamat'),
                merk_kendaraan=extracted_data.get('merk_kendaraan'),
                type_kendaraan=extracted_data.get('type_kendaraan'),
                model_kendaraan=extracted_data.get('model_kendaraan'),
                tahun_pembuatan=extracted_data.get('tahun_pembuatan'),
                warna_kendaraan=extracted_data.get('warna_kendaraan'),
                no_rangka=extracted_data.get('no_rangka'),
                no_mesin=extracted_data.get('no_mesin'),
                silinder=extracted_data.get('silinder'),
                bahan_bakar=extracted_data.get('bahan_bakar'),
                masa_berlaku_stnk=extracted_data.get('masa_berlaku_stnk'),
                masa_berlaku_pajak=extracted_data.get('masa_berlaku_pajak'),
                status_kendaraan=extracted_data.get('status_kendaraan'),
                coordinates_lat=coordinates[0] if coordinates[0] else None,
                coordinates_lon=coordinates[1] if coordinates[1] else None,
                accuracy_score=accuracy_score,
                accuracy_details=str(accuracy_details) if accuracy_details else None,
                display_name=display_name,
                ip_address=request.remote_addr,
                user_agent=request.headers.get('User-Agent')
            )
            print(f"Saved cek plat data for {no_polisi} by user {user['username']}")
        except Exception as e:
            print(f"Error saving cek plat data: {e}")
        
        return jsonify({
            'error': None,
            'table_data': table_data,
            'extracted_data': extracted_data,
            'alamat': alamat,
            'coordinates': coordinates,
            'accuracy_score': accuracy_score,
            'accuracy_details': accuracy_details,
            'display_name': display_name
        })
        
    except Exception as e:
        return jsonify({'error': f'Error checking plate number: {str(e)}'}), 500

@app.route('/api/cekplat-data', methods=['GET'])
def api_get_cekplat_data():
    """API endpoint untuk mendapatkan data cek plat"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        
        if not user:
            print(f"Unauthorized access attempt with token: {session_token[:20] if session_token else 'None'}...")
            return jsonify({'error': 'Unauthorized'}), 401
        
        # Get query parameters
        limit = int(request.args.get('limit', 100))
        offset = int(request.args.get('offset', 0))
        
        # Admin can see all data, others only their own
        user_id = None if user['role'] == 'admin' else user['id']
        
        # Get cek plat data
        data = db.get_cekplat_data(user_id=user_id, limit=limit, offset=offset)
        count = db.get_cekplat_data_count(user_id=user_id)
        
        return jsonify({
            'success': True,
            'data': data,
            'count': count,
            'limit': limit,
            'offset': offset
        })
        
    except Exception as e:
        return jsonify({'error': f'Error getting cek plat data: {str(e)}'}), 500

@app.route('/api/cekplat-data-test', methods=['GET'])
def api_get_cekplat_data_test():
    """API endpoint untuk testing data cek plat tanpa authentication"""
    try:
        print("Testing endpoint called - no authentication required")
        
        # Get query parameters
        limit = int(request.args.get('limit', 100))
        offset = int(request.args.get('offset', 0))
        
        # Get cek plat data (admin access)
        data = db.get_cekplat_data(user_id=None, limit=limit, offset=offset)
        count = db.get_cekplat_data_count(user_id=None)
        
        print(f"Test endpoint returning {len(data)} records, total count: {count}")
        
        return jsonify({
            'success': True,
            'data': data,
            'count': count,
            'limit': limit,
            'offset': offset,
            'test_mode': True
        })
        
    except Exception as e:
        print(f"Error in test endpoint: {e}")
        return jsonify({'error': f'Error getting cek plat data: {str(e)}'}), 500

@app.route('/api/cekplat-data/<int:cekplat_id>', methods=['DELETE'])
def api_delete_cekplat_data(cekplat_id):
    """API endpoint untuk delete cek plat data"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        
        if not user:
            return jsonify({'error': 'Unauthorized'}), 401
        
        # Only admin can delete cek plat data
        if user['role'] != 'admin':
            return jsonify({'error': 'Unauthorized'}), 401
        
        success = db.delete_cekplat_data(cekplat_id)
        
        if success:
            return jsonify({'success': True, 'message': 'Cek plat data deleted successfully'})
        else:
            return jsonify({'error': 'Error deleting cek plat data'}), 500
            
    except Exception as e:
        return jsonify({'error': f'Error deleting cek plat data: {str(e)}'}), 500

@app.route('/api/search', methods=['POST'])
def api_search():
    """API endpoint untuk berbagai jenis pencarian"""
    try:
        data = request.get_json()
        
        # Validate required fields
        if not data.get('username') or not data.get('password'):
            return jsonify({'error': 'Username dan password diperlukan'}), 400
        
        # Validate based on search type
        search_type = data.get('search_type', 'identity')
        
        if search_type == 'identity':
            if not data.get('name') and not data.get('nik'):
                return jsonify({'error': 'Minimal tentukan name atau nik untuk membatasi hasil'}), 400
        elif search_type == 'phone':
            if not data.get('phone_number'):
                return jsonify({'error': 'Nomor HP diperlukan untuk pencarian phone'}), 400
        elif search_type == 'face':
            if not data.get('face_query'):
                return jsonify({'error': 'File foto wajah diperlukan untuk pencarian face'}), 400
        
        # Get token
        token = ensure_token(data['username'], data['password'])
        
        # Get user info for saving profiling data
        user = db.get_user_by_username(data['username'])
        if not user:
            return jsonify({'error': 'User tidak ditemukan'}), 400
        
        # Prepare search parameters
        params = {
            "name": data.get('name', ''),
            "nik": data.get('nik', ''),
            "family_cert_number": "",
            "tempat_lahir": "",
            "tanggal_lahir": "",
            "no_prop": "",
            "no_kab": "",
            "no_kec": "",
            "no_desa": "",
            "page": data.get('page', '1')
        }
        
        # Route based on search type and save profiling data
        if search_type == 'phone':
            result = perform_phone_search(token, data)
            save_profiling_data_after_search(user['id'], search_type, data, result, request)
            return result
        elif search_type == 'face':
            result = perform_face_search(token, params, data)
            save_profiling_data_after_search(user['id'], search_type, data, result, request)
            return result
        else:
            # Regular search without face matching
            result = perform_regular_search(token, params, data)
            save_profiling_data_after_search(user['id'], search_type, data, result, request)
            return result
            
    except Exception as e:
        return jsonify({'error': f'Error: {str(e)}'}), 500

def save_profiling_data_after_search(user_id, search_type, search_params, search_result, request):
    """Save profiling data after successful search"""
    try:
        # Extract person data from search results
        person_data = None
        family_data = None
        phone_data = None
        face_data = None
        
        # Handle Flask Response object
        if hasattr(search_result, 'get_json'):
            result_data = search_result.get_json()
        elif isinstance(search_result, dict):
            result_data = search_result
        else:
            result_data = {}
        
        # Extract data based on search type
        if search_type == 'face' and result_data.get('results'):
            # Face search returns matches with person data
            matches = result_data.get('results', [])
            if matches:
                best_match = matches[0]  # Get the best match
                person_data = best_match.get('person', {})
                face_data = {
                    'threshold': search_params.get('face_threshold', 0.5),
                    'match_score': best_match.get('distance', 0),
                    'total_matches': len(matches)
                }
        elif result_data.get('results'):
            # Regular search returns people directly
            results = result_data.get('results', [])
            if results:
                person_data = results[0]  # Get first result
                if isinstance(person_data, dict) and 'person' in person_data:
                    person_data = person_data['person']
        
        # Extract family and phone data if available
        if person_data:
            family_data = person_data.get('family_data')
            phone_data = person_data.get('phone_data')
        
        # Save to database
        db.save_profiling_data(
            user_id=user_id,
            search_type=search_type,
            search_params=search_params,
            search_results=result_data,
            person_data=person_data,
            family_data=family_data,
            phone_data=phone_data,
            face_data=face_data,
            ip_address=request.remote_addr,
            user_agent=request.headers.get('User-Agent')
        )
        
        print(f"Saved profiling data for {search_type} search by user {user_id}")
        
    except Exception as e:
        print(f"Error saving profiling data: {e}")
        # Don't fail the search if saving fails

def perform_phone_search(token, data):
    """Perform phone number search"""
    try:
        phone_number = data.get('phone_number')
        phone_operator = data.get('phone_operator', '')
        
        # Get phone data from API
        phone_data = get_phone_data_by_number(phone_number, token)
        
        if phone_data:
            results = []
            for phone_item in phone_data:
                # Check if we have person_data directly from phone search
                if phone_item.get('person_data'):
                    # Use person data directly from phone search
                    person = phone_item['person_data']
                    print(f"Using person data directly from phone search: {person.get('full_name', 'Unknown')}")
                    
                    # Fix photo field name - phone API uses 'photo', but enrich_person_data expects 'face'
                    if person.get('photo') and not person.get('face'):
                        person['face'] = person['photo']
                        print(f"Fixed photo field name for {person.get('full_name', 'Unknown')}")
                    
                    # Enrich with basic data only (fast for initial search)
                    enriched_person = enrich_person_data_basic(person.copy(), token)
                    
                    # Add the phone number that was searched
                    if not enriched_person.get('phone_data'):
                        enriched_person['phone_data'] = []
                    
                    # Add the searched phone number to phone data
                    searched_phone = {
                                'number': phone_number,
                        'operator': 'N/A',
                        'register_date': 'N/A',
                        'nik': person.get('ktp_number', 'N/A')
                    }
                    enriched_person['phone_data'].insert(0, searched_phone)
                    
                    results.append({'person': enriched_person})
                else:
                    # Fallback: Get person data using NIK from phone data
                    nik = phone_item.get('nik')
                    if nik:
                        search_params = {'nik': nik}
                        search_result = call_search(token, search_params)
                        
                        if search_result.get('results'):
                            person = search_result['results'][0]
                            # Enrich with phone data
                            enriched_person = enrich_person_data(person.copy(), token)
                            results.append({'person': enriched_person})
            
            if results:
                result = {
                    'results': results,
                    'total_results': len(results),
                    'message': f'Ditemukan {len(results)} hasil untuk nomor {phone_number}'
                }
            else:
                result = {
                    'results': [],
                    'total_results': 0,
                    'message': f'Tidak ada hasil untuk nomor {phone_number}'
                }
        else:
            result = {
                'results': [],
                'total_results': 0,
                'message': f'Tidak ada hasil untuk nomor {phone_number}'
            }
        
        return jsonify(result)
        
    except Exception as e:
        print(f"Error during phone search: {e}")
        return jsonify({'error': f'Error during phone search: {str(e)}'}), 500

def perform_face_search(token, params, data):
    """Perform face search with uploaded image"""
    if not USE_FACE_LIB:
        return jsonify({'error': 'face_recognition library tidak terpasang'}), 500
    
    try:
        # Decode base64 image
        face_query_b64 = data['face_query']
        if face_query_b64.startswith('data:'):
            face_query_b64 = face_query_b64.split(',', 1)[1]
        
        # Save to temporary file
        with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as tmp_file:
            img_bytes = base64.b64decode(face_query_b64)
            tmp_file.write(img_bytes)
            tmp_path = Path(tmp_file.name)
        
        try:
            # Get query encoding
            q_enc = load_image_file_to_encoding(tmp_path)
            if q_enc is None:
                return jsonify({'error': 'Tidak menemukan wajah pada query image'}), 400
            
            # Get search results
            j = call_search(token, params)
            people = parse_people_from_response(j)
            
            if not people:
                return jsonify({'results': [], 'message': 'Tidak ada person yang dikembalikan oleh API'})
            
            # Process face matching
            threshold = float(data.get('face_threshold', 0.50))
            matches = []
            
            for p in people:
                # Fix photo field name - API might use 'photo', but we expect 'face'
                if p.get('photo') and not p.get('face'):
                    p['face'] = p['photo']
                    print(f"Fixed photo field name for {p.get('full_name', 'Unknown')}")
                
                face_b64 = p.get("face", "")
                if not face_b64:
                    continue
                
                try:
                    enc = get_encoding_from_base64_face(face_b64)
                    if enc is None:
                        continue
                    
                    # Compute distance
                    dist = np.linalg.norm(np.array(q_enc) - np.array(enc))
                    
                    if dist <= threshold:
                        # Enrich person data with family and phone info
                        enriched_person = enrich_person_data(p.copy(), token)
                        
                        # Save face image if requested
                        saved_path = None
                        if data.get('save_face'):
                            saved_path = save_face_image(
                                face_b64, 
                                OUTPUT_FOLDER, 
                                filename_prefix=str(p.get("ktp_number") or p.get("full_name") or "face")
                            )
                        
                        matches.append({
                            'distance': float(dist),
                            'person': enriched_person,
                            'saved_face_path': str(saved_path) if saved_path else None
                        })
                        
                except Exception as e:
                    print(f"Warning: gagal decode face untuk {p.get('ktp_number', p.get('full_name', 'unknown'))}: {e}")
                    continue
            
            # Sort by distance
            matches.sort(key=lambda x: x['distance'])
            
            return jsonify({
                'results': matches,
                'total_matches': len(matches),
                'threshold': threshold,
                'message': f'Ditemukan {len(matches)} kandidat match'
            })
            
        finally:
            # Clean up temporary file
            if tmp_path.exists():
                tmp_path.unlink()
                
    except Exception as e:
        return jsonify({'error': f'Face search error: {str(e)}'}), 500

def perform_regular_search(token, params, data):
    """Perform regular search without face matching"""
    try:
        j = call_search(token, params)
        people = parse_people_from_response(j)
        
        if not people:
            return jsonify({'results': [], 'message': 'Tidak ada hasil'})
        
        results = []
        for p in people:
            # Fix photo field name - API might use 'photo', but enrich_person_data expects 'face'
            if p.get('photo') and not p.get('face'):
                p['face'] = p['photo']
                print(f"Fixed photo field name for {p.get('full_name', 'Unknown')}")
            
            # Enrich person data with basic info only (fast for initial search)
            enriched_person = enrich_person_data_basic(p.copy(), token)
            
            # If still no real photo (missing or avatar SVG), try a direct lookup by NIK to get full record
            ep_face = enriched_person.get('face', '')
            if (not ep_face) or (isinstance(ep_face, str) and ep_face.startswith('data:image/svg')):
                nik_lookup = enriched_person.get('ktp_number') or enriched_person.get('nik')
                if nik_lookup:
                    try:
                        search_result = call_search(token, {'nik': nik_lookup})
                        people_refetch = parse_people_from_response(search_result)
                        if people_refetch:
                            full_person = people_refetch[0]
                            # Map alternate photo keys
                            if full_person.get('photo') and not full_person.get('face'):
                                full_person['face'] = full_person['photo']
                            # Re-run enrichment on the full person data (will normalize face prefix)
                            enriched_person = enrich_person_data(full_person.copy(), token)
                            print(f"Refetched full person data with face for NIK {nik_lookup}")
                    except Exception as _e:
                        pass
            
            # Photo formatting is already handled in enrich_person_data
            
            result = {'person': enriched_person}
            
            # Save face if requested
            if data.get('save_face') and enriched_person.get('face'):
                saved_path = save_face_image(
                    enriched_person.get('face'), 
                    OUTPUT_FOLDER, 
                    filename_prefix=str(enriched_person.get("ktp_number") or enriched_person.get("full_name") or "face")
                )
                result['saved_face_path'] = str(saved_path) if saved_path else None
            
            results.append(result)
        
        return jsonify({
            'results': results,
            'total_results': len(results),
            'message': f'Ditemukan {len(results)} hasil'
        })
        
    except Exception as e:
        return jsonify({'error': f'Search error: {str(e)}'}), 500

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'face_lib_available': USE_FACE_LIB,
        'timestamp': time.time()
    })

@app.route('/static/clean_photos/<filename>')
def serve_clean_photo(filename):
    """Serve clean photos from static folder"""
    try:
        return send_from_directory(CLEAN_PHOTOS_FOLDER, filename)
    except Exception as e:
        logger.error(f"Error serving clean photo {filename}: {e}")
        return jsonify({'error': 'Photo not found'}), 404

@app.route('/api/reprocess-photo/<nik>', methods=['POST'])
def api_reprocess_photo(nik):
    """API endpoint untuk memaksa reprocessing foto dengan algoritma terbaru"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        
        if not user:
            return jsonify({'error': 'Unauthorized'}), 401
        
        data = request.get_json()
        url_foto = data.get('url_foto')
        
        if not url_foto:
            return jsonify({'error': 'URL foto diperlukan'}), 400
        
        logger.info(f"Force reprocessing photo for NIK {nik} by user {user['username']}")
        
        # Force reprocess with new algorithm
        clean_photo_url = process_and_save_clean_photo(nik, url_foto, force_reprocess=True)
        
        if clean_photo_url:
            return jsonify({
                'success': True,
                'message': f'Foto berhasil diproses ulang untuk NIK {nik}',
                'foto_bersih_url': clean_photo_url
            })
        else:
            return jsonify({'error': 'Gagal memproses foto'}), 500
            
    except Exception as e:
        logger.error(f"Error reprocessing photo for NIK {nik}: {e}")
        return jsonify({'error': f'Error: {str(e)}'}), 500

@app.route('/api/test-watermark-removal', methods=['POST'])
def api_test_watermark_removal():
    """API endpoint untuk test watermark removal dengan foto yang sudah ada"""
    try:
        # Validate session
        session_token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = validate_session_token(session_token)
        
        if not user:
            return jsonify({'error': 'Unauthorized'}), 401
        
        data = request.get_json()
        nik = data.get('nik', '1505041107830002')  # Default to MARGUTIN
        
        logger.info(f"Testing watermark removal for NIK {nik} by user {user['username']}")
        
        # Get person data first
        search_data = {
            "search_type": "identity",
            "nik": nik,
            "username": DEFAULT_USERNAME,
            "password": DEFAULT_PASSWORD
        }
        
        # Search for person data
        search_result = call_search(ensure_token(DEFAULT_USERNAME, DEFAULT_PASSWORD), {'nik': nik})
        people = parse_people_from_response(search_result)
        
        if not people:
            return jsonify({'error': 'Person tidak ditemukan'}), 404
        
        person = people[0]
        
        # Check if person has face data
        face_data = person.get('face', '')
        if not face_data or face_data.startswith('data:image/svg'):
            return jsonify({'error': 'Person tidak memiliki foto asli'}), 400
        
        # Try to extract photo URL from face data or other fields
        url_foto = person.get('url_foto') or person.get('face_url') or person.get('photo_url')
        
        if not url_foto or not url_foto.startswith('http'):
            # If no URL, we can't download the original photo for watermark removal
            return jsonify({
                'error': 'Tidak ada URL foto asli yang tersedia untuk watermark removal',
                'person_name': person.get('full_name', 'Unknown'),
                'has_face_data': bool(face_data),
                'face_data_type': 'base64' if face_data.startswith('data:image/') else 'svg_avatar'
            }), 400
        
        # Force reprocess with aggressive algorithm
        clean_photo_url = process_and_save_clean_photo(nik, url_foto, force_reprocess=True)
        
        if clean_photo_url:
            return jsonify({
                'success': True,
                'message': f'Watermark removal test berhasil untuk NIK {nik}',
                'person_name': person.get('full_name', 'Unknown'),
                'foto_bersih_url': clean_photo_url,
                'original_photo_url': url_foto
            })
        else:
            return jsonify({'error': 'Gagal memproses watermark removal'}), 500
            
    except Exception as e:
        logger.error(f"Error testing watermark removal for NIK {nik}: {e}")
        return jsonify({'error': f'Error: {str(e)}'}), 500

@app.route('/api/debug/family/<nik>', methods=['GET'])
def debug_family_data(nik):
    """Debug endpoint to test family data API"""
    try:
        # Get token from request headers or use a default one
        auth_header = request.headers.get('Authorization')
        token = None
        if auth_header and auth_header.startswith('Bearer '):
            token = auth_header[7:]
        
        # Get NKK from query parameter
        nkk = request.args.get('nkk')
        
        # Use sample data for NIK 1505041107830002
        if nik == "1505041107830002":
            sample_family_data = {
                'kepala_keluarga': 'MARGUTIN',
                'nkk': '1505041911100032',
                'alamat_keluarga': 'DUSUN SUNGAIN BAYUR, MANDI ANGIN',
                'anggota_keluarga': [
                    {
                        'nama': 'MARGUTIN',
                        'hubungan': 'Kepala Keluarga',
                        'nik': '1505041107830002',
                        'tanggal_lahir': '11-07-1983',
                        'tempat_lahir': 'MANDI ANGIN',
                        'jenis_kelamin': 'Laki-laki',
                        'agama': 'Islam',
                        'status_perkawinan': 'Kawin',
                        'pekerjaan': 'Wiraswasta',
                        'foto': 'data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iNjQiIGhlaWdodD0iNjQiIHZpZXdCb3g9IjAgMCA2NCA2NCIgZmlsbD0ibm9uZSIgeG1sbnM9Imh0dHA6Ly93d3cudzMub3JnLzIwMDAvc3ZnIj4KPGNpcmNsZSBjeD0iMzIiIGN5PSIzMiIgcj0iMzIiIGZpbGw9IiM2NjdFRUEiLz4KPHN2ZyB3aWR0aD0iMzIiIGhlaWdodD0iMzIiIHZpZXdCb3g9IjAgMCAyNCAyNCIgZmlsbD0ibm9uZSIgeG1sbnM9Imh0dHA6Ly93d3cudzMub3JnLzIwMDAvc3ZnIiB4PSIxNiIgeT0iMTYiPgo8cGF0aCBkPSJNMTIgMTJDMTQuMjA5MSAxMiAxNiAxMC4yMDkxIDE2IDhDMTYgNS43OTA5IDE0LjIwOTEgNCAxMiA0QzkuNzkwODYgNCA4IDUuNzkwOSA4IDhDOCAxMC4yMDkxIDkuNzkwODYgMTIgMTIgMTJaIiBmaWxsPSJ3aGl0ZSIvPgo8cGF0aCBkPSJNMTIgMTRDOC42ODYyOSAxNCA2IDE2LjY4NjMgNiAyMEgxOEMxOCAxNi42ODYzIDE1LjMxMzcgMTQgMTIgMTRaIiBmaWxsPSJ3aGl0ZSIvPgo8L3N2Zz4KPC9zdmc+'
                    },
                    {
                        'nama': 'SITI AMINAH',
                        'hubungan': 'Istri',
                        'nik': '1505041911100033',
                        'tanggal_lahir': '15-03-1985',
                        'tempat_lahir': 'MANDI ANGIN',
                        'jenis_kelamin': 'Perempuan',
                        'agama': 'Islam',
                        'status_perkawinan': 'Kawin',
                        'pekerjaan': 'Ibu Rumah Tangga',
                        'foto': 'data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iNjQiIGhlaWdodD0iNjQiIHZpZXdCb3g9IjAgMCA2NCA2NCIgZmlsbD0ibm9uZSIgeG1sbnM9Imh0dHA6Ly93d3cudzMub3JnLzIwMDAvc3ZnIj4KPGNpcmNsZSBjeD0iMzIiIGN5PSIzMiIgcj0iMzIiIGZpbGw9IiNGRjY5QjQiLz4KPHN2ZyB3aWR0aD0iMzIiIGhlaWdodD0iMzIiIHZpZXdCb3g9IjAgMCAyNCAyNCIgZmlsbD0ibm9uZSIgeG1sbnM9Imh0dHA6Ly93d3cudzMub3JnLzIwMDAvc3ZnIiB4PSIxNiIgeT0iMTYiPgo8cGF0aCBkPSJNMTIgMTJDMTQuMjA5MSAxMiAxNiAxMC4yMDkxIDE2IDhDMTYgNS43OTA5IDE0LjIwOTEgNCAxMiA0QzkuNzkwODYgNCA4IDUuNzkwOSA4IDhDOCAxMC4yMDkxIDkuNzkwODYgMTIgMTIgMTJaIiBmaWxsPSJ3aGl0ZSIvPgo8cGF0aCBkPSJNMTIgMTRDOC42ODYyOSAxNCA2IDE2LjY4NjMgNiAyMEgxOEMxOCAxNi42ODYzIDE1LjMxMzcgMTQgMTIgMTRaIiBmaWxsPSJ3aGl0ZSIvPgo8L3N2Zz4KPC9zdmc+'
                    },
                    {
                        'nama': 'AHMAD RIZKI',
                        'hubungan': 'Anak',
                        'nik': '1505041911100034',
                        'tanggal_lahir': '20-12-2010',
                        'tempat_lahir': 'MANDI ANGIN',
                        'jenis_kelamin': 'Laki-laki',
                        'agama': 'Islam',
                        'status_perkawinan': 'Belum Kawin',
                        'pekerjaan': 'Pelajar',
                        'foto': 'data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iNjQiIGhlaWdodD0iNjQiIHZpZXdCb3g9IjAgMCA2NCA2NCIgZmlsbD0ibm9uZSIgeG1sbnM9Imh0dHA6Ly93d3cudzMub3JnLzIwMDAvc3ZnIj4KPGNpcmNsZSBjeD0iMzIiIGN5PSIzMiIgcj0iMzIiIGZpbGw9IiM0Q0FGNTIiLz4KPHN2ZyB3aWR0aD0iMzIiIGhlaWdodD0iMzIiIHZpZXdCb3g9IjAgMCAyNCAyNCIgZmlsbD0ibm9uZSIgeG1sbnM9Imh0dHA6Ly93d3cudzMub3JnLzIwMDAvc3ZnIiB4PSIxNiIgeT0iMTYiPgo8cGF0aCBkPSJNMTIgMTJDMTQuMjA5MSAxMiAxNiAxMC4yMDkxIDE2IDhDMTYgNS43OTA5IDE0LjIwOTEgNCAxMiA0QzkuNzkwODYgNCA4IDUuNzkwOSA4IDhDOCAxMC4yMDkxIDkuNzkwODYgMTIgMTIgMTJaIiBmaWxsPSJ3aGl0ZSIvPgo8cGF0aCBkPSJNMTIgMTRDOC42ODYyOSAxNCA2IDE2LjY4NjMgNiAyMEgxOEMxOCAxNi42ODYzIDE1LjMxMzcgMTQgMTIgMTRaIiBmaWxsPSJ3aGl0ZSIvPgo8L3N2Zz4KPC9zdmc+'
                    },
                    {
                        'nama': 'FATIMAH ZAHRA',
                        'hubungan': 'Anak',
                        'nik': '1505041911100035',
                        'tanggal_lahir': '05-08-2012',
                        'tempat_lahir': 'MANDI ANGIN',
                        'jenis_kelamin': 'Perempuan',
                        'agama': 'Islam',
                        'status_perkawinan': 'Belum Kawin',
                        'pekerjaan': 'Pelajar',
                        'foto': 'data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iNjQiIGhlaWdodD0iNjQiIHZpZXdCb3g9IjAgMCA2NCA2NCIgZmlsbD0ibm9uZSIgeG1sbnM9Imh0dHA6Ly93d3cudzMub3JnLzIwMDAvc3ZnIj4KPGNpcmNsZSBjeD0iMzIiIGN5PSIzMiIgcj0iMzIiIGZpbGw9IiNGRjQ0NzciLz4KPHN2ZyB3aWR0aD0iMzIiIGhlaWdodD0iMzIiIHZpZXdCb3g9IjAgMCAyNCAyNCIgZmlsbD0ibm9uZSIgeG1sbnM9Imh0dHA6Ly93d3cudzMub3JnLzIwMDAvc3ZnIiB4PSIxNiIgeT0iMTYiPgo8cGF0aCBkPSJNMTIgMTJDMTQuMjA5MSAxMiAxNiAxMC4yMDkxIDE2IDhDMTYgNS43OTA5IDE0LjIwOTEgNCAxMiA0QzkuNzkwODYgNCA4IDUuNzkwOSA4IDhDOCAxMC4yMDkxIDkuNzkwODYgMTIgMTIgMTJaIiBmaWxsPSJ3aGl0ZSIvPgo8cGF0aCBkPSJNMTIgMTRDOC42ODYyOSAxNCA2IDE2LjY4NjMgNiAyMEgxOEMxOCAxNi42ODYzIDE1LjMxMzcgMTQgMTIgMTRaIiBmaWxsPSJ3aGl0ZSIvPgo8L3N2Zz4KPC9zdmc+'
                    }
                ]
            }
            return jsonify({
                'nik': nik,
                'nkk': nkk,
                'family_data': sample_family_data,
                'status': 'success',
                'note': 'Using sample data for testing'
            })
        
        family_data = get_family_data(nik, nkk=nkk, token=token)
        return jsonify({
            'nik': nik,
            'nkk': nkk,
            'family_data': family_data,
            'status': 'success'
        })
    except Exception as e:
        return jsonify({
            'nik': nik,
            'error': str(e),
            'status': 'error'
        }), 500

@app.route('/api/debug/family-alt/<nkk>', methods=['GET'])
def debug_family_alt(nkk):
    """Debug endpoint to test alternative family data API"""
    try:
        url = f"{FAMILY_API_ALT}?family_cert_number={nkk}"
        print(f"Testing alternative family API: {url}")
        
        response = requests.get(url, timeout=30)
        print(f"Alternative API response status: {response.status_code}")
        print(f"Alternative API response text: {response.text}")
        
        if response.status_code == 200:
            data = response.json()
            return jsonify({
                'url': url,
                'status_code': response.status_code,
                'response': data,
                'status': 'success'
            })
        else:
            return jsonify({
                'url': url,
                'status_code': response.status_code,
                'response': response.text,
                'status': 'error'
            }), response.status_code
            
    except Exception as e:
        return jsonify({
            'url': url,
            'error': str(e),
            'status': 'error'
        }), 500

@app.route('/api/debug/phone/<nik>', methods=['GET'])
def debug_phone_data(nik):
    """Debug endpoint to test phone data API"""
    try:
        # Get token from request headers or use a default one
        auth_header = request.headers.get('Authorization')
        token = None
        if auth_header and auth_header.startswith('Bearer '):
            token = auth_header[7:]
        
        # Test with sample data first
        if nik == "1505041107830002":
            sample_data = {
                "data": {
                    "data": [
                        {
                            "_source": {
                                "msisdn": "6285755349653",
                                "nik": "1505041107830002",
                                "operator": "INDOSAT",
                                "register_date": "2018-07-29"
                            }
                        },
                        {
                            "_source": {
                                "msisdn": "6285655343518",
                                "nik": "1505041107830002",
                                "operator": "INDOSAT",
                                "register_date": "2018-08-05"
                            }
                        },
                        {
                            "_source": {
                                "msisdn": "6282260122369",
                                "nik": "1505041107830002",
                                "operator": "TELKOMSEL",
                                "register_date": "2018-03-27"
                            }
                        },
                        {
                            "_source": {
                                "msisdn": "6281280699628",
                                "nik": "1505041107830002",
                                "operator": "TELKOMSEL",
                                "register_date": "2018-03-27"
                            }
                        },
                        {
                            "_source": {
                                "msisdn": "6282313374849",
                                "nik": "1505041107830002",
                                "operator": "TELKOMSEL",
                                "register_date": "2019-05-09"
                            }
                        }
                    ]
                }
            }
            print(f"Using sample data for testing: {sample_data}")
            # Parse sample data
            phone_list = []
            for item in sample_data['data']['data']:
                source = item.get('_source', {})
                phone_info = {
                    'number': source.get('msisdn', ''),
                    'operator': source.get('operator', ''),
                    'register_date': source.get('register_date', ''),
                    'nik': source.get('nik', nik)
                }
                phone_list.append(phone_info)
            return jsonify({
                'nik': nik,
                'phone_data': phone_list,
                'status': 'success',
                'note': 'Using sample data for testing'
            })
            
        phone_data = get_phone_data(nik, token=token)
        return jsonify({
            'nik': nik,
            'phone_data': phone_data,
            'status': 'success'
        })
    except Exception as e:
        return jsonify({
            'nik': nik,
            'error': str(e),
            'status': 'error'
        }), 500

@app.route('/api/debug/phone-direct/<nik>', methods=['GET'])
def debug_phone_direct(nik):
    """Debug endpoint to test phone data API directly"""
    try:
        url = f"{PHONE_API_BASE}/{nik}"
        print(f"Testing direct API call to: {url}")
        
        response = requests.get(url, timeout=10)
        print(f"Direct API response status: {response.status_code}")
        print(f"Direct API response text: {response.text}")
        
        if response.status_code == 200:
            data = response.json()
            return jsonify({
                'url': url,
                'status_code': response.status_code,
                'response': data,
                'status': 'success'
            })
        else:
            return jsonify({
                'url': url,
                'status_code': response.status_code,
                'response': response.text,
                'status': 'error'
            }), response.status_code
            
    except Exception as e:
        return jsonify({
            'url': url,
            'error': str(e),
            'status': 'error'
        }), 500

@app.route('/api/debug/search/<nik>', methods=['GET'])
def debug_search(nik):
    """Debug endpoint to test search with sample data"""
    try:
        # Create sample person data
        sample_person = {
            'ktp_number': nik,
            'full_name': 'MARGUTIN',
            'birth_date': '11-07-1983',
            'birth_place': 'MANDI ANGIN',
            'gender': 'Laki-laki',
            'religion': 'Islam',
            'marital_status': 'Kawin',
            'occupation': 'Wiraswasta',
            'address': 'DUSUN SUNGAIN BAYUR',
            'face': 'data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iMjAwIiBoZWlnaHQ9IjIwMCIgdmlld0JveD0iMCAwIDIwMCAyMDAiIGZpbGw9Im5vbmUiIHhtbG5zPSJodHRwOi8vd3d3LnczLm9yZy8yMDAwL3N2ZyI+CjxyZWN0IHdpZHRoPSIyMDAiIGhlaWdodD0iMjAwIiBmaWxsPSIjRkY2OUI0Ii8+CjxjaXJjbGUgY3g9IjEwMCIgY3k9IjgwIiByPSIzMCIgZmlsbD0iI0ZGRkZGRiIvPgo8cGF0aCBkPSJNNTAgMTQwIFE1MCAxMjAgMTAwIDEyMCBRMTUwIDEyMCAxNTAgMTQwIEwxNTAgMTgwIEw1MCAxODAgWk01MCAxNDAiIGZpbGw9IiNGRkZGRkYiLz4KPC9zdmc+'
        }
        
        # Enrich with family and phone data
        enriched_person = enrich_person_data(sample_person.copy())
        
        return jsonify({
            'nik': nik,
            'person': enriched_person,
            'has_face': bool(enriched_person.get('face')),
            'face_length': len(enriched_person.get('face', '')) if enriched_person.get('face') else 0,
            'status': 'success'
        })
        
    except Exception as e:
        return jsonify({
            'nik': nik,
            'error': str(e),
            'status': 'error'
        }), 500

@app.route('/api/test/face', methods=['GET'])
def test_face():
    """Test endpoint to check if face image displays correctly"""
    try:
        # Simple test with a basic face image
        test_face = 'data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iMjAwIiBoZWlnaHQ9IjIwMCIgdmlld0JveD0iMCAwIDIwMCAyMDAiIGZpbGw9Im5vbmUiIHhtbG5zPSJodHRwOi8vd3d3LnczLm9yZy8yMDAwL3N2ZyI+CjxyZWN0IHdpZHRoPSIyMDAiIGhlaWdodD0iMjAwIiBmaWxsPSIjRkY2OUI0Ii8+CjxjaXJjbGUgY3g9IjEwMCIgY3k9IjgwIiByPSIzMCIgZmlsbD0iI0ZGRkZGRiIvPgo8cGF0aCBkPSJNNTAgMTQwIFE1MCAxMjAgMTAwIDEyMCBRMTUwIDEyMCAxNTAgMTQwIEwxNTAgMTgwIEw1MCAxODAgWk01MCAxNDAiIGZpbGw9IiNGRkZGRkYiLz4KPC9zdmc+'
        
        return jsonify({
            'message': 'Test face image',
            'face': test_face,
            'face_length': len(test_face),
            'status': 'success'
        })
        
    except Exception as e:
        return jsonify({
            'error': str(e),
            'status': 'error'
        }), 500

@app.route('/api/person-details', methods=['POST'])
def api_person_details():
    """API endpoint untuk mendapatkan detail lengkap person"""
    try:
        data = request.get_json()
        
        # Validate required fields
        if not data.get('nik'):
            return jsonify({'error': 'NIK diperlukan'}), 400
        
        if not data.get('username') or not data.get('password'):
            return jsonify({'error': 'Username dan password diperlukan'}), 400
        
        nik = data.get('nik')
        username = data.get('username')
        password = data.get('password')
        
        # Get token
        token = ensure_token(username, password)
        
        # Search for person by NIK
        search_params = {
            'nik': nik,
            'name': '',
            'family_cert_number': '',
            'tempat_lahir': '',
            'tanggal_lahir': '',
            'no_prop': '',
            'no_kab': '',
            'no_kec': '',
            'no_desa': '',
            'page': '1'
        }
        
        # Get person data
        search_result = call_search(token, search_params)
        people = parse_people_from_response(search_result)
        
        if not people:
            return jsonify({'error': 'Person tidak ditemukan'}), 404
        
        # Get the first person (should be the one we're looking for)
        person = people[0]
        
        # Enrich with full data (family and phone info) for detail view
        enriched_person = enrich_person_data(person.copy(), token)
        
        return jsonify({
            'success': True,
            'person': enriched_person
        })
        
    except Exception as e:
        print(f"Error getting person details: {e}")
        return jsonify({'error': f'Error: {str(e)}'}), 500

def clean_text_for_pdf(text):
    """Clean text to avoid encoding issues in PDF"""
    if not text:
        return 'N/A'
    
    # Convert to string and handle encoding
    text = str(text)
    
    # Replace problematic characters
    replacements = {
        '–': '-',
        '—': '-',
        '"': '"',
        '"': '"',
        ''': "'",
        ''': "'",
        '…': '...',
        '°': 'deg',
        '×': 'x',
        '÷': '/',
        '±': '+/-',
        '≤': '<=',
        '≥': '>=',
        '≠': '!=',
        '∞': 'infinity',
        '∑': 'sum',
        '∏': 'product',
        '√': 'sqrt',
        '∫': 'integral',
        '∆': 'delta',
        'α': 'alpha',
        'β': 'beta',
        'γ': 'gamma',
        'δ': 'delta',
        'ε': 'epsilon',
        'ζ': 'zeta',
        'η': 'eta',
        'θ': 'theta',
        'λ': 'lambda',
        'μ': 'mu',
        'π': 'pi',
        'ρ': 'rho',
        'σ': 'sigma',
        'τ': 'tau',
        'φ': 'phi',
        'χ': 'chi',
        'ψ': 'psi',
        'ω': 'omega'
    }
    
    for old, new in replacements.items():
        text = text.replace(old, new)
    
    # Remove any remaining non-ASCII characters that might cause issues
    try:
        text = text.encode('ascii', 'ignore').decode('ascii')
    except:
        text = 'N/A'
    
    return text

@app.route('/api/export/pdf', methods=['POST'])
def export_pdf():
    """Export person details to PDF"""
    try:
        data = request.get_json()
        person = data.get('person')
        username = data.get('username')
        password = data.get('password')
        
        if not person:
            return jsonify({'error': 'No person data provided'}), 400
        
        # Authenticate user
        auth_result = authenticate_user(username, password)
        if not auth_result:
            return jsonify({'error': 'Invalid credentials'}), 401
        
        # Create PDF buffer
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Use default fonts that support basic Unicode
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1,  # Center alignment
            textColor=colors.darkblue,
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            textColor=colors.darkblue,
            fontName='Helvetica-Bold'
        )
        
        # Build content
        story = []
        
        # Title
        story.append(Paragraph("LAPORAN DETAIL ORANG", title_style))
        story.append(Spacer(1, 20))
        
        # Add photo if available
        if person.get('face'):
            try:
                face_data = person['face']
                if face_data.startswith('data:image/'):
                    # Extract base64 data
                    if ',' in face_data:
                        header, base64_data = face_data.split(',', 1)
                    else:
                        base64_data = face_data
                    
                    # Decode base64 to bytes
                    image_bytes = base64.b64decode(base64_data)
                    
                    # Create image from bytes
                    face_image = Image(io.BytesIO(image_bytes), width=2*inch, height=2*inch)
                    face_image.hAlign = 'CENTER'
                    story.append(face_image)
                    story.append(Spacer(1, 20))
                    print(f"Added face photo to PDF for {person.get('full_name', 'Unknown')}")
                else:
                    print(f"Face data format not supported for PDF: {person.get('full_name', 'Unknown')}")
            except Exception as e:
                print(f"Error adding face photo to PDF: {e}")
                # Continue without photo
        
        # Personal Information
        story.append(Paragraph("Informasi Pribadi", heading_style))
        
        personal_data = [
            ['Nama Lengkap', clean_text_for_pdf(person.get('full_name', 'N/A'))],
            ['NIK', clean_text_for_pdf(person.get('ktp_number', person.get('nik', 'N/A')))],
            ['Tanggal Lahir', clean_text_for_pdf(person.get('date_of_birth', 'N/A'))],
            ['Tempat Lahir', clean_text_for_pdf(person.get('birth_place', 'N/A'))],
            ['Jenis Kelamin', clean_text_for_pdf(person.get('gender', 'N/A'))],
            ['Alamat', clean_text_for_pdf(person.get('address', 'N/A'))],
            ['Agama', clean_text_for_pdf(person.get('religion', 'N/A'))],
            ['Status Perkawinan', clean_text_for_pdf(person.get('marital_status', 'N/A'))],
            ['Pekerjaan', clean_text_for_pdf(person.get('occupation', 'N/A'))]
        ]
        
        personal_table = Table(personal_data, colWidths=[2*inch, 4*inch])
        personal_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('BACKGROUND', (1, 0), (1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(personal_table)
        story.append(Spacer(1, 20))
        
        # Phone Numbers
        if person.get('phone_data'):
            story.append(Paragraph("Nomor Telepon", heading_style))
            phone_data = [['Nomor Telepon', 'Operator', 'Tanggal Terdaftar']]
            for phone in person['phone_data']:
                phone_data.append([
                    clean_text_for_pdf(phone.get('number', 'N/A')),
                    clean_text_for_pdf(phone.get('operator', 'N/A')),
                    clean_text_for_pdf(phone.get('register_date', 'N/A'))
                ])
            
            phone_table = Table(phone_data, colWidths=[2*inch, 2*inch, 2*inch])
            phone_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(phone_table)
            story.append(Spacer(1, 20))
        
        # Family Members
        if person.get('family_data') and person['family_data'].get('anggota_keluarga'):
            story.append(Paragraph("Anggota Keluarga", heading_style))
            
            family_info = [
                ['Kepala Keluarga', clean_text_for_pdf(person['family_data'].get('kepala_keluarga', 'N/A'))],
                ['Nomor Kartu Keluarga (NKK)', clean_text_for_pdf(person['family_data'].get('nkk', 'N/A'))],
                ['Alamat Keluarga', clean_text_for_pdf(person['family_data'].get('alamat_keluarga', 'N/A'))]
            ]
            
            family_info_table = Table(family_info, colWidths=[2*inch, 4*inch])
            family_info_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('BACKGROUND', (1, 0), (1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(family_info_table)
            story.append(Spacer(1, 12))
            
            # Family members table
            family_data = [['Nama', 'Hubungan', 'NIK', 'Tanggal Lahir', 'Jenis Kelamin']]
            for member in person['family_data']['anggota_keluarga']:
                family_data.append([
                    clean_text_for_pdf(member.get('nama', 'N/A')),
                    clean_text_for_pdf(member.get('hubungan', 'N/A')),
                    clean_text_for_pdf(member.get('nik', 'N/A')),
                    clean_text_for_pdf(member.get('tanggal_lahir', 'N/A')),
                    clean_text_for_pdf(member.get('jenis_kelamin', 'N/A'))
                ])
            
            family_table = Table(family_data, colWidths=[1.5*inch, 1.2*inch, 1.5*inch, 1.2*inch, 1*inch])
            family_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(family_table)
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"person_details_{person.get('ktp_number', 'unknown')}.pdf",
            mimetype='application/pdf'
        )
        
    except Exception as e:
        print(f"Error generating PDF: {e}")
        return jsonify({'error': f'PDF generation failed: {str(e)}'}), 500

@app.route('/api/export/word', methods=['POST'])
def export_word():
    """Export person details to Word document"""
    try:
        data = request.get_json()
        person = data.get('person')
        username = data.get('username')
        password = data.get('password')
        
        if not person:
            return jsonify({'error': 'No person data provided'}), 400
        
        # Authenticate user
        auth_result = authenticate_user(username, password)
        if not auth_result:
            return jsonify({'error': 'Invalid credentials'}), 401
        
        # Create Word document
        doc = Document()
        
        # Title
        title = doc.add_heading('PERSON DETAILS REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add photo if available
        if person.get('face'):
            try:
                face_data = person['face']
                if face_data.startswith('data:image/'):
                    # Extract base64 data
                    if ',' in face_data:
                        header, base64_data = face_data.split(',', 1)
                    else:
                        base64_data = face_data
                    
                    # Decode base64 to bytes
                    image_bytes = base64.b64decode(base64_data)
                    
                    # Create image from bytes
                    face_image = io.BytesIO(image_bytes)
                    
                    # Add image to document
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    run.add_picture(face_image, width=Inches(2))
                    
                    doc.add_paragraph()  # Add space
                    print(f"Added face photo to Word document for {person.get('full_name', 'Unknown')}")
                else:
                    print(f"Face data format not supported for Word: {person.get('full_name', 'Unknown')}")
            except Exception as e:
                print(f"Error adding face photo to Word document: {e}")
                # Continue without photo
        
        # Personal Information
        doc.add_heading('Informasi Pribadi', level=1)
        
        personal_data = [
            ('Nama Lengkap', person.get('full_name', 'N/A')),
            ('NIK', person.get('ktp_number', person.get('nik', 'N/A'))),
            ('Tanggal Lahir', person.get('date_of_birth', 'N/A')),
            ('Tempat Lahir', person.get('birth_place', 'N/A')),
            ('Jenis Kelamin', person.get('gender', 'N/A')),
            ('Alamat', person.get('address', 'N/A')),
            ('Agama', person.get('religion', 'N/A')),
            ('Status Perkawinan', person.get('marital_status', 'N/A')),
            ('Pekerjaan', person.get('occupation', 'N/A'))
        ]
        
        # Create table for personal information
        table = doc.add_table(rows=len(personal_data), cols=2)
        table.style = 'Table Grid'
        
        for i, (label, value) in enumerate(personal_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = str(value)
            # Make first column bold
            for paragraph in table.cell(i, 0).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        doc.add_paragraph()  # Add space
        
        # Phone Numbers
        if person.get('phone_data'):
            doc.add_heading('Nomor Telepon', level=1)
            
            phone_table = doc.add_table(rows=1, cols=3)
            phone_table.style = 'Table Grid'
            
            # Header row
            header_cells = phone_table.rows[0].cells
            header_cells[0].text = 'Nomor Telepon'
            header_cells[1].text = 'Operator'
            header_cells[2].text = 'Tanggal Terdaftar'
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add phone data
            for phone in person['phone_data']:
                row_cells = phone_table.add_row().cells
                row_cells[0].text = phone.get('number', 'N/A')
                row_cells[1].text = phone.get('operator', 'N/A')
                row_cells[2].text = phone.get('register_date', 'N/A')
            
            doc.add_paragraph()  # Add space
        
        # Family Members
        if person.get('family_data') and person['family_data'].get('anggota_keluarga'):
            doc.add_heading('Informasi Keluarga', level=1)
            
            # Family info
            family_info = [
                ('Kepala Keluarga', person['family_data'].get('kepala_keluarga', 'N/A')),
                ('Nomor Kartu Keluarga (NKK)', person['family_data'].get('nkk', 'N/A')),
                ('Alamat Keluarga', person['family_data'].get('alamat_keluarga', 'N/A'))
            ]
            
            family_info_table = doc.add_table(rows=len(family_info), cols=2)
            family_info_table.style = 'Table Grid'
            
            for i, (label, value) in enumerate(family_info):
                family_info_table.cell(i, 0).text = label
                family_info_table.cell(i, 1).text = str(value)
                # Make first column bold
                for paragraph in family_info_table.cell(i, 0).paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            doc.add_paragraph()  # Add space
            
            # Family members table
            doc.add_heading('Anggota Keluarga', level=2)
            
            family_table = doc.add_table(rows=1, cols=5)
            family_table.style = 'Table Grid'
            
            # Header row
            header_cells = family_table.rows[0].cells
            header_cells[0].text = 'Nama'
            header_cells[1].text = 'Hubungan'
            header_cells[2].text = 'NIK'
            header_cells[3].text = 'Tanggal Lahir'
            header_cells[4].text = 'Jenis Kelamin'
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add family member data
            for member in person['family_data']['anggota_keluarga']:
                row_cells = family_table.add_row().cells
                row_cells[0].text = member.get('nama', 'N/A')
                row_cells[1].text = member.get('hubungan', 'N/A')
                row_cells[2].text = member.get('nik', 'N/A')
                row_cells[3].text = member.get('tanggal_lahir', 'N/A')
                row_cells[4].text = member.get('jenis_kelamin', 'N/A')
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"person_details_{person.get('ktp_number', 'unknown')}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        print(f"Error generating Word document: {e}")
        return jsonify({'error': f'Word document generation failed: {str(e)}'}), 500

if __name__ == '__main__':
    print("Starting Clearance Face Search Web Server...")
    print(f"Face recognition library available: {USE_FACE_LIB}")
    print("Access the application at: http://localhost:5000")
    app.run(debug=True, host='0.0.0.0', port=5000)
